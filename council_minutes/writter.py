
import dateparser
from docx import Document
from docx.shared import RGBColor, Pt
from .models import Request
from .cases.case_utils import header
from docx.enum.style import WD_STYLE_TYPE


class UnifiedWritter():

    def __init__(self):
        self.document = Document()
        self.filename = 'public/'
        styles = self.document.styles
        h2a = styles.add_style('Heading 2 Ancizar', WD_STYLE_TYPE.PARAGRAPH)
        h2a.base_style = styles['Heading 2']
        h2a = styles.add_style('Heading 3 Ancizar', WD_STYLE_TYPE.PARAGRAPH)
        h2a.base_style = styles['Heading 3']
        hls = styles.add_style('List Hyperlink', WD_STYLE_TYPE.PARAGRAPH)
        hls.base_style = self.document.styles['List Bullet']
        for style in styles:
            if style.name != 'No List':
                self.document.styles[style.name].font.name = 'Ancizar Sans'
                self.document.styles[style.name].font.size = Pt(11)
                self.document.styles[style.name].font.color.rgb = RGBColor(
                    0x00, 0x00, 0x00)
        hls.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)
        hls.font.underline = True
        self.case_count = 0

    def generate_case_example_by_id(self, caseid, pre):
        case = Request.get_case_by_id(caseid)
        if case is None:
            raise KeyError
        if pre:
            self.filename += 'pcm' + caseid + '.docx'
            self.__write_case_pcm(case)
        else:
            self.filename += 'cm' + caseid + '.docx'
            self.__write_case_cm(case)
        self.__generate()

    def generate_document_by_querie(self, query, precm):
        cases = Request.get_cases_by_query(query).order_by(
            'academic_program', '_cls')
        casespre = [
            case for case in cases if case.is_pre()]
        casespos = [
            case for case in cases if not case.is_pre()]
        self.__write_case_collection(casespre, True, precm)
        self.__write_case_collection(casespos, False, precm)
        self.__generate()

    def __write_case_cm(self, case):
        case.cm(self.document)

    def __write_case_pcm(self, case):
        header(case, self.document)
        case.pcm(self.document)

    def __write_document_header(self, precm):
        run = self.document.add_paragraph(style='Heading 1').add_run(
            '{}. ASUNTOS ESTUDIANTILES DE {}'.format(
                9 if precm else 10,
                'PREGRADO' if precm else 'POSGRADO'))
        run.font.bold = True
        run.font.size = Pt(11)

    def __write_case_type_header(self, case_type_name):
        run = self.document.add_paragraph(
            style='Heading 2 Ancizar').add_run(case_type_name.upper())
        run.font.bold = True
        run.font.size = Pt(11)

    def __write_academic_program_header(self, academic_program):
        run = self.document.add_paragraph(
            style='Heading 2 Ancizar').add_run(academic_program)
        run.font.bold = True
        run.font.size = Pt(11)

    def __write_case_collection(self, cases, pre, pcm):
        list_level_1 = 9 if pre else 10
        list_level_2 = 0
        list_level_3 = 0
        actual_case = 'dummy'
        actual_academic_program = 'dummy'
        for request in cases:
            if actual_academic_program != request.academic_program:
                list_level_2 = list_level_2 + 1
                actual_academic_program = request.academic_program
                self.__write_academic_program_header(
                    '{}.{} {}'.format(
                        list_level_1,
                        list_level_2,
                        request.get_academic_program_display().upper()))
                list_level_3 = 0
            if actual_case != request.full_name:
                actual_case = request.full_name
                self.__write_case_type_header(request.full_name)
            para = self.document.add_paragraph(style='Heading 3 Ancizar')
            list_level_3 = list_level_3 + 1
            run = para.add_run('{}.{}.{} {}\tDNI. {}'.format(
                list_level_1, list_level_2, list_level_3, request.student_name,
                request.student_dni))
            run.font.bold = True
            run.font.size = Pt(11)
            try:
                if pcm:
                    self.__write_case_pcm(request)
                else:
                    self.__write_case_cm(request)
            except NotImplementedError:
                self.document.add_paragraph()
                self.document.add_paragraph(
                    'Not Implemented case {}'.format(request.full_name))
                self.document.add_paragraph()
            except Exception as err:  # pylint: disable=broad-except
                self.document.add_paragraph()
                self.document.add_paragraph(
                    'Error en el acta {}'.format(request.id))
                self.document.add_paragraph('Trace: {}'.format(err))
                self.document.add_paragraph()

    def __generate(self):
        self.document.save(self.filename)
