import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Pt
from ..models import Request


def add_hyperlink(paragraph_, text, url):
    """
    A function that places a hyperlink within a paragraph object.
    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: A string containing the required url
    :param text: The text displayed for the url
    :return: The hyperlink object
    """

    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph_.part
    r_id = part.relate_to(
        url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    # pylint: disable=protected-access
    paragraph_._p.append(hyperlink)
    # paragraph_.style.font.underline = True

    return hyperlink


def add_analysis_paragraph(docx_, analysis_list, header=True):
    """
     A function that adds the analysis paragraph within a docx object.
     : param docx_: The docx we are adding the analysis to.
     : param analysis_list: A list of the analysis to be added
     """
    paragraph = docx_.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_after = Pt(0)
    if header:
        paragraph.add_run('Analisis: ').font.bold = True
    add_analysis_list(docx_, analysis_list)


def add_analysis_list(docx_, analysis_list):
    """
    A function that adds several analysis within a docx object.
    : param docx_: The docx we are adding the analysis to.
    : param analysis_list: A list of the analysis to be added
    """
    for analysis in analysis_list:
        add_analysis(docx_, analysis)


def add_analysis(docx_, analysis):
    """
    A function that adds an unique analysis within a docx object.
    : param docx_: The docx we are adding the analysis to.
    : param analysis: The analysis to be added
    """
    paragraph = docx_.add_paragraph()
    paragraph.style = 'List Bullet'
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.add_run(analysis)


def string_to_date(string):
    ret = string[8:10]
    ret += num_to_month(string[5:7])
    ret += string[0:4]
    return ret


def get_academic_program(cod_program):
    for p in Request.PLAN_CHOICES:
        if p[0] == cod_program:
            return p[1]


def num_to_month(month):
    if int(month) == 1:
        return ' de enero de '
    elif int(month) == 2:
        return ' de febrero de '
    elif int(month) == 3:
        return ' de marzo de '
    elif int(month) == 4:
        return ' de abril de '
    elif int(month) == 5:
        return ' de mayo de '
    elif int(month) == 6:
        return ' de junio de '
    elif int(month) == 7:
        return ' de julio de '
    elif int(month) == 8:
        return ' de agosto de '
    elif int(month) == 9:
        return ' de septiembre de '
    elif int(month) == 10:
        return ' de octubre de '
    elif int(month) == 11:
        return ' de nomviembre de '
    elif int(month) == 12:
        return ' de diciembre de '


def header(request, docx_):
    para = docx_.add_paragraph()
    para.add_run('Tipo de solicitud:\t{}\n'.format(request.full_name))
    para.add_run('Justificación:\t\t{}\n'.format(
        request.student_justification))
    para.add_run('Soportes:\t\t{}\n'.format(request.supports))
    month = num_to_month(request.date.strftime('%m'))
    para.add_run('Fecha radicación:\t{}\n'.format(
        request.date.strftime("%d{}%Y".format(month))))
    para.add_run('Normatividad:')
    para.paragraph_format.space_after = Pt(0)
    for regulation in request.regulation_list:
        para = docx_.add_paragraph(style='List Hyperlink')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.space_after = Pt(0)
        add_hyperlink(
            para, request.regulations[regulation][0], request.regulations[regulation][1])
    para.paragraph_format.space_after = Pt(0)


def table_general_data(general_data, case, docx_):
    '''
    Add a generated table with general data
        Params:
        1.  general_data (list of list):    A list with the general data of the student,
        in the first field of the sublist goes the name of the information,
        and in the second field goes the value of the information, e.g.:
            [["Nombre Estudiante", "Juan Pérez"], ['DNI', '1111111'],
             ['Plan de estudios', 'Ingeniería de Sistemas'],
             ['Código del plan de estudios', '2879'],
             ['Fecha de la Solicitud', '29 de abril del 2019']]
        2.  case (string):  DOBLE TITULACIÓN or REINGRESO or TRASLADO
        3.  docx_ (docx_):  The document to which the table will be added
    '''
    table = docx_.add_table(rows=len(general_data) + 1,
                            cols=3, style='Table Grid')
    table.style.font.size = Pt(8)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.columns[0].width = 400000
    table.columns[1].width = 2400000
    table.columns[2].width = 2400000
    for cell in table.columns[0].cells:
        cell.width = 400000
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for cell in table.columns[1].cells:
        cell.width = 2400000
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for cell in table.columns[2].cells:
        cell.width = 2400000
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cellp = table.cell(0, 0).merge(table.cell(0, 2)).paragraphs[0]
    cellp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cellp.add_run(case + '\n').font.bold = True
    cellp.runs[0].font.size = Pt(8)
    if case == "DOBLE TITULACIÓN":
        cellp.add_run(
            'Normativa Asociada: Articulo 47 al 50 del Acuerdo 008 '
            'de 2008 del CSU y Acuerdo 155 de 2014 del CSU').font.size = Pt(8)
    elif case == "REINGRESO":
        cellp.add_run(
            'Normativa Asociada: Articulo 46 del Acuerdo 008 de 2008'
            ' del CSU y Resolución 012 de 2014 de VRA').font.size = Pt(8)
    elif case == "TRASLADO":
        cellp.add_run(
            'Normativa Asociada: Articulo 39 del Acuerdo 008 de 2008 '
            'del CSU y Acuerdo 089 de 2014 del C.A.').font.size = Pt(8)
    # pylint: disable=consider-using-enumerate
    for i in range(len(general_data)):
        table.cell(
            i + 1, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(i + 1, 0).paragraphs[0].add_run(str(i + 1)).font.bold = True
        table.cell(i + 1, 0).paragraphs[0].runs[0].font.size = Pt(8)
        for j in range(0, len(general_data[i])):
            table.cell(
                i+1, j+1).paragraphs[0].add_run(general_data[i][j]).font.size = Pt(8)


def table_subjects(docx_, data):
    '''Add a generated table with approvals subjects
        Params:
            docx_ (docx_): The document to which the table will be added
            subjects (list): A list of list with the subjects in table,
            each list must be a list with following data:
            [0]: Subject's SIA code
            [1]: Subject's SIA name
            [2]: Subject's SIA group
            [3]: Subject's SIA tipology
            [4]: Subject's SIA credits
        Raises:
            IndexError: All lists must have same size

    '''
    table = docx_.add_table(rows=len(data)+1, cols=5)
    for column in table.columns:
        for cell in column.cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table.style = 'Table Grid'
    table.style.font.size = Pt(9)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.columns[0].width = 700000
    table.columns[1].width = 2250000
    table.columns[2].width = 600000
    table.columns[3].width = 1050000
    table.columns[4].width = 600000
    for cell in table.columns[0].cells:
        cell.width = 700000
    for cell in table.columns[1].cells:
        cell.width = 2250000
    for cell in table.columns[2].cells:
        cell.width = 600000
    for cell in table.columns[3].cells:
        cell.width = 1050000
    for cell in table.columns[4].cells:
        cell.width = 600000
    table.cell(0, 0).paragraphs[0].add_run('Código').font.bold = True
    table.cell(0, 1).paragraphs[0].add_run('Asignatura').font.bold = True
    table.cell(0, 2).paragraphs[0].add_run('Grupo').font.bold = True
    table.cell(0, 3).paragraphs[0].add_run('Tipología').font.bold = True
    table.cell(0, 4).paragraphs[0].add_run('Créditos').font.bold = True
    for i in range(5):
        table.cell(0, i).paragraphs[0].runs[0].font.size = Pt(8)
    index = 1
    for _ in data:
        table.cell(index, 0).paragraphs[0].add_run(
            data[index-1][0]).font.size = Pt(8)
        table.cell(index, 1).paragraphs[0].add_run(
            data[index-1][1]).font.size = Pt(8)
        table.cell(index, 2).paragraphs[0].add_run(
            data[index-1][2]).font.size = Pt(8)
        table.cell(index, 3).paragraphs[0].add_run(
            data[index-1][3]).font.size = Pt(8)
        table.cell(index, 4).paragraphs[0].add_run(
            data[index-1][4]).font.size = Pt(8)
        index = index + 1


def table_english(docx_, subjects, details):
    '''Add a generated table with approvals subjects

        Params:
            docx_ (docx_): The document to which the table will be added
            subjects (list): A list of list with the subjects in table,
            each list must be a list with following data:
            [0]: Subject's SIA code
            [1]: Subject's SIA name
            [2]: Subject's SIA credits
            [3]: Subject's SIA tipology
            details (list): A list with the datails of homologation,
            must be contains the following data:
            [0]: Exam or institution's name
            [1]: Grade obtained in the institution
            [2]: Student's name
            [3]: Student's DNI
            [4]: Currucular program's name
            [5]: Curricular programs's code

        Raises:
            IndexError: All lists must have same size

    '''
    table = docx_.add_table(rows=len(subjects)+5, cols=7)
    table.style = 'Table Grid'
    table.style.font.size = Pt(8)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.columns[0].width = 600000
    table.columns[1].width = 1800000
    table.columns[2].width = 300000
    table.columns[3].width = 300000
    table.columns[4].width = 400000
    table.columns[5].width = 1400000
    table.columns[6].width = 400000
    for cell in table.columns[0].cells:
        cell.width = 600000
    for cell in table.columns[1].cells:
        cell.width = 1800000
    for cell in table.columns[2].cells:
        cell.width = 300000
    for cell in table.columns[3].cells:
        cell.width = 300000
    for cell in table.columns[4].cells:
        cell.width = 400000
    for cell in table.columns[5].cells:
        cell.width = 1400000
    for cell in table.columns[6].cells:
        cell.width = 400000
    cell = table.cell(0, 0).merge(table.cell(0, 6)).paragraphs[0]
    cell.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    cell.add_run(details[2] + '\t\tDNI. ' + details[3]).font.bold = True
    cell = table.cell(1, 0).merge(table.cell(1, 4)).paragraphs[0]
    str_prog = 'Asignaturas a homologar en el plan de estudios {} ({})'
    cell.add_run(str_prog.format(details[4], details[5])).font.bold = True

    cellp = table.cell(1, 5).merge(table.cell(2, 5)).paragraphs[0]
    cellp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(1, 5).merge(table.cell(2, 5)).paragraphs[0].add_run(
        'Examen de inglés presentado').font.bold = True
    table.cell(1, 5).merge(table.cell(2, 5)
                           ).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    cellp = table.cell(1, 6).merge(table.cell(2, 6)).paragraphs[0]
    cellp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(1, 6).merge(table.cell(2, 6)).paragraphs[0].add_run(
        'Nota').font.bold = True
    table.cell(1, 6).merge(table.cell(2, 6)
                           ).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    cell = table.cell(3, 5).merge(table.cell(
        len(subjects)+2, 5)).paragraphs[0]
    cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.add_run(details[0])
    table.cell(3, 5).merge(table.cell(len(
        subjects)+2, 5)).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    cell = table.cell(3, 6).merge(table.cell(
        len(subjects)+2, 6)).paragraphs[0]
    cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.add_run(details[1])
    table.cell(3, 6).merge(table.cell(len(
        subjects)+2, 6)).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    table.cell(2, 0).paragraphs[0].add_run('Código').font.bold = True
    table.cell(2, 1).paragraphs[0].add_run('Asignatura').font.bold = True
    table.cell(2, 2).paragraphs[0].add_run('C').font.bold = True
    table.cell(2, 3).paragraphs[0].add_run('T').font.bold = True
    table.cell(2, 4).paragraphs[0].add_run('Nota').font.bold = True
    index = 0
    credits_sum = 0
    for subject in subjects:
        credits_sum = credits_sum + int(3)
        table.cell(index+3, 0).paragraphs[0].add_run(subject[0])
        table.cell(index+3, 1).paragraphs[0].add_run(subject[1])
        table.cell(index+3, 2).paragraphs[0].add_run(subject[2])
        table.cell(index+3, 3).paragraphs[0].add_run(subject[3])
        table.cell(index+3, 4).paragraphs[0].add_run(subject[4])
        index = index + 1
    cellp = table.cell(index+3, 0).merge(table.cell(index+3, 3)).paragraphs[0]
    cellp.add_run('Créditos homologados P')
    cellp = table.cell(index+3, 4).merge(table.cell(index+3, 6)).paragraphs[0]
    cellp.add_run(str(credits_sum))
    cellp = table.cell(index+4, 0).merge(table.cell(index+4, 3)).paragraphs[0]
    cellp.add_run('Total créditos que se homologan')
    cellp = table.cell(index+4, 4).merge(table.cell(index+4, 6)).paragraphs[0]
    cellp.add_run(str(credits_sum))


def table_approvals(docx_, subjects, details):
    '''Add a generated table with approvals subjects

    Params:
        docx_ (docx_): The document to which the table will be added
        subjects (list): A list of list with the subjects in table,
        each list must be a list with following data:
        [0]: Subject's period
        [1]: Subject's SIA code
        [2]: Subject's SIA name
        [3]: Subject's SIA credits
        [4]: Subject's SIA tipology
        [5]: Subject's SIA grade
        [6]: Subject's old name
        [7]: Subject's old grade
        details (list): A list with the datails of homologation,
        must be contains the following data:
        [0]: Student's name
        [1]: Student's DNI
        [2]: Student's academica plan code
        [3]: Source institution


    Raises:
        IndexError: All lists must have same size


    '''
    tipology = {}
    periods = []
    for asign in subjects:
        if asign[4] in tipology:
            tipology.update(
                {asign[4]: tipology[asign[4]] + int(asign[3])})
        else:
            tipology.update({asign[4]: int(asign[3])})
        if asign[0] not in periods:
            periods.append(asign[0])
    asign_number = len(subjects)
    tipology_number = len(tipology)
    table = docx_.add_table(
        rows=(3+asign_number), cols=8, style='Table Grid')
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.columns[0].width = 500000
    table.columns[1].width = 550000
    table.columns[2].width = 1350000
    table.columns[3].width = 300000
    table.columns[4].width = 300000
    table.columns[5].width = 400000
    table.columns[6].width = 1400000
    table.columns[7].width = 400000
    for cell in table.columns[0].cells:
        cell.width = 500000
    for cell in table.columns[1].cells:
        cell.width = 550000
    for cell in table.columns[2].cells:
        cell.width = 1350000
    for cell in table.columns[3].cells:
        cell.width = 300000
    for cell in table.columns[4].cells:
        cell.width = 300000
    for cell in table.columns[5].cells:
        cell.width = 400000
    for cell in table.columns[6].cells:
        cell.width = 1400000
    for cell in table.columns[7].cells:
        cell.width = 400000
    cellp = table.cell(0, 0).merge(table.cell(0, 7)).paragraphs[0]
    cellp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cellp.add_run('{}\t\t\tDNI.{}'.format(
        details[0], details[1])).font.bold = True
    cellp.runs[0].font.size = Pt(8)
    cellp = table.cell(1, 0).merge(table.cell(1, 5)).paragraphs[0]
    cellp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cellp.add_run('Asignaturas a homologar en el plan de estudios de {} ({})'.format(
        get_academic_program(details[2]), details[2])).font.bold = True
    cellp.runs[0].font.size = Pt(8)
    table.cell(1, 0).merge(table.cell(1, 5)
                           ).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cellp = table.cell(1, 6).merge(table.cell(1, 7)).paragraphs[0]
    cellp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cellp.add_run('Asignaturas cursadas en {}'.format(
        details[3])).font.bold = True
    cellp.runs[0].font.size = Pt(8)
    for i in range(2, asign_number + 3):
        for j in range(8):
            table.cell(
                i, j).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(2, 0).paragraphs[0].add_run('Periodo').font.bold = True
    table.cell(2, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table.cell(2, 1).paragraphs[0].add_run('Código').font.bold = True
    table.cell(2, 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table.cell(2, 2).paragraphs[0].add_run('Asignatura').font.bold = True
    table.cell(2, 2).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table.cell(2, 3).paragraphs[0].add_run('C').font.bold = True
    table.cell(2, 3).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table.cell(2, 4).paragraphs[0].add_run('T').font.bold = True
    table.cell(2, 4).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table.cell(2, 5).paragraphs[0].add_run('Nota').font.bold = True
    table.cell(2, 5).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table.cell(2, 6).paragraphs[0].add_run('Asignatura').font.bold = True
    table.cell(2, 6).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table.cell(2, 7).paragraphs[0].add_run('Nota').font.bold = True
    table.cell(2, 7).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for i in range(8):
        table.cell(2, i).paragraphs[0].runs[0].font.size = Pt(8)
    subjects.sort(key=lambda s: (s[0], s[2]))
    summary_subjects_right = {}
    summary_subjects_left = {}
    for sbj in subjects:
        if sbj[6] in summary_subjects_right:
            summary_subjects_right[sbj[6]] += [sbj]
        else:
            if sbj[6] == '':
                summary_subjects_right[list(
                    summary_subjects_right.keys())[-1]] += [sbj]
            else:
                summary_subjects_right.update({sbj[6]: [sbj]})
        if sbj[1] in summary_subjects_left:
            summary_subjects_left[sbj[1]] += [sbj]
        else:
            if sbj[1] == '':
                summary_subjects_left[list(
                    summary_subjects_left.keys())[-1]] += [sbj]
            else:
                summary_subjects_left.update({sbj[1]: [sbj]})
        index = 3
    if len(summary_subjects_right) > len(summary_subjects_left):
        for item in summary_subjects_left:
            for k in range(6):
                mg_c = table.cell(index, k).merge(table.cell(
                    index + len(summary_subjects_left[item]) - 1, k)).paragraphs[0]
                table.cell(index, k).merge(table.cell(
                    index + len(summary_subjects_left[
                        item]) - 1, k)).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                mg_c.add_run(
                    str(summary_subjects_left[item][0][k])).font.size = Pt(8)
            index += len(summary_subjects_left[item])
        count = 3
        for subject in subjects:
            table.cell(count, 6).paragraphs[0].add_run(
                subject[6]).font.size = Pt(8)
            table.cell(count, 6).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            table.cell(count, 7).paragraphs[0].add_run(
                subject[7]).font.size = Pt(8)
            table.cell(count, 7).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            count += 1
    elif len(summary_subjects_right) < len(summary_subjects_left):
        if len(summary_subjects_right) == 1:
            for k in (6, 7):
                mg_c = table.cell(index, k).merge(table.cell(
                    index + len(subjects) - 1, k)).paragraphs[0]
                table.cell(index, k).merge(table.cell(
                    index + len(subjects) - 1, k)).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                mg_c.add_run(summary_subjects_right[list(
                    summary_subjects_right.keys())[-1]][0][k]).font.size = Pt(8)
            index += len(subjects)
        else:
            for item in summary_subjects_right:
                for k in (6, 7):
                    mg_c = table.cell(index, k).merge(table.cell(
                        index + len(summary_subjects_right[item]) - 1, k)).paragraphs[0]
                    table.cell(index, k).merge(table.cell(
                        index + len(
                            summary_subjects_right[
                                item]) - 1, k)).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    mg_c.add_run(
                        summary_subjects_right[item][0][k]).font.size = Pt(8)
                index += len(summary_subjects_right[item])
        count = 3
        for subject in subjects:
            table.cell(count, 0).paragraphs[0].add_run(
                subject[0]).font.size = Pt(8)
            table.cell(count, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            table.cell(count, 1).paragraphs[0].add_run(
                subject[1]).font.size = Pt(8)
            table.cell(count, 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            table.cell(count, 2).paragraphs[0].add_run(
                subject[2]).font.size = Pt(8)
            table.cell(count, 2).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            table.cell(count, 3).paragraphs[0].add_run(
                str(subject[3])).font.size = Pt(8)
            table.cell(count, 3).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            table.cell(count, 4).paragraphs[0].add_run(
                subject[4]).font.size = Pt(8)
            table.cell(count, 4).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            table.cell(count, 5).paragraphs[0].add_run(
                subject[5]).font.size = Pt(8)
            table.cell(count, 5).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            count += 1
    else:
        count = 3
        for subject in subjects:
            table.cell(count, 0).paragraphs[0].add_run(
                subject[0]).font.size = Pt(8)
            table.cell(count, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            table.cell(count, 1).paragraphs[0].add_run(
                subject[1]).font.size = Pt(8)
            table.cell(count, 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            table.cell(count, 2).paragraphs[0].add_run(
                subject[2]).font.size = Pt(8)
            table.cell(count, 2).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            table.cell(count, 3).paragraphs[0].add_run(
                str(subject[3])).font.size = Pt(8)
            table.cell(count, 3).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            table.cell(count, 4).paragraphs[0].add_run(
                subject[4]).font.size = Pt(8)
            table.cell(count, 4).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            table.cell(count, 5).paragraphs[0].add_run(
                subject[5]).font.size = Pt(8)
            table.cell(count, 5).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            table.cell(count, 6).paragraphs[0].add_run(
                subject[6]).font.size = Pt(8)
            table.cell(count, 6).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            table.cell(count, 7).paragraphs[0].add_run(
                subject[7]).font.size = Pt(8)
            table.cell(count, 7).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            count += 1
    total_homologated = 0
    table = docx_.add_table(
        rows=(tipology_number+1), cols=8, style='Table Grid')
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.columns[0].width = 500000
    table.columns[1].width = 550000
    table.columns[2].width = 1350000
    table.columns[3].width = 300000
    table.columns[4].width = 300000
    table.columns[5].width = 400000
    table.columns[6].width = 1400000
    table.columns[7].width = 400000
    for cell in table.columns[0].cells:
        cell.width = 500000
    for cell in table.columns[1].cells:
        cell.width = 550000
    for cell in table.columns[2].cells:
        cell.width = 1350000
    for cell in table.columns[3].cells:
        cell.width = 300000
    for cell in table.columns[4].cells:
        cell.width = 300000
    for cell in table.columns[5].cells:
        cell.width = 400000
    for cell in table.columns[6].cells:
        cell.width = 1400000
    for cell in table.columns[7].cells:
        cell.width = 400000
    count = 0
    for tip in tipology:
        text = 'Créditos homologados ' + str(tip)
        table.cell(count, 2).paragraphs[0].add_run(text).font.size = Pt(8)
        table.cell(
            count, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(count, 3).paragraphs[0].add_run(
            str(tipology[tip])).font.size = Pt(8)
        table.cell(
            count, 3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(
            count, 6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        total_homologated += int(tipology[tip])
        count += 1
    table.cell(count, 2).paragraphs[0].add_run(
        'Total créditos que se homologan').font.size = Pt(8)
    table.cell(count, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(count, 3).paragraphs[0].add_run(
        str(total_homologated)).font.size = Pt(8)
    table.cell(count, 3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(0, 0).merge(table.cell(count, 1))
    table.cell(0, 4).merge(table.cell(count, 7))


def table_repprovals(docx_, subjects, details):
    '''Add a generated table with repprovals subjects

    Params:
        docx_ (docx_): The document to which the table will be added
        subjects (list): A list of list with the subjects in table,
        each list must be a list with following data:
        [0]: Subject's period
        [1]: Subject's SIA name
        [2]: Subject's old name
        [3]: Subject's justification
        [4]: Subject's credits
        [5]: Subject's grade
        details (list): A list with the datails of homologation,
        must be contains the following data:
        [0]: Student's name
        [1]: Student's DNI
        [2]: Student's academica plan code
        [3]: Source institution


    Raises:
        IndexError: All lists must have same size


    '''
    asign_number = len(subjects)
    table = docx_.add_table(
        rows=(3+asign_number), cols=6, style='Table Grid')
    table.style.font.size = Pt(8)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.columns[0].width = 500000
    table.columns[1].width = 850000
    table.columns[2].width = 1400000
    table.columns[3].width = 1850000
    table.columns[4].width = 300000
    table.columns[5].width = 300000
    for cell in table.columns[0].cells:
        cell.width = 500000
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for cell in table.columns[1].cells:
        cell.width = 850000
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for cell in table.columns[2].cells:
        cell.width = 1400000
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for cell in table.columns[3].cells:
        cell.width = 1850000
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for cell in table.columns[4].cells:
        cell.width = 300000
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for cell in table.columns[5].cells:
        cell.width = 300000
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cellp = table.cell(0, 0).merge(table.cell(0, 5)).paragraphs[0]
    cellp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cellp.add_run('{}\t\t\tDNI.{}'.format(
        details[0], details[1])).font.bold = True
    cellp.runs[0].font.size = Pt(8)
    cellp = table.cell(1, 0).merge(table.cell(1, 5)).paragraphs[0]
    cellp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cellp.add_run('Asignaturas que no se homologan en el plan de estudios {} ({})'.format(
        get_academic_program(details[2]), details[2])).font.bold = True
    cellp.runs[0].font.size = Pt(8)
    for i in range(2, asign_number + 3):
        for j in range(6):
            table.cell(
                i, j).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(2, 0).paragraphs[0].add_run('Periodo').font.bold = True
    table.cell(2, 0).paragraphs[0].runs[0].font.size = Pt(8)
    table.cell(2, 1).paragraphs[0].add_run(
        'Asignatura Universidad Nacional de Colombia - ({})'.format(details[2])).font.bold = True
    table.cell(2, 1).paragraphs[0].runs[0].font.size = Pt(8)
    table.cell(2, 2).paragraphs[0].add_run(
        'Asignatura cursada en {}'.format(details[3])).font.bold = True
    table.cell(2, 2).paragraphs[0].runs[0].font.size = Pt(8)
    table.cell(2, 3).paragraphs[0].add_run('Justificación').font.bold = True
    table.cell(2, 3).paragraphs[0].runs[0].font.size = Pt(8)
    table.cell(2, 4).paragraphs[0].add_run('C').font.bold = True
    table.cell(2, 4).paragraphs[0].runs[0].font.size = Pt(8)
    table.cell(2, 5).paragraphs[0].add_run('Nota').font.bold = True
    table.cell(2, 5).paragraphs[0].runs[0].font.size = Pt(8)
    count = 3
    for subject in subjects:
        table.cell(count, 0).paragraphs[0].add_run(subject[0])
        table.cell(count, 0).paragraphs[0].runs[0].font.size = Pt(8)
        table.cell(count, 1).paragraphs[0].add_run(subject[1])
        table.cell(count, 1).paragraphs[0].runs[0].font.size = Pt(8)
        table.cell(count, 2).paragraphs[0].add_run(
            subject[2])
        table.cell(count, 2).paragraphs[0].runs[0].font.size = Pt(8)
        table.cell(count, 3).paragraphs[0].add_run(str(subject[3]))
        table.cell(count, 3).paragraphs[0].runs[0].font.size = Pt(8)
        table.cell(count, 4).paragraphs[0].add_run(str(subject[4]))
        table.cell(count, 4).paragraphs[0].runs[0].font.size = Pt(8)
        table.cell(count, 5).paragraphs[0].add_run(subject[5])
        table.cell(count, 5).paragraphs[0].runs[0].font.size = Pt(8)
        count += 1


def table_credits_summary(docx_, credits_, case):
    '''Add a generated table with credits summary
    Params:
        1.  credits: A list of list of 3 rows and 5 columns with the credits in table,
        each list must be a list with following data:
        [0][0]: mandatory required credits of typology B
        [0][1]: optional required credits of typology B
        [0][2]: mandatory required credits of typology C
        [0][3]: optional required credits of typology C
        [0][4]: required credits of typology L
        [1][0]: mandatory equivalent or validated credits  of typology B
        [1][1]: optional equivalent or validated credits of typology B
        [1][2]: mandatory equivalent or validated credits of typology C
        [1][3]: optional equivalent or validated credits of typology C
        [1][4]: equivalent or validated credits of typology L
        [2][0]: mandatory outstanding credits of typology B
        [2][1]: optional outstanding credits of typology B
        [2][2]: mandatory outstanding credits of typology B
        [2][3]: optional outstanding credits of typology B
        [2][4]: outstanding credits of typology B
        2.  case: DOBLE TITULACIÓN or REINGRESO or TRASLADO
        3.  docx_ (docx_): The document to which the table will be added
    Raises:
        IndexError: All lists must have same size
    '''
    table = docx_.add_table(rows=5, cols=7, style='Table Grid')
    table.style.font.size = Pt(8)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for cell in table.columns[0].cells:
        cell.width = 1610000
    for cell in table.columns[1].cells:
        cell.width = 690000
    for cell in table.columns[2].cells:
        cell.width = 610000
    for cell in table.columns[3].cells:
        cell.width = 690000
    for cell in table.columns[4].cells:
        cell.width = 610000
    for cell in table.columns[5].cells:
        cell.width = 675000
    for cell in table.columns[6].cells:
        cell.width = 375000
    table.columns[0].width = 1610000
    table.columns[1].width = 690000
    table.columns[2].width = 610000
    table.columns[3].width = 690000
    table.columns[4].width = 610000
    table.columns[5].width = 675000
    table.columns[6].width = 375000
    for column in table.columns:
        for cell in column.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cellp = table.cell(0, 0).merge(table.cell(1, 0)).paragraphs[0]
    cellp.add_run('Créditos').font.size = Pt(8)
    cellp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(2, 0).paragraphs[0].add_run('Exigidos*').font.size = Pt(8)
    if case == "DOBLE TITULACIÓN":
        table.cell(3, 0).paragraphs[0].add_run(
            'Convalidados/equivalentes**').font.size = Pt(8)
    elif case == 'REINGRESO':
        table.cell(3, 0).paragraphs[0].add_run(
            'Aprobados del plan de estudios').font.size = Pt(8)
    else:
        table.cell(3, 0).paragraphs[0].add_run(
            'Convalidados/equivalentes').font.size = Pt(8)
    table.cell(4, 0).paragraphs[0].add_run('Pendientes').font.size = Pt(8)
    cellp = table.cell(0, 1).merge(table.cell(0, 2)).paragraphs[0]
    cellp.add_run('Fundamentación (B)').font.size = Pt(8)
    cellp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(1, 1).paragraphs[0].add_run('Obligatorios').font.size = Pt(8)
    table.cell(1, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(1, 2).paragraphs[0].add_run('Optativos').font.size = Pt(8)
    table.cell(1, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cellp = table.cell(0, 3).merge(table.cell(0, 4)).paragraphs[0]
    cellp.add_run('Disciplinar (C)').font.size = Pt(8)
    cellp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(1, 3).paragraphs[0].add_run('Obligatorios').font.size = Pt(8)
    table.cell(1, 3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(1, 4).paragraphs[0].add_run('Optativos').font.size = Pt(8)
    table.cell(1, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cellp = table.cell(0, 5).merge(table.cell(1, 5)).paragraphs[0]
    cellp.add_run('Libre Elección (L)').font.size = Pt(8)
    cellp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cellp = table.cell(0, 6).merge(table.cell(1, 6)).paragraphs[0]
    cellp.add_run('Total').font.size = Pt(8)
    cellp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i in range(0, 3):
        suma = 0
        for j in range(0, 5):
            table.cell(
                2+i, j+1).paragraphs[0].add_run(str(credits_[i][j])).font.size = Pt(8)
            table.cell(
                2+i, j+1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            suma += credits_[i][j]
        table.cell(2+i, 6).paragraphs[0].add_run(str(suma)).font.size = Pt(8)
        table.cell(2+i, 6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


def table_recommend(docx_, details):
    '''Add a generated table with approvals subjects
    Params:
        docx_(docx_): The document to which the table will be added
        details(list): A list with the datails of homologation,
        must be contains the following data:
        [0]: Comite's name
        [1]: Comite's date(string)(DD-MM-YYYY)
        [2]: Comite's acta number
        [3]: Comite's acta year
        [4]: Recommend(boolean)
    '''
    table = docx_.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    table.style.font.size = Pt(8)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for cell in table.columns[0].cells:
        cell.width = 3000000
    for cell in table.columns[1].cells:
        cell.width = 800000
    for cell in table.columns[2].cells:
        cell.width = 300000
    for cell in table.columns[3].cells:
        cell.width = 800000
    for cell in table.columns[4].cells:
        cell.width = 300000
    table.columns[0].width = 3000000
    table.columns[1].width = 800000
    table.columns[2].width = 300000
    table.columns[3].width = 800000
    table.columns[4].width = 300000

    table.cell(0, 0).paragraphs[0].add_run(
        'El ' + details[0] + ' en sesión del día ').font.size = Pt(8)
    table.cell(0, 0).paragraphs[0].add_run(
        str(details[1])[0:2] + num_to_month(int(str(details[1])[4:5])) + str(details[1])[6:10]).font.size = Pt(8)
    table.cell(0, 0).paragraphs[0].add_run(
        '. Acta ' + str(details[2]) + ' de ' + str(details[3]) + '.').font.size = Pt(8)
    table.cell(0, 1).paragraphs[0].add_run('Recomienda').font.size = Pt(8)
    table.cell(0, 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(0, 3).paragraphs[0].add_run('No Recomienda').font.size = Pt(8)
    table.cell(0, 3).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table.cell(0, 3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if details[4]:
        table.cell(0, 2).paragraphs[0].add_run('X').font.size = Pt(8)
        table.cell(0, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        table.cell(0, 4).paragraphs[0].add_run('X').font.size = Pt(8)
        table.cell(0, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(0, 2).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table.cell(0, 4).vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def table_change_typology(docx_, subjects):
    '''Add a generated table with approvals subjects
        Params:
            docx_ (docx_): The document to which the table will be added
            subjects (list): A list of list with the subjects in table,
            each list must be a list with following data:
            [0]: Subject's SIA code
            [1]: Subject's SIA name
            [2]: Subject's SIA grade
            [4]: Subject's SIA old component
            [5]: Subject's SIA new component
        Raises:
            IndexError: All lists must have same size
    '''
    table = docx_.add_table(rows=len(subjects)+1, cols=5)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.style = 'Table Grid'
    table.style.font.size = Pt(9)
    table.columns[0].width = 700000
    table.columns[1].width = 2000000
    table.columns[2].width = 600000
    table.columns[3].width = 1050000
    table.columns[4].width = 1050000
    for cell in table.columns[0].cells:
        cell.width = 700000
    for cell in table.columns[1].cells:
        cell.width = 2000000
    for cell in table.columns[2].cells:
        cell.width = 600000
    for cell in table.columns[3].cells:
        cell.width = 1050000
    for cell in table.columns[4].cells:
        cell.width = 1050000
    table.cell(0, 0).paragraphs[0].add_run('Código').font.bold = True
    table.cell(0, 1).paragraphs[0].add_run('Asignatura').font.bold = True
    table.cell(0, 2).paragraphs[0].add_run('Nota').font.bold = True
    table.cell(0, 3).paragraphs[0].add_run(
        'Componente Registrado').font.bold = True
    table.cell(0, 4).paragraphs[0].add_run('Nuevo Componente').font.bold = True
    for i in range(5):
        table.cell(0, i).paragraphs[0].runs[0].font.size = Pt(8)
    index = 0
    for subject in subjects:
        table.cell(
            index+1, 0).paragraphs[0].add_run(subject[0]).font.size = Pt(8)
        table.cell(
            index+1, 1).paragraphs[0].add_run(subject[1]).font.size = Pt(8)
        table.cell(
            index+1, 2).paragraphs[0].add_run(subject[2]).font.size = Pt(8)
        table.cell(
            index+1, 3).paragraphs[0].add_run(subject[3]).font.size = Pt(8)
        table.cell(
            index+1, 4).paragraphs[0].add_run(subject[4]).font.size = Pt(8)
        index = index + 1
