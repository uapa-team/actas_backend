from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from mongoengine import StringField
from ..models import Request
from .case_utils import add_analysis_paragraph


class IATE(Request):

    full_name = 'Informe de avance de tesis'
    decision_maker = Request.decision_makers[2]
    in_cm = False
    in_pcm = False

    subject = StringField(
        required=True, choices=Request.GRADE_OPTION_CHOICES, display='Nombre Asignatura',
        default=Request.GRADE_OPTION_TESIS_MAESTRIA)
    code = StringField(required=True, display='Código Asignatura', default='')

    regulation_list = []

    str_cm = [
        'calificación avance satisfactorio (AS) en la asignatura {} ({}).'
    ]
    str_pcm = []

    def cm(self, docx):
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.str_council_header + ' ')
        self.cm_answer(paragraph)

    def cm_answer(self, paragraph):
        paragraph.add_run(
            # pylint: disable=no-member
            self.get_approval_status_display().upper() + ' ').font.bold = True
        self.add_text(paragraph)

    def pcm(self, docx):
        add_analysis_paragraph(docx, self.extra_analysis)
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.str_answer + ' ').font.bold = True
        paragraph.add_run(self.str_comittee_header + ' ')
        self.pcm_answer(paragraph)

    def pcm_answer(self, paragraph):
        paragraph.add_run(
            # pylint: disable=no-member
            self.get_advisor_response_display().upper() + ' ').font.bold = True
        self.add_text(paragraph)

    def add_text(self, paragraph):
        paragraph.add_run(self.str_cm[0].format(
            # pylint: disable=no-member
            self.get_subject_display(), self.code))

    def resource_analysis(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.pcm_answer(last_paragraph)
    
    def resource_pre_answer(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.pcm_answer(last_paragraph)

    def resource_answer(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.cm_answer(last_paragraph)
