import datetime
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from mongoengine import StringField, DateField, BooleanField
from mongoengine import EmbeddedDocumentListField
from ..models import Request, Professor
from .case_utils import add_analysis_paragraph

class DJCT(Request):

    full_name = 'Designación de jurados calificadores de tesis de maestría ' + \
        'o evaluadores de trabajo final de maestría'

    node = StringField(
        display='Perfil', choices=Request.PROFILE_CHOICES, default=Request.PROFILE_INVE)
    grade_option = StringField(
        required=True, choices=Request.GRADE_OPTION_CHOICES, display='Opción de grado',
        default=Request.GRADE_OPTION_TESIS_MAESTRIA)
    advisor = StringField(required=True, display='Director', default='')
    title = StringField(
        requiered=True, display='Título de Tesis/Trabajo Final', default='')
    date_approval = DateField(required=True, display='Fecha de Aprobación',
                              default=datetime.date.today)
    proposal_jury = BooleanField(
        required=True, display='¿Jurados Propuestos?', default=False)
    proffesors = EmbeddedDocumentListField(Professor, display='Docentes')

    regulation_list = ['040|2017|CFA', '056|2012|CSU']

    str_cm = [
        'designar como {} de {}, cuyo título es ',
        'al(los) profesor(es): '
    ]

    str_pcm = [
        'designar como {} de {}, cuyo título es ',
        'al(los) profesor(es): '
    ]

    names = ['jurado calificador', 'evaluador']

    str_pcm_mag = [
        'SIA: Perfil de {}. El estudiante tiene la asignatura {} ({}).',
        'Concepto motivado acerca del trabajo por parte del director {} (Artículo 43).',
        'Propuesta de tesis aprobada: {}: {}.',
        'Copia impresa y versión electrónica en formato PDF (Artículo 43)',
        'Solicitud de nombramiento de jurados (Artículo 44)',
        'Uno o más evaluadores para los trabajos finales, dos o más jurados para las ' +
        'tesis de Maestría y cuatro jurados para tesis de Doctorado (Artículo 44).'
    ]

    str_pcm_doc = [
        'El estudiante tiene la asignatura {} ({}).',
        'Copia impresa y versión electrónica en formato PDF (Artículo 43)',
        'El proyecto de tesis debe inscribirse y entregarse, antes de alcanzar ' +
        'el 50% de la duración establecida para el programa (Artículo 33)',
        'El documento Proyecto de Tesis de Doctorado será evaluado por un grupo ' +
        'de evaluadores conformado por mínimo tres integrantes, designados por ' +
        'el Comité Asesor de Posgrado. (Artículo 36). Jurados propuestos: {}',
        'El estudiante deberá realizar una sustentación pública de su Proyecto ' +
        'de Tesis de Doctorado ante los evaluadores. En la sustentación deberán ' +
        'participar, presencialmente o mediante video conferencia, el estudiante, ' +
        'los evaluadores, el profesor tutor del estudiante y un profesor activo ' +
        'de la Universidad Nacional de Colombia delegado por el Comité Asesor de ' +
        'Posgrado, quien hará las veces de coordinador de la sustentación (Artículo 37).'
    ]

    def cm(self, docx):
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.str_council_header + ' ')
        self.cm_answer(paragraph)

    def cm_answer(self, paragraph):
        paragraph.add_run(
            # pylint: disable=no-member
            self.get_approval_status_display().upper() + ' ').font.bold = True
        self.add_text(paragraph)

    def pcm(self, docx):
        self.pcm_analysis_handler(docx)
        self.pcm_answer_handler(docx)

    def pcm_answer(self, paragraph):
        paragraph.add_run(
            # pylint: disable=no-member
            self.get_advisor_response_display().upper() + ' ').font.bold = True
        self.add_text(paragraph)

    def add_text(self, paragraph):
        if self.grade_option == Request.GRADE_OPTION_TRABAJO_FINAL_MAESTRIA:
            name = self.names[1]
        else:
            name = self.names[0]
        paragraph.add_run(self.str_cm[0].format(
            # pylint: disable=no-member
            name, self.get_grade_option_display()))
        paragraph.add_run('"{}" '.format(self.title)).font.italic = True
        paragraph.add_run(self.str_pcm[1])
        self.add_proffesors(paragraph)

    def add_proffesors(self, paragraph):
        for i in range(len(self.proffesors)):
            if self.proffesors[i].department not in (self.DP_EMPTY, self.DP_EXTERNO_FACULTAD):
                mod = self.proffesors[i].get_department_display()
            else:
                mod = self.proffesors[i].institution
                if self.proffesors[i].country != '':
                    mod += ' ({})'.format(self.proffesors[i].country)
            end = ', ' if i + 1 < len(self.proffesors) else '.'
            paragraph.add_run(
                '{} - {}{}'.format(self.proffesors[i].name, mod, end))

    def pcm_answer_handler(self, docx):
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.str_answer + ':\n').font.bold = True
        paragraph.add_run(self.str_comittee_header + ' ')
        self.pcm_answer(paragraph)

    def pcm_analysis_handler(self, docx):
        if self.grade_option != Request.GRADE_OPTION_TESIS_DOCTORADO:
            analysis = self.pcm_analysis_magister()
        else:
            analysis = self.pcm_analysis_phd()
        add_analysis_paragraph(docx, analysis + self.extra_analysis)

    def pcm_analysis_magister(self):
        # pylint: disable=no-member
        return [
            self.str_pcm_mag[0].format(
                self.get_node_display(), self.get_grade_option_display(),
                self.get_academic_program_display()),
            self.str_pcm_mag[1].format(self.advisor),
            self.str_pcm_mag[2].format(
                self.date_approval.strftime('%d/%m/%Y '), self.title)
        ] + self.str_pcm_mag[3:]

    def pcm_analysis_phd(self):
        proposed = 'SI' if self.proposal_jury else 'NO'
        return [
            self.str_pcm_doc[0].format(
                # pylint: disable=no-member
                self.get_grade_option_display(), self.get_academic_program_display()),
            self.str_pcm_doc[1],
            self.str_pcm_doc[2],
            self.str_pcm_doc[3].format(proposed),
            self.str_pcm_doc[4]
        ]

    def resource_analysis(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.pcm_answer(last_paragraph)
    
    def resource_pre_answer(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.pcm_answer(last_paragraph)

    def resource_answer(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.cm_answer(last_paragraph)
