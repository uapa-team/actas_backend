import datetime
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from mongoengine import StringField, BooleanField, DateTimeField, IntField
from mongoengine import EmbeddedDocumentListField, FloatField, EmbeddedDocument
from ..models import Request, Subject
from .case_utils import add_analysis_paragraph, table_general_data, string_to_date
from .case_utils import table_credits_summary, table_recommend, indent_table


class TRASPRE(Request):

    class HomologatedSubject(Subject):
        TIP_CHOICES = (
            (Subject.TIP_PRE_FUND_OBLIGATORIA, 'Disciplinar Obligatoria'),
            (Subject.TIP_PRE_FUND_OPTATIVA, 'Disciplinar Optativa'),
            (Subject.TIP_PRE_DISC_OBLIGATORIA, 'Fundamentación Obligatoria'),
            (Subject.TIP_PRE_DISC_OPTATIVA, 'Fundamentación Optativa'),
            (Subject.TIP_PRE_LIBRE_ELECCION, 'Libre Elección'),
        )
        name2 = StringField(required=True, display='Nuevo Nombre Asignatura', default='')
        code2 = StringField(required=True, display='Nuevo Código', default='')
        group = StringField(required=True, display='Agrupación', default='')
        grade = StringField(required=True, display='Nota', default='')
        period = StringField(required=True, display='Periodo',
                choices=Request.PERIOD_CHOICES, default=Request.PERIOD_DEFAULT)

    class PendingSubject(Subject):
        TIP_FUNDAMENTACION = 'B'
        TIP_DISCIPLINAR = 'C'
        TIP_CHOICES = (
            (TIP_FUNDAMENTACION, 'Fundamentación'),
            (TIP_DISCIPLINAR, 'Disciplinar'),
        )
        group = StringField(required=True, display='Agrupación', default='')
        tipology = StringField(
            required=True, choices=TIP_CHOICES, display='Tipología', default=TIP_FUNDAMENTACION)

    class Optative(EmbeddedDocument):
        TIP_FUNDAMENTACION = 'B'
        TIP_DISCIPLINAR = 'C'
        TIP_CHOICES = (
            (TIP_FUNDAMENTACION, 'Fundamentación'),
            (TIP_DISCIPLINAR, 'Disciplinar'),
        )
        group = StringField(required=True, display='Agrupación', default='')
        tipology = StringField(
            required=True, choices=TIP_CHOICES, display='Tipología', default=TIP_FUNDAMENTACION)
        required_creds = IntField(min_value=0, required=True,
                                  display='Créditos requeridos', default=0)
        pending_creds = IntField(min_value=0, required=True,
                                 display='Créditos pendientes', default=0)

    full_name = 'Traslado de programa curricular (Pregrado)'

    TT_INTERCAMPUS = 'TTIC'
    TT_INTERFACULTY = 'TTIF'
    TT_INTRAFACULTY = 'TTRF'
    TT_CHOICES = (
        (TT_INTERCAMPUS, 'Traslado Intersede'),
        (TT_INTERFACULTY, 'Traslado Interfacultad'),
        (TT_INTRAFACULTY, 'Traslado Intrafacultad'),
    )
    TC_BOGOTA = 'BOG'
    TC_MEDELLIN = 'MED'
    TC_MANIZALES = 'MAN'
    TC_PALMIRA = 'PAL'
    TC_LAPAZ = 'PAZ'
    TC_CHOICES = (
        (TC_BOGOTA, 'Bogotá'),
        (TC_MEDELLIN, 'Medellín'),
        (TC_MANIZALES, 'Manizales'),
        (TC_PALMIRA, 'Palmira'),
        (TC_LAPAZ, 'La Paz'),
    )

    offer_regulation = {
        '2541': ('030', 2016),
        '2542': ('027', 2015),
        '2544': ('034', 2016),
        '2545': ('087', 2014),
        '2546': ('024', 2014),
        '2547': ('068', 2018),
        '2548': ('018', 2014),
        '2549': ('002', 2013),
        '2879': ('026', 2014),
    }

    at_least_one_period = BooleanField(
        required=True, default=True,
        display='¿Ha cursado por lo menos un periodo académico del primer plan de estudios?')
    finish_first_plan = BooleanField(
        required=True, default=False,
        display='¿Ha culminado el primer plan de estudios?')
    have_entitled_to_enrrol = BooleanField(
        required=True, default=True,
        display='¿Tiene derecho a renovar matrícula?')
    enroled_number = IntField(
        required=True, min_value=1, default=1,
        display='Número de matrículas')
    currently_studying_double_degree = BooleanField(
        required=True, default=False,
        display='¿Está cursando actualmente doble titulación?')
    available_quota_for_transit = BooleanField(
        required=True, default=True,
        display='El cupo de créditos para traslado es suficiente?')
    availabe_quota_number = IntField(
        required=True, default=0, min_value=0,
        display='Número de cupos ofertados para traslado')
    campus_origin = StringField(
        required=True, default=TC_BOGOTA, choices=TC_CHOICES,
        display='Sede de origen')
    transit_type = StringField(
        required=True, default=TT_INTRAFACULTY, choices=TT_CHOICES,
        display='Tipo de traslado')
    admission_period = StringField(
        required=True, display='Periodo de admisión del estudiante', default='')
    same_degree = BooleanField(
        required=True, default=False,
        display='¿Estos planes de estudios conducen al mismo título?')
    origin_program_code = StringField(
        required=True,
        display='Código del plan de estudios de origen', default='')
    origin_program_name = StringField(
        required=True,
        display='Nombre del plan de estudios de origen', default='')
    enrroled = BooleanField(
        required=True, default=True,
        display='¿Se encuentra matriculado en el semestre de presentar la solicitud?')
    prev_plan = BooleanField(
        required=True, default=False,
        display='¿Tuvo calidad de estudiante en el plan de estudios destino?')
    completion_percentage = FloatField(
        required=True, default=0.0, min_value=0.0, max_value=100.0,
        display='Porcentaje de créditos aprobados en el plan de estudios origen')
    student_admission_score = FloatField(
        default=600.0, min_value=0.0,
        display='Puntaje de admisión del solicitante')
    last_admitted_score = FloatField(
        default=600.0, min_value=0.0,
        display='Puntaje de admisión del último admitido')
    PAPA = FloatField(
        default=3.0, min_value=0.0, max_value=5.0,
        display='P.A.P.A.')
    PAPA_in_threshold = BooleanField(
        default=True,
        display='¿El P.A.P.A. se encuentra dentro de la franja del 30% de los mejores?')
    creds_miunus_remaining = IntField(
        required=True, default=0, min_value=0,
        display='Cupo de créditos menos pendientes en el plan de origen')
    creds_for_transit = IntField(
        required=True, default=0, min_value=0,
        display='Cupo de créditos para traslado')
    advisor_meeting_date = DateTimeField(
        display='Fecha de reunión del comité', default=datetime.date.today)
    council_number_advisor = IntField(
        required=True, default=1, min_value=1,
        display='Número del acta del comité')
    council_year_advisor = IntField(
        required=True, default=2000, min_value=2000,
        display='Año del acta del comité')
    exiged_b_ob = IntField(
        min_value=0, default=0, required=True,
        display='Créditos exigidos fundamentación obligatorios')
    exiged_b_op = IntField(
        min_value=0, default=0, required=True,
        display='Créditos exigidos fundamentación optativos')
    exiged_c_ob = IntField(
        min_value=0, default=0, required=True,
        display='Créditos exigidos disciplinares obligatorios')
    exiged_c_op = IntField(
        min_value=0, default=0, required=True,
        display='Créditos exigidos disciplinares optativos')
    exiged_l = IntField(
        min_value=0, default=0, required=True,
        display='Créditos exigidos libre elección')
    equivalence = EmbeddedDocumentListField(
        HomologatedSubject,
        display="Asignaturas a homologar en el segundo plan de estudios")
    remaining = EmbeddedDocumentListField(
        PendingSubject,
        display="Asignaturas pendientes por cursar en el segundo plan de estudios")
    optative_remaining = EmbeddedDocumentListField(
        Optative,
        display='Agrupaciones de las asignaturas pendientes optativas')

    regulation_list = ['008|2008|CSU', '089|2014|CAC']  # List of regulations

    str_cm = ['traslado {} del programa {} ({}) - Sede {}, al programa {} ({}) - Sede ' +
              '{}, en el periodo académico {}', ', condicionado a conservar la ' +
              'calidad de estudiante al finalizar el periodo académico {}. (Artículo 39 ' +
              'del {} y {}).', 'debido a que', 'Comité Asesor de {}']

    srt_titles = ['I) Datos Generales', 'II) Información Académica']

    list_analysis = ['Viene del plan {} de la sede {}.',
                     'a tenido calidad de estudiante en ese programa previamente ' +
                     '(Parágrafo 1. Artículo 2, {}). Universitas: OK.',
                     'a Culminado el primer plan de estudios.',
                     'iene derecho a renovar la matrícula. Universitas: OK.',
                     'a cursado por lo menos un periodo académico del primer plan ' +
                     'de estudios (Artículo 39, {}). SIA: OK.', 'Ha cursado {} periodos ' +
                     'académicos desde {}.', 'stá cursando doble titulación (Artículo ' +
                     '7. {}). SIA: OK.', 'ay cupos disponibles en el plan de estudios ' +
                     'del programa curricular solicitado (Estipulados por Consejo de ' +
                     'Facultad).', 'El estudiante {}cuenta con el suficiente cupo ' +
                     'de créditos para inscribir las asignaturas pendientes de ' +
                     'aprobación en el nuevo plan (Artículo 3, {}).', 'iene puntaje de ' +
                     'admisión igual o superior al puntaje del útimo admitido regular ' +
                     'al plan de estudios de destino (Artículo 3, {}).', 'e encuentra ' +
                     'dentro de la franja del 30% de los mejores promedios en el plan ' +
                     'de estudios origen.', 'La oferta de asignatura optativas en cada ' +
                     'una de las agrupaciones y componentes del plan de estudios del ' +
                     'programa de {}, la encuentra el acuerdo No. {} del año {}, ' +
                     'expedido por el Consejo de Facultad.']

    str_table = ['Estudiante', 'DNI', 'Plan de estudios de origen (1er plan) - Sede {}',
                 'Código del plan de estudios de origen (1er plan)',
                 'Plan de estudios de destino (2° plan) - Sede {}',
                 'Código del plan de estudios de destino (2° plan)',
                 'Fecha de la solicitud', '¿Estos planes de estudios conducen al mismo título?',
                 'Periodo para el cual fue admitido', '¿El solicitante se encuentra matriculado' +
                 ' en el semestre de presentar la solicitud?', '¿El solicitante tuvo calidad' +
                 ' de estudiante en el plan de estudios de destino (2° plan)?',
                 'Porcentaje de créditos aprobados en el plan de estudios origen (1er plan)',
                 'CUADRO EQUIVALENCIAS Y CONVALIDACIONES DE ASIGNATURAS CURSADAS Y {}APROBADAS' +
                 ' HASTA LA FECHA DE PRESENTACIÓN DE LA SOLICITUD POR PARTE DEL ESTUDIANTE.',
                 'Universidad Nacional de Colombia plan de estudios de {}',
                 'La oferta de asignaturas en cada una de las agrupaciones y componentes del' +
                 ' plan de estudios del programa de {} - perfil {}, la encuentra en el Acuerdo' +
                 ' No. {} del año {}, expedido por Consejo de Facultad de Ingeniería.',
                 '¿Cuál fue el puntaje de admisión del solicitante?', 'Puntaje de admisión del ' +
                 'último admitido regular al plan destino (2° plan) en la misma prueba de ' +
                 'ingreso del solicitante* ', 'P.A.P.A. a la fecha de la solicitud', '¿El ' +
                 'PAPA se encuentra en la franja del 30 % de los mejores promedios en el plan' +
                 ' de estudios origen (1er plan)?', 'Estudio de créditos', 'Cupo de créditos ' +
                 'menos créditos pendientes en el plan de estudios origen (1er plan):', 'Cupo ' +
                 'de créditos para traslado (literal d del artículo 3 del {})', '¿El cupo de ' +
                 'créditos para traslado es igual o mayor al número de créditos pendientes de ' +
                 'aprobación en el plan de estudios destino (2° plan)?', '* en caso que el ' +
                 'plan destino sea de convocatoria anual el puntaje será con la anterior ' +
                 'convocatoria.', 'III) Resumen general de créditos del segundo plan de estudios:',
                 '*Sin incluir los créditos correspondientes al cumplimiento del requisito de ' +
                 'suficiencia en idioma extranjero ', 'PLAN DE ESTUDIOS ({})', 'Periodo', 'Código',
                 'Asignatura', 'T*', 'Agrupación', 'C*', 'Nota', '*T: tipología (C/T/B/O/L). ' +
                 'C*: créditos', 'ASIGNATURAS PENDIENTES POR CURSAR EN EL SEGUNDO PLAN DE ESTUDIOS',
                 'Componente de Fundamentación (B)', 'Obligatorias', 'Créditos Asignatura',
                 'Créditos pendientes por cursar por el estudiante', 'Total créditos pendientes',
                 'Nombre de la Agrupación', 'Créditos requeridos', 'Componente Disciplinar/' +
                 'Profesional (C)', 'Componente de Libre Elección (L) (Créditos Pendientes)',
                 'Optativas', 'La oferta de asignaturas optativas en cada una de las ' +
                 'agrupaciones y componentes del plan de estudios del programa curricular de {}' +
                 ', la encuentra en el Acuerdo No. {} del año {}, expedido por el Consejo de ' +
                 'la Facultad de Ingeniería.', 'Total créditos que se equivalen/convalidan',
                 'Bogotá']

    def cm(self, docx):
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.str_council_header + ' ')
        self.cm_answer(paragraph)
        self.add_tables(docx)

    def cm_answer(self, paragraph):
        # pylint: disable=no-member
        paragraph.add_run(
            self.get_approval_status_display().upper() + ' ').font.bold = True
        paragraph.add_run(
            self.str_cm[0].format(
                self.get_transit_type_display().split(
                    ' ')[1].lower(), self.origin_program_name, self.origin_program_code,
                self.get_campus_origin_display(), self.get_academic_program_display(),
                self.academic_program, self.str_table[48],
                self.get_next_period(self.academic_period)))
        if self.is_affirmative_response_approval_status():
            self.cm_af(paragraph)
        else:
            self.cm_ng(paragraph)

    def pcm(self, docx):
        self.pcm_analysis(docx)
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.str_answer + ': ').font.bold = True
        paragraph.add_run(self.str_comittee_header + ' ')
        self.pcm_answer(paragraph)
        self.add_tables(docx)

    def pcm_answer(self, paragraph):
        # pylint: disable=no-member
        paragraph.add_run(
            self.get_advisor_response_display().upper() + ' ').font.bold = True
        paragraph.add_run(
            self.str_cm[0].format(
                self.get_transit_type_display().split(
                    ' ')[1].lower(), self.origin_program_name, self.origin_program_code,
                self.get_campus_origin_display(), self.get_academic_program_display(),
                self.academic_program, self.str_table[48],
                self.get_next_period(self.academic_period)))
        if self.is_affirmative_response_approval_status():
            self.cm_af(paragraph)
        else:
            self.cm_ng(paragraph)

    def get_next_period(self, actual_period):
        year = int(actual_period[0:4])
        semester = int(actual_period[5])
        if semester == 1:
            return str(year) + '-' + str(semester + 1) + 'S'
        elif semester == 2:
            return str(year + 1) + '-1S'

    def cm_af(self, paragraph):
        paragraph.add_run(self.str_cm[1].format(
            self.academic_period, Request.regulations['008|2008|CSU'][0],
            Request.regulations['089|2014|CAC'][0]))

    def cm_ng(self, paragraph):
        paragraph.add_run(
            ', ' + self.str_cm[2] + ' ' + self.council_decision + '.')

    def pcm_analysis(self, docx):
        # pylint: disable=no-member
        final_analysis = []
        final_analysis += [self.list_analysis[0].format(
            self.origin_program_name, self.get_campus_origin_display())]
        aux_str = 'H' if self.prev_plan else 'No h'
        final_analysis += [aux_str + self.list_analysis[1]
                           .format(Request.regulations['089|2014|CAC'][0])]
        aux_str = 'H' if self.finish_first_plan else 'No h'
        final_analysis += [aux_str + self.list_analysis[2]]
        aux_str = 'T' if self.have_entitled_to_enrrol else 'No t'
        final_analysis += [aux_str + self.list_analysis[3]]
        aux_str = 'H' if self.at_least_one_period else 'No h'
        final_analysis += [aux_str + self.list_analysis[4]
                           .format(Request.regulations['008|2008|CSU'][0])]
        final_analysis += [self.list_analysis[5].format(
            self.enroled_number, self.admission_period)]
        aux_str = 'H' if self.availabe_quota_number > 0 else 'No h'
        final_analysis += [aux_str +
                           self.list_analysis[6].format(self.availabe_quota_number)]
        aux_str = 'H' if self.available_quota_for_transit else 'No h'
        final_analysis += [aux_str + self.list_analysis[7].format(
            Request.regulations['089|2014|CAC'][0])]
        for extra_a in self.extra_analysis:
            final_analysis += [extra_a]
        add_analysis_paragraph(docx, final_analysis)

    def add_tables(self, docx):
        paragraph = docx.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(self.srt_titles[0])
        run.font.bold = True
        run.font.size = Pt(8)
        # pylint: disable=no-member
        general_data = [
            [self.str_table[0], self.student_name],
            [self.str_table[1], self.student_dni],
            [self.str_table[2].format(
                self.get_campus_origin_display()), self.origin_program_name],
            [self.str_table[3], self.origin_program_code],
            [self.str_table[4].format(
                self.str_table[48]), self.get_academic_program_display()],
            [self.str_table[5], self.academic_program],
            [self.str_table[6], string_to_date(str(self.date))],
            [self.str_table[7], 'Sí' if self.same_degree else 'No'],
        ]
        table_general_data(general_data, 'TRASLADO', docx)
        paragraph = docx.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = paragraph.add_run(self.srt_titles[1])
        run.font.bold = True
        run.font.size = Pt(8)
        table = docx.add_table(rows=4, cols=2, style='Table Grid')
        table.alignment = WD_ALIGN_PARAGRAPH.LEFT
        indent_table(table, 963)
        table.columns[0].width = 4350000
        table.columns[1].width = 850000
        for cell in table.columns[0].cells:
            cell.width = 4350000
        for cell in table.columns[1].cells:
            cell.width = 850000
        for i in range(4):
            table.cell(
                i, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(0, 0).paragraphs[0].add_run(
            self.str_table[8]).font.size = Pt(8)
        table.cell(0, 1).paragraphs[0].add_run(
            self.admission_period).font.size = Pt(8)
        table.cell(1, 0).paragraphs[0].add_run(
            self.str_table[9]).font.size = Pt(8)
        table.cell(1, 1).paragraphs[0].add_run(
            'Sí' if self.enrroled else 'No').font.size = Pt(8)
        table.cell(2, 0).paragraphs[0].add_run(
            self.str_table[10]).font.size = Pt(8)
        table.cell(2, 1).paragraphs[0].add_run(
            'Sí' if self.prev_plan else 'No').font.size = Pt(8)
        table.cell(3, 0).paragraphs[0].add_run(
            self.str_table[11]).font.size = Pt(8)
        table.cell(3, 1).paragraphs[0].add_run(
            str(self.completion_percentage) + '%').font.size = Pt(8)
        table = docx.add_table(rows=2, cols=2, style='Table Grid')
        table.alignment = WD_ALIGN_PARAGRAPH.LEFT
        indent_table(table, 963)
        table.columns[0].width = 4350000
        table.columns[1].width = 850000
        for cell in table.columns[0].cells:
            cell.width = 4350000
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for cell in table.columns[1].cells:
            cell.width = 850000
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        creds_study = True
        if self.completion_percentage < 30.0:
            table.cell(0, 0).paragraphs[0].add_run(
                self.str_table[15]).font.size = Pt(8)
            table.cell(0, 1).paragraphs[0].add_run(
                str(self.student_admission_score)).font.size = Pt(8)
            table.cell(1, 0).paragraphs[0].add_run(
                self.str_table[16]).font.size = Pt(8)
            table.cell(1, 1).paragraphs[0].add_run(
                str(self.last_admitted_score)).font.size = Pt(8)
            creds_study = creds_study and self.student_admission_score > self.last_admitted_score
        else:
            table.cell(0, 0).paragraphs[0].add_run(
                self.str_table[17]).font.size = Pt(8)
            table.cell(0, 1).paragraphs[0].add_run(
                str(self.PAPA)).font.size = Pt(8)
            table.cell(1, 0).paragraphs[0].add_run(
                self.str_table[18]).font.size = Pt(8)
            table.cell(1, 1).paragraphs[0].add_run(
                'Sí' if self.PAPA_in_threshold else 'No').font.size = Pt(8)
            table.cell(
                1, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            creds_study = creds_study and self.PAPA_in_threshold
        for i in range(2):
            table.cell(
                i, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if not creds_study:
            paragraph = docx.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.add_run(' ').font.size = Pt(8)
            paragraph.runs[0].font.size = Pt(8)
            details = [self.str_cm[3].format(
                self.get_academic_program_display()), self.advisor_meeting_date.strftime(
                    '%d/%m/%Y '), self.council_number_advisor,
                self.council_year_advisor,
                self.is_affirmative_response_advisor_response()]
            table_recommend(docx, details)
        else:
            paragraph = docx.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(' ').font.size = Pt(8)
            table = docx.add_table(rows=4, cols=3, style='Table Grid')
            table.style.font.size = Pt(8)
            table.alignment = WD_ALIGN_PARAGRAPH.LEFT
            indent_table(table, 963)
            table.columns[0].width = 200000
            table.columns[1].width = 4150000
            table.columns[2].width = 850000
            for cell in table.columns[0].cells:
                cell.width = 200000
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for cell in table.columns[1].cells:
                cell.width = 4150000
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for cell in table.columns[2].cells:
                cell.width = 850000
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            mg_cell = table.cell(0, 0).merge(table.cell(
                0, 2)).paragraphs[0].add_run(self.str_table[19])
            mg_cell.font.bold = True
            mg_cell.font.size = Pt(8)
            table.cell(1, 0).paragraphs[0].add_run('1').font.size = Pt(8)
            table.cell(1, 1).paragraphs[0].add_run(
                self.str_table[20]).font.size = Pt(8)
            table.cell(1, 2).paragraphs[0].add_run(
                str(self.creds_miunus_remaining)).font.size = Pt(8)
            table.cell(2, 0).paragraphs[0].add_run('2').font.size = Pt(8)
            table.cell(2, 1).paragraphs[0].add_run(
                self.str_table[21].format(Request.regulations['089|2014|CAC'][0])).font.size = Pt(8)
            table.cell(2, 2).paragraphs[0].add_run(
                str(self.creds_for_transit)).font.size = Pt(8)
            table.cell(3, 0).paragraphs[0].add_run('3').font.size = Pt(8)
            table.cell(3, 1).paragraphs[0].add_run(
                self.str_table[22]).font.size = Pt(8)
            table.cell(3, 2).paragraphs[0].add_run(
                'Sí' if self.creds_for_transit >= self.creds_miunus_remaining else 'No').font.size = Pt(8)
            for i in range(1, 4):
                table.cell(
                    i, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph = docx.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph.add_run(self.str_table[23]).font.size = Pt(8)
            paragraph = docx.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = paragraph.add_run(self.str_table[24])
            run.font.size = Pt(8)
            run.font.bold = True
            equivalence_creds = Subject.creds_summary(
                self.equivalence)
            pending_creds = [self.exiged_b_ob - equivalence_creds[0], self.exiged_b_op -
                             equivalence_creds[1], self.exiged_c_ob -
                             equivalence_creds[2],
                             self.exiged_c_op - equivalence_creds[3], self.exiged_l -
                             equivalence_creds[4]]
            table_credits_summary(docx, [[self.exiged_b_ob, self.exiged_b_op, self.exiged_c_ob,
                                          self.exiged_c_op, self.exiged_l], equivalence_creds,
                                         pending_creds], 'TRASLADO')
            paragraph = docx.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph.add_run(self.str_table[25]).font.size = Pt(8)
            paragraph = docx.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.add_run(' ').font.size = Pt(8)
            details = [self.str_cm[3].format(
                self.get_academic_program_display()),
                self.advisor_meeting_date.strftime('%d/%m/%Y '),
                self.council_number_advisor, self.council_year_advisor,
                self.is_affirmative_response_advisor_response]
            table_recommend(docx, details)
            paragraph = docx.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.add_run(' ').font.size = Pt(8)
            paragraph = docx.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(self.str_table[12].format(''))
            run.font.underline = True
            run.font.bold = True
            run.font.size = Pt(8)
            paragraph = docx.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.add_run(' ').font.size = Pt(8)
            reproved = 0
            for sbj in self.equivalence:
                try:
                    if float(sbj.grade) < 3.0:
                        reproved += 1
                except ValueError:
                    if sbj.grade == 'RE':
                        reproved += 1
            table = docx.add_table(
                rows=(len(self.equivalence) + 3 - reproved), cols=9, style='Table Grid')
            table.style.font.size = Pt(8)
            table.alignment = WD_ALIGN_PARAGRAPH.LEFT
            indent_table(table, 963)
            table.columns[0].width = 550000
            table.columns[1].width = 550000
            table.columns[2].width = 900000
            table.columns[3].width = 500000
            table.columns[4].width = 900000
            table.columns[5].width = 250000
            table.columns[6].width = 900000
            table.columns[7].width = 250000
            table.columns[8].width = 400000
            for cell in table.columns[0].cells:
                cell.width = 550000
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for cell in table.columns[1].cells:
                cell.width = 550000
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for cell in table.columns[2].cells:
                cell.width = 900000
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for cell in table.columns[3].cells:
                cell.width = 500000
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for cell in table.columns[4].cells:
                cell.width = 900000
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for cell in table.columns[5].cells:
                cell.width = 250000
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for cell in table.columns[6].cells:
                cell.width = 900000
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for cell in table.columns[7].cells:
                cell.width = 250000
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for cell in table.columns[8].cells:
                cell.width = 400000
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cellm = table.cell(0, 0).merge(table.cell(0, 2)).paragraphs[0]
            cellm.add_run(
                self.str_table[26].format('1')).font.bold = True
            cellm.runs[0].font.size = Pt(8)
            cellm.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cellm = table.cell(0, 3).merge(table.cell(0, 8)).paragraphs[0]
            cellm.add_run(
                self.str_table[26].format('2')).font.bold = True
            cellm.runs[0].font.size = Pt(8)
            cellm.alignment = WD_ALIGN_PARAGRAPH.CENTER
            table.cell(1, 0).paragraphs[0].add_run(
                self.str_table[27]).font.bold = True
            table.cell(1, 1).paragraphs[0].add_run(
                self.str_table[28]).font.bold = True
            table.cell(1, 2).paragraphs[0].add_run(
                self.str_table[29]).font.bold = True
            table.cell(1, 3).paragraphs[0].add_run(
                self.str_table[28]).font.bold = True
            table.cell(1, 4).paragraphs[0].add_run(
                self.str_table[29]).font.bold = True
            table.cell(1, 5).paragraphs[0].add_run(
                self.str_table[30]).font.bold = True
            table.cell(1, 6).paragraphs[0].add_run(
                self.str_table[31]).font.bold = True
            table.cell(1, 7).paragraphs[0].add_run(
                self.str_table[32]).font.bold = True
            table.cell(1, 8).paragraphs[0].add_run(
                self.str_table[33]).font.bold = True
            for i in range(9):
                table.cell(
                    1, i).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                table.cell(1, i).paragraphs[0].runs[0].font.size = Pt(8)
            index = 2
            total_creds = 0
            for sbj in self.equivalence:
                try:
                    if float(sbj.grade) < 3.0:
                        continue
                except ValueError:
                    if sbj.grade == 'RE':
                        continue
                table.cell(index, 0).paragraphs[0].add_run(
                    sbj.period).font.size = Pt(8)
                table.cell(index, 1).paragraphs[0].add_run(
                    sbj.code).font.size = Pt(8)
                table.cell(index, 2).paragraphs[0].add_run(
                    sbj.name).font.size = Pt(8)
                table.cell(index, 3).paragraphs[0].add_run(
                    sbj.code2).font.size = Pt(8)
                table.cell(index, 4).paragraphs[0].add_run(
                    sbj.name2).font.size = Pt(8)
                table.cell(index, 5).paragraphs[0].add_run(
                    sbj.tipology[-1]).font.size = Pt(8)
                table.cell(index, 6).paragraphs[0].add_run(
                    sbj.group).font.size = Pt(8)
                table.cell(index, 7).paragraphs[0].add_run(
                    str(sbj.credits)).font.size = Pt(8)
                table.cell(index, 8).paragraphs[0].add_run(
                    str(sbj.grade)).font.size = Pt(8)
                total_creds += sbj.credits
                for i in range(9):
                    table.cell(
                        index, i).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                index += 1
            mg_cll = table.cell(index, 0).merge(table.cell(index, 6)).paragraphs[0].add_run(
                self.str_table[47])
            mg_cll.font.bold = True
            mg_cll.font.size = Pt(8)
            table.cell(index, 0).merge(table.cell(index, 6)
                                       ).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            mg_cll = table.cell(index, 7).merge(table.cell(index, 8)).paragraphs[0].add_run(
                str(total_creds))
            mg_cll.font.bold = True
            mg_cll.font.size = Pt(8)
            table.cell(index, 7).merge(table.cell(index, 8)
                                       ).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if reproved > 0:
                paragraph = docx.add_paragraph()
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.add_run(' ').font.size = Pt(8)
                paragraph = docx.add_paragraph()
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run(self.str_table[12].format('NO '))
                run.font.underline = True
                run.font.bold = True
                run.font.size = Pt(8)
                paragraph = docx.add_paragraph()
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.add_run(' ').font.size = Pt(8)
                table = docx.add_table(
                    rows=(reproved + 3), cols=9, style='Table Grid')
                table.style.font.size = Pt(8)
                table.alignment = WD_ALIGN_PARAGRAPH.LEFT
                indent_table(table, 963)
                table.columns[0].width = 550000
                table.columns[1].width = 550000
                table.columns[2].width = 900000
                table.columns[3].width = 500000
                table.columns[4].width = 900000
                table.columns[5].width = 250000
                table.columns[6].width = 900000
                table.columns[7].width = 250000
                table.columns[8].width = 400000
                for cell in table.columns[0].cells:
                    cell.width = 550000
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for cell in table.columns[1].cells:
                    cell.width = 550000
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for cell in table.columns[2].cells:
                    cell.width = 900000
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for cell in table.columns[3].cells:
                    cell.width = 500000
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for cell in table.columns[4].cells:
                    cell.width = 900000
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for cell in table.columns[5].cells:
                    cell.width = 250000
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for cell in table.columns[6].cells:
                    cell.width = 900000
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for cell in table.columns[7].cells:
                    cell.width = 250000
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for cell in table.columns[8].cells:
                    cell.width = 400000
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                cellm = table.cell(0, 0).merge(table.cell(0, 2)).paragraphs[0]
                cellm.add_run(
                    self.str_table[26].format('1')).font.bold = True
                cellm.runs[0].font.size = Pt(8)
                cellm.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cellm = table.cell(0, 3).merge(table.cell(0, 8)).paragraphs[0]
                cellm.add_run(
                    self.str_table[26].format('2')).font.bold = True
                cellm.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cellm.runs[0].font.size = Pt(8)
                table.cell(1, 0).paragraphs[0].add_run(
                    self.str_table[27]).font.bold = True
                table.cell(1, 1).paragraphs[0].add_run(
                    self.str_table[28]).font.bold = True
                table.cell(1, 2).paragraphs[0].add_run(
                    self.str_table[29]).font.bold = True
                table.cell(1, 3).paragraphs[0].add_run(
                    self.str_table[28]).font.bold = True
                table.cell(1, 4).paragraphs[0].add_run(
                    self.str_table[29]).font.bold = True
                table.cell(1, 5).paragraphs[0].add_run(
                    self.str_table[30]).font.bold = True
                table.cell(1, 6).paragraphs[0].add_run(
                    self.str_table[31]).font.bold = True
                table.cell(1, 7).paragraphs[0].add_run(
                    self.str_table[32]).font.bold = True
                table.cell(1, 8).paragraphs[0].add_run(
                    self.str_table[33]).font.bold = True
                for i in range(9):
                    table.cell(
                        1, i).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    table.cell(1, i).paragraphs[0].runs[0].font.size = Pt(8)
                index = 2
                total_creds = 0
                for sbj in self.equivalence:
                    try:
                        if float(sbj.grade) >= 3.0:
                            continue
                    except ValueError:
                        if sbj.grade in ('AP', 'AS'):
                            continue
                    table.cell(index, 0).paragraphs[0].add_run(
                        sbj.period).font.size = Pt(8)
                    table.cell(index, 1).paragraphs[0].add_run(
                        sbj.code).font.size = Pt(8)
                    table.cell(index, 2).paragraphs[0].add_run(
                        sbj.name).font.size = Pt(8)
                    table.cell(index, 3).paragraphs[0].add_run(
                        sbj.code2).font.size = Pt(8)
                    table.cell(index, 4).paragraphs[0].add_run(
                        sbj.name2).font.size = Pt(8)
                    table.cell(index, 5).paragraphs[0].add_run(
                        sbj.tipology[-1]).font.size = Pt(8)
                    table.cell(index, 6).paragraphs[0].add_run(
                        sbj.group).font.size = Pt(8)
                    table.cell(index, 7).paragraphs[0].add_run(
                        str(sbj.credits)).font.size = Pt(8)
                    table.cell(index, 8).paragraphs[0].add_run(
                        str(sbj.grade)).font.size = Pt(8)
                    total_creds += sbj.credits
                    for i in range(9):
                        table.cell(
                            index, i).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    index += 1
                mg_cll = table.cell(index, 0).merge(table.cell(index, 6)).paragraphs[0].add_run(
                    self.str_table[47])
                mg_cll.font.bold = True
                mg_cll.font.size = Pt(8)
                table.cell(index, 0).merge(table.cell(index, 6)
                                           ).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                mg_cll = table.cell(index, 7).merge(table.cell(index, 8)).paragraphs[0].add_run(
                    str(total_creds))
                mg_cll.font.bold = True
                mg_cll.font.size = Pt(8)
                table.cell(index, 7).merge(table.cell(index, 8)
                                           ).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph = docx.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.add_run(' ').font.size = Pt(8)
            paragraph = docx.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = paragraph.add_run(self.str_table[34])
            run.font.size = Pt(8)
            run.font.underline = True
            run.font.italic = True
            paragraph = docx.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.add_run(' ').font.size = Pt(8)
            paragraph = docx.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(self.str_table[35])
            run.font.size = Pt(8)
            run.font.underline = True
            run.font.bold = True
            paragraph = docx.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.add_run(' ').font.size = Pt(8)
            disciplinare = 0
            for sbj in self.remaining:
                if sbj.tipology == self.PendingSubject.TIP_DISCIPLINAR:
                    disciplinare += 1
            table = docx.add_table(
                rows=(len(self.remaining) - disciplinare + 4), cols=5, style='Table Grid')
            table.style.font.size = Pt(8)
            table.alignment = WD_ALIGN_PARAGRAPH.LEFT
            indent_table(table, 963)
            table.columns[0].width = 1000000
            table.columns[1].width = 600000
            table.columns[2].width = 1700000
            table.columns[3].width = 700000
            table.columns[4].width = 1200000
            for cell in table.columns[0].cells:
                cell.width = 1000000
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for cell in table.columns[1].cells:
                cell.width = 600000
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for cell in table.columns[2].cells:
                cell.width = 1700000
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for cell in table.columns[3].cells:
                cell.width = 700000
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for cell in table.columns[4].cells:
                cell.width = 1200000
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            mg_cll = table.cell(0, 0).merge(table.cell(0, 4)).paragraphs[0].add_run(
                self.str_table[36])
            mg_cll.font.bold = True
            mg_cll.font.size = Pt(8)
            mg_cll = table.cell(1, 0).merge(table.cell(1, 4)).paragraphs[0].add_run(
                self.str_table[37])
            mg_cll.font.bold = True
            mg_cll.font.size = Pt(8)
            table.cell(0, 0).merge(table.cell(
                0, 4)).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            table.cell(1, 0).merge(table.cell(
                1, 4)).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            table.cell(2, 0).paragraphs[0].add_run(
                self.str_table[31]).font.size = Pt(8)
            table.cell(2, 1).paragraphs[0].add_run(
                self.str_table[28]).font.size = Pt(8)
            table.cell(2, 2).paragraphs[0].add_run(
                self.str_table[29]).font.size = Pt(8)
            table.cell(2, 3).paragraphs[0].add_run(
                self.str_table[38]).font.size = Pt(8)
            table.cell(2, 4).paragraphs[0].add_run(
                self.str_table[39]).font.size = Pt(8)
            for i in range(5):
                table.cell(
                    2, i).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            index = 3
            total_creds = 0
            proceced_remaining = {}
            proceced_creds = {}
            for sbj in self.remaining:
                if sbj.tipology == self.PendingSubject.TIP_DISCIPLINAR:
                    continue
                total_creds += sbj.credits
                if sbj.group in proceced_remaining.keys():
                    proceced_remaining[sbj.group].append(sbj)
                    proceced_creds[sbj.group] += sbj.credits
                else:
                    proceced_remaining[sbj.group] = [sbj]
                    proceced_creds[sbj.group] = sbj.credits
            ordered_list = []
            for i in proceced_remaining:
                fc = table.cell(index, 0).merge(table.cell(
                    index + len(proceced_remaining[i]) - 1, 0)).paragraphs[0]
                fc.alignment = WD_ALIGN_PARAGRAPH.CENTER
                fc.add_run(i).font.size = Pt(8)
                sc = table.cell(index, 4).merge(table.cell(
                    index + len(proceced_remaining[i]) - 1, 4)).paragraphs[0]
                sc.alignment = WD_ALIGN_PARAGRAPH.CENTER
                sc.add_run(str(proceced_creds[i])).font.size = Pt(8)
                index += len(proceced_remaining[i])
                ordered_list += proceced_remaining[i]
            index = 3
            for sbj in ordered_list:
                table.cell(index, 1).paragraphs[0].add_run(
                    sbj.code).font.size = Pt(8)
                table.cell(index, 2).paragraphs[0].add_run(
                    sbj.name).font.size = Pt(8)
                table.cell(index, 3).paragraphs[0].add_run(
                    str(sbj.credits)).font.size = Pt(8)
                for i in range(1, 4):
                    table.cell(
                        index, i).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                index += 1
            mg_cll = table.cell(index, 0).merge(
                table.cell(index, 2)).paragraphs[0].add_run(self.str_table[40])
            mg_cll.font.bold = True
            mg_cll.font.size = Pt(8)
            table.cell(index, 0).merge(
                table.cell(index, 2)).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            table.cell(index, 3).merge(
                table.cell(index, 4)).paragraphs[0].add_run(str(total_creds)).font.size = Pt(8)
            table.cell(index, 3).merge(
                table.cell(index, 4)).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            fundam = 0
            for agr in self.optative_remaining:
                if agr.tipology == self.Optative.TIP_DISCIPLINAR:
                    continue
                fundam += 1
            table = docx.add_table(
                rows=(fundam + 3), cols=3, style='Table Grid')
            table.style.font.size = Pt(8)
            table.alignment = WD_ALIGN_PARAGRAPH.LEFT
            indent_table(table, 963)
            table.columns[0].width = 2000000
            table.columns[1].width = 2000000
            table.columns[2].width = 1200000
            for cell in table.columns[0].cells:
                cell.width = 2000000
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for cell in table.columns[1].cells:
                cell.width = 2000000
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for cell in table.columns[2].cells:
                cell.width = 1200000
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell = table.cell(0, 0).merge(table.cell(0, 2)).paragraphs[0]
            cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.add_run(self.str_table[45]).font.bold = True
            cell.runs[0].font.size = Pt(8)
            cell = table.cell(1, 0).paragraphs[0]
            cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.add_run(self.str_table[41]).font.bold = True
            cell.runs[0].font.size = Pt(8)
            cell = table.cell(1, 1).paragraphs[0]
            cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.add_run(self.str_table[39]).font.bold = True
            cell.runs[0].font.size = Pt(8)
            cell = table.cell(1, 2).paragraphs[0]
            cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.add_run(self.str_table[41]).font.bold = True
            cell.runs[0].font.size = Pt(8)
            index = 2
            total = 0
            for agr in self.optative_remaining:
                if agr.tipology == self.Optative.TIP_DISCIPLINAR:
                    continue
                cell = table.cell(index, 0).paragraphs[0]
                cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.add_run(agr.group).font.size = Pt(8)
                cell = table.cell(index, 1).paragraphs[0]
                cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.add_run(str(agr.required_creds)).font.size = Pt(8)
                cell = table.cell(index, 2).paragraphs[0]
                cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.add_run(str(agr.pending_creds)).font.size = Pt(8)
                total += agr.pending_creds
                index += 1
            cell = table.cell(index, 0).merge(
                table.cell(index, 1)).paragraphs[0]
            cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.add_run(self.str_table[40]).font.bold = True
            cell.runs[0].font.size = Pt(8)
            cell = table.cell(index, 2).paragraphs[0]
            cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.add_run(str(total)).font.size = Pt(8)
            paragraph = docx.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.add_run(' ').font.size = Pt(8)
            fundam = 0
            for sbj in self.remaining:
                if sbj.tipology == self.PendingSubject.TIP_FUNDAMENTACION:
                    fundam += 1
            table = docx.add_table(
                rows=(len(self.remaining) - fundam + 4), cols=5, style='Table Grid')
            table.style.font.size = Pt(8)
            table.alignment = WD_ALIGN_PARAGRAPH.LEFT
            indent_table(table, 963)
            table.columns[0].width = 1000000
            table.columns[1].width = 600000
            table.columns[2].width = 1700000
            table.columns[3].width = 700000
            table.columns[4].width = 1200000
            for cell in table.columns[0].cells:
                cell.width = 1000000
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for cell in table.columns[1].cells:
                cell.width = 600000
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for cell in table.columns[2].cells:
                cell.width = 1700000
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for cell in table.columns[3].cells:
                cell.width = 700000
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for cell in table.columns[4].cells:
                cell.width = 1200000
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            mg_cll = table.cell(0, 0).merge(table.cell(0, 4)).paragraphs[0].add_run(
                self.str_table[43])
            mg_cll.font.bold = True
            mg_cll.font.size = Pt(8)
            mg_cll = table.cell(1, 0).merge(table.cell(1, 4)).paragraphs[0].add_run(
                self.str_table[37])
            mg_cll.font.bold = True
            mg_cll.font.size = Pt(8)
            table.cell(0, 0).merge(table.cell(
                0, 4)).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            table.cell(1, 0).merge(table.cell(
                1, 4)).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            table.cell(2, 0).paragraphs[0].add_run(
                self.str_table[31]).font.size = Pt(8)
            table.cell(2, 1).paragraphs[0].add_run(
                self.str_table[28]).font.size = Pt(8)
            table.cell(2, 2).paragraphs[0].add_run(
                self.str_table[29]).font.size = Pt(8)
            table.cell(2, 3).paragraphs[0].add_run(
                self.str_table[38]).font.size = Pt(8)
            table.cell(2, 4).paragraphs[0].add_run(
                self.str_table[39]).font.size = Pt(8)
            for i in range(5):
                table.cell(
                    2, i).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            index = 3
            total_creds = 0
            proceced_remaining = {}
            proceced_creds = {}
            for sbj in self.remaining:
                if sbj.tipology == self.PendingSubject.TIP_FUNDAMENTACION:
                    continue
                total_creds += sbj.credits
                if sbj.group in proceced_remaining.keys():
                    proceced_remaining[sbj.group].append(sbj)
                    proceced_creds[sbj.group] += sbj.credits
                else:
                    proceced_remaining[sbj.group] = [sbj]
                    proceced_creds[sbj.group] = sbj.credits
            ordered_list = []
            for i in proceced_remaining:
                fc = table.cell(index, 0).merge(table.cell(
                    index + len(proceced_remaining[i]) - 1, 0)).paragraphs[0]
                fc.alignment = WD_ALIGN_PARAGRAPH.CENTER
                fc.add_run(i).font.size = Pt(8)
                sc = table.cell(index, 4).merge(table.cell(
                    index + len(proceced_remaining[i]) - 1, 4)).paragraphs[0]
                sc.alignment = WD_ALIGN_PARAGRAPH.CENTER
                sc.add_run(str(proceced_creds[i])).font.size = Pt(8)
                index += len(proceced_remaining[i])
                ordered_list += proceced_remaining[i]
            index = 3
            for sbj in ordered_list:
                table.cell(index, 1).paragraphs[0].add_run(
                    sbj.code).font.size = Pt(8)
                table.cell(index, 2).paragraphs[0].add_run(
                    sbj.name).font.size = Pt(8)
                table.cell(index, 3).paragraphs[0].add_run(
                    str(sbj.credits)).font.size = Pt(8)
                for i in range(1, 4):
                    table.cell(
                        index, i).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                index += 1
            mg_cll = table.cell(index, 0).merge(
                table.cell(index, 2)).paragraphs[0].add_run(self.str_table[40])
            mg_cll.font.bold = True
            mg_cll.font.size = Pt(8)
            table.cell(index, 0).merge(
                table.cell(index, 2)).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            table.cell(index, 3).merge(
                table.cell(index, 4)).paragraphs[0].add_run(str(total_creds)).font.size = Pt(8)
            table.cell(index, 3).merge(
                table.cell(index, 4)).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            disc = 0
            for agr in self.optative_remaining:
                if agr.tipology == self.Optative.TIP_FUNDAMENTACION:
                    continue
                disc += 1
            table = docx.add_table(
                rows=(disc + 3), cols=3, style='Table Grid')
            table.style.font.size = Pt(8)
            table.alignment = WD_ALIGN_PARAGRAPH.LEFT
            indent_table(table, 963)
            table.columns[0].width = 2000000
            table.columns[1].width = 2000000
            table.columns[2].width = 1200000
            for cell in table.columns[0].cells:
                cell.width = 2000000
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for cell in table.columns[1].cells:
                cell.width = 2000000
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for cell in table.columns[2].cells:
                cell.width = 1200000
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell = table.cell(0, 0).merge(table.cell(0, 2)).paragraphs[0]
            cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.add_run(self.str_table[45]).font.bold = True
            cell.runs[0].font.size = Pt(8)
            cell = table.cell(1, 0).paragraphs[0]
            cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.add_run(self.str_table[41]).font.bold = True
            cell.runs[0].font.size = Pt(8)
            cell = table.cell(1, 1).paragraphs[0]
            cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.add_run(self.str_table[39]).font.bold = True
            cell.runs[0].font.size = Pt(8)
            cell = table.cell(1, 2).paragraphs[0]
            cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.add_run(self.str_table[41]).font.bold = True
            cell.runs[0].font.size = Pt(8)
            index = 2
            total = 0
            for agr in self.optative_remaining:
                if agr.tipology == self.Optative.TIP_FUNDAMENTACION:
                    continue
                cell = table.cell(index, 0).paragraphs[0]
                cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.add_run(agr.group).font.size = Pt(8)
                cell = table.cell(index, 1).paragraphs[0]
                cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.add_run(str(agr.required_creds)).font.size = Pt(8)
                cell = table.cell(index, 2).paragraphs[0]
                cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.add_run(str(agr.pending_creds)).font.size = Pt(8)
                total += agr.pending_creds
                index += 1
            cell = table.cell(index, 0).merge(
                table.cell(index, 1)).paragraphs[0]
            cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.add_run(self.str_table[40]).font.bold = True
            cell.runs[0].font.size = Pt(8)
            cell = table.cell(index, 2).paragraphs[0]
            cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.add_run(str(total)).font.size = Pt(8)
            paragraph = docx.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.add_run(' ').font.size = Pt(8)
            table = docx.add_table(rows=1, cols=2, style='Table Grid')
            table.style.font.size = Pt(8)
            table.alignment = WD_ALIGN_PARAGRAPH.LEFT
            indent_table(table, 963)
            table.columns[0].width = 4350000
            table.columns[1].width = 850000
            for cell in table.columns[0].cells:
                cell.width = 4350000
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for cell in table.columns[1].cells:
                cell.width = 850000
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            table.cell(0, 0).paragraphs[0].add_run(
                self.str_table[44]).font.bold = True
            table.cell(0, 0).paragraphs[0].runs[0].font.size = Pt(8)
            table.cell(0, 1).paragraphs[0].add_run(
                str(pending_creds[4])).font.size = Pt(8)
            table.cell(
                0, 1).paragraphs[0].alignment = WD_ALIGN_VERTICAL.CENTER
            paragraph = docx.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.add_run(' ').font.size = Pt(8)
            paragraph = docx.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(self.str_table[46].format(
                self.get_academic_program_display(), *self.offer_regulation[
                    self.academic_program]))
            run.font.size = Pt(8)
            run.font.italic = True
            run.font.underline = True

    def resource_analysis(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.pcm_answer(last_paragraph)

    def resource_pre_answer(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.pcm_answer(last_paragraph)

    def resource_answer(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.cm_answer(last_paragraph)
