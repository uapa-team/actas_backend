from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from mongoengine import StringField
from ..models import Request
from .case_utils import add_analysis_paragraph


class AAUT(Request):

    full_name = 'Admisión automática al posgrado'

    regulation_list = ['008|2008|CSU', '070|2009|CAC']  # List of regulations

    academic_profile = StringField(
        default=Request.PROFILE_INVE, choices=Request.PROFILE_CHOICES,
        display='Perfil de programa curricular')
    places_resolution = StringField(
        default='',
        display='Resolución de cupos para admisión automática.')
    last_academic_program = StringField(
        min_length=4, max_length=4, choices=Request.PLAN_CHOICES,
        required=True, display='Anterior Programa Académico',
        default=Request.PI_CIVIL)
    admission_academic_period = StringField(max_length=10,
                                            display='Periodo de ingreso al nuevo programa',
                                            choices=Request.PERIOD_CHOICES,
                                            default=Request.PERIOD_DEFAULT)

    str_cm = [
        'admisión automática al programa {} ({}), a partir del periodo académico {}.',
        'Debido a que {}justifica debidamente la solicitud.'
    ]

    str_pcm = [
        'El estudiante completó plan de estudios en el plan curricular {} ({}).',
        'Cupo de admisión automática en resolución {}.',
        'Solicita admisión al plan de estudios {} ({}) - perfil de {}.',
        'admisión automática al programa {} ({}) en el plan de estudios de {} a partir del period' +
        'o académico {}.'
    ]

    def cm(self, docx):
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        self.cm_answer(paragraph)
        paragraph.add_run(self.str_cm[1].format(
            '' if self.is_affirmative_response_approval_status() else 'no ') + '. ')
        paragraph.add_run('({}. {}). '.format(
            self.regulations[self.regulation_list[0]][0],
            self.regulations[self.regulation_list[1]][0]))

    def cm_answer(self, paragraph):
        paragraph.add_run(self.str_council_header + ' ')
        paragraph.add_run(
            # pylint: disable=no-member
            self.get_approval_status_display().upper() + ' ').font.bold = True
        paragraph.add_run(self.str_cm[0].format(
            # pylint: disable=no-member
            self.get_academic_program_display(),
            self.academic_program,
            self.academic_period,
            '' if self.is_affirmative_response_approval_status() else 'no ') + ' ')

    def pcm(self, docx):
        self.pcm_analysis(docx)
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.str_answer + ': ').bold = True
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.str_comittee_header + ' ')
        self.pcm_answer(paragraph)

    def pcm_analysis(self, docx):
        analysis_list = []
        analysis_list += [self.str_pcm[0].format(
            # pylint: disable=no-member
            self.get_last_academic_program_display(),
            self.last_academic_program)]
        analysis_list += [self.str_pcm[1].format(self.places_resolution)]
        analysis_list += [self.str_pcm[2].format(
            # pylint: disable=no-member
            self.get_academic_program_display(),
            self.academic_program,
            self.get_academic_profile_display())]
        analysis_list += self.extra_analysis
        add_analysis_paragraph(docx, analysis_list)

    def pcm_answer(self, paragraph):
        paragraph.add_run(
            # pylint: disable=no-member
            self.get_advisor_response_display().upper()).font.bold = True
        paragraph.add_run(' ' + self.str_pcm[3].format(
            # pylint: disable=no-member
            self.get_academic_program_display(),
            self.academic_program,
            self.get_last_academic_program_display(),
            self.last_academic_program,
            self.admission_academic_period
        ))
        paragraph.add_run(' ({}. {}).'.format(
            self.regulations['070|2009|CAC'][0], self.regulations['008|2008|CSU'][0]))
