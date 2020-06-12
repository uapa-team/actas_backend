import datetime
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from mongoengine import IntField, DateField
from ..models import Request
from .case_utils import add_analysis_paragraph, num_to_month


class EBAP(Request):

    full_name = 'Eliminación de la historia académica BAPI'

    commite_cm = IntField(default=1, display='Acta de comité')
    commite_cm_date = DateField(
        display='Fecha acta de comité', default=datetime.date.today)

    regulation_list = ['008|2008|CSU']  # List of regulations

    str_cm = [
        'eliminar la historia académica BAPI, debido a que {}.',
    ]

    str_pcm = [
        'Modalidad de trabajo de grado: Asignaturas de posgrado. Acta de comité {}, del {} de {} ' +
        'del {}.'
    ]

    def cm(self, docx):
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.str_council_header + ' ')
        self.cm_answer(paragraph)
        paragraph.add_run(' ({}). '.format(
            self.regulations[self.regulation_list[0]][0]))

    def cm_answer(self, paragraph):
        paragraph.add_run(
            # pylint: disable=no-member
            self.get_approval_status_display().upper() + ' ').font.bold = True
        paragraph.add_run(self.str_cm[0].format(
            '' if self.is_affirmative_response_approval_status() else 'no ') + '.')

    def pcm(self, docx):
        self.pcm_analysis(docx)
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.str_answer + ': ').bold = True
        paragraph.add_run(self.str_comittee_header + ' ')
        self.pcm_answer(paragraph)

    def pcm_analysis(self, docx):
        analysis_list = []
        analysis_list += [self.str_pcm[0].format(
            # pylint: disable=no-member
            self.commite_cm,
            self.commite_cm_date.day,
            num_to_month(self.commite_cm_date.month),
            self.commite_cm_date.year
        )]
        analysis_list += self.extra_analysis
        add_analysis_paragraph(docx, analysis_list)

    def pcm_answer(self, paragraph):
        paragraph.add_run(
            # pylint: disable=no-member
            self.get_advisor_response_display().upper() + ' ').font.bold = True
        paragraph.add_run(self.str_cm[0].format(self.council_decision))

    def resource_analysis(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.pcm_answer(last_paragraph)
    
    def resource_pre_answer(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.pcm_answer(last_paragraph)

    def resource_answer(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.cm_answer(last_paragraph)
