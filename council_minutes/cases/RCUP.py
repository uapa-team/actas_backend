from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from mongoengine import IntField,StringField
from ..models import Request
from .case_utils import add_analysis_paragraph


class RCUP(Request):

    full_name = 'Reserva de cupo adicional'

    CJT_ANSWER_JUST_DEB = 'JD'
    CJT_ANSWER_OTRO = 'OT'
    CJT_ANSWER_N_DEB = 'ND'

    CJT_ANSWER_CHOICES = (
        (CJT_ANSWER_JUST_DEB,'Justifica debidamente su solicitud'),
        (CJT_ANSWER_N_DEB,'No justifica debidamente su solicitud'),
        (CJT_ANSWER_OTRO, 'Otro')
    )

    CJT_ANSWERS_DICT = {
        CJT_ANSWER_JUST_DEB : 'justifica debidamente su solicitud',
        CJT_ANSWER_N_DEB : 'no justifica debidamente su solicitud',
        CJT_ANSWER_OTRO : 'existen otros factores que lo justifican'
    }

    council_decision = StringField(
        max_length=255, choices=CJT_ANSWER_CHOICES,
        default=CJT_ANSWER_JUST_DEB, display='Justificación del Consejo')
    index = IntField(min_value=0, default=1,
                     display='Reservas adicionales aprobadas')

    regulation_list = ['008|2008|CSU']

    str_cm = [
        'reserva de cupo adicional en el periodo académico {}, debido a que {}.',
        'justifica debidamente la solicitud',
        '(Artículo 20 del {}).'
    ]

    str_pcm_aff = [
        'reserva de cupo adicional en el periodo académico {}, debido a que ' +
        'justifica debidamente la solicitud. (Artículo 20 del {}).'
    ]

    str_pcm_neg = [
        'reserva de cupo adicional en el periodo académico {}, teniendo ' +
        'en cuenta que esta posibilidad es viable a continuación de la segunda reserva ' +
        'de cupo automática. (Artículo 20 del {}).'
    ]

    analysis = [
        'El comité de {} considera que la situación personal está debidamente justificada.',
        'Se le han aprobado {} reservas de cupo adicionales.'
    ]

    def cm(self, docx):
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.str_council_header + ' ')
        self.cm_answer(paragraph)

    def cm_answer(self, paragraph):
        # pylint: disable=no-member
        paragraph.add_run(
            self.get_approval_status_display().upper() + ' ').font.bold = True
        paragraph.add_run(self.str_cm[0].format(
            self.academic_period, self.CJT_ANSWERS_DICT[self.council_decision]))
        paragraph.add_run(self.str_cm[2].format(
            self.regulations[self.regulation_list[0]][0]))

    def pcm(self, docx):
        add_analysis_paragraph(docx, self.add_analysis())
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.str_answer + ' ').font.bold = True
        paragraph.add_run(self.str_comittee_header + ' ')
        self.pcm_answer(paragraph)

    def pcm_answer(self, paragraph):
        paragraph.add_run(
            # pylint: disable=no-member
            self.get_advisor_response_display().upper() + ' ').font.bold = True

        if self.is_affirmative_response_advisor_response():
            value = self.str_pcm_aff[0]
        else:
            value = self.str_pcm_neg[0]
        paragraph.add_run(value.format(
            self.academic_period, self.regulations[self.regulation_list[0]][0]))

    def add_analysis(self):
        # pylint: disable=no-member
        modifier = self.get_academic_program_display()
        if not self.is_affirmative_response_advisor_response():
            modifier += ' no'
        return [
            self.analysis[0].format(modifier),
            self.analysis[1].format(self.index)
        ] + self.extra_analysis

    def resource_analysis(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.pcm_answer(last_paragraph)
    
    def resource_pre_answer(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.pcm_answer(last_paragraph)

    def resource_answer(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.cm_answer(last_paragraph)
