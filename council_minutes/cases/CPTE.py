from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from mongoengine import StringField, BooleanField
from ..models import Request
from .case_utils import add_analysis_paragraph


class CPTE(Request):

    full_name = 'Cambio de proyecto de tesis de maestría/doctorado o ' + \
        'propuesta de trabajo final de maestría'

    GO_TRABAJO_FINAL_MAESTRIA = 'TFM'
    GO_TESIS_MAESTRIA = 'TSM'
    GO_TESIS_DOCTORADO = 'TSD'
    GO_CHOICES = (
        (GO_TRABAJO_FINAL_MAESTRIA, 'Trabajo Final de Maestría'),
        (GO_TESIS_MAESTRIA, 'Tesis de Maestría'),
        (GO_TESIS_DOCTORADO, 'Tesis de Doctorado')
    )

    title = StringField(
        required=True, display='Nuevo título de la tesis/trabajo final', default='')
    grade_option = StringField(
        required=True, display='Tipo de tesis/trabajo final',
        choices=GO_CHOICES, default=GO_TESIS_MAESTRIA)
    new_advisor = StringField(
        required=True, display='Nuevo director de tesis/trabajo final', default='')
    old_advisor = StringField(
        display='Antiguo director de tesis/trabajo final', default='')
    new_co_advisor = StringField(
        display='Nuevo codirector de tesis/trabajo final', default='')
    old_co_advisor = StringField(
        display='Antiguo codirector de tesis/trabajo final', default='')
    inst_new_advisor = StringField(choices=Request.DP_CHOICES,
                                   display='Departamento de adscripción del nuevo director',
                                   default=Request.DP_EMPTY)
    inst_new_co_advisor = StringField(choices=Request.DP_CHOICES,
                                      display='Departamento de adscripción del nuevo codirector',
                                      default=Request.DP_EMPTY)
    inst_old_co_advisor = StringField(choices=Request.DP_CHOICES,
                                      display='Departamento de adscripción del antiguo codirector',
                                      default=Request.DP_EMPTY)
    inst_old_advisor = StringField(choices=Request.DP_CHOICES,
                                   display='Departamento de adscripción del antiguo director',
                                   default=Request.DP_EMPTY)
    enrolled_thesis = BooleanField(required=True, default=False,
                                   display='¿Tiene inscrita la asignatura tesis/trabajo final?')
    have_signature = BooleanField(required=True, default=False,
                                  display='¿Tiene la firma del (los) director(es)?')

    regulation_list = ['002|2011|CFA', '056|2012|CSU']  # List of regulations

    str_cm = ['cambiar el título de {} del programa {} a: ',
              '"{}"',
              'ratificar director',
              'designar nuevo director',
              'al profesor',
              'en reemplazo del profesor',
              'Designar nuevo codirector',
              'ratificar director',
              ' del ',
              'debido a que',
              'Ratificar nuevo codirector']

    list_analysis = ['Perfil de {}.',
                     'El estudiante {}tiene inscrita la asignatura {}.',
                     'iene la firma de los directores de tesis/trabajo final: {}']

    def cm(self, docx):
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.str_council_header + ' ')
        self.cm_answer(paragraph)

    def cm_answer(self, paragraph):
        paragraph.add_run(
            # pylint: disable=no-member
            self.get_approval_status_display().upper() + ' ').font.bold = True
        # pylint: disable=no-member
        paragraph.add_run(self.str_cm[0].format(
            self.get_grade_option_display(), self.get_academic_program_display()))
        paragraph.add_run(self.str_cm[1].format(self.title)).font.italic = True
        if self.is_affirmative_response_approval_status():
            self.cm_af(paragraph)
        else:
            self.cm_ng(paragraph)

    def pcm(self, docx):
        self.pcm_analysis(docx)
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.str_answer + ': ').font.bold = True
        paragraph.add_run(self.str_comittee_header + ' ')
        self.pcm_answer(paragraph)

    def pcm_answer(self, paragraph):
        paragraph.add_run(
            # pylint: disable=no-member
            self.get_advisor_response_display().upper()).font.bold = True
        # pylint: disable=no-member
        paragraph.add_run(' ' + self.str_cm[0].format(
            self.get_grade_option_display(), self.get_academic_program_display()))
        paragraph.add_run(self.str_cm[1].format(self.title)).font.italic = True
        if self.is_affirmative_response_approval_status():
            self.cm_af(paragraph)
        else:
            self.cm_ng(paragraph)

    def cm_af(self, paragraph):
        if self.old_advisor == self.new_advisor or self.old_advisor == '':
            paragraph.add_run(
                ', ' + self.str_cm[2] + ' ' + self.str_cm[4] + ' ')
            paragraph.add_run(self.new_advisor)
            if self.inst_new_advisor != self.DP_EXTERNO_FACULTAD and \
                    self.inst_new_advisor != '':
                # pylint: disable=no-member
                paragraph.add_run(
                    self.str_cm[8] + self.get_inst_new_advisor_display())
        else:
            paragraph.add_run(
                ', ' + self.str_cm[3] + ' ' + self.str_cm[4] + ' ')
            paragraph.add_run(self.new_advisor)
            if self.inst_new_advisor != self.DP_EXTERNO_FACULTAD and \
                    self.inst_new_advisor != '':
                # pylint: disable=no-member
                paragraph.add_run(
                    self.str_cm[8] + self.get_inst_new_advisor_display())
            paragraph.add_run(' ' + self.str_cm[5] + ' ' + self.old_advisor)
            # pylint: disable=no-member
            paragraph.add_run(
                self.str_cm[8] + self.get_inst_old_advisor_display())
        paragraph.add_run('.')
        if self.new_co_advisor != '':
            if self.old_co_advisor == self.new_co_advisor or\
                    self.old_co_advisor == '':
                paragraph.add_run(
                    ' ' + self.str_cm[10] + ' ' + self.str_cm[4] + ' ')
                paragraph.add_run(self.new_co_advisor)
                if self.inst_new_co_advisor != self.DP_EXTERNO_FACULTAD and \
                        self.inst_new_co_advisor != '':
                    # pylint: disable=no-member
                    paragraph.add_run(
                        self.str_cm[8] + self.get_inst_new_co_advisor_display())
            else:
                paragraph.add_run(
                    ' ' + self.str_cm[6] + ' ' + self.str_cm[4] + ' ')
                paragraph.add_run(self.new_co_advisor)
                if self.inst_new_co_advisor != self.DP_EXTERNO_FACULTAD and \
                        self.inst_new_co_advisor != '':
                    # pylint: disable=no-member
                    paragraph.add_run(
                        self.str_cm[8] + self.get_inst_new_co_advisor_display())
                paragraph.add_run(' ' + self.str_cm[5] + ' ')
                paragraph.add_run(self.old_co_advisor)
                if self.inst_old_co_advisor != self.DP_EXTERNO_FACULTAD and \
                        self.inst_old_co_advisor != '':
                    # pylint: disable=no-member
                    paragraph.add_run(
                        self.str_cm[8] + self.get_inst_old_co_advisor_display())
            paragraph.add_run('.')

    def cm_ng(self, paragraph):
        paragraph.add_run(
            ' ' + self.str_cm[9] + ' ' + self.council_decision + '.')

    def pcm_analysis(self, docx):
        if self.grade_option in [self.GO_TESIS_MAESTRIA, self.GO_TESIS_DOCTORADO]:
            profile = 'investigación'
        else:
            profile = 'profundización'
        final_analysis = []
        final_analysis += [self.list_analysis[0].format(profile)]
        ets = '' if self.enrolled_thesis else 'no '
        # pylint: disable=no-member
        final_analysis += [self.list_analysis[1].format(
            ets, self.get_grade_option_display())]
        hss = 'T' if self.have_signature else 'No t'
        final_analysis += [hss +
                           self.list_analysis[2].format(self.new_advisor)]
        for extra_a in self.extra_analysis:
            final_analysis += [extra_a]
        add_analysis_paragraph(docx, final_analysis)

    def resource_analysis(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.pcm_answer(last_paragraph)
    
    def resource_pre_answer(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.pcm_answer(last_paragraph)

    def resource_answer(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.cm_answer(last_paragraph)
