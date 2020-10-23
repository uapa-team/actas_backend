from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from mongoengine import StringField, BooleanField, ListField
from ..models import Request
from .case_utils import add_analysis_paragraph


class MOTP(Request):

    full_name = 'Cambio de objetivos de tesis o trabajo final'
    decision_maker = Request.decision_makers[1]
    in_cm = False

    GO_TRABAJO_FINAL_MAESTRIA = 'TFM'
    GO_TESIS_MAESTRIA = 'TSM'
    GO_TESIS_DOCTORADO = 'TSD'
    GO_CHOICES = (
        (GO_TRABAJO_FINAL_MAESTRIA, 'Trabajo Final de Maestría'),
        (GO_TESIS_MAESTRIA, 'Tesis de Maestría'),
        (GO_TESIS_DOCTORADO, 'Tesis de Doctorado')
    )

    title = StringField(
        required=True, display='Título de la tesis/trabajo final', default=''),
    grade_option = StringField(
        required=True, display='Tipo de tesis/trabajo final',
        choices=GO_CHOICES, default=GO_TESIS_MAESTRIA)
    general_objetive = StringField(
        required=True, display='Objetivo general', default='')
    specific_objetives = ListField(StringField(),
                                   display='Objetivos específicos')
    enrolled_thesis = BooleanField(
        required=True, default=False,
        display='¿Tiene inscrita la asignatura tesis/trabajo final?')

    regulation_list = ['040|2017|CFA', '056|2012|CSU']  # List of regulations

    str_cm = ['cambiar los objetivos de ',
              '"{}"',
              'a:',
              'debido a que']

    list_analysis = ['Título:',
                     'Objetivo general:',
                     'Objetivos específicos:',
                     'Perfil de {}.',
                     'El estudiante {}tiene inscrita la asignatura {}.']

    def cm(self, docx):
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.str_council_header + ' ')
        self.cm_answer(paragraph)
        self.cm_ob(docx)

    def cm_answer(self, paragraph):
        # pylint: disable=no-member
        paragraph.add_run(
            self.get_approval_status_display().upper() + ' ').font.bold = True
        paragraph.add_run(
            self.str_cm[0] + self.get_grade_option_display() + ' ')
        paragraph.add_run(self.str_cm[1].format(self.title)).font.italic = True
        if self.is_affirmative_response_approval_status():
            self.cm_af(paragraph)
        else:
            self.cm_ng(paragraph)

    def pcm(self, docx):
        self.pcm_analysis(docx)
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.str_answer + ': ').font.bold = True
        paragraph.add_run(self.str_comittee_header + ' ')
        self.pcm_answer(paragraph)
        self.cm_ob(docx)

    def pcm_answer(self, paragraph):
        paragraph.add_run(
            # pylint: disable=no-member
            self.get_advisor_response_display().upper()).font.bold = True
        # pylint: disable=no-member
        paragraph.add_run(' ' + self.str_cm[0].format(
            self.get_grade_option_display(), self.get_academic_program_display()))
        paragraph.add_run(self.str_cm[1].format(self.title)).font.italic = True
        if self.is_affirmative_response_approval_status():
            self.cm_af(paragraph)
        else:
            self.cm_ng(paragraph)

    def cm_af(self, paragraph):
        paragraph.add_run(' ' + self.str_cm[2])

    def cm_ng(self, paragraph):
        paragraph.add_run(
            ' ' + self.str_cm[3] + ' ' + self.council_decision + '.')

    def cm_ob(self, docx):
        if not self.is_affirmative_response_approval_status():
            return
        paragraph = docx.add_paragraph()
        paragraph.style = 'List Bullet'
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.list_analysis[1]).font.bold = True
        paragraph = docx.add_paragraph()
        paragraph.style = 'List Bullet 2'
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.general_objetive + '.')
        paragraph = docx.add_paragraph()
        paragraph.style = 'List Bullet'
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.list_analysis[2]).font.bold = True
        for spec in self.specific_objetives:
            paragraph = docx.add_paragraph()
            paragraph.style = 'List Bullet 2'
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.add_run(spec + '.')

    def pcm_analysis(self, docx):
        if self.grade_option in [self.GO_TESIS_MAESTRIA, self.GO_TESIS_DOCTORADO]:
            profile = 'investigación'
        else:
            profile = 'profundización'
        final_analysis = []
        final_analysis += [self.list_analysis[3].format(profile)]
        ets = '' if self.enrolled_thesis else 'no '
        # pylint: disable=no-member
        final_analysis += [self.list_analysis[4].format(
            ets, self.get_grade_option_display())]
        for extra_a in self.extra_analysis:
            final_analysis += [extra_a]
        add_analysis_paragraph(docx, final_analysis)
        paragraph = docx.add_paragraph()
        paragraph.style = 'List Bullet'
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.list_analysis[0] + ' ').font.bold = True
        paragraph.add_run(self.title + '.').font.italic = True

    def resource_analysis(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.pcm_answer(last_paragraph)
    
    def resource_pre_answer(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.pcm_answer(last_paragraph)

    def resource_answer(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.cm_answer(last_paragraph)
