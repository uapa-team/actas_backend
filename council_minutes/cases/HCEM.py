from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from mongoengine import (StringField, BooleanField, IntField,
                         EmbeddedDocumentListField, EmbeddedDocument)
from ..models import Request, Subject
from .case_utils import table_approvals_cases, table_repprovals_cases, add_analysis_paragraph 


class HCEM(Request):

    class HomologatedSubject(Subject):
        HT_HOMOLOGACION = 'H'
        HT_CONVALIDACION = 'C'
        HT_EQUIVALENCIA = 'E'
        HT_ANDES = 'A'
        HT_INTERNACIONAL = 'I'
        HT_CHOICES = (
            (HT_HOMOLOGACION, 'Homologación'),
            (HT_CONVALIDACION, 'Convalidación'),
            (HT_EQUIVALENCIA, 'Equivalencia'),
            (HT_ANDES, 'Homologación conv. Uniandes'),
            (HT_INTERNACIONAL, 'Homologación conv. internacional'),
        )
        old_credits = IntField(default=3, min_value=0, required=True,
                               display='Créditos de la asignatura en la anterior institución')
        old_name = StringField(
            required=True, display='Nombre Asignatura en la anterior institución', default='')
        old_grade = StringField(
            required=True, default='', display='Calificación anterior del estudiante')
        grade = StringField(
            required=True, default='', display='Nueva calificación del estudiante')
        period = StringField(display='Periodo', choices=Request.PERIOD_CHOICES,
                default=Request.PERIOD_DEFAULT)
        approved = BooleanField(
            default=True, required=True, display='¿Fue aprobada la homologación?')
        reason = StringField(
            default='', display='Razón por la cuál no fue aprobada')
        h_type = StringField(required=True, default=HT_HOMOLOGACION,
                             choices=HT_CHOICES, display='Tipo de homologación')

    class MobilitySubject(Subject):
        GD_AP = 'AP'
        GD_NA = 'NA'
        HT_CHOICES = (
            (GD_AP, 'aprobada'),
            (GD_NA, 'reprobada'),
        )
        period = StringField(display='Periodo', 
                choices=Request.PERIOD_CHOICES, default=Request.PERIOD_DEFAULT)
        code = StringField(display='Código de la asignatura', default='')
        grade = StringField(display='Calificación',
                            default=GD_AP, choices=HT_CHOICES)

    full_name = 'Homologación, convalidación o equivalencia'

    institution_origin = StringField(
        required=True, default='Universidad Nacional de Colombia',
        display='Institución donde cursó las asignaturas')
    origin_plan = StringField(
        default='',
        display='Plan de estudios donde cursó las asignaturas')
    homologated_subjects = EmbeddedDocumentListField(
        HomologatedSubject, display='Asignaturas a homologar')
    mobility_subject = EmbeddedDocumentListField(MobilitySubject,
                                                 display='Asignaturas de movilidad')
    subject_accomplish_pr = BooleanField(
        default=True, display='¿Las asignaturas a homologar cumplen con los prerrequisitos?')
    greatger_than_50 = BooleanField(
        default=False, display='¿Se homologan/convalidan más del 50% de créditos del plan?')
    prev_hcem = BooleanField(
        default=False, display='¿Ha tenido homologaciones/convalidaciones anteriores.?')

    regulation_list = ['008|2008|CSU']  # List of regulations

    homologable_subjects = {
        '2011183': 'Intercambio Académico Internacional',
        '2014269': 'Intercambio Académico Internacional Prórroga',
        '2026630': 'Intercambio académico internacional – II',
        '2026631': 'Intercambio académico internacional - II Prórroga',
        '2024944': 'Asignatura por convenio con Universidad de los Andes I - POSGRADO',
        '2011302': 'Asignatura por convenio con Universidad de los Andes I - PREGRADO',
        '2012698': 'Asignatura por convenio con Universidad de los Andes II - PREGRADO',
    }

    verbs = {
        HomologatedSubject.HT_CONVALIDACION: 'convalidar',
        HomologatedSubject.HT_EQUIVALENCIA: 'equivaler',
        HomologatedSubject.HT_HOMOLOGACION: 'homologar',
        HomologatedSubject.HT_ANDES: 'homologar',
        HomologatedSubject.HT_INTERNACIONAL: 'homologar'}

    str_cm = [
        '{} la(s) siguiente(s) asignatura(s) cursada(s) en', 'el programa {} de la {}',
        'el intercambio académico internacional en la institución', 'el convenio con la ' +
        'Universidad de los Andes', 'de la siguiente manera', 'por la(s) siguiente(s) razon(es)',
        'calificar', 'la asignatura {} - {}, en el periodo {}']

    list_analysis = ['Solicitud de homologación de {} asignaturas del programa {} de' +
                     ' la institución {}.', 'Las asignaturas a homologar {}cumplen' +
                     ' con los prerrequisitos.', '{}e homologan/convalidan más' +
                     ' del 50% de créditos del plan (Artículo 38, {}).',
                     '{}a tenido homologaciones/convalidaciones anteriores.']

    srt_status = [['NO APROBAR', 'APROBAR'], ['NO APRUEBA', 'APRUEBA']]

    def counter(self):
        summary = [0, 0]
        types = {self.HomologatedSubject.HT_CONVALIDACION: 0,
                 self.HomologatedSubject.HT_EQUIVALENCIA: 0,
                 self.HomologatedSubject.HT_HOMOLOGACION: 0,
                 self.HomologatedSubject.HT_ANDES: 0,
                 self.HomologatedSubject.HT_INTERNACIONAL: 0, }
        for sbj in self.homologated_subjects:
            summary[sbj.approved] += 1
            types[sbj.h_type] += 1
        counter = 0
        if summary[0] == 0:
            counter += 1
        if summary[1] == 0:
            counter += 1
        if types[self.HomologatedSubject.HT_CONVALIDACION] == 0:
            counter += 1
        if types[self.HomologatedSubject.HT_EQUIVALENCIA] == 0:
            counter += 1
        if types[self.HomologatedSubject.HT_HOMOLOGACION] == 0:
            counter += 1
        if types[self.HomologatedSubject.HT_ANDES] == 0:
            counter += 1
        if types[self.HomologatedSubject.HT_INTERNACIONAL] == 0:
            counter += 1
        if self.mobility_subject == []:
            counter += 1
        return counter

    def cm(self, docx):
        if self.counter() == 6:
            paragraph = docx.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.add_run(self.str_council_header + ' ')
            self.cm_answer(paragraph)
            self.add_single_table(docx)
        else:
            self.add_composite_hcem(docx, True)

    def cm_answer(self, paragraph):
        # pylint: disable=no-member
        paragraph.add_run(
            self.get_approval_status_display().upper() + ' ').font.bold = True
        paragraph.add_run(self.str_cm[0].format(
            self.verbs[self.homologated_subjects[0].h_type]))
        paragraph.add_run(
            ' ' + self.str_cm[1].format(self.origin_plan, self.institution_origin))
        if self.is_affirmative_response_advisor_response():
            paragraph.add_run(' ' + self.str_cm[4] + ':')
        else:
            paragraph.add_run(' ' + self.str_cm[5] + ':')

    def add_analysis(self, docx):
        final_analysis = []
        final_analysis += [self.list_analysis[0].format(
            str(len(self.homologated_subjects)), self.origin_plan, self.institution_origin)]
        aux = '' if self.subject_accomplish_pr else 'no '
        final_analysis += [self.list_analysis[1].format(aux)]
        aux = 'S' if self.greatger_than_50 else 'No s'
        final_analysis += [self.list_analysis[2].format(
            aux, self.regulations['008|2008|CSU'][0])]
        aux = 'H' if self.prev_hcem else 'No h'
        final_analysis += [self.list_analysis[3].format(
            aux, self.regulations['008|2008|CSU'][0])]
        for extra_a in self.extra_analysis:
            final_analysis += [extra_a]
        add_analysis_paragraph(docx, final_analysis)

    def add_composite_hcem(self, docx, pre):
        # pylint: disable=consider-using-enumerate
        if not pre:
            paragraph = docx.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.add_run(self.str_answer + ': ').font.bold = True
        types = {self.HomologatedSubject.HT_CONVALIDACION: [[], []],
                 self.HomologatedSubject.HT_EQUIVALENCIA: [[], []],
                 self.HomologatedSubject.HT_HOMOLOGACION: [[], []],
                 self.HomologatedSubject.HT_ANDES: [[], []],
                 self.HomologatedSubject.HT_INTERNACIONAL: [[], []], }
        for sbj in self.homologated_subjects:
            types[sbj.h_type][sbj.approved].append(sbj)
        details = [self.student_name, self.student_dni,
                   self.academic_program, self.str_cm[1].format(
                       self.origin_plan, self.institution_origin)]
        for i in range(len(types)):
            for j in range(len(types[list(types.keys())[i]]) - 1, -1, -1):
                if len(types[list(types.keys())[i]][j]) != 0:
                    paragraph = docx.add_paragraph()
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.style = 'List Bullet'
                    if not pre:
                        paragraph.add_run(self.str_comittee_header + ' ')
                    else:
                        paragraph.add_run(self.str_council_header + ' ')
                    paragraph.add_run(
                        self.srt_status[pre][j] + ' ').font.bold = True
                    paragraph.add_run(self.str_cm[0].format(
                        self.verbs[types[list(types.keys())[i]][j][0].h_type]))
                    paragraph.add_run(
                        ' ' + self.str_cm[1].format(self.origin_plan, self.institution_origin))
                    if j != 0:
                        paragraph.add_run(' ' + self.str_cm[4] + ':')
                        data = []
                        for sbj in types[list(types.keys())[i]][j]:
                            data.append([sbj.period, sbj.code, sbj.name, sbj.credits,
                                         sbj.tipology[-1], sbj.grade, sbj.old_name, sbj.old_grade, sbj.old_credits])
                        idxSaved = []
                        # #Calculate grade of subjects many to one 
                        for a in range (len(data)-1):
                            if a in idxSaved:
                                continue
                            else:
                                auxArr = []
                                auxArr.append(a)
                                for b in range( a + 1 , len(data) ):
                                    if data[a][2] == data[b][2]:
                                        auxArr.append(b)
                                    
                                auxGrades = [] 
                                auxCredits = 0

                                # Calculate final grade, multiple subjects to one
                                for c in range(len(auxArr)):
                                    gradeSub = float(data[auxArr[c]][7]) * int(data[auxArr[c]][8])
                                    auxGrades.append(round(gradeSub,1))
                                    auxCredits = auxCredits + int(data[auxArr[c]][8])

                                if(len(auxArr) == 1):
                                    finalNote = round(auxGrades[0] / auxCredits, 1)
                                else:
                                    finalNote = round(sum(auxGrades) / auxCredits, 1)

                                # Update all grades of that subject
                                for a in range(len(auxArr)):
                                    data[auxArr[a]][5] = finalNote

                                idxSaved.extend(auxArr)

                        if  len(data)-1 not in idxSaved:
                            data[len(data)-1][5] = data[len(data)-1][7]
                        
                        table_approvals_cases(docx, data, details, types[list(types.keys())[i]][j][0].h_type)
                    else:
                        paragraph.add_run(' ' + self.str_cm[5] + ':')
                        data = []
                        for sbj in types[list(types.keys())[i]][j]:
                            data.append([sbj.period, sbj.name, sbj.old_name, sbj.reason,
                                         sbj.old_credits, sbj.old_grade])
                        table_repprovals_cases(docx, data, details,types[list(types.keys())[i]][j][0].h_type)
        if self.mobility_subject != []:
            for sbj in self.mobility_subject:
                paragraph = docx.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.style = 'List Bullet'
                if not pre:
                    paragraph.add_run(self.str_comittee_header + ' ')
                else:
                    paragraph.add_run(self.str_council_header + ' ')
                paragraph.add_run(
                    self.srt_status[pre][1] + ' ').font.bold = True
                paragraph.add_run(self.str_cm[6] + ' ')
                paragraph.add_run('{} ({})'.format(
                    sbj.get_grade_display(), sbj.grade) + ' ')
                try:
                    paragraph.add_run(self.str_cm[7].format(
                        sbj.code, self.homologable_subjects[sbj.code], sbj.period) + '.')
                except KeyError as e:
                    print(e)

    def pcm(self, docx):
        self.add_analysis(docx)
        if self.counter() == 6:
            paragraph = docx.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.add_run(self.str_answer + ': ').font.bold = True
            paragraph.add_run(self.str_comittee_header + ' ')
            self.pcm_answer(paragraph)
            self.add_single_table(docx)
        else:
            self.add_composite_hcem(docx, False)

    def pcm_answer(self, paragraph):
        # pylint: disable=no-member
        paragraph.add_run(
            self.get_advisor_response_display().upper() + ' ').font.bold = True
        paragraph.add_run(self.str_cm[0].format(
            self.verbs[self.homologated_subjects[0].h_type]))
        paragraph.add_run(
            ' ' + self.str_cm[1].format(self.origin_plan, self.institution_origin))
        if self.is_affirmative_response_advisor_response():
            paragraph.add_run(' ' + self.str_cm[4] + ':')
        else:
            paragraph.add_run(' ' + self.str_cm[5] + ':')

    def add_single_table(self, docx):
        data = []
        for sbj in self.homologated_subjects:
            data.append([sbj.period, sbj.code, sbj.name, sbj.credits,
                         sbj.tipology[-1], sbj.grade, sbj.old_name, sbj.old_grade, sbj.old_credits])
        
        idxSaved = []
        
        #Calculate grade of subjects many to one 
        for i in range (len(data)-1):
            if i in idxSaved:
                continue
            else:
                auxArr = []
                auxArr.append(i)
                for j in range( i + 1 , len(data) ):
                    if data[i][2] == data[j][2]:
                        auxArr.append(j)
                    
                auxGrades = [] 
                auxCredits = 0

                # Calculate final grade, multiple subjects to one
                for c in range(len(auxArr)):
                    gradeSub = float(data[auxArr[c]][7]) * int(data[auxArr[c]][8])
                    print(float(data[auxArr[c]][7]) * int(data[auxArr[c]][8]))
                    auxGrades.append(round(gradeSub,1))
                    auxCredits = auxCredits + int(data[auxArr[c]][8])

                if(len(auxArr) == 1):
                    finalNote = round(auxGrades[0] / auxCredits, 1)
                else:
                    finalNote = round(sum(auxGrades) / auxCredits, 1)

                # Update all grades of that subject
                for a in range(len(auxArr)):
                    data[auxArr[a]][5] = finalNote

                idxSaved.extend(auxArr)

        if  len(data)-1 not in idxSaved:
            data[len(data)-1][5] = data[len(data)-1][7]
        
        table_approvals_cases(docx, data, [self.student_name, self.student_dni,
                                     self.academic_program, self.str_cm[1].format(
                                         self.origin_plan, self.institution_origin)], 'H')

    def resource_analysis(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.pcm_answer(last_paragraph)

    def resource_pre_answer(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.pcm_answer(last_paragraph)

    def resource_answer(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.cm_answer(last_paragraph)
