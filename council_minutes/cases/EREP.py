import datetime
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from mongoengine import DateField, BooleanField
from ..models import Request
from .case_utils import string_to_date, add_analysis_paragraph


class EREP(Request):

    full_name = 'Expedición de recibo de matrícula'
    decision_maker = Request.decision_makers[3]

    ah_active = BooleanField(
        required=True, display='¿Tiene activa la historia académica?', default=False)
    payment_date = DateField(
        display='Fecha límite de pago', default=datetime.date.today)

    regulation_list = ['051|2003|CSU']  # List of regulations

    str_cm = ['presentar con concepto positivo al Comité de Matrículas de la Sede Bogotá, ' +
              'la expedición de un único recibo correspondiente a los derechos académicos y ' +
              'administrativos para el periodo académico {}', ' y se le concede como fecha de ' +
              'pago el {} ({}), teniendo en cuenta el estado de pago por parte de {}.']

    list_analysis = ['El estudiante {}tiene la historia académica activa.']

    def cm(self, docx):
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.str_council_header + ' ')
        self.cm_answer(paragraph)

    def cm_answer(self, paragraph):
        paragraph.add_run(
            # pylint: disable=no-member
            self.get_approval_status_display().upper() + ' ').font.bold = True
        paragraph.add_run(self.str_cm[0].format(self.academic_period))
        if self.is_affirmative_response_approval_status():
            self.cm_af(paragraph)
        else:
            self.cm_ng(paragraph)

    def pcm(self, docx):
        self.pcm_analysis(docx)
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.str_answer + ': ').font.bold = True
        paragraph.add_run(self.str_comittee_header + ' ')
        self.pcm_answer(paragraph)

    def pcm_answer(self, paragraph):
        paragraph.add_run(
            # pylint: disable=no-member
            self.get_advisor_response_display().upper()).font.bold = True
        paragraph.add_run(' ' + self.str_cm[0].format(self.academic_period))
        if self.is_affirmative_response_advisor_response():
            self.pcm_answers_af(paragraph)
        else:
            self.pcm_answers_ng(paragraph)

    def cm_af(self, paragraph):
        paragraph.add_run(self.str_cm[1].format(
            self.payment_date, string_to_date(str(self.payment_date)), self.student_name))

    def cm_ng(self, paragraph):
        paragraph.add_run(self.council_decision + '.')

    def pcm_analysis(self, docx):
        active = '' if self.ah_active else 'no '
        final_analysis = []
        final_analysis += [self.list_analysis[0].format(active)]
        for extra_a in self.extra_analysis:
            final_analysis += [extra_a]
        add_analysis_paragraph(docx, final_analysis)

    def pcm_answers_af(self, paragraph):
        paragraph.add_run(self.str_cm[1].format(
            self.payment_date, string_to_date(str(self.payment_date)), self.student_name))

    def pcm_answers_ng(self, paragraph):
        paragraph.add_run(self.council_decision + '.')

    def resource_analysis(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.pcm_answer(last_paragraph)
    
    def resource_pre_answer(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.pcm_answer(last_paragraph)

    def resource_answer(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.cm_answer(last_paragraph)
