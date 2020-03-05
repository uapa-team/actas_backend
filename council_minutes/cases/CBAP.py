from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from mongoengine import IntField
from ..models import Request
from .case_utils import add_analysis_paragraph


class CBAP(Request):

    full_name = 'Créditos excedentes BAPI'

    credits = IntField(required=True, display='Créditos excedentes', default=0)

    regulation_list = []

    str_cm = [
        'trasladar ',
        ' crédito(s) aprobado(s) en ',
        ' debido a que ',
        ' exigido(s) por la asignatura Trabajo de Grado, que se asumirá(n)' +
        ' como crédito(s) inscrito(s) y aprobado(s) del componente de libre elección,' +
        ' si en este componente aún hay créditos por ser aprobados. ',
        '(Artículo 16 del '
    ]

    def cm(self, docx):
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.str_council_header + ' ')
        self.cm_answer(paragraph)

    def cm_answer(self, paragraph):
        # pylint: disable=no-member
        paragraph.add_run(
            self.get_approval_status_display().upper() + ' ').font.bold = True
        paragraph.add_run(self.str_cm[0] + str(self.credits) + self.str_cm[1])
        paragraph.add_run(self.get_academic_program_display())
        if self.approval_status != 'AP':
            paragraph.add_run(self.str_cm[2] + self.student_justification)
            return
        paragraph.add_run(self.str_cm[3])
        paragraph.add_run(
            self.str_cm[4] + self.regulations['026|2012|CSU'][0] + ').')

    def pcm_analysis(self, docx):
        # pylint: disable=no-member
        analysis_list = []
        analysis_list += self.extra_analysis
        add_analysis_paragraph(docx, analysis_list)

    # PCM for CBAP not found. Using the same of cm.
    def pcm(self, docx):
        self.pcm_analysis(docx)
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph = docx.add_paragraph()
        paragraph.add_run(self.str_answer + ': ').bold = True
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.str_comittee_header + ' ')
        self.pcm_answer(paragraph)

    def pcm_answer(self, paragraph):
        # pylint: disable=no-member
        paragraph.add_run(
            self.get_advisor_response_display().upper() + ' ').font.bold = True
        paragraph.add_run(self.str_cm[0] + str(self.credits) + self.str_cm[1])
        paragraph.add_run(self.get_academic_program_display())
        if self.advisor_response != Request.ARCR_APROBAR:
            paragraph.add_run(self.str_cm[2] + self.council_decision)
            return
        paragraph.add_run(self.str_cm[3])
        paragraph.add_run(
            self.str_cm[4] + self.regulations['026|2012|CSU'][0] + ').')

    def resource_analysis(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.pcm_answer(last_paragraph)
    
    def resource_pre_answer(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.pcm_answer(last_paragraph)

    def resource_answer(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.cm_answer(last_paragraph)
