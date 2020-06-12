from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from mongoengine import StringField, IntField
from ..models import Request
from .case_utils import add_analysis_paragraph


class CGRU(Request):

    full_name = 'Cambio de grupo'

    name = StringField(required=True, display='Nombre Asignatura', default='')
    code = StringField(required=True, display='Código', default='')
    group = StringField(required=True, display='Grupo', default='')
    new_group = StringField(required=True, display='Nuevo Grupo', default='')
    free_places = IntField(required=True, min_value=0,
                           default=0, display='Cupos disponibles')
    professor = StringField(required=True, display='Profesor', default='')
    department = StringField(required=True, choices=Request.DP_CHOICES,
                             default=Request.DP_EMPTY, display='Departamento')

    regulation_list = ['008|2008|CSU']

    str_cm = [
        'cambio de grupo de la asignatura {} ({}) al grupo {} en el periodo {}, ' +
        'debido a que {}.',
        'justifica debidamente la solicitud'
    ]

    str_pcm = [
        'El grupo {} de la asignatura {} ({}) cuenta con {} cupos.',
        'cambio de grupo de la asignatura/actividad {}, código {}, ' +
        'inscrita en el periodo {}, del grupo {} al grupo {} con ' +
        'el profesor {} del {}, debido a que {}.'
    ]

    def cm(self, docx):
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.str_council_header + ' ')
        self.cm_answer(paragraph)

    def cm_answer(self, paragraph):
        # pylint: disable=no-member
        paragraph.add_run(
            self.get_approval_status_display().upper() + ' ').font.bold = True
        if self.council_decision == Request.council_decision.default or len(self.council_decision) == 0:
            modifier = self.str_cm[1]
        else:
            modifier = self.council_decision
        paragraph.add_run(self.str_cm[0].format(
            self.name, self.code, self.new_group, self.academic_period, modifier
        ))

    def pcm(self, docx):
        add_analysis_paragraph(docx, self.create_analysis())
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.str_answer + ' ').font.bold = True
        paragraph.add_run(self.str_comittee_header + ' ')
        self.pcm_answer(paragraph)

    def pcm_answer(self, paragraph):
        # pylint: disable=no-member
        paragraph.add_run(
            self.get_advisor_response_display().upper() + ' ').font.bold = True
        paragraph.add_run(self.str_pcm[1].format(
            self.name, self.code, self.academic_period, self.group,
            self.new_group, self.professor, self.get_department_display(), self.council_decision
        ))

    def create_analysis(self):
        return [
            self.str_pcm[0].format(self.group, self.name,
                                   self.code, self.free_places)
        ] + self.extra_analysis

    def resource_analysis(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.pcm_answer(last_paragraph)
    
    def resource_pre_answer(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.pcm_answer(last_paragraph)

    def resource_answer(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.cm_answer(last_paragraph)
