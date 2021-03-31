from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from mongoengine import EmbeddedDocumentListField, BooleanField, ListField, StringField
from ..models import Request, Subject
from .case_utils import table_subjects, add_analysis_paragraph


class IASISubject(Subject):
    offered = BooleanField(display='Ofrecida para el plan de estudios', default=True)
    overlap = BooleanField(display='Materia cruzada', default=False)
    approved = BooleanField(display='Aprobado inscribir materia', default=True)


class IASI(Request):

    full_name = 'Inscripción de Asignaturas'

    CJT_ANSWER_DEFAULT = 'DF'
    CJT_ANSWER_PROC_ACT = 'PA'
    CJT_ANSWER_HIST_BAPI = 'HB'
    CJT_ANSWER_PROC_MAS = 'PM'
    CJT_ANSWER_OTRO = 'OT'
    CJT_ANSWER_N_DEB = 'ND'

    CJT_ANSWER_CHOICES = (
        (CJT_ANSWER_DEFAULT,''),
        (CJT_ANSWER_PROC_ACT,'Proceso de actualización'),
        (CJT_ANSWER_HIST_BAPI,'Inscripción en la historia académica BAPI'),
        (CJT_ANSWER_PROC_MAS,'Falta de registro en el proceso masivo'),
        (CJT_ANSWER_N_DEB,'No justifica debidamente su solicitud'),
        (CJT_ANSWER_OTRO, 'Otro')
    )

    subjects = EmbeddedDocumentListField(
        IASISubject, display='Asignaturas')

    council_decision = StringField(
        max_length=255, choices=CJT_ANSWER_CHOICES,
        default=CJT_ANSWER_DEFAULT, display='Justificación del Consejo')

    str_cm = [
        'inscribir la(s) siguiente(s) asignatura(s) del programa {} ({}), en el periodo académico' +
        ' {}, debido a que {}.',
    ]

    str_pcm = [
        'Se solicita inscribir la asignatura {} ({}). La materia {}es ofrecida para el plan de es' +
        'tudios {} ({}) y {}tiene cruces con el horario actual del estudiante.'
    ]

    regulation_list = ['008|2008|CSU']  # List of regulations

    def cm(self, docx):
        sapproved = []
        snotapproved = []
        for subject in self.subjects:
            if subject.approved:
                sapproved.append(subject)
            else:
                snotapproved.append(subject)
        if len(sapproved) > 0:
            paragraph = docx.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.add_run(self.str_council_header + ' ')
            self.cm_answer_approved(paragraph)
            table_subjects(docx, Subject.subjects_to_array(sapproved))
        if len(snotapproved) > 0:
            paragraph = docx.add_paragraph()
            paragraph = docx.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.add_run(self.str_council_header + ' ')
            self.cm_answer_not_approved(paragraph)
            table_subjects(docx, Subject.subjects_to_array(snotapproved))

    def cm_answer(self, paragraph):
        paragraph.add_run(self.str_council_header + ' ')
        paragraph.add_run(
            # pylint: disable=no-member
            self.get_approval_status_display().upper() + ' ').font.bold = True
        paragraph.add_run(self.str_cm[0].format(
            # pylint: disable=no-member
            self.get_academic_program_display(),
            self.academic_program,
            self.academic_period,
            self.council_decision))
        paragraph.add_run('({}).'.format(self.regulations['008|2008|CSU'][0]))

    def cm_answer_approved(self, paragraph):
        paragraph.add_run(
            # pylint: disable=no-member
            'APRUEBA ').font.bold = True
        paragraph.add_run(self.str_cm[0].format(
            # pylint: disable=no-member
            self.get_academic_program_display(),
            self.academic_program,
            self.academic_period,
            self.council_decision))
        paragraph.add_run('({}).'.format(self.regulations['008|2008|CSU'][0]))

    def cm_answer_not_approved(self, paragraph):
        paragraph.add_run(
            # pylint: disable=no-member
            'NO APRUEBA ').font.bold = True
        paragraph.add_run(self.str_cm[0].format(
            # pylint: disable=no-member
            self.get_academic_program_display(),
            self.academic_program,
            self.academic_period,
            self.council_decision))
        paragraph.add_run('({}).'.format(self.regulations['008|2008|CSU'][0]))

    def pcm(self, docx):
        self.pcm_analysis(docx)
        paragraph = docx.add_paragraph()
        paragraph.add_run(self.str_answer + ': ').bold = True
        sapproved = []
        snotapproved = []
        for subject in self.subjects:
            if subject.approved:
                sapproved.append(subject)
            else:
                snotapproved.append(subject)
        if len(sapproved) > 0:
            paragraph = docx.add_paragraph()
            paragraph.add_run(self.str_comittee_header)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph.paragraph_format.space_after = Pt(0)
            self.pcm_answer_approved(paragraph)
            table_subjects(docx, Subject.subjects_to_array(sapproved))
        if len(snotapproved) > 0:
            paragraph = docx.add_paragraph()
            paragraph.add_run(self.str_comittee_header)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph.paragraph_format.space_after = Pt(0)
            self.pcm_answer_not_approved(paragraph)
            table_subjects(docx, Subject.subjects_to_array(snotapproved))

    def pcm_analysis(self, docx):
        # pylint: disable=no-member
        analysis_list = []
        for subject in self.subjects:
            analysis_list += [self.str_pcm[0].format(
                subject.name,
                subject.code,
                '' if subject.offered else 'no ',
                self.get_academic_program_display(),
                self.academic_program,
                '' if subject.overlap else 'no ')]
        analysis_list += self.extra_analysis
        add_analysis_paragraph(docx, analysis_list)

    def pcm_answer(self, paragraph):
        paragraph.add_run(
            # pylint: disable=no-member
            ' ' + self.get_advisor_response_display().upper() + ' ').font.bold = True
        paragraph.add_run(self.str_cm[0].format(
            # pylint: disable=no-member
            self.get_academic_program_display(),
            self.academic_program,
            self.academic_period,
            self.council_decision))

    def pcm_answer_approved(self, paragraph):
        paragraph.add_run(
            # pylint: disable=no-member
            ' APROBAR ').font.bold = True
        paragraph.add_run(self.str_cm[0].format(
            # pylint: disable=no-member
            self.get_academic_program_display(),
            self.academic_program,
            self.academic_period,
            self.council_decision))

    def pcm_answer_not_approved(self, paragraph):
        paragraph.add_run(
            # pylint: disable=no-member
            ' NO APROBAR ').font.bold = True
        paragraph.add_run(self.str_cm[0].format(
            # pylint: disable=no-member
            self.get_academic_program_display(),
            self.academic_program,
            self.academic_period,
            self.council_decision))

    def resource_analysis(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.pcm_answer(last_paragraph)
    
    def resource_pre_answer(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.pcm_answer(last_paragraph)

    def resource_answer(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.cm_answer(last_paragraph)
