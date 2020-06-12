# pylint: disable=no-name-in-module
import datetime
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from mongoengine import StringField, DateField
from ..models import Request
from .case_utils import add_analysis_paragraph, string_to_date


class PEAC(Request):

    full_name = 'Permiso Académico'

    reason_permision = StringField(required=True, default='Con el objetivo de ...',
                                   display='Razón del permiso académico')
    from_date = DateField(
        required=True, display='Fecha de inicio del permiso', default=datetime.date.today)
    to_date = DateField(
        required=True, display='Fecha de fin del permiso', default=datetime.date.today)

    str_cm = [
        'otorgar permiso académico desde el {} hasta el {} con el objetivo de {}',
        ' debido a que {}.',
    ]

    regulation_list = ['070|2012|CSU']  # List of regulations

    def cm(self, docx):
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.str_council_header + ' ')
        self.cm_answer(paragraph)
        paragraph.add_run(' ({}). '.format(
            self.regulations[self.regulation_list[0]][0]))

    def cm_answer(self, paragraph):
        # pylint: disable=no-member
        paragraph.add_run(
            self.get_approval_status_display().upper() + ' ').font.bold = True
        paragraph.add_run(self.str_cm[0].format(
            string_to_date(str(self.from_date)),
            string_to_date(str(self.to_date)),
            str(self.reason_permision)[0].lower() +
            str(self.reason_permision)[1:]
        ))
        if not self.is_affirmative_response_approval_status():
            paragraph.add_run(self.str_cm[1].format(self.council_decision))
        else:
            paragraph.add_run('.')

    def pcm(self, docx):
        self.pcm_analysis(docx)
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.str_answer + ': ').bold = True
        paragraph.add_run(self.str_comittee_header + ' ')
        self.pcm_answer(paragraph)

    def pcm_analysis(self, docx):
        analysis_list = []
        analysis_list += self.extra_analysis
        add_analysis_paragraph(docx, analysis_list)

    def pcm_answer(self, paragraph):
        paragraph.add_run(
            # pylint: disable=no-member
            self.get_advisor_response_display().upper() + ' ').font.bold = True
        paragraph.add_run(self.str_cm[0].format(
            string_to_date(str(self.from_date)),
            string_to_date(str(self.to_date)),
            str(self.reason_permision)[0].lower() +
            str(self.reason_permision)[1:]
        ))
        if not self.is_affirmative_response_advisor_response():
            paragraph.add_run(self.str_cm[1].format(self.council_decision))
        else:
            paragraph.add_run('.')

    def resource_analysis(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.pcm_answer(last_paragraph)

    def resource_pre_answer(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.pcm_answer(last_paragraph)

    def resource_answer(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.cm_answer(last_paragraph)
