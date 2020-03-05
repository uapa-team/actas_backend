from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from mongoengine import StringField, EmbeddedDocumentListField
from ..models import Request, Subject
from .case_utils import add_analysis_paragraph, table_change_typology, table_approvals


class ChangeTipologySubject(Subject):
    new_tipology = StringField(
        required=True, choices=Subject.TIP_CHOICES,
        default='', display='Nueva tipología')
    grade = StringField(display='Nota obtenida', default='')


class HomologationSubject(Subject):
    period = StringField(display='Periodo de vista la materia')
    grade = StringField(display='Nota obtenida')
    new_grade = StringField(display='Nota a homologar')
    new_name = StringField(display='Nombre de la materia a homologar')


class CPER(Request):

    full_name = 'Cambio de perfil'

    origin_profile = StringField(
        choices=Request.PROFILE_CHOICES, display='Perfil origen')
    destin_profile = StringField(
        choices=Request.PROFILE_CHOICES, display='Perfil destino')
    subjects_change_tipology = EmbeddedDocumentListField(
        ChangeTipologySubject, display='Asignaturas cambio de componente')
    subjects_homologations = EmbeddedDocumentListField(
        HomologationSubject, display='Asignaturas equivalencia')

    regulation_list = ['032|2010|CSU', '1416|2013|RE']  # List of regulations

    str_cm = [
        'traslado intrafacultad del estudiante de {} ({}) en el perfil de {} al plan de estudios ' +
        '{} ({}) en el perfil de {}, debido a que {}justifica adecuadamente su solicitud.',
        'cambiar de componente las siguientes asignaturas del programa {} ({}):',
        'Equivaler en el programa {} ({}) perfil de {}, las siguientes asignaturas cursadas en el' +
        ' programa {} ({}) perfil de {}:'
    ]

    str_pcm = [
        'traslado intrafacultad del estudiante de {} ({}) en el perfil de {} al plan de estudios ' +
        '{} ({}) en el perfil de {}.',
        'Se cursa la materia {} ({}), con la tipología {} ({}) y se solicita hacer el cambio a la' +
        ' tipología {} ({}).',
        'Se cursa la materia {} ({}), en el periodo {}, y se solicita equivaler a la materia {} (' +
        '{}), con nota obtenida {} y nota equivalente {}.'
    ]

    def cm(self, docx):
        affirmative = self.is_affirmative_response_approval_status()
        has_subjects = len(self.subjects_change_tipology) + \
            len(self.subjects_homologations) > 0
        if affirmative and has_subjects:
            self.answer_subjects(docx, self.str_council_header + ': ')
        else:
            self.answer_no_subjects(docx, self.str_council_header + ': ')

    def cm_answer(self, paragraph):
        paragraph.add_run(
            # pylint: disable=no-member
            self.get_approval_status_display().upper() + ' ').font.bold = True
        paragraph.add_run(self.str_cm[0].format(
            # pylint: disable=no-member
            self.get_academic_program_display(),
            self.academic_program,
            self.get_origin_profile_display(),
            self.get_academic_program_display(),
            self.academic_program,
            self.get_destin_profile_display(),
            '' if self.is_affirmative_response_approval_status() else 'no'))

    def answer_no_subjects(self, docx, str_in):
        paragraph = self.answer_paragraph_normal(docx)
        paragraph.add_run(str_in)
        self.cm_answer(paragraph)

    def answer_paragraph(self, docx, style):
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.style = style
        return paragraph

    def answer_paragraph_normal(self, docx):
        return self.answer_paragraph(docx, 'Normal')

    def answer_add_bullet(self, docx):
        return self.answer_paragraph(docx, 'List Bullet')

    def answer_subjects(self, docx, str_in):
        paragraph = self.answer_paragraph_normal(docx)
        paragraph.add_run(str_in)
        paragraph = self.answer_add_bullet(docx)
        self.cm_answer(paragraph)
        if len(self.subjects_change_tipology) != 0:
            paragraph = self.answer_add_bullet(docx)
            paragraph.add_run(
                # pylint: disable=no-member
                self.get_approval_status_display().upper() + ' ').font.bold = True
            paragraph.add_run(self.str_cm[1].format(
                # pylint: disable=no-member
                self.get_academic_program_display(), self.academic_program))
            self.add_subjects_change_tipology_table(docx)
        if len(self.subjects_homologations) != 0:
            paragraph = self.answer_add_bullet(docx)
            paragraph.add_run(
                # pylint: disable=no-member
                self.get_approval_status_display().upper() + ' ').font.bold = True
            paragraph.add_run(self.str_cm[2].format(
                # pylint: disable=no-member
                self.get_academic_program_display(),
                self.academic_program,
                self.get_origin_profile_display(),
                self.get_academic_program_display(),
                self.academic_program,
                self.get_destin_profile_display()))
            self.add_subjects_homologation_table(docx)

    def add_subjects_change_tipology_table(self, docx):
        subjects = []
        for subject in self.subjects_change_tipology:
            subjects.append([
                subject.code,
                subject.name,
                subject.grade,
                subject.tipology[1],
                subject.new_tipology[1]
            ])
        table_change_typology(docx, subjects)

    def add_subjects_homologation_table(self, docx):
        subjects = []
        for subject in self.subjects_homologations:
            subjects.append([
                subject.period,
                subject.code,
                subject.name,
                str(subject.credits),
                subject.tipology[1],
                subject.new_grade,
                subject.new_name,
                subject.grade
            ])
        table_approvals(docx, subjects, [
            self.student_name,
            self.student_dni,
            self.academic_program,
            'perfil de {}'.format(
                # pylint: disable=no-member
                self.get_destin_profile_display()
            )
        ])

    def pcm(self, docx):
        self.pcm_analysis(docx)
        affirmative = self.is_affirmative_response_advisor_response()
        has_subjects = len(self.subjects_change_tipology) + \
            len(self.subjects_homologations) > 0
        if affirmative and has_subjects:
            self.answer_subjects(docx, self.str_comittee_header + ': ')
        else:
            self.answer_no_subjects(docx, self.str_comittee_header + ': ')

    def pcm_answer(self, paragraph):
        paragraph.add_run(
            # pylint: disable=no-member
            self.get_advisor_response_display().upper() + ' ').font.bold = True
        paragraph.add_run(self.str_cm[0].format(
            # pylint: disable=no-member
            self.get_academic_program_display(),
            self.academic_program,
            self.get_origin_profile_display(),
            self.get_academic_program_display(),
            self.academic_program,
            self.get_destin_profile_display(),
            '' if self.is_affirmative_response_advisor_response() else 'no'))

    def pcm_analysis(self, docx):
        analysis_list = []
        analysis_list += [self.str_pcm[0].format(
            # pylint: disable=no-member
            self.get_academic_program_display(),
            self.academic_program,
            self.get_origin_profile_display(),
            self.get_academic_program_display(),
            self.academic_program,
            self.get_destin_profile_display(),
        )]
        for subject in self.subjects_change_tipology:
            analysis_list += [self.str_pcm[1].format(
                subject.name,
                subject.code,
                subject.get_tipology_display(),
                subject.tipology[1],
                subject.get_new_tipology_display(),
                subject.new_tipology[1],
            )]
        for subject in self.subjects_homologations:
            analysis_list += [self.str_pcm[2].format(
                subject.name,
                subject.code,
                subject.period,
                subject.new_name,
                subject.code,
                subject.grade,
                subject.new_grade
            )]
        analysis_list += self.extra_analysis
        add_analysis_paragraph(docx, analysis_list)

    def resource_analysis(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.pcm_answer(last_paragraph)
    
    def resource_pre_answer(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.pcm_answer(last_paragraph)

    def resource_answer(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.cm_answer(last_paragraph)
