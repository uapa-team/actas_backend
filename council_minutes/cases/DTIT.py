# pylint: disable=no-name-in-module
import datetime
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Pt
from mongoengine import StringField, BooleanField, EmbeddedDocument, FloatField, DateTimeField, IntField, EmbeddedDocumentListField
from ..models import Request,Subject
from .case_utils import add_analysis_paragraph
from .case_utils import string_to_date, table_general_data
from council_minutes.cases.case_utils import table_credits_summary, table_recommend
from ..helpers import get_period_choices

class DTIT_subject(Subject):
        PERIOD_CHOICES = get_period_choices()
        PERIOD_DEFAULT = PERIOD_CHOICES[1][0] if datetime.date.today(
        ).month <= 6 else PERIOD_CHOICES[0][0]
        academic_period = StringField(
            max_length=10, display='Periodo Académico en el cual fue visto', choices=PERIOD_CHOICES, default=PERIOD_DEFAULT)
        grade = FloatField(required = True, display='Nota', default=0.0)
        group = StringField(required = True, display='Agrupación', default='')
        
        def subjects_to_array(self, subjects):
            """
            A function that converts a List of Subjects into a classic array.
            : param subjects: EmbeddedDocumentListField of Subjects to be converted
            """
            data = []
            for subject in subjects:
                data.append([
                    subject.academic_period,
                    subject.code,
                    subject.name,
                    subject.tipology[-1],
                    subject.group,
                    str(subject.credits),
                    str(subject.grade)
                ])
            return data

class DTIT_pending_subject(Subject):
    #Class of pending subjets
        group = StringField(required = True, display='Agrupación', default='')
        
        def subjects_to_array(self, subjects):
            """
            A function that converts a List of Subjects into a classic array.
            : param subjects: EmbeddedDocumentListField of Subjects to be converted
            """
            data = []
            for subject in subjects:
                data.append([
                    subject.code,
                    subject.name,
                    subject.tipology[-1],
                    subject.group,
                    str(subject.credits),
                ])
            return data

class DTIT_pending_group(EmbeddedDocument):
    #Class of pending groups

    group = StringField(required = True, display='Agrupación', default='')
    tipology = StringField( required=True, choices=Subject.TIP_CHOICES
        , display='Tipología', default=Subject.TIP_PRE_FUND_OBLIGATORIA)
    credits = IntField(required = True, display='Créditos', default='0')

    def groups_to_array(self, groups):
        """
        A function that converts a List of Groups into a classic array.
        : param groups: EmbeddedDocumentListField of Groups to be converted
        """
        data = []
        for group in groups:
            data.append([
                group.group,
                group.credits,
                group.tipology[-1],
            ])
        return data
class DTIT(Request):

    full_name = 'Doble Titulación'

    second_plan = StringField(min_length=4, max_length=4, choices=Request.PLAN_CHOICES,
        required=True, display='¿Cuál es el segundo plan?', default=Request.PI_AGRICOLA)
    is_graduated = BooleanField(
        required=True, display='¿Es estudiante de posgrado?', default=False)
    is_enrolled = BooleanField(
        required=True, display='¿El estudiante se encuentra matriculado?', default=True)
    was_student = BooleanField(
        required=True, display='¿Ha sido estudiante en el segundo plan?', default=False)
    has_credits = BooleanField(
        required=True, display='¿Dispone del cupo suficiente de créditos?', default=False)
    wasnt_student = BooleanField(
        required=True, display='¿Ha perdido calidad de estudiante?', default=False)
    papa = FloatField(required=True, display='P.A.P.A', default=0.0)
    quota_credits = IntField(display = 'Cupo de créditos menos créditos pendientes del primer plan',
        default = 0)

    subjects = EmbeddedDocumentListField(
        DTIT_subject, display='Asignaturas cursadas')
    pending_subjects = EmbeddedDocumentListField(
        DTIT_pending_subject, display='Asignaturas no cursadas')
    groups = EmbeddedDocumentListField(
        DTIT_pending_group, display='Agrupaciones optativas faltantes') 
    ob_fund_credit = IntField(display = 'Créditos necesarios de fundamentación obligatorios del segundo plan',
        default = 0)
    op_fund_credit = IntField(display = 'Créditos necesarios de fundamentación optativos del segundo plan',
        default = 0)
    ob_disc_credit = IntField(display = 'Créditos necesarios disciplinares obligatorios del segundo plan',
        default = 0)
    op_disc_credit = IntField(display = 'Créditos necesarios disciplinares optativos del segundo plan',
        default = 0)
    free_elect_credit = IntField(display = 'Créditos necesarios de libre elección del segundo plan',
        default = 0)
    
    equivalent_credits_subjects = []
    pending_credits_subjects = []


    regulation_list = ['008|2008|CSU', '155|2014|CSU']  # List of regulations

    str_cm = [
        ' recomendar al Consejo de Sede que formalice la admisión y ubicación en ' + 
        'el programa de Ingeniería {} - ({}), ',
        ' Teniendo en cuenta que el estudiante tiene un Promedio Académico' + 
        ' Ponderado Acumulado superior o igual a 4.3. ({})',
        ' Teniendo en cuenta que el estudiante cuenta con un cupo de créditos '
        'suficiente para culminar el segundo plan de estudios. ({})'
    ]
    str_cm_false = [
        ' recomendar al Consejo de Sede que formalice la admisión y ubicación en el' +
        ' programa de pregrado Ingeniería {} – ({}), debido a que tiene ' +
        'un PAPA de {} y no cuenta con el cupo suficiente de créditos para culminar el' +
        'segundo plan. ({})'
    ]

    str_pcm = [
        ' recomendar al Consejo de Sede que formalice la admisión y ubicación en ' + 
        'el programa de Ingeniería {} - ({}), ',
        ' Teniendo en cuenta que el estudiante tiene un Promedio Académico' + 
        ' Ponderado Acumulado superior o igual a 4.3. ({})',
        ' Teniendo en cuenta que el estudiante cuenta con un cupo de créditos '
        'suficiente para culminar el segundo plan de estudios. ({})',
        '1. Datos Generales:',
        '2. Información Académica:',
        '3. Cuadro equivalencia y convalidaciones de asignaturas cursadas y aprobadas '+
        'hasta la fecha de presentación de la solicitud por parte del estudiante:',
        '4. Asignaturas pendientes por cursar en el segundo plan de estudios:',
        '5. Resumen general de créditos del segundo plan de estudios:',
        '*Sin incluir los créditos correspondientes al cumplimiento del requisito de' +
        ' suficiencia en idioma extranjero.',
        '**Aprobados del plan de estudios, sin excedentes.'
    ]

    str_pcm_false = [
        ' recomendar al Consejo de Sede que formalice la admisión y ubicación en el' +
        ' programa de pregrado Ingeniería {} – ({}), debido a que tiene ' +
        'un PAPA de {} y no cuenta con el cupo suficiente de créditos para culminar el' +
        'segundo plan. ({})'
    ]

    str_analysis = [
        '{} estudiante de posgrado (Artículo 49 {}). Universitas y SIA: .',
        '{} está matriculado al momento de la solicitud (Artículo 1, {}). Universitas y SIA; .',
        '{} ha tenido calidad de estudiante en el plan de estudios de doble titulación (Artículo 4, ' +
        '{}).Universitas: .',
        '{} dispone del cupo de créditos necesario para optar por el segundo título luego de convalidar ' +
        'o hacer equivaler todas las asignaturas pertinentes cursadas y aprobadas en el primer plan de ' +
        'estudios (parágrafo 1, {}).',
        'Régimen de convalidaciones y equivalencias PERTINENTES entre el primero y el segundo plan de ' +
        'estudios (Artículo 2, {}).',
        '{} ha perdido la calidad de estudiante por las causales 2, 3, 4 o 5 del Artículo 44 del '+
        '{} (Artículo 7, {}).Universitas: .'
    ]


    def cm(self, docx):
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.str_council_header + ' ')
        self.cm_answer(paragraph)
        if self.is_affirmative_response_approval_status():
            # is_affirmative_response_advisor_response
            self.cm_adds(docx)
            self.pcm_adds(docx)
            self.dtit_general_data_table(docx)
            self.dtit_academic_info_table(docx)
            self.table_subjects(docx, DTIT_subject.subjects_to_array(self, self.subjects))
            self.table_pending_subjects(docx, DTIT_pending_subject.subjects_to_array(self, self.pending_subjects),
                    DTIT_pending_group.groups_to_array(self, self.groups))
            self.dtit_recommend_table(docx)

    def cm_answer(self, paragraph):
        # pylint: disable=no-member
        paragraph.add_run(
            self.get_approval_status_display().upper()).font.bold = True
        if self.is_affirmative_response_approval_status():
            paragraph.add_run(self.str_cm[0].format(self.get_academic_program_display(),
                self.academic_program))
        else:
            paragraph.add_run(self.str_cm_false[0].format(self.get_academic_program_display(), 
                self.academic_program, self.papa, self.regulations[self.regulation_list[1]][0]))

    def cm_adds(self, docx):
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.style = 'List Bullet'
        paragraph.add_run(self.str_cm[1].format(self.regulations[self.regulation_list[1]][0]))

        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.style = 'List Bullet'
        paragraph.add_run(self.str_cm[2].format(self.regulations[self.regulation_list[1]][0]))


    def pcm(self, docx):
        self.pcm_analysis(docx)
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.str_answer + ': ').bold = True
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.str_comittee_header + ' ')
        self.pcm_answer(paragraph)
        if self.is_affirmative_response_advisor_response():
            self.pcm_adds(docx)
            self.dtit_general_data_table(docx)
            self.dtit_academic_info_table(docx)
            self.table_subjects(docx, DTIT_subject.subjects_to_array(self, self.subjects))
            self.table_pending_subjects(docx, DTIT_pending_subject.subjects_to_array(self, self.pending_subjects),
                    DTIT_pending_group.groups_to_array(self, self.groups))
            self.dtit_recommend_table(docx)


    def pcm_adds(self, docx):
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.style = 'List Bullet'
        paragraph.add_run(self.str_pcm[1].format(self.regulations[self.regulation_list[1]][0]))

        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.style = 'List Bullet'
        paragraph.add_run(self.str_pcm[2].format(self.regulations[self.regulation_list[1]][0]))

    def pcm_analysis(self, docx):
        analysis_list = [
            self.str_analysis[0].format('Es' if self.is_graduated else 'No es',
                    self.regulations[self.regulation_list[0]][0]),
            self.str_analysis[1].format('Si' if self.is_enrolled else 'No',
                    self.regulations[self.regulation_list[1]][0]),
            self.str_analysis[2].format('Si' if self.was_student else 'No', 
                    self.regulations[self.regulation_list[1]][0]),
            self.str_analysis[3].format('Si' if self.has_credits else 'No', 
                    self.regulations[self.regulation_list[0]][0]),
            self.str_analysis[4].format(self.regulations[self.regulation_list[1]][0]),
            self.str_analysis[5].format('Si' if self.wasnt_student else 'No', 
                    self.regulations[self.regulation_list[0]][0],
                    self.regulations[self.regulation_list[1]][0])
        ]
        
        add_analysis_paragraph(docx, analysis_list)

    def pcm_answer(self, paragraph):
        # pylint: disable=no-member
        paragraph.add_run(
            self.get_advisor_response_display().upper()).font.bold = True
        if self.is_affirmative_response_advisor_response():
            paragraph.add_run(self.str_pcm[0].format(self.get_academic_program_display(),
                self.academic_program))
        else:
            paragraph.add_run(self.str_pcm_false[0].format(self.get_academic_program_display(), 
                self.academic_program, self.papa, self.regulations[self.regulation_list[1]][0]))


    def resource_analysis(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.pcm_answer(last_paragraph)
    
    def resource_pre_answer(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.pcm_answer(last_paragraph)

    def resource_answer(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.cm_answer(last_paragraph)

    # Table General data
    def dtit_general_data_table(self, docx):
        # pylint: disable=no-member
        general_data = [['Nombre del estudiante', self.student_name],
                        ['DNI', self.student_dni],
                        ['Plan de estudios origen (1er plan) - Sede', self.get_academic_program_display()],
                        ['Código del plan de estudios origen', self.academic_program],
                        ['Plan de estudios doble titulación (2° plan)', self.get_second_plan_display()],
                        ['Código del plan de estudios doble titulación', self.second_plan],
                        ['Fecha de la solicitud a través del SIA', string_to_date(str(self.date))]]

        case = 'DOBLE TITULACIÓN'

        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        bullet = paragraph.add_run(self.str_pcm[3])
        bullet.font.bold = True
        bullet.font.size = Pt(8)

        table_general_data(general_data, case, docx)
        paragraph = docx.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run('\n').font.size = Pt(8)

    # Table Academic Information
    def dtit_academic_info_table(self, docx):
        '''
        Add a generated table with general data
            Params:
            1.  general_data (list of list):    A list with the general data of the student,
            in the first field of the sublist goes the name of the information,
            and in the second field goes the value of the information, e.g.:
                [["Nombre Estudiante", "Juan Pérez"], ['DNI', '1111111'],
                ['Plan de estudios', 'Ingeniería de Sistemas'],
                ['Código del plan de estudios', '2879'],
                ['Fecha de la Solicitud', '29 de abril del 2019']]
            2.  docx_ (docx_):  The document to which the table will be added
        '''
        # pylint: disable=no-member

        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        bullet = paragraph.add_run(self.str_pcm[4])
        bullet.font.bold = True
        bullet.font.size = Pt(8)

        table = docx.add_table(rows=4, cols=5, style='Table Grid')
        table.style.font.size = Pt(8)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.columns[0].width = 2600000
        table.columns[1].width = 650000
        table.columns[2].width = 650000
        table.columns[3].width = 650000
        table.columns[4].width = 650000
        for cell in table.columns[0].cells:
            cell.width = 2600000
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for cell in table.columns[1].cells:
            cell.width = 650000
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for cell in table.columns[2].cells:
            cell.width = 650000
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for cell in table.columns[3].cells:
            cell.width = 650000
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for cell in table.columns[4].cells:
            cell.width = 650000
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        table.cell(0, 0).paragraphs[0].add_run('¿Tuvo calidad de estudiante en el ' +
                            '2° plan?').font.size = Pt(8)
        table.cell(0, 1).paragraphs[0].add_run('Sí').font.size = Pt(8)
        table.cell(0, 3).paragraphs[0].add_run('No').font.size = Pt(8)

        table.cell(1, 0).paragraphs[0].add_run('Se encuentra matriculado al momento ' +
                            'de presentar la solicitud').font.size = Pt(8)
        table.cell(1, 1).paragraphs[0].add_run('Sí').font.size = Pt(8)
        table.cell(1, 3).paragraphs[0].add_run('No').font.size = Pt(8)

        table.cell(2, 0).paragraphs[0].add_run('PAPA en el primer ' +
                            'plan de estudio').font.size = Pt(8)    
        cellp = table.cell(2, 1).merge(table.cell(2, 2)).paragraphs[0]
        cellp = table.cell(2, 2).merge(table.cell(2, 3)).paragraphs[0]
        cellp = table.cell(2, 3).merge(table.cell(2, 4)).paragraphs[0]
        cellp.add_run(str(self.papa)).font.size = Pt(8)

        table.cell(3, 0).paragraphs[0].add_run('Cupo de créditos menos créditos ' +
                            'pendientes del primer plan:').font.size = Pt(8)
        cellp = table.cell(3, 1).merge(table.cell(3, 2)).paragraphs[0]
        cellp = table.cell(3, 2).merge(table.cell(3, 3)).paragraphs[0]
        cellp = table.cell(3, 3).merge(table.cell(3, 4)).paragraphs[0]
        cellp.add_run(str( self.quota_credits )).font.size = Pt(8)

        if self.was_student:
            table.cell(0, 2).paragraphs[0].add_run('X').font.size = Pt(8)
            table.cell(0, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            table.cell(0, 4).paragraphs[0].add_run('X').font.size = Pt(8)
            table.cell(0, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if self.is_enrolled:
            table.cell(1, 2).paragraphs[0].add_run('X').font.size = Pt(8)
            table.cell(1, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            table.cell(1, 4).paragraphs[0].add_run('X').font.size = Pt(8)
            table.cell(1, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        paragraph = docx.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run('\n').font.size = Pt(8)


    #Third table of docx
    def table_subjects(self, docx, data):
        '''Add a generated table with approvals subjects
            Params:
                docx_ (docx_): The document to which the table will be added
                subjects (list): A list of list with the subjects in table,
                each list must be a list with following data:
                [0]: Subject's academic period
                [1]: Subject's SIA code
                [2]: Subject's SIA name
                [3]: Subject's SIA tipology
                [4]: Subject's group
                [5]: Subject's credits
                [6]: Subject's grade
            Raises:
                IndexError: All lists must have same size
        '''
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        bullet = paragraph.add_run(self.str_pcm[5])
        bullet.font.bold = True
        bullet.font.size = Pt(8)

        # first table
        cont_fund = 0
        sum_total = 0
        fund_index = []
        for i in range (len(data)):
            if data[i][3] == 'B' or data[i][3] == 'O':
                cont_fund = cont_fund + 1
                fund_index.append(i)

        table = docx.add_table(rows=cont_fund+2, cols=9)
        for column in table.columns:
            for cell in column.cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        table.style = 'Table Grid'
        table.style.font.size = Pt(9)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.columns[0].width = 500000
        table.columns[1].width = 600000
        table.columns[2].width = 1150000
        table.columns[3].width = 600000
        table.columns[4].width = 1100000
        table.columns[5].width = 250000
        table.columns[6].width = 650000
        table.columns[7].width = 250000
        table.columns[8].width = 400000
        for cell in table.columns[0].cells:
            cell.width = 500000
        for cell in table.columns[1].cells:
            cell.width = 600000
        for cell in table.columns[2].cells:
            cell.width = 1150000
        for cell in table.columns[3].cells:
            cell.width = 600000
        for cell in table.columns[4].cells:
            cell.width = 1100000
        for cell in table.columns[5].cells:
            cell.width = 250000
        for cell in table.columns[6].cells:
            cell.width = 650000
        for cell in table.columns[7].cells:
            cell.width = 250000
        for cell in table.columns[8].cells:
            cell.width = 400000
        
        cellp = table.cell(0, 0).merge(table.cell(0, 1)).paragraphs[0]
        cellp = table.cell(0, 1).merge(table.cell(0, 2)).paragraphs[0]
        cellp.add_run('PLAN DE ESTUDIOS (1)').font.bold = True
        cellp.runs[0].font.size = Pt(8)
        cellp = table.cell(0, 3).merge(table.cell(0, 4)).paragraphs[0]
        cellp = table.cell(0, 4).merge(table.cell(0, 5)).paragraphs[0]
        cellp = table.cell(0, 5).merge(table.cell(0, 6)).paragraphs[0]
        cellp = table.cell(0, 6).merge(table.cell(0, 7)).paragraphs[0]
        cellp = table.cell(0, 7).merge(table.cell(0, 8)).paragraphs[0]
        cellp.add_run('PLAN DE ESTUDIOS (2)' + '\n' + 'COMPONENTE ' + 
            'DE FUNDAMENTACIÓN (B y O)').font.bold = True
        cellp.runs[0].font.size = Pt(8)
        table.cell(1, 0).paragraphs[0].add_run('Periodo').font.bold = True
        table.cell(1, 1).paragraphs[0].add_run('Código').font.bold = True
        table.cell(1, 2).paragraphs[0].add_run('Asignatura').font.bold = True
        table.cell(1, 3).paragraphs[0].add_run('Código').font.bold = True
        table.cell(1, 4).paragraphs[0].add_run('Asignatura').font.bold = True
        table.cell(1, 5).paragraphs[0].add_run('T*').font.bold = True
        table.cell(1, 6).paragraphs[0].add_run('Agrupación').font.bold = True
        table.cell(1, 7).paragraphs[0].add_run('C*').font.bold = True
        table.cell(1, 8).paragraphs[0].add_run('NOTA').font.bold = True
        for i in range(9):
            table.cell(1, i).paragraphs[0].runs[0].font.size = Pt(8)
        idx = 0
        for index in fund_index:
            table.cell(idx + 2, 0).paragraphs[0].add_run(
                data[index][0]).font.size = Pt(8)
            table.cell(idx + 2, 1).paragraphs[0].add_run(
                data[index][1]).font.size = Pt(8)
            table.cell(idx + 2, 2).paragraphs[0].add_run(
                data[index][2]).font.size = Pt(8)
            table.cell(idx + 2, 3).paragraphs[0].add_run(
                data[index][1]).font.size = Pt(8)
            table.cell(idx + 2, 4).paragraphs[0].add_run(
                data[index][2]).font.size = Pt(8)
            table.cell(idx + 2, 5).paragraphs[0].add_run(
                data[index][3]).font.size = Pt(8)
            table.cell(idx + 2, 6).paragraphs[0].add_run(
                data[index][4]).font.size = Pt(8)
            table.cell(idx + 2, 7).paragraphs[0].add_run(
                data[index][5]).font.size = Pt(8)
            table.cell(idx + 2, 8).paragraphs[0].add_run(
                data[index][6]).font.size = Pt(8)
            sum_total = sum_total + int(data[index][5])
            idx = idx + 1

        cont_B = 0
        cont_O = 0
        for i in range (len(data)):
            if data[i][3] == 'B':
                cont_B = cont_B + int(data[i][5])
            elif data[i][3] == 'O':
                cont_O = cont_O + int(data[i][5])
            else:
                continue

        self.equivalent_credits_subjects.append(cont_B)
        self.equivalent_credits_subjects.append(cont_O)
        
        table = docx.add_table(rows=1, cols=2)
        for column in table.columns:
            for cell in column.cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        table.style = 'Table Grid'
        table.style.font.size = Pt(9)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.columns[0].width = 2500000
        table.columns[1].width = 750000
        for cell in table.columns[0].cells:
            cell.width = 2500000
        for cell in table.columns[1].cells:
            cell.width = 750000
        
        table.cell(0, 0).paragraphs[0].add_run('Total créditos ' + 
                'convalidados/equivalentes en el componente').font.size = Pt(8)
        table.cell(0, 1).paragraphs[0].add_run(str( sum_total )).font.size = Pt(8)
        
        paragraph = docx.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        bullet = paragraph.add_run('Ob: obligatoria. Op: optativa. C: créditos')
        bullet.font.size = Pt(8)
        bullet.font.italic = True
        bullet.font.underline = True

        paragraph = docx.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run('\n').font.size = Pt(8)


        # Table #2
        cont_disc = 0
        sum_total = 0
        disc_index = []
        for i in range (len(data)):
            if data[i][3] == 'C' or data[i][3] == 'T':
                cont_disc = cont_disc + 1
                disc_index.append(i)
    

        table = docx.add_table(rows=cont_disc+2, cols=9)
        for column in table.columns:
            for cell in column.cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        table.style = 'Table Grid'
        table.style.font.size = Pt(9)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.columns[0].width = 500000
        table.columns[1].width = 600000
        table.columns[2].width = 1150000
        table.columns[3].width = 600000
        table.columns[4].width = 1100000
        table.columns[5].width = 250000
        table.columns[6].width = 650000
        table.columns[7].width = 250000
        table.columns[8].width = 400000
        for cell in table.columns[0].cells:
            cell.width = 500000
        for cell in table.columns[1].cells:
            cell.width = 600000
        for cell in table.columns[2].cells:
            cell.width = 1150000
        for cell in table.columns[3].cells:
            cell.width = 600000
        for cell in table.columns[4].cells:
            cell.width = 1100000
        for cell in table.columns[5].cells:
            cell.width = 250000
        for cell in table.columns[6].cells:
            cell.width = 650000
        for cell in table.columns[7].cells:
            cell.width = 250000
        for cell in table.columns[8].cells:
            cell.width = 400000

        cellp = table.cell(0, 0).merge(table.cell(0, 1)).paragraphs[0]
        cellp = table.cell(0, 1).merge(table.cell(0, 2)).paragraphs[0]
        cellp.add_run('PLAN DE ESTUDIOS (1)').font.bold = True
        cellp.runs[0].font.size = Pt(8)
        cellp = table.cell(0, 3).merge(table.cell(0, 4)).paragraphs[0]
        cellp = table.cell(0, 4).merge(table.cell(0, 5)).paragraphs[0]
        cellp = table.cell(0, 5).merge(table.cell(0, 6)).paragraphs[0]
        cellp = table.cell(0, 6).merge(table.cell(0, 7)).paragraphs[0]
        cellp = table.cell(0, 7).merge(table.cell(0, 8)).paragraphs[0]
        cellp.add_run('PLAN DE ESTUDIOS (2)' + '\n' + 'COMPONENTE ' + 
            'DISCIPLINAR/PROFESIONAL (C y T)').font.bold = True
        cellp.runs[0].font.size = Pt(8)
        table.cell(1, 0).paragraphs[0].add_run('Periodo').font.bold = True
        table.cell(1, 1).paragraphs[0].add_run('Código').font.bold = True
        table.cell(1, 2).paragraphs[0].add_run('Asignatura').font.bold = True
        table.cell(1, 3).paragraphs[0].add_run('Código').font.bold = True
        table.cell(1, 4).paragraphs[0].add_run('Asignatura').font.bold = True
        table.cell(1, 5).paragraphs[0].add_run('T*').font.bold = True
        table.cell(1, 6).paragraphs[0].add_run('Agrupación').font.bold = True
        table.cell(1, 7).paragraphs[0].add_run('C*').font.bold = True
        table.cell(1, 8).paragraphs[0].add_run('NOTA').font.bold = True
        for i in range(9):
            table.cell(1, i).paragraphs[0].runs[0].font.size = Pt(8)
        
        idx = 0
        for index in disc_index:
            table.cell(idx + 2, 0).paragraphs[0].add_run(
                data[index][0]).font.size = Pt(8)
            table.cell(idx + 2, 1).paragraphs[0].add_run(
                data[index][1]).font.size = Pt(8)
            table.cell(idx + 2, 2).paragraphs[0].add_run(
                data[index][2]).font.size = Pt(8)
            table.cell(idx + 2, 3).paragraphs[0].add_run(
                data[index][1]).font.size = Pt(8)
            table.cell(idx + 2, 4).paragraphs[0].add_run(
                data[index][2]).font.size = Pt(8)
            table.cell(idx + 2, 5).paragraphs[0].add_run(
                data[index][3]).font.size = Pt(8)
            table.cell(idx + 2, 6).paragraphs[0].add_run(
                data[index][4]).font.size = Pt(8)
            table.cell(idx + 2, 7).paragraphs[0].add_run(
                data[index][5]).font.size = Pt(8)
            table.cell(idx + 2, 8).paragraphs[0].add_run(
                data[index][6]).font.size = Pt(8)
            sum_total = sum_total + int(data[index][5])
            idx = idx +1
            
        cont_C = 0
        cont_T = 0
        for i in range (len(data)):
            if data[i][3] == 'C':
                cont_C = cont_C + int(data[i][5])
            elif data[i][3] == 'T':
                cont_T = cont_T + int(data[i][5])
            else:
                continue

        self.equivalent_credits_subjects.append(cont_C)
        self.equivalent_credits_subjects.append(cont_T)

        table = docx.add_table(rows=1, cols=2)
        for column in table.columns:
            for cell in column.cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        table.style = 'Table Grid'
        table.style.font.size = Pt(9)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.columns[0].width = 2500000
        table.columns[1].width = 750000
        for cell in table.columns[0].cells:
            cell.width = 2500000
        for cell in table.columns[1].cells:
            cell.width = 750000
        
        table.cell(0, 0).paragraphs[0].add_run('Total créditos ' + 
                'convalidados/equivalentes en el componente').font.size = Pt(8)
        table.cell(0, 1).paragraphs[0].add_run(str( sum_total )).font.size = Pt(8)
        
        paragraph = docx.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        bullet = paragraph.add_run('Ob: obligatoria. Op: optativa. C: créditos')
        bullet.font.size = Pt(8)
        bullet.font.italic = True
        bullet.font.underline = True

        paragraph = docx.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run('\n').font.size = Pt(8)

        # Free election table 
        cont_free = 0
        cont_credits = 0
        sum_total = 0
        free_index = []
        for i in range (len(data)):
            if data[i][3] == 'L':
                cont_free = cont_free + 1
                free_index.append(i)
                cont_credits = cont_credits + int(data[i][5])

        table = docx.add_table(rows=cont_free+2, cols=7)
        for column in table.columns:
            for cell in column.cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        table.style = 'Table Grid'
        table.style.font.size = Pt(9)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.columns[0].width = 500000
        table.columns[1].width = 550000
        table.columns[2].width = 1600000
        table.columns[3].width = 550000
        table.columns[4].width = 1200000
        table.columns[5].width = 400000
        table.columns[6].width = 400000
        for cell in table.columns[0].cells:
            cell.width = 500000
        for cell in table.columns[1].cells:
            cell.width = 550000
        for cell in table.columns[2].cells:
            cell.width = 1600000
        for cell in table.columns[3].cells:
            cell.width = 550000
        for cell in table.columns[4].cells:
            cell.width = 1200000
        for cell in table.columns[5].cells:
            cell.width = 400000
        for cell in table.columns[6].cells:
            cell.width = 400000
        cellp = table.cell(0, 0).merge(table.cell(0, 1)).paragraphs[0]
        cellp = table.cell(0, 1).merge(table.cell(0, 2)).paragraphs[0]
        cellp.add_run('PLAN DE ESTUDIOS (1)').font.bold = True
        cellp.runs[0].font.size = Pt(8)
        cellp = table.cell(0, 3).merge(table.cell(0, 4)).paragraphs[0]
        cellp = table.cell(0, 4).merge(table.cell(0, 5)).paragraphs[0]
        cellp = table.cell(0, 5).merge(table.cell(0, 6)).paragraphs[0]
        cellp.add_run('PLAN DE ESTUDIOS (2)' + '\n' + 'COMPONENTE ' + 
            'DE LIBRE ELECCIÓN (L)').font.bold = True
        cellp.runs[0].font.size = Pt(8)
        table.cell(1, 0).paragraphs[0].add_run('Periodo').font.bold = True
        table.cell(1, 1).paragraphs[0].add_run('Código').font.bold = True
        table.cell(1, 2).paragraphs[0].add_run('Asignatura').font.bold = True
        table.cell(1, 3).paragraphs[0].add_run('Código').font.bold = True
        table.cell(1, 4).paragraphs[0].add_run('Asignatura').font.bold = True
        table.cell(1, 5).paragraphs[0].add_run('C*').font.bold = True
        table.cell(1, 6).paragraphs[0].add_run('NOTA').font.bold = True
        for i in range(7):
            table.cell(1, i).paragraphs[0].runs[0].font.size = Pt(8)
        
        idx = 0
        for index in free_index:
            table.cell(idx + 2, 0).paragraphs[0].add_run(
                data[index][0]).font.size = Pt(8)
            table.cell(idx + 2, 1).paragraphs[0].add_run(
                data[index][1]).font.size = Pt(8)
            table.cell(idx + 2, 2).paragraphs[0].add_run(
                data[index][2]).font.size = Pt(8)
            table.cell(idx + 2, 3).paragraphs[0].add_run(
                data[index][1]).font.size = Pt(8)
            table.cell(idx + 2, 4).paragraphs[0].add_run(
                data[index][2]).font.size = Pt(8)
            table.cell(idx + 2, 5).paragraphs[0].add_run(
                data[index][5]).font.size = Pt(8)
            table.cell(idx + 2, 6).paragraphs[0].add_run(
                data[index][6]).font.size = Pt(8)
            sum_total = sum_total + int(data[index][5])
            idx = idx +1
        
        self.equivalent_credits_subjects.append(cont_credits)

        table = docx.add_table(rows=1, cols=2)
        for column in table.columns:
            for cell in column.cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        table.style = 'Table Grid'
        table.style.font.size = Pt(9)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.columns[0].width = 2500000
        table.columns[1].width = 750000
        for cell in table.columns[0].cells:
            cell.width = 2500000
        for cell in table.columns[1].cells:
            cell.width = 750000
        
        table.cell(0, 0).paragraphs[0].add_run('Total créditos ' + 
                'convalidados/equivalentes en el componente').font.size = Pt(8)
        table.cell(0, 1).paragraphs[0].add_run(str( sum_total )).font.size = Pt(8)
        
        paragraph = docx.add_paragraph()
        paragraph.add_run('C: créditos').font.size = Pt(8)
      
        paragraph = docx.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run('\n').font.size = Pt(8)


    #Forth table of docx
    def table_pending_subjects(self, docx, data, group_data):
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        bullet = paragraph.add_run(self.str_pcm[6])
        bullet.font.bold = True
        bullet.font.size = Pt(8)
        
        cont_fund_OB = 0
        fund_OB_index = []
        aux_group = []
        for i in range (len(data)):
            if data[i][2] == 'B':
                cont_fund_OB = cont_fund_OB + 1
                fund_OB_index.append(i)    

        table = docx.add_table(rows=cont_fund_OB+4, cols=5)
        for column in table.columns:
            for cell in column.cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        table.style = 'Table Grid'
        table.style.font.size = Pt(9)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.columns[0].width = 1250000
        table.columns[1].width = 550000
        table.columns[2].width = 1850000
        table.columns[3].width = 550000
        table.columns[4].width = 1000000
        for cell in table.columns[0].cells:
            cell.width = 1250000
        for cell in table.columns[1].cells:
            cell.width = 550000
        for cell in table.columns[2].cells:
            cell.width = 1850000
        for cell in table.columns[3].cells:
            cell.width = 550000
        for cell in table.columns[4].cells:
            cell.width = 1000000
        cellp = table.cell(0, 0).merge(table.cell(0, 1)).paragraphs[0]
        cellp = table.cell(0, 1).merge(table.cell(0, 2)).paragraphs[0]
        cellp = table.cell(0, 2).merge(table.cell(0, 3)).paragraphs[0]
        cellp = table.cell(0, 3).merge(table.cell(0, 4)).paragraphs[0]
        cellp.add_run('Componente de Fundamentación(B)').font.bold = True
        cellp.runs[0].font.size = Pt(8)
        
        cellp = table.cell(1, 0).merge(table.cell(1, 1)).paragraphs[0]
        cellp = table.cell(1, 1).merge(table.cell(1, 2)).paragraphs[0]
        cellp = table.cell(1, 2).merge(table.cell(1, 3)).paragraphs[0]
        cellp = table.cell(1, 3).merge(table.cell(1, 4)).paragraphs[0]
        cellp.add_run('Obligatorias').font.bold = True
        cellp.runs[0].font.size = Pt(8)
        
        table.cell(2, 0).paragraphs[0].add_run('Agrupación').font.bold = True
        table.cell(2, 1).paragraphs[0].add_run('Código').font.bold = True
        table.cell(2, 2).paragraphs[0].add_run('Asignatura').font.bold = True
        table.cell(2, 3).paragraphs[0].add_run('Créditos asignatura').font.bold = True
        table.cell(2, 4).paragraphs[0].add_run('Créditos pendientes por cursar por ' +
            'el estudiante').font.bold = True
        for i in range(5):
            table.cell(2, i).paragraphs[0].runs[0].font.size = Pt(8)
        
        #Required to build table
        idx = 0
        sum_group = 0
        sum_total = 0
        cont_subject = 0
        merge_subjects = []
        credits_subjects = []

        for i in range(cont_fund_OB):
            if len(aux_group) == 0:
                var = data[fund_OB_index[i]][3]
                aux_group.append(var)
                table.cell(idx + 3, 1).paragraphs[0].add_run(
                    data[fund_OB_index[i]][0]).font.size = Pt(8)
                table.cell(idx + 3, 2).paragraphs[0].add_run(
                    data[fund_OB_index[i]][1]).font.size = Pt(8)
                table.cell(idx + 3, 3).paragraphs[0].add_run(
                    data[fund_OB_index[i]][4]).font.size = Pt(8)
                sum_group = sum_group + int(data[fund_OB_index[i]][4])
                cont_subject = cont_subject + 1
                idx = idx +1

                for j in range (i+1, cont_fund_OB):
                    if var == data[j][3]:
                        table.cell(idx + 3, 1).paragraphs[0].add_run(
                            data[fund_OB_index[j]][0]).font.size = Pt(8)
                        table.cell(idx + 3, 2).paragraphs[0].add_run(
                            data[fund_OB_index[j]][1]).font.size = Pt(8)
                        table.cell(idx + 3, 3).paragraphs[0].add_run(
                            data[fund_OB_index[j]][4]).font.size = Pt(8)
                        sum_group = sum_group + int(data[fund_OB_index[j]][4])
                        cont_subject = cont_subject + 1
                        idx = idx +1
                merge_subjects.append(cont_subject)
                credits_subjects.append(sum_group)
                sum_total = sum_total + sum_group
                cont_subject = 0
                sum_group = 0
            elif data[i][3] not in aux_group:
                var = data[fund_OB_index[i]][3]
                aux_group.append(var)
                table.cell(idx + 3, 1).paragraphs[0].add_run(
                    data[fund_OB_index[i]][0]).font.size = Pt(8)
                table.cell(idx + 3, 2).paragraphs[0].add_run(
                    data[fund_OB_index[i]][1]).font.size = Pt(8)
                table.cell(idx + 3, 3).paragraphs[0].add_run(
                    data[fund_OB_index[i]][4]).font.size = Pt(8)
                sum_group = sum_group + int(data[fund_OB_index[i]][4])
                cont_subject = cont_subject + 1
                idx = idx +1

                for j in range (i+1, cont_fund_OB):
                    if var == data[j][3]:
                        table.cell(idx + 3, 1).paragraphs[0].add_run(
                            data[fund_OB_index[j]][0]).font.size = Pt(8)
                        table.cell(idx + 3, 2).paragraphs[0].add_run(
                            data[fund_OB_index[j]][1]).font.size = Pt(8)
                        table.cell(idx + 3, 3).paragraphs[0].add_run(
                            data[fund_OB_index[j]][4]).font.size = Pt(8)
                        sum_group = sum_group + int(data[fund_OB_index[j]][4])
                        cont_subject = cont_subject + 1
                        idx = idx +1
                merge_subjects.append(cont_subject)
                credits_subjects.append(sum_group)
                sum_total = sum_total + sum_group
                cont_subject = 0
                sum_group = 0
            else:
                continue
        
        cont = 0
        aux = 0
        for i in merge_subjects:
            if i == 1:
                table.cell(cont + 3, 0).paragraphs[0].add_run(data[fund_OB_index[cont]][3]).font.size = Pt(8)
                table.cell(cont + 3, 4).paragraphs[0].add_run(str(credits_subjects[aux])).font.size = Pt(8)
                cont = cont + i 
                aux = aux + 1
            else:
                for a in range(i-1):
                    cellp3 = table.cell(a + cont + 3, 0).merge(table.cell(a + cont + 4, 0)).paragraphs[0]
                    cellp4 = table.cell(a + cont + 3, 4).merge(table.cell(a + cont + 4, 4)).paragraphs[0]
                cellp3.add_run(data[fund_OB_index[cont]][3]).font.size = Pt(8)
                cellp4.add_run(str(credits_subjects[aux])).font.size = Pt(8)
                aux = aux + 1
                cont = cont + i 

        self.pending_credits_subjects.append(sum_total)
        cellp = table.cell(cont_fund_OB+3, 0).merge(table.cell(cont_fund_OB+3, 1)).paragraphs[0]
        cellp = table.cell(cont_fund_OB+3, 1).merge(table.cell(cont_fund_OB+3, 2)).paragraphs[0]
        cellp = table.cell(cont_fund_OB+3, 2).merge(table.cell(cont_fund_OB+3, 3)).paragraphs[0]
        cellp.add_run('Total créditos pendientes').font.bold = True
        cellp.runs[0].font.size = Pt(8)
        cellp.runs[0].font.highlight_color = WD_COLOR_INDEX.YELLOW
        table.cell(cont_fund_OB+3, 4).paragraphs[0].add_run(str(sum_total)).font.size = Pt(8)


        # Second part of table
        cont_fund_OP = 0
        fund_OP_index = []
        for i in range (len(group_data)):
            if group_data[i][2] == 'O':
                cont_fund_OP = cont_fund_OP + 1
                fund_OP_index.append(i)

        table = docx.add_table(rows=cont_fund_OP+3, cols=3)
        for column in table.columns:
            for cell in column.cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        table.style = 'Table Grid'
        table.style.font.size = Pt(9)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.columns[0].width = 3100000
        table.columns[1].width = 1100000
        table.columns[2].width = 1000000
        for cell in table.columns[0].cells:
            cell.width = 3100000
        for cell in table.columns[1].cells:
            cell.width = 1100000
        for cell in table.columns[2].cells:
            cell.width = 1000000
        cellp = table.cell(0, 0).merge(table.cell(0, 1)).paragraphs[0]
        cellp = table.cell(0, 1).merge(table.cell(0, 2)).paragraphs[0]
        cellp.add_run('Optativas').font.bold = True
        cellp.runs[0].font.size = Pt(8)
        
        table.cell(1, 0).paragraphs[0].add_run('Nombre de la Agrupación').font.size = Pt(8)
        table.cell(1, 1).paragraphs[0].add_run('Créditos Requeridos').font.size = Pt(8)
        table.cell(1, 2).paragraphs[0].add_run('Créditos pendientes por cursar ' +
            'por el estudiante').font.size = Pt(8)
        
        for i in range(3):
            table.cell(1, i).paragraphs[0].runs[0].font.bold = True

        if cont_fund_OP != 0:
            idx = 0
            sum_total = 0

            for index in fund_OP_index:
                table.cell(idx + 2, 0).paragraphs[0].add_run(
                    group_data[index][0]).font.size = Pt(8)
                table.cell(idx + 2, 1).paragraphs[0].add_run(
                    str(group_data[index][1])).font.size = Pt(8)

                sum_total = sum_total + group_data[index][1]
                idx = idx +1

            if cont_fund_OP == 1:
                table.cell(2, 2).paragraphs[0].add_run(str( sum_total )).font.size = Pt(8)
            else:
                for index in range(cont_fund_OP - 1):
                    cellp = table.cell(index + 2, 2).merge(table.cell(index + 3, 2)).paragraphs[0]    
                cellp.add_run(str( sum_total )).font.size = Pt(8)

        self.pending_credits_subjects.append(sum_total)
        cellp = table.cell(cont_fund_OP+2, 0).merge(table.cell(cont_fund_OP+2, 1)).paragraphs[0]
        cellp.add_run('Total créditos pendientes').font.bold = True
        cellp.runs[0].font.size = Pt(8)
        cellp.runs[0].font.highlight_color = WD_COLOR_INDEX.YELLOW
        table.cell(cont_fund_OP+2, 2).paragraphs[0].add_run(str( sum_total )).font.size = Pt(8)


        paragraph = docx.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run('\n').font.size = Pt(8)

        # Disciplinar table
        cont_disc_OB = 0
        disc_OB_index = []
        aux_group = []
        for i in range (len(data)):
            if data[i][2] == 'C':
                cont_disc_OB = cont_disc_OB + 1
                disc_OB_index.append(i)

        table = docx.add_table(rows=cont_disc_OB+4, cols=5)
        for column in table.columns:
            for cell in column.cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        table.style = 'Table Grid'
        table.style.font.size = Pt(9)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.columns[0].width = 1250000
        table.columns[1].width = 550000
        table.columns[2].width = 1850000
        table.columns[3].width = 550000
        table.columns[4].width = 1000000
        for cell in table.columns[0].cells:
            cell.width = 1250000
        for cell in table.columns[1].cells:
            cell.width = 550000
        for cell in table.columns[2].cells:
            cell.width = 1850000
        for cell in table.columns[3].cells:
            cell.width = 550000
        for cell in table.columns[4].cells:
            cell.width = 1000000
        cellp = table.cell(0, 0).merge(table.cell(0, 1)).paragraphs[0]
        cellp = table.cell(0, 1).merge(table.cell(0, 2)).paragraphs[0]
        cellp = table.cell(0, 2).merge(table.cell(0, 3)).paragraphs[0]
        cellp = table.cell(0, 3).merge(table.cell(0, 4)).paragraphs[0]
        cellp.add_run('Componente Disciplinar/Profesional (C)').font.bold = True
        cellp.runs[0].font.size = Pt(8)
        
        cellp = table.cell(1, 0).merge(table.cell(1, 1)).paragraphs[0]
        cellp = table.cell(1, 1).merge(table.cell(1, 2)).paragraphs[0]
        cellp = table.cell(1, 2).merge(table.cell(1, 3)).paragraphs[0]
        cellp = table.cell(1, 3).merge(table.cell(1, 4)).paragraphs[0]
        cellp.add_run('Obligatorias').font.bold = True
        cellp.runs[0].font.size = Pt(8)
        
        table.cell(2, 0).paragraphs[0].add_run('Agrupación').font.bold = True
        table.cell(2, 1).paragraphs[0].add_run('Código').font.bold = True
        table.cell(2, 2).paragraphs[0].add_run('Asignatura').font.bold = True
        table.cell(2, 3).paragraphs[0].add_run('Créditos asignatura').font.bold = True
        table.cell(2, 4).paragraphs[0].add_run('Créditos pendientes por cursar por ' +
            'el estudiante').font.bold = True
        for i in range(5):
            table.cell(2, i).paragraphs[0].runs[0].font.size = Pt(8)
        
        #Required to build table
        idx = 0
        sum_group = 0
        sum_total = 0
        cont_subject = 0
        merge_subjects = []
        credits_subjects = []

        for i in range(cont_disc_OB):
            if len(aux_group) == 0:
                var = data[disc_OB_index[i]][3]
                aux_group.append(var)
                table.cell(idx + 3, 1).paragraphs[0].add_run(
                    data[disc_OB_index[i]][0]).font.size = Pt(8)
                table.cell(idx + 3, 2).paragraphs[0].add_run(
                    data[disc_OB_index[i]][1]).font.size = Pt(8)
                table.cell(idx + 3, 3).paragraphs[0].add_run(
                    data[disc_OB_index[i]][4]).font.size = Pt(8)
                sum_group = sum_group + int(data[disc_OB_index[i]][4])
                cont_subject = cont_subject + 1
                idx = idx +1

                for j in range (i+1, cont_disc_OB):
                    if var == data[j][3]:
                        table.cell(idx + 3, 1).paragraphs[0].add_run(
                            data[disc_OB_index[j]][0]).font.size = Pt(8)
                        table.cell(idx + 3, 2).paragraphs[0].add_run(
                            data[disc_OB_index[j]][1]).font.size = Pt(8)
                        table.cell(idx + 3, 3).paragraphs[0].add_run(
                            data[disc_OB_index[j]][4]).font.size = Pt(8)
                        sum_group = sum_group + int(data[disc_OB_index[j]][4])
                        cont_subject = cont_subject + 1
                        idx = idx +1
                merge_subjects.append(cont_subject)
                credits_subjects.append(sum_group)
                sum_total = sum_total + sum_group
                cont_subject = 0
                sum_group = 0
            elif data[i][3] not in aux_group:
                var = data[disc_OB_index[i]][3]
                aux_group.append(var)
                table.cell(idx + 3, 1).paragraphs[0].add_run(
                    data[disc_OB_index[i]][0]).font.size = Pt(8)
                table.cell(idx + 3, 2).paragraphs[0].add_run(
                    data[disc_OB_index[i]][1]).font.size = Pt(8)
                table.cell(idx + 3, 3).paragraphs[0].add_run(
                    data[disc_OB_index[i]][4]).font.size = Pt(8)
                sum_group = sum_group + int(data[disc_OB_index[i]][4])
                cont_subject = cont_subject + 1
                idx = idx +1

                for j in range (i+1, cont_disc_OB):
                    if var == data[j][3]:
                        table.cell(idx + 3, 1).paragraphs[0].add_run(
                            data[disc_OB_index[j]][0]).font.size = Pt(8)
                        table.cell(idx + 3, 2).paragraphs[0].add_run(
                            data[disc_OB_index[j]][1]).font.size = Pt(8)
                        table.cell(idx + 3, 3).paragraphs[0].add_run(
                            data[disc_OB_index[j]][4]).font.size = Pt(8)
                        sum_group = sum_group + int(data[disc_OB_index[j]][4])
                        cont_subject = cont_subject + 1
                        idx = idx +1
                merge_subjects.append(cont_subject)
                credits_subjects.append(sum_group)
                sum_total = sum_total + sum_group
                cont_subject = 0
                sum_group = 0
            else:
                continue
        
        cont = 0
        aux = 0
        for i in merge_subjects:
            if i == 1:
                table.cell(cont + 3, 0).paragraphs[0].add_run(data[disc_OB_index[cont]][3]).font.size = Pt(8)
                table.cell(cont + 3, 4).paragraphs[0].add_run(str(credits_subjects[aux])).font.size = Pt(8)
                cont = cont + i 
                aux = aux + 1
            else:
                for a in range(i-1):
                    cellp3 = table.cell(a + cont + 3, 0).merge(table.cell(a + cont + 4, 0)).paragraphs[0]
                    cellp4 = table.cell(a + cont + 3, 4).merge(table.cell(a + cont + 4, 4)).paragraphs[0]
                cellp3.add_run(data[disc_OB_index[cont]][3]).font.size = Pt(8)
                cellp4.add_run(str(credits_subjects[aux])).font.size = Pt(8)
                aux = aux + 1
                cont = cont + i 

        self.pending_credits_subjects.append(sum_total)
        cellp = table.cell(cont_disc_OB+3, 0).merge(table.cell(cont_disc_OB+3, 1)).paragraphs[0]
        cellp = table.cell(cont_disc_OB+3, 1).merge(table.cell(cont_disc_OB+3, 2)).paragraphs[0]
        cellp = table.cell(cont_disc_OB+3, 2).merge(table.cell(cont_disc_OB+3, 3)).paragraphs[0]
        cellp.add_run('Total créditos pendientes').font.bold = True
        cellp.runs[0].font.size = Pt(8)
        cellp.runs[0].font.highlight_color = WD_COLOR_INDEX.YELLOW
        table.cell(cont_disc_OB+3, 4).paragraphs[0].add_run(str(sum_total)).font.size = Pt(8)


        # Second part of table
        cont_disc_OP = 0
        disc_OP_index = []
        for i in range (len(group_data)):
            if group_data[i][2] == 'T':
                cont_disc_OP = cont_disc_OP + 1
                disc_OP_index.append(i)

        table = docx.add_table(rows=cont_disc_OP+3, cols=3)
        for column in table.columns:
            for cell in column.cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        table.style = 'Table Grid'
        table.style.font.size = Pt(9)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.columns[0].width = 3100000
        table.columns[1].width = 1100000
        table.columns[2].width = 1000000
        for cell in table.columns[0].cells:
            cell.width = 3100000
        for cell in table.columns[1].cells:
            cell.width = 1100000
        for cell in table.columns[2].cells:
            cell.width = 1000000
        cellp = table.cell(0, 0).merge(table.cell(0, 1)).paragraphs[0]
        cellp = table.cell(0, 1).merge(table.cell(0, 2)).paragraphs[0]
        cellp.add_run('Optativas').font.bold = True
        cellp.runs[0].font.size = Pt(8)
        
        table.cell(1, 0).paragraphs[0].add_run('Nombre de la Agrupación').font.size = Pt(8)
        table.cell(1, 1).paragraphs[0].add_run('Créditos Requeridos').font.size = Pt(8)
        table.cell(1, 2).paragraphs[0].add_run('Créditos pendientes por cursar ' +
            'por el estudiante').font.size = Pt(8)
        
        for i in range(3):
            table.cell(1, i).paragraphs[0].runs[0].font.bold = True
        
        if cont_disc_OP != 0:
            idx = 0
            sum_total = 0

            for index in disc_OP_index:
                table.cell(idx + 2, 0).paragraphs[0].add_run(
                    group_data[index][0]).font.size = Pt(8)
                table.cell(idx + 2, 1).paragraphs[0].add_run(
                    str(group_data[index][1])).font.size = Pt(8)

                sum_total = sum_total + group_data[index][1]
                idx = idx +1

            if cont_fund_OP == 1:
                table.cell(2, 2).paragraphs[0].add_run(str( sum_total )).font.size = Pt(8)
            else:
                for index in range(cont_disc_OP - 1):
                    cellp = table.cell(index + 2, 2).merge(table.cell(index + 3, 2)).paragraphs[0]    
                cellp.add_run(str( sum_total )).font.size = Pt(8)

        self.pending_credits_subjects.append(sum_total)
        cellp = table.cell(cont_disc_OP+2, 0).merge(table.cell(cont_disc_OP+2, 1)).paragraphs[0]
        cellp.add_run('Total créditos pendientes').font.bold = True
        cellp.runs[0].font.size = Pt(8)
        cellp.runs[0].font.highlight_color = WD_COLOR_INDEX.YELLOW
        table.cell(cont_disc_OP+2, 2).paragraphs[0].add_run(str( sum_total )).font.size = Pt(8)
        paragraph = docx.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run('\n').font.size = Pt(8)

        # Free subjects
        table = docx.add_table(rows= 1, cols= 2)
        for column in table.columns:
            for cell in column.cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        table.style = 'Table Grid'
        table.style.font.size = Pt(9)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.columns[0].width = 4100000
        table.columns[1].width = 1100000
        for cell in table.columns[0].cells:
            cell.width = 4100000
        for cell in table.columns[1].cells:
            cell.width = 1100000

        table.cell(0, 0).paragraphs[0].add_run("Componente de Libre Elección (L) (Créditos pendientes)").font.size = Pt(8)
        table.cell(0, 1).paragraphs[0].add_run( str(self.free_elect_credit - self.equivalent_credits_subjects[4]) ).font.size = Pt(8)
        paragraph = docx.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run('\n').font.size = Pt(8)


    #Fifth table of docx
    def dtit_recommend_table(self, docx):
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        bullet = paragraph.add_run(self.str_pcm[7])
        bullet.font.bold = True
        bullet.font.size = Pt(8)
        credits_data = [[self.ob_fund_credit, self.op_fund_credit, self.ob_disc_credit, 
                            self.op_disc_credit, self.free_elect_credit],
                        [self.equivalent_credits_subjects[0], self.equivalent_credits_subjects[1], 
                            self.equivalent_credits_subjects[2], self.equivalent_credits_subjects[3], 
                                self.equivalent_credits_subjects[4]],
                        [self.pending_credits_subjects[0], self.pending_credits_subjects[1], 
                            self.pending_credits_subjects[2],self.pending_credits_subjects[3], 
                                self.free_elect_credit - self.equivalent_credits_subjects[4]]]
        case = 'DOBLE TITULACIÓN'
        table_credits_summary(docx, credits_data, case)

        
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        bullet = paragraph.add_run(self.str_pcm[8])
        bullet.font.size = Pt(8)
        
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        bullet = paragraph.add_run(self.str_pcm[9])
        bullet.font.size = Pt(8)

        details = []
        consFac = "Consejo de la Facultad de " + self.get_academic_program_display() 
        details.append(
            # pylint: disable=no-member
            consFac
            )

        # Migrate to case_utils?
        details.append(str( self.received_date.day ) + '-' + str(self.received_date.month) + '-' + str(self.received_date.year))
        details.append(self.consecutive_minute)
        details.append(self.year)
        if self.advisor_response == self.ARCR_APROBAR:
            details.append(True)
        else:
            details.append(False)

        table_recommend(docx,details)