import datetime
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from mongoengine import FloatField, BooleanField, DateField
from ..models import Request
from .case_utils import add_analysis_paragraph


class EMSP(Request):

    full_name = 'Exención resultados destacados en SABER PRO'

    regulation_list = ['002|2011|CFA']  # List of regulations

    percentaje = FloatField(
        display='Porcentaje de exención de matrícula', default=0.0)
    is_renovation = BooleanField(
        display='Es renovación de beca', default=False)
    is_best = BooleanField(
        display='¿Es el mejor puntaje?', default=False)
    date_presentation = DateField(display='Fecha de presentación del examen',
                                  default=datetime.date.today)

    str_cm = [
        'BECA EXCENCIÓN DE DERECHOS ACADÉMICOS del programa {} ({}) por obtener un excelente resu' +
        'ltado en el exámen de estado SABER-PRO.'
    ]

    str_pcm_modifiers = [
        'la renovación de ',
        'al mejor puntaje',
        'a los 10 mejores puntajes',
    ]

    str_pcm = [
        'la exención del {}% del pago de los derechos académicos y renovación de matrícula {} de' +
        'l Examen de Estado de la Calidad de la Educación Superior (SABER PRO) {} - {} para el pe' +
        'riodo académico {}{}. (Literal b, Artículo 16 del {}).', 'debido a que '
    ]

    def cm(self, docx):
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.str_council_header + ' ')
        self.cm_answer(paragraph)

    def cm_answer(self, paragraph):
        paragraph.add_run(
            # pylint: disable=no-member
            self.get_approval_status_display().upper() + ' ').font.bold = True
        paragraph.add_run(
            self.str_pcm[0].format(
                # pylint: disable=no-member
                self.percentaje,
                self.str_pcm_modifiers[1] if self.is_best else self.str_pcm_modifiers[2],
                self.date_presentation.year,
                self.get_academic_program_display(),
                self.academic_period,
                '' if self.is_affirmative_response_approval_status() else (
                    ' ' + self.str_pcm[1] + self.council_decision),
                self.regulations['002|2011|CFA'][0]
            ))

    def pcm(self, docx):
        self.pcm_analysis(docx)
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.str_answer + ': ').bold = True
        paragraph.add_run(self.str_comittee_header + ' ')
        self.pcm_answer(paragraph)

    def pcm_answer(self, paragraph):
        paragraph.add_run(
            # pylint: disable=no-member
            self.get_advisor_response_display().upper() + ' ').font.bold = True
        paragraph.add_run(
            self.str_pcm_modifiers[0] if self.is_renovation else '')
        paragraph.add_run(
            self.str_pcm[0].format(
                # pylint: disable=no-member
                self.percentaje,
                self.str_pcm_modifiers[1] if self.is_best else self.str_pcm_modifiers[2],
                self.date_presentation.year,
                self.get_academic_program_display(),
                self.academic_period,
                '' if self.is_affirmative_response_advisor_response() else (
                    ' ' + self.str_pcm[1] + self.council_decision),
                self.regulations['002|2011|CFA'][0]
            ))

    def pcm_analysis(self, docx):
        analysis_list = []
        analysis_list += self.extra_analysis
        add_analysis_paragraph(docx, analysis_list)

    def resource_analysis(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.pcm_answer(last_paragraph)

    def resource_pre_answer(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.pcm_answer(last_paragraph)

    def resource_answer(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.cm_answer(last_paragraph)
