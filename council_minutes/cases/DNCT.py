import datetime
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from mongoengine import StringField, DateField, BooleanField
from mongoengine import EmbeddedDocumentListField
from ..models import Request, Professor
from .case_utils import add_analysis_paragraph

class DNCT(Request):

    full_name = 'Designación de jurados para propuesta de tesis y de examen de calificación doctoral'
    decision_maker = Request.decision_makers[1]
    in_cm = False

    student_justification = StringField(
        default='El estudiante pide designación de jurados.', display='Justificación del Estudiante')
    supports = StringField(
        default='Copia digital de proyecto de tesis, solicitud de nombramiento de jurados.', display='Soportes')

    has_thesis_project = BooleanField(
        default=False, display='Tiene inscrita la asignatura Proyecto de Tesis de Doctorado')
    thesis_name = StringField(
        default='', display='Título de la tesis')

    code_exam = StringField(
        default='', display='Código de asignatura Examen de calificación doctoral')
    code_thesis = StringField(
        default='', display='Código de asignatura Tesis de doctorado')

    proffesors_thesis = EmbeddedDocumentListField(
        Professor, display='Docentes jurados evaluadores del proyecto de tesis')

    regulation_list = ['056|2012|CSU', '040|2017|CFA']

    str_cm = []

    str_pcm = [
        'El estudiante {}tiene inscrita la asignatura Proyecto de tesis de Doctorado ({}).',
        'Versión electrónica en formato pdf (Artículo 35).',
        'El documento Proyecto de Tesis de Doctorado será evaluado por un grupo ' +
        'de evaluadores conformado por mínimo tres integrantes, designados por el ' +
        'Comité Asesor de Posgrado. (Artículo 36). Jurados propuestos: NO ',
        'El estudiante deberá realizar una sustentación pública de su Proyecto de '
        'Tesis de Doctorado ante los evaluadores. En la sustentación deberán ' +
        'participar, presencialmente o mediante video conferencia el estudiante, ' +
        'los evaluadores, el profesor tutor del estudiante y un profesor activo de la ' +
        'Universidad Nacional de Colombia delegado por el Comité Asesor de ' +
        'Posgrado, quien hará las veces de coordinador de la sustentación (Artículo ' +
        '37).',
        'En el jurado evaluador del proyecto de tesis del {} con título: “{}”, a los '
        'profesores {}.',
        'En el jurado evaluador del examen de calificación del {}, al (a los) profesor(es) {}',
    ]

    def cm(self, docx):
        raise NotImplementedError

    def cm_answer(self, paragraph):
        raise NotImplementedError

    def pcm(self, docx):
        self.pcm_analysis_handler(docx)
        self.pcm_answer_handler(docx)

    def pcm_answer(self, paragraph):
        paragraph.add_run(self.str_answer + ':\n').font.bold = True
        paragraph.add_run('El Comité Asesor ')
        paragraph.add_run('DESIGNA:').font.bold = True

    def add_proffesors(self, professors):
        answer = ''
        for index in range(len(professors)-2):
            answer = answer + self.add_proffesor(professors[index])
            answer += ', '
        if len(professors) > 1:
            answer = answer + self.add_proffesor(professors[-2])
            answer += ' y el profesor(a) '
        if len(professors) != 0:
            answer = answer + self.add_proffesor(professors[-1])
        return answer
        
    def add_proffesor(self, professor):
        answer = ''
        answer = answer + str(professor.name)
        answer = answer + ' del departamento ' + str(professor.get_department_display())
        answer = answer + ' de la institución ' + str(professor.institution)
        return answer

    def pcm_answer_handler(self, docx):
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        self.pcm_answer(paragraph)
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.str_pcm[4].format(
            self.get_academic_program_display(), 
            self.thesis_name,
            self.add_proffesors(self.proffesors_thesis)
        ))
        paragraph.style = 'List Bullet'
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.str_pcm[5].format(
            self.get_academic_program_display(), 
            self.add_proffesors(self.proffesors_thesis)
        ))
        paragraph.style = 'List Bullet'

    def pcm_analysis_handler(self, docx):
        analysis = self.pcm_analysis_phd()
        add_analysis_paragraph(docx, analysis + self.extra_analysis)

    def pcm_analysis_phd(self):
        return [
            self.str_pcm[0].format('' if self.has_thesis_project else 'no ', self.code_thesis),
            self.str_pcm[1],
            self.str_pcm[2],
            self.str_pcm[3]
        ]

    def resource_analysis(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.pcm_answer(last_paragraph)
    
    def resource_pre_answer(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.pcm_answer(last_paragraph)

    def resource_answer(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.cm_answer(last_paragraph)
