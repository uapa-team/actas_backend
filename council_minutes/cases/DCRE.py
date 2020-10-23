from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from mongoengine import EmbeddedDocumentListField
from ..models import Request, Subject
from .case_utils import add_analysis_paragraph, table_subjects


class DCRE(Request):

    full_name = 'Reintegro de Créditos'

    subjects = EmbeddedDocumentListField(Subject, display='Asignaturas')

    regulation_list = ['001|2019|VSB', '230|2016|CSU']

    str_cm = [
        'reintegrar al cupo el total de {} céditos descontados por la cancelación de la(s) ' +
        'siguiente(s) asignatura(s) en el periodo académico {}.'
    ]

    str_pcm = [
        'reintegrar al cupo el total de {} céditos descontados por la cancelación de la(s) ' +
        'siguiente(s) asignatura(s) en el periodo académico {}.'
    ]

    def cm(self, docx):
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.str_council_header + ' ')
        self.cm_answer(paragraph)
        table_subjects(docx, Subject.subjects_to_array(self.subjects))

    def cm_answer(self, paragraph):
        # pylint: disable=no-member
        paragraph.add_run(
            self.get_approval_status_display().upper() + ' ').font.bold = True
        paragraph.add_run(self.str_cm[0].format(
            self.total_credits(), self.academic_period))
        paragraph.add_run(' ({}, {}).'.format(
            self.regulations[self.regulation_list[0]][0],
            self.regulations[self.regulation_list[1]][0]
        ))

    def pcm(self, docx):
        add_analysis_paragraph(docx, self.extra_analysis)
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.str_answer + ': ').font.bold = True
        paragraph.add_run(self.str_comittee_header + ' ')
        self.pcm_answer(paragraph)
        table_subjects(docx, Subject.subjects_to_array(self.subjects))

    def pcm_answer(self, paragraph):
        paragraph.add_run(
            # pylint: disable=no-member
            self.get_advisor_response_display().upper() + ' ').font.bold = True
        paragraph.add_run(self.str_pcm[0].format(
            self.total_credits(), self.academic_period))
        paragraph.add_run(' ({}, {}).'.format(
            self.regulations[self.regulation_list[0]][0],
            self.regulations[self.regulation_list[1]][0]
        ))

    def total_credits(self):
        total = 0
        for subject in self.subjects:
            total += subject.credits
        return total

    def resource_analysis(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.pcm_answer(last_paragraph)
    
    def resource_pre_answer(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.pcm_answer(last_paragraph)

    def resource_answer(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.cm_answer(last_paragraph)
