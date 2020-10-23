from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from ..models import Request
from .case_utils import add_analysis_paragraph


class DEHA(Request):

    full_name = 'Desbloquear historia académica'

    str_cm = [
        'Desbloqueo de la historia académica en el programa {} ({}).',
        'Restablecimiento de la carga académica de la historia académica del programa {} ({}).'
    ]

    regulation_list = ['070|2012|CSU']  # List of regulations

    def cm(self, docx):
        # pylint: disable=no-member
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.str_council_header + ' ')
        paragraph.add_run(
            self.get_approval_status_display().upper() + ':').font.bold = True
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.style = 'List Bullet'
        self.cm_answer(paragraph)
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.style = 'List Bullet'
        self.cm_retrieve(paragraph)

    def cm_answer(self, paragraph):
        # pylint: disable=no-member
        paragraph.add_run(self.str_cm[0].format(
            self.get_academic_program_display(),
            self.academic_program))

    def cm_retrieve(self, paragraph):
        # pylint: disable=no-member
        paragraph.add_run(self.str_cm[1].format(
            self.get_academic_program_display(),
            self.academic_program))

    def pcm(self, docx):
        # pylint: disable=no-member
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.str_comittee_header + ' ')
        paragraph.add_run(
            self.get_advisor_response_display().upper() + ':').font.bold = True
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.style = 'List Bullet'
        self.pcm_answer(paragraph)
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.style = 'List Bullet'
        self.pcm_retrieve(paragraph)

    def pcm_analysis(self, docx):
        analysis_list = []
        analysis_list += self.extra_analysis
        add_analysis_paragraph(docx, analysis_list)

    def pcm_answer(self, paragraph):
        self.cm_answer(paragraph)

    def pcm_retrieve(self, paragraph):
        self.cm_retrieve(paragraph)

    def resource_analysis(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.pcm_answer(last_paragraph)
    
    def resource_pre_answer(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.pcm_answer(last_paragraph)

    def resource_answer(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.cm_answer(last_paragraph)
