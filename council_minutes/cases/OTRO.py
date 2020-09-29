from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from mongoengine import ListField, StringField
from ..models import Request


class OTRO(Request):

    full_name = 'Otros casos'

    regulation_list = []  # List of regulations

    comitee_answer = StringField(display='Texto respuesta Comité', default='')

    comitee_aux = ListField(
        StringField(), display='Parrafos adicionales Comité Asesor')

    council_answer = StringField(display='Texto respuesta Consejo', default='')

    council_aux = ListField(
        StringField(), display='Parrafos adicionales Consejo Facultad')

    def cm(self, docx):
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.str_council_header + ' ')
        self.cm_answer(paragraph)
        
        for p in self.council_aux:
            paragraph = docx.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.add_run(p)

    def cm_answer(self, paragraph):
        paragraph.add_run(
            # pylint: disable=no-member
            self.get_approval_status_display().upper() + ' ').font.bold = True
        paragraph.add_run(' ' + self.council_answer)

    def pcm(self, docx):
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.str_answer + ': ').bold = True
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.str_comittee_header + ' ')
        self.pcm_answer(paragraph)

        for p in self.comitee_aux:
            paragraph = docx.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.add_run(p)

    def pcm_answer(self, paragraph):
        paragraph.add_run(
            # pylint: disable=no-member
            self.get_advisor_response_display().upper()).font.bold = True
        paragraph.add_run(' ' + self.comitee_answer)
