from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from mongoengine import StringField, EmbeddedDocumentListField
from ..models import Request, Subject
from .case_utils import table_change_typology, add_analysis_paragraph


class ChangeTipologySubject(Subject):
    new_tipology = StringField(
        required=True, choices=Subject.TIP_CHOICES, display='Nuevo componente',
        default=Subject.TIP_PRE_FUND_OBLIGATORIA)
    grade = StringField(display='Nota obtenida', default='')


class CTIP(Request):

    full_name = 'Cambio de componente'

    subjects_change_tipology = EmbeddedDocumentListField(
        ChangeTipologySubject, display='Asignaturas')

    regulation_list = ['008|2008|CSU']


    CJT_ANSWER_JUST_DEB = 'JD'
    CJT_ANSWER_OTRO = 'OT'
    CJT_ANSWER_N_DEB = 'ND'

    CJT_ANSWER_CHOICES = (
        (CJT_ANSWER_JUST_DEB,'Justifica debidamente su solicitud'),
        (CJT_ANSWER_N_DEB,'No justifica debidamente su solicitud'),
        (CJT_ANSWER_OTRO, 'Otro')
    )

    CJT_ANSWERS_DICT = {
        CJT_ANSWER_JUST_DEB : 'justifica debidamente su solicitud',
        CJT_ANSWER_N_DEB : 'no justifica debidamente su solicitud',
        CJT_ANSWER_OTRO : 'existen otros factores que lo justifican'
    }

    council_decision = StringField(
        max_length=255, choices=CJT_ANSWER_CHOICES,
        default=CJT_ANSWER_N_DEB, display='Justificación del Consejo')

    str_cm = [
        'cambiar de componente la(s) siguiente(s) asignatura(s) del programa {} ({}), cursada en ' +
        'el periodo académico {}',
        'debido a que {}.'
    ]

    str_pcm = [
        'Se solicita cambiar la tipología de la asignatura {} ({}). Tipología original: {}. Tipo' +
        'logía destino: {}. Nota obtenida: {}.'
    ]

    def cm(self, docx):
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.str_council_header + ' ')
        self.cm_answer(paragraph)
        paragraph.add_run(', ' + self.str_cm[1].format(self.CJT_ANSWERS_DICT[self.council_decision]))
        if self.is_affirmative_response_approval_status():
            self.add_subjects_change_tipology_table(docx)

    def cm_answer(self, paragraph):
        paragraph.add_run(
            # pylint: disable=no-member
            self.get_approval_status_display().upper() + ' ').font.bold = True
        paragraph.add_run(self.str_cm[0].format(
            # pylint: disable=no-member
            self.get_academic_program_display(),
            self.academic_program,
            self.academic_period
        ))

    def pcm(self, docx):
        self.pcm_analysis(docx)
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.str_comittee_header + ' ')
        self.pcm_answer(paragraph)
        paragraph.add_run(', ' + self.str_cm[1].format(self.CJT_ANSWERS_DICT[self.council_decision]))
        if self.is_affirmative_response_advisor_response() or self.is_waiting_response_advisor_response():
            self.add_subjects_change_tipology_table(docx)

    def pcm_answer(self, paragraph):
        paragraph.add_run(
            # pylint: disable=no-member
            self.get_advisor_response_display().upper() + ' ').font.bold = True
        paragraph.add_run(self.str_cm[0].format(
            # pylint: disable=no-member
            self.get_academic_program_display(),
            self.academic_program,
            self.academic_period
        ))

    def pcm_analysis(self, docx):
        analysis_list = []
        for subject in self.subjects_change_tipology:
            analysis_list += [self.str_pcm[0].format(
                subject.name,
                subject.code,
                subject.tipology[1],
                subject.new_tipology[1],
                subject.grade
            )]
        analysis_list += self.extra_analysis
        add_analysis_paragraph(docx, analysis_list)

    def add_subjects_change_tipology_table(self, docx):
        subjects = []
        for subject in self.subjects_change_tipology:
            subjects.append([
                subject.code,
                subject.name,
                subject.grade,
                subject.tipology[1],
                subject.new_tipology[1]
            ])
        table_change_typology(docx, subjects)

    def resource_analysis(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.pcm_answer(last_paragraph)
    
    def resource_pre_answer(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.pcm_answer(last_paragraph)

    def resource_answer(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.cm_answer(last_paragraph)
