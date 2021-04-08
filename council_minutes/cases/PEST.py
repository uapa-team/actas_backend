from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from mongoengine import StringField, IntField, FloatField, BooleanField
from ..models import Request, Subject
from .case_utils import add_analysis_paragraph, table_subjects


class PEST(Request):

    full_name = 'Inscripción práctica estudiantil'

    SUB_P1 = 'P1'
    SUB_P2 = 'P2'
    SUB_P3 = 'P3'
    SUB_P4 = 'P4'
    SUB_P5 = 'P5'
    SUB_P6 = 'P6'
    SUBJECT_CHOICES = (
        (SUB_P1, 'Práctica Estudiantil I'),
        (SUB_P2, 'Práctica Estudiantil II'),
        (SUB_P3, 'Práctica Estudiantil III'),
        (SUB_P4, 'Práctica Colombia I'),
        (SUB_P5, 'Práctica Colombia II'),
        (SUB_P6, 'Práctica Colombia III')
    )
    SUBJECT_INFO = {
        SUB_P1: ('2016762', 3,'Práctica Estudiantil I'),
        SUB_P2: ('2016763', 6,'Práctica Estudiantil II'),
        SUB_P3: ('2016764', 9,'Práctica Estudiantil III')
    }

    institution = StringField(required=True, display='Institución/Empresa', default='')
    is_intern = BooleanField(required=True, display='¿Es práctica interna?', default=False)
    professor = StringField(required=True, display='Profesor', default='')
    ins_person = StringField(required=True, display='Encargado Institución', default='')
    subject = StringField(required=True, choices=SUBJECT_CHOICES,
                          default=SUB_P1, display='Asignatura')
    group = StringField(required=True, display='Grupo', default='0')
    advance = FloatField(required=True, min_value=0, display='Avance SIA', default=0.0)
    another_practice = BooleanField(
        required=True, display='¿Primera práctica?', default=False)  
    another_practice_institution = StringField(required=False,display='Empresas de prácticas anteriores (si aplica)',
        default='')
    hours = IntField(required=True, min_value=0, display='Horas Semana', default=0)
    duration = StringField(required=True, display='Duración', default='')
    documentation = BooleanField(
        required=True, display='¿Documentación Completa?', default=True)
    creditsDisc = IntField(required=True, min_value=0, 
        display='Créditos Disciplinares (Únicamente Ingeniería de Sistemas e Industrial)', default=0)

    regulation_list = ['008|2008|CSU', '102|2013|CSU', '016|2011|CAC']

    str_cm = [
        'inscribir la asignatura ',
        'en el periodo académico {}, a desarrollar en la empresa {}, a cargo del docente ' +
        '{} por parte de la Universidad Nacional de Colombia',
        ' y {} por parte de la entidad',
        'debido a que {} ({}).'
    ]

    str_analysis = [
        'El estudiante {}cumple con el requisito de haber aprobado el ' +
        '70% de los créditos del plan de estudios. SIA: {:0.1f}% de avance en ' +
        'los créditos exigidos del plan de estudios.',
        'El estudiante {}ha cursado otra de las asignaturas con ' +
        'el nombre Práctica Estudiantil.',
        'Requisitos: Pertinencia, objetivos, alcance, empresa {}, duración: {} ' +
        'horas/semana durante {}, costos, descripción de actividades ' +
        'a cargo de un profesor de la Facultad: {}, porcentajes de evaluación definidos ' +
        '(Artículo 3 del {}).',
        'Documentación {}cumple con requisitos: Formato está completamente diligenciado, ' +
        'adjunta copia del Acuerdo firmado, ' +
        'adjunta el recibido de la carta de presentación de la Universidad, ' +
        'están fijados los porcentajes de evaluación.'
    ]

    # Analysis to Sistemas & Industrial
    str_system_industrial_analysis = [
        'El estudiante {}cumple con el requisito de haber aprobado los ' +
        'créditos del plan de estudios. Créditos: {} aprobados ' +
        'del componente disciplinar exigidos del plan de estudios.',
        'El estudiante {}ha cursado otra de las asignaturas con ' +
        'el nombre Práctica Estudiantil.',
        'Requisitos: Pertinencia, objetivos, alcance, empresa {}, duración: {} ' +
        'horas/semana durante {}, costos, descripción de actividades ' +
        'a cargo de un profesor de la Facultad: {}, porcentajes de evaluación definidos ' +
        '(Artículo 3 del {}).',
        'Documentación {}cumple con requisitos: Formato está completamente diligenciado, ' +
        'adjunta copia del Acuerdo firmado, ' +
        'adjunta el recibido de la carta de presentación de la Universidad, ' +
        'están fijados los porcentajes de evaluación.'
    ]

    str_pcm = []

    def cm(self, docx):
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.str_council_header + ' ')
        self.cm_answer(paragraph)
        self.add_table(docx)

    def cm_answer(self, paragraph):
        # pylint: disable=no-member
        paragraph.add_run(
            self.get_approval_status_display().upper() + ' ').font.bold = True
        self.add_text(
            paragraph, self.is_affirmative_response_approval_status())

    def pcm(self, docx):
        if ((self.academic_program == str(2546)) or   
                (self.academic_program == str(2879))):
            add_analysis_paragraph(docx, self.add_system_industrial_analysis())
        else:
            add_analysis_paragraph(docx, self.add_analysis())
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.str_answer + ' ').font.bold = True
        paragraph.add_run(self.str_comittee_header + ' ')
        self.pcm_answer(paragraph)
        #self.add_table(docx)

    def pcm_answer(self, paragraph):
        paragraph.add_run(
            # pylint: disable=no-member
            self.get_advisor_response_display().upper() + ' ').font.bold = True
        self.add_text(
            paragraph, self.is_affirmative_response_advisor_response())

    def add_text(self, paragraph, affirmative):
        code, _credits, name = self.SUBJECT_INFO[self.subject]
        # pylint: disable=no-member
        paragraph.add_run(self.str_cm[0]+str(name)+' ')

        if affirmative:
            paragraph.add_run(self.str_cm[1].format(
                self.academic_period, self.institution, self.professor
            ))
            if not self.is_intern:
                paragraph.add_run(self.str_cm[2].format(self.ins_person))
            paragraph.add_run('.')
        else:
            paragraph.add_run(self.str_cm[3].format(
                self.council_decision,
                self.regulations[self.regulation_list[1]][0]
            ))
    
    def add_table(self, docx):
        code, _credits, name = self.SUBJECT_INFO[self.subject]
        table_subjects(docx, [[
            code, self.get_subject_display(), self.group, 'L', str(_credits)
        ]])

    def add_analysis(self):
        analysis = []
        # Avance por porcentaje
        modifier = '' if self.advance >= 70 else 'no '

        analysis.append(self.str_analysis[0].format(modifier, self.advance))

        modifier = '' if not self.another_practice else 'no '
        analysis.append(self.str_analysis[1].format(modifier))

        analysis.append(self.str_analysis[2].format(
            self.institution, self.hours, self.duration, self.professor,
            self.regulations[self.regulation_list[2]][0]))

        modifier = '' if self.documentation else 'no '
        analysis.append(self.str_analysis[3].format(modifier))

        return analysis + self.extra_analysis

    def add_system_industrial_analysis(self):
        analysis = []

        # Avance por creditos disciplinares
        if ((self.creditsDisc >= 45 and  self.academic_program == str(2546)) or 
                (self.creditsDisc >= 40 and  self.academic_program == str(2879))):
            modifier = ''
        else:
            modifier = 'no '
                
        analysis.append(self.str_system_industrial_analysis[0].format(modifier, self.creditsDisc))

        modifier = '' if self.another_practice else 'no '
        analysis.append(self.str_system_industrial_analysis[1].format(modifier))

        analysis.append(self.str_system_industrial_analysis[2].format(
            self.institution, self.hours, self.duration, self.proffesor,
            self.regulations[self.regulation_list[2]][0]))

        modifier = '' if self.documentation else 'no '
        analysis.append(self.str_system_industrial_analysis[3].format(modifier))

        return analysis + self.extra_analysis

    def resource_analysis(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.pcm_answer(last_paragraph)
    
    def resource_pre_answer(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.pcm_answer(last_paragraph)

    def resource_answer(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.cm_answer(last_paragraph)
