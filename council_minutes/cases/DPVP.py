from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from mongoengine import StringField, FloatField, BooleanField
from ..models import Request
from .case_utils import add_analysis_paragraph


class DPVP(Request):

    full_name = 'Devolución proporcional del valor pagado por concepto de derechos de matrícula'

    percentage = FloatField(required=True, display='Porcentaje devuelto', default=0.0)
    cancel = BooleanField(
        required=True, display='¿Cancelación fue aprobada?', default=True)
    council_number = StringField(
        required=True, max_length=2, default='00', display='# Acta de cancelación')
    council_year = StringField(
        required=True, min_length=4, max_length=4, display='Año del Acta', default='0000')

    regulation_list = ['032|2010|CSU', '1416|2013|RE']

    str_cm = [
        'devolución proporcional del {:0.2f} % del valor pagado por concepto de ' +
        'derechos de matricula del periodo {}, teniendo en cuenta la fecha de ' +
        'presentación de la solicitud y que {}le fue aprobada la cancelación ' +
        'de periodo en Acta {} de {} de Consejo de Facultad.'
    ]
    str_pcm = []

    def cm(self, docx):
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.str_council_header + ' ')
        self.cm_answer(paragraph)

    def cm_answer(self, paragraph):
        paragraph.add_run(
            # pylint: disable=no-member
            self.get_approval_status_display().upper() + ' ').font.bold = True
        self.add_text(paragraph)

    def pcm(self, docx):
        add_analysis_paragraph(docx, self.extra_analysis)
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.str_answer + ' ').font.bold = True
        paragraph.add_run(self.str_comittee_header + ' ')
        self.pcm_answer(paragraph)

    def pcm_answer(self, paragraph):
        paragraph.add_run(
            # pylint: disable=no-member
            self.get_advisor_response_display().upper() + ' ').font.bold = True
        self.add_text(paragraph)

    def add_text(self, paragraph):
        modifier = '' if self.cancel else 'no '
        paragraph.add_run(self.str_cm[0].format(
            self.percentage,
            self.academic_period,
            modifier,
            self.council_number,
            self.council_year
        ))

    def resource_analysis(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.pcm_answer(last_paragraph)
    
    def resource_pre_answer(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.pcm_answer(last_paragraph)

    def resource_answer(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.cm_answer(last_paragraph)
