from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from mongoengine import StringField, BooleanField, DateField, IntField
from mongoengine import EmbeddedDocumentListField, FloatField, EmbeddedDocument
from ..models import Request, Subject
from .case_utils import add_analysis_paragraph


class TRAS(Request):

    class HomologatedSubject(Subject):
        group = StringField()
        name2 = StringField(required=True, display='Nuevo Nombre Asignatura')
        code2 = StringField(required=True, display='Nuevo Código')
        agroup = StringField(required=True, display='Agrupación')
        grade = FloatField(min_value=0.0, required=True, display='Nota')
        period = StringField(required=True, display='Periodo')

    class PendingSubject(EmbeddedDocument):
        TIP_FUNDAMENTACION = 'B'
        TIP_DISCIPLINAR = 'C'
        TIP_CHOICES = (
            (TIP_FUNDAMENTACION, 'Fundamentación'),
            (TIP_DISCIPLINAR, 'Disciplinar'),
        )
        group = StringField(required=True, display='Agrupación')
        name = StringField(required=True, display='Nombre Asignatura')
        code = StringField(required=True, display='Código')
        credits = IntField(required=True, display='Créditos')
        tipology = StringField(
            required=True, choices=TIP_CHOICES, display='Tipología')

    full_name = 'Traslado de programa curricular'

    TT_INTERCAMPUS = 'TTIC'
    TT_INTERFACULTY = 'TTIF'
    TT_INTRAFACULTY = 'TTRF'
    TT_CHOICES = (
        (TT_INTERCAMPUS, 'Traslado Intersede'),
        (TT_INTERFACULTY, 'Traslado Interfacultad'),
        (TT_INTRAFACULTY, 'Traslado Intrafacultad'),
    )
    TC_BOGOTA = 'BOG'
    TC_MEDELLIN = 'MED'
    TC_MANIZALES = 'MAN'
    TC_PALMIRA = 'PAL'
    TC_LAPAZ = 'PAZ'
    TC_CHOICES = (
        (TC_BOGOTA, 'Bogotá'),
        (TC_MEDELLIN, 'Medellín'),
        (TC_MANIZALES, 'Manizales'),
        (TC_PALMIRA, 'Palmira'),
        (TC_LAPAZ, 'La Paz'),
    )

    at_least_one_period = BooleanField(
        required=True, default=True,
        display='¿Ha cursado por lo menos un periodo académico del primer plan de estudios?')
    finish_first_plan = BooleanField(
        required=True, default=False,
        display='¿Ha culminado el primer plan de estudios?')
    have_entitled_to_enrrol = BooleanField(
        required=True, default=True,
        display='¿Tiene derecho a renovar matrícula?')
    enroled_number = IntField(
        required=True, min_value=1, default=1,
        display='Número de matrículas')
    currently_studying_double_degree = BooleanField(
        required=True, default=False,
        display='¿Está cursando actualmente doble titulación?')
    available_quota_for_transit = BooleanField(
        required=True, default=True,
        display='El cupo de créditos para traslado es suficiente?')
    availabe_quota_number = IntField(
        required=True, default=0, min_value=0,
        display='Número de cupos ofertados para traslado')
    campus_origin = StringField(
        required=True, default=TC_BOGOTA, choices=TC_CHOICES,
        display='Sede de origen')
    transit_type = StringField(
        required=True, default=TT_INTRAFACULTY, choices=TT_CHOICES,
        display='Tipo de traslado')
    admission_period = StringField(
        required=True, display='Periodo de admisión del estudiante')
    same_degree = BooleanField(
        required=True, default=False,
        display='¿Estos planes de estudios conducen al mismo título?')
    transit_program_code = StringField(
        required=True,
        display='Código del plan de estudios de destino')
    transit_program_name = StringField(
        required=True,
        display='Nombre del plan de estudios de destino')
    enrroled = BooleanField(
        required=True, default=True,
        display='¿Se encuentra matriculado en el semestre de presentar la solicitud?')
    prev_plan = BooleanField(
        required=True, default=False,
        display='¿Tuvo calidad de estudiante en el plan de estudios destino?')
    completion_percentage = FloatField(
        required=True, default=0.0, min_value=0.0, max_value=100.0,
        display='Porcentaje de créditos aprobados en el plan de estudios origen')
    student_admission_score = FloatField(
        default=600.0, min_value=0.0,
        display='Puntaje de admisión del solicitante')
    last_admitted_score = FloatField(
        default=600.0, min_value=0.0,
        display='Puntaje de admisión del último admitido')
    PAPA = FloatField(
        default=3.0, min_value=0.0, max_value=5.0,
        display='P.A.P.A.')
    PAPA_in_threshold = BooleanField(
        default=True,
        display='¿El P.A.P.A. se encuentra dentro de la franja del 30% de los mejores?')
    creds_miunus_remaining = IntField(
        required=True, default=0, min_value=0,
        display='Cupo de créditos menos pendientes en el plan de origen')
    creds_for_transit = IntField(
        required=True, default=0, min_value=0,
        display='Cupo de créditos para traslado')
    advisor_meeting_date = DateField(
        display='Fecha de reunión del comité')
    exiged_b_ob = IntField(
        min_value=0, default=0, required=True,
        display='Créditos exigidos fundamentación obligatorios')
    exiged_b_op = IntField(
        min_value=0, default=0, required=True,
        display='Créditos exigidos fundamentación optativos')
    exiged_c_ob = IntField(
        min_value=0, default=0, required=True,
        display='Créditos exigidos disciplinares obligatorios')
    exiged_c_op = IntField(
        min_value=0, default=0, required=True,
        display='Créditos exigidos disciplinares optativos')
    exiged_l = IntField(
        min_value=0, default=0, required=True,
        display='Créditos exigidos libre elección')
    equivalence = EmbeddedDocumentListField(
        HomologatedSubject, required=True,
        display="Asignaturas a homologar en el segundo plan de estudios")
    remaining = EmbeddedDocumentListField(
        PendingSubject, required=True,
        display="Asignaturas pendientes por cursar en el segundo plan de estudios")
    free_choice_pending = IntField(min_value=0, required=True,
                                   display='Créditos pendientes de libre elección')

    regulation_list = ['008|2008|CSU', '089|2014|CAC']  # List of regulations

    str_cm = ['']

    list_analysis = ['']

    def cm(self, docx):
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        self.cm_answer(paragraph)

    def cm_answer(self, paragraph):
        paragraph.add_run(self.str_council_header + ' ')
        # pylint: disable=no-member
        paragraph.add_run(
            self.get_approval_status_display().upper() + ' ').font.bold = True
        paragraph.add_run(
            self.str_cm[0] + self.get_grade_option_display() + ' ')
        paragraph.add_run(self.str_cm[1].format(self.title)).font.italic = True
        if self.is_affirmative_response_approval_status():
            self.cm_af(paragraph)
        else:
            self.cm_ng(paragraph)

    def pcm(self, docx):
        self.pcm_analysis(docx)
        self.pcm_answer(docx)

    def pcm_answer(self, docx):
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.str_answer + ': ').font.bold = True
        paragraph.add_run(self.str_comittee_header + ' ')
        paragraph.add_run(
            # pylint: disable=no-member
            self.get_advisor_response_display().upper()).font.bold = True
        # pylint: disable=no-member
        paragraph.add_run(' ' + self.str_cm[0].format(
            self.get_grade_option_display(), self.get_academic_program_display()))
        paragraph.add_run(self.str_cm[1].format(self.title)).font.italic = True
        if self.is_affirmative_response_approval_status():
            self.cm_af(paragraph)
        else:
            self.cm_ng(paragraph)

    def cm_af(self, paragraph):
        paragraph.add_run(' ' + self.str_cm[2])

    def cm_ng(self, paragraph):
        paragraph.add_run(
            ' ' + self.str_cm[3] + ' ' + self.council_decision + '.')

    def pcm_analysis(self, docx):
        final_analysis = []
        final_analysis += [self.list_analysis[3]]
        ets = ''
        # pylint: disable=no-member
        final_analysis += [self.list_analysis[4].format(
            ets, self.get_grade_option_display())]
        for extra_a in self.extra_analysis:
            final_analysis += [extra_a]
        add_analysis_paragraph(docx, final_analysis)
        paragraph = docx.add_paragraph()
        paragraph.style = 'List Bullet'
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.list_analysis[0] + ' ').font.bold = True
        paragraph.add_run(self.title + '.').font.italic = True
