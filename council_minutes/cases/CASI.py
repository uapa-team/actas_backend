from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from mongoengine import StringField, IntField, FloatField, EmbeddedDocumentListField
from ..models import Request, Subject
from .case_utils import table_subjects, add_analysis_paragraph


class CASI(Request):

    full_name = 'Cancelación de Asignaturas'

    CN_ANSWER_NO_DILIGENTE = 'ND'
    CN_ANSWER_MOTIVOS_LABORALES = 'ML'
    CN_ANSWER_INFORMACION_FALSA = 'TR'
    CN_ANSWER_SOPORTES_NO_SOPORTAN = 'SN'
    CN_ANSWER_FALTA_DE_CONOCIMIENTO = 'FC'
    CN_ANSWER_ARGUMENTOS_INSUFICIENTES = 'AI'
    CN_ANSWER_INCOHERENTE_O_INCONSECUENTE = 'II'
    CN_ANSWER_OTRO = 'OT'
    CN_ANSWER_CHOICES = (
        (CN_ANSWER_NO_DILIGENTE, 'No diligente'),
        (CN_ANSWER_MOTIVOS_LABORALES, 'Motivos Laborales'),
        (CN_ANSWER_INFORMACION_FALSA, 'Información Falsa'),
        (CN_ANSWER_SOPORTES_NO_SOPORTAN, 'Argumento cuando los soportes no soportan'),
        (CN_ANSWER_FALTA_DE_CONOCIMIENTO, 'Falta de conocimiento'),
        (CN_ANSWER_ARGUMENTOS_INSUFICIENTES, 'Argumentos insuficientes'),
        (CN_ANSWER_INCOHERENTE_O_INCONSECUENTE, 'Incoherente o no consecuente'),
        (CN_ANSWER_OTRO, 'Otro')
    )

    subjects = EmbeddedDocumentListField(
        Subject, display='Asignaturas')
    advance = FloatField(required=True, display='% de Avance', default=0.0)
    enrolled_academic_periods = IntField(
        required=True, display='# Periodos Matriculados', default=0)
    papa = FloatField(required=True, display='PAPA', default=0.0)
    available_credits = IntField(
        required=True, display='Créditos Disponibles', default=0)
    current_credits = IntField(
        required=True, display='Créditos Inscritos', default=0)
    nrc_answer = StringField(choices=CN_ANSWER_CHOICES, default=CN_ANSWER_NO_DILIGENTE,
                             display='Motivo de rechazo')

    regulation_list = ['008|2008|CSU']  # List of regulations

    str_cm = [
        'cancelar la(s) siguiente(s) asignatura(s) inscrita(s) del periodo académico {}',
        'porque {}justifica debidamente la solicitud.'
    ]

    str_pcm = [
        'SIA: Porcentaje de avance en el plan: {}. Número de matrículas: {}. PAPA: {}.',
        'SIA: Créditos disponibles: {}.',
        'SIA: Al aprobar la cancelación de la asignatura {} ({}) el estudiante quedaría con {} cr' +
        'éditos inscritos.'
    ]

    str_pcmap = [
        ' cancelar la(s) siguiente(s) asignatura(s) inscrita(s) del periodo académico {}, ',
    ]

    str_pcma_cap = [
        'porque se justifica debidamente la solicitud.'
    ]

    str_pcma_cna = [
        'porque no existe coherencia entre la documentación y justificación que presenta.',
        'porque lo expuesto es un hecho de su conocimiento desde el inicio del periodo académico' +
        '; tuvo la oportunidad de resolverlo oportunamente hasta el 50% del periodo académico, p' +
        'or tanto, no constituye causa extraña que justifique la cancelación de la(s) asignatura' +
        '(s).',
        'porque de acuerdo con la documentación que presenta, su situación laboral no le impide ' +
        'asistir a las clases y tiene el tiempo suficiente para responder por las actividades ac' +
        'adémicas de la(s) asignatura(s). ',
        'porque verificada la información de los soportes, se encontró que el contenido de los m' +
        'ismos no coincide con lo que en ellos se afirma.',
        'poque es responsabilidad del estudiante indagar sobre el conocimiento requerido y la pr' +
        'eparación necesaria para cursar la(s) asignatura(s) antes de inscribir.',
        'porque lo expuesto no es un hecho que constituya causa extraña que justifique la cancel' +
        'ación de la(s) asignatura(s).',
        'porque de la documentación aportada, se tiene que no hay justificación para acceder a l' +
        'o pedido. '
    ]

    def cm(self, docx):
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.str_council_header + ' ')
        self.cm_answer(paragraph)
        table_subjects(docx, Subject.subjects_to_array(self.subjects))

    def cm_answer(self, paragraph):
        paragraph.add_run(
            # pylint: disable=no-member
            self.get_approval_status_display().upper() + ' ').font.bold = True
        paragraph.add_run(self.str_cm[0].format(self.academic_period) + ', ')
        if self.is_affirmative_response_approval_status():
            self.cm_ap(paragraph)
        else:
            self.cm_na(paragraph)
        paragraph.add_run('({}).'.format(self.regulations['008|2008|CSU'][0]))

    def pcm(self, docx):
        self.pcm_analysis(docx)
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.str_answer + ': ').bold = True
        paragraph.add_run(self.str_comittee_header)
        self.pcm_answer(paragraph)
        table_subjects(docx, Subject.subjects_to_array(self.subjects))

    def pcm_answer(self, paragraph):
        paragraph.add_run(
            # pylint: disable=no-member
            ' ' + self.get_advisor_response_display().upper()).font.bold = True
        paragraph.add_run(self.str_pcmap[0].format(self.academic_period))
        if self.is_affirmative_response_advisor_response():
            self.pcm_answers_cr(paragraph)
        else:
            self.pcm_answers_cn(paragraph)

    def cm_ap(self, paragraph):
        if self.council_decision == Request.council_decision.default or len(self.council_decision) == 0:
            paragraph.add_run(self.str_cm[1].format('') + ' ')    
        else:
            paragraph.add_run(self.council_decision + ' ')

    def cm_na(self, paragraph):
        if self.council_decision == Request.council_decision.default or len(self.council_decision) == 0:
            paragraph.add_run(self.str_cm[1].format('no ') + ' ')
        else:
            paragraph.add_run(self.council_decision + ' ')

    def pcm_analysis(self, docx):
        analysis_list = []
        analysis_list += [self.str_pcm[0].format(
            self.advance, self.enrolled_academic_periods, self.papa)]
        analysis_list += [self.str_pcm[1].format(self.available_credits)]
        analysis_list += self.pcm_analysis_subject_list()
        analysis_list += self.extra_analysis
        add_analysis_paragraph(docx, analysis_list)

    def pcm_analysis_subject_list(self):
        analysis_subject_list = []
        for subject in self.subjects:
            current_credits = self.current_credits
            subject_credits = subject.credits
            subject_info = {
                'remaining': int(current_credits) - int(subject_credits),
                'code': subject.code,
                'name': subject.name
            }
            analysis_subject_list += [self.str_pcm[2].format(
                subject_info['name'], subject_info['code'], subject_info['remaining'])]
        return analysis_subject_list

    def pcm_answers_cr(self, paragraph):
        paragraph.add_run(self.str_pcma_cap[0])
        paragraph.add_run(' ({}).'.format(self.regulations['008|2008|CSU'][0]))

    def pcm_answers_cn(self, paragraph):
        if self.nrc_answer == self.CN_ANSWER_INCOHERENTE_O_INCONSECUENTE:
            paragraph.add_run(self.str_pcma_cna[0] + ' ')
        elif self.nrc_answer == self.CN_ANSWER_NO_DILIGENTE:
            paragraph.add_run(self.str_pcma_cna[1] + ' ')
        elif self.nrc_answer == self.CN_ANSWER_MOTIVOS_LABORALES:
            paragraph.add_run(self.str_pcma_cna[2] + ' ')
        elif self.nrc_answer == self.CN_ANSWER_INFORMACION_FALSA:
            paragraph.add_run(self.str_pcma_cna[3] + ' ')
        elif self.nrc_answer == self.CN_ANSWER_FALTA_DE_CONOCIMIENTO:
            paragraph.add_run(self.str_pcma_cna[4] + ' ')
        elif self.nrc_answer == self.CN_ANSWER_ARGUMENTOS_INSUFICIENTES:
            paragraph.add_run(self.str_pcma_cna[5] + ' ')
        elif self.nrc_answer == self.CN_ANSWER_SOPORTES_NO_SOPORTAN:
            paragraph.add_run(self.str_pcma_cna[6] + ' ')
        elif self.nrc_answer == self.CN_ANSWER_OTRO:
            paragraph.add_run(self.str_pcma_cna[6] + ' ')
        else:
            raise AssertionError(
                self.assertionerror['CHOICES'].format('NRC_answer'))
        paragraph.add_run(' ({}).'.format(self.regulations['008|2008|CSU'][0]))

    def resource_analysis(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.pcm_answer(last_paragraph)
        table_subjects(docx, Subject.subjects_to_array(self.subjects))
    
    def resource_pre_answer(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.pcm_answer(last_paragraph)
        table_subjects(docx, Subject.subjects_to_array(self.subjects))

    def resource_answer(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.cm_answer(last_paragraph)
        table_subjects(docx, Subject.subjects_to_array(self.subjects))
