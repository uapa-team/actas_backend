import datetime
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from mongoengine import StringField, BooleanField, DateField
from ..models import Request
from .case_utils import add_analysis_paragraph, string_to_date


class APAS(Request):

    full_name = 'Aprobación de pasantía'

    GO_TRABAJO_FINAL_MAESTRIA = 'TFM'
    GO_TESIS_MAESTRIA = 'TSM'
    GO_TESIS_DOCTORADO = 'TSD'
    GO_CHOICES = (
        (GO_TRABAJO_FINAL_MAESTRIA, 'Trabajo Final de Maestría'),
        (GO_TESIS_MAESTRIA, 'Tesis de Maestría'),
        (GO_TESIS_DOCTORADO, 'Tesis de Doctorado')
    )

    date_start = DateField(
        required=True, display='Fecha de inicio de la pasantía', default=datetime.date.today)
    date_finish = DateField(
        required=True, display='Fecha de finalización de la pasantía', default=datetime.date.today)
    place = StringField(
        required=True, display='Lugar: Ciudad, País', default='')
    format_present = BooleanField(
        required=True, default=False,
        display='¿Presenta el formato de solicitud de movilidad saliente?')
    institut = StringField(
        required=True, display='Institución donde se va a desarrollar la pasantía', default='')
    grade_option = StringField(
        required=True, display='Tipo de tesis/trabajo final', choices=GO_CHOICES,
        default=GO_TESIS_MAESTRIA)
    internship_period = StringField(
        required=True, display='Periodo en el que se va a desarrollar la pasantía',
        choices=Request.PERIOD_CHOICES, default=Request.PERIOD_DEFAULT)
    enrolled_thesis = BooleanField(
        required=True, default=False, display='¿Tiene inscrita la asignatura tesis/trabajo final?')

    regulation_list = []  # List of regulations

    str_cm = ['realizar pasantía en la institución {} en {}, como parte del desarrollo de su {}',
              ', desde el {} hasta el {}. El estudiante deberá estar debidamente matriculado' +
              ' en el periodo {}.',
              'debido a que']

    list_analysis = ['Perfil de {}',
                     'El estudiante {}tiene inscrita la asignatura {}.',
                     'Estancia de investigación del {} al {}.',
                     'Lugar: {}.',
                     'El estudiante {}presenta el formato de solicitud de movilidad saliente.']

    def cm(self, docx):
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.str_council_header + ' ')
        self.cm_answer(paragraph)

    def cm_answer(self, paragraph):
        # pylint: disable=no-member
        paragraph.add_run(
            self.get_approval_status_display().upper() + ' ').font.bold = True
        paragraph.add_run(self.str_cm[0].format(
            self.institut, self.place, self.get_grade_option_display()))
        if self.is_affirmative_response_approval_status():
            self.cm_af(paragraph)
        else:
            self.cm_ng(paragraph)

    def pcm(self, docx):
        self.pcm_analysis(docx)
        self.pcm_answer(docx)

    def pcm_answer(self, docx):
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.str_answer + ': ').font.bold = True
        paragraph.add_run(self.str_comittee_header + ' ')
        # pylint: disable=no-member
        paragraph.add_run(
            self.get_advisor_response_display().upper() + ' ').font.bold = True
        paragraph.add_run(self.str_cm[0].format(
            self.institut, self.place, self.get_grade_option_display()))
        if self.is_affirmative_response_approval_status():
            self.cm_af(paragraph)
        else:
            self.cm_ng(paragraph)

    def cm_af(self, paragraph):
        paragraph.add_run(self.str_cm[1].format(
            string_to_date(str(self.date_start)), string_to_date(
                str(self.date_finish)),
            self.internship_period,))

    def cm_ng(self, paragraph):
        paragraph.add_run(
            ' ' + self.str_cm[2] + ' ' + self.council_decision + '.')

    def pcm_analysis(self, docx):
        if self.grade_option in [self.GO_TESIS_MAESTRIA, self.GO_TESIS_DOCTORADO]:
            profile = 'investigación'
        else:
            profile = 'profundización'
        final_analysis = []
        final_analysis += [self.list_analysis[0].format(profile)]
        ets = '' if self.enrolled_thesis else 'no '
        # pylint: disable=no-member
        final_analysis += [self.list_analysis[1].format(
            ets, self.get_grade_option_display())]
        final_analysis += [self.list_analysis[2].format(
            string_to_date(str(self.date_start)),
            string_to_date(str(self.date_finish)))]
        final_analysis += [self.list_analysis[3].format(
            self.place)]
        pss = '' if self.format_present else 'no '
        final_analysis += [self.list_analysis[4].format(pss)]
        for extra_a in self.extra_analysis:
            final_analysis += [extra_a]
        add_analysis_paragraph(docx, final_analysis)

    def resource_analysis(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.pcm_answer(last_paragraph)
    
    def resource_pre_answer(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.pcm_answer(last_paragraph)

    def resource_answer(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.cm_answer(last_paragraph)
