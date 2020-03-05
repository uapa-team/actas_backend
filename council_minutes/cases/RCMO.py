from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from mongoengine import StringField, EmbeddedDocumentListField
from ..models import Request, Subject
from .case_utils import table_approvals, add_analysis_paragraph


class SubjectMovility(Subject):
    name_origin = StringField(
        required=True, display='Nombre Asignatura Origen')
    grade_origin = StringField(required=True, display='Nota obtenida')
    grade = StringField(required=True, display='Nota homologada')
    min_grade_origin = StringField(
        required=True, display='Mínima nota', default='0.0')
    max_grade_origin = StringField(
        required=True, display='Máxima nota', default='5.0')
    approval_grade_origin = StringField(
        required=True, display='Mínima nota aprobatoria', default='3.0')

    @staticmethod
    def subjects_to_table_array(subjects, period):
        """
        A function that converts a List of Subjects into a classic array.
        : param subjects: EmbeddedDocumentListField of Subjects to be converted
        """
        data = []
        for subject in subjects:
            data.append([
                period,
                subject.code,
                subject.name,
                str(subject.credits),
                subject.tipology[1],
                subject.grade,
                subject.name_origin,
                subject.grade_origin
            ])
        return data


class RCMO(Request):

    full_name = 'Registro de asignatura de movilidad'

    CALIFICATION_AP = 'AP'
    CALIFICATION_NA = 'NA'
    CALIFICATION_CHOICES = (
        (CALIFICATION_AP, 'Aprobada'),
        (CALIFICATION_NA, 'No aprobada')
    )

    calification = StringField(
        choices=CALIFICATION_CHOICES, display='Calificación Movilidad', default=CALIFICATION_AP)
    institution = StringField(display='Institución origen', default='')
    subject_code = StringField(display='Código asignatura', default='')
    subject_name = StringField(display='Nombre asignatura', default='')
    subject_period = StringField(
        display='Periodo asignatura', 
        choices=Request.PERIOD_CHOICES, default=Request.PERIOD_DEFAULT)
    subjects = EmbeddedDocumentListField(
        SubjectMovility, display='Asignaturas')

    str_cm = [
        'calificar {} la asignatura {} ({}) en el periodo {}.',
        'Homologar en el periodo académico {}, la(s) siguiente(s) asignatura(s) cursada(s) bajo l' +
        'a asignatura {}.',
    ]
    str_pcm = [
        'El estudiante realizó movilidad bajo la modalidad de la materia {} ({}) en el periodo {}' +
        ', en la institución {}.',
        'El estudiante cursó la asignatura {}, obteniendo una calificación {}, con nota mínima ap' +
        'robatoria {} en el rango de notas {} a {}, y se solicita homologar bajo el contenido de ' +
        'la materia {} ({}), con una nota {}; la cantidad de créditos {} y tipología {}.',
        'calificar {} la asignatura {} ({}) en el periodo {} y {} homologar en el periodo académi' +
        'co {}, la(s) asignatura(s) cursada(s) bajo la asignatura.',
    ]
    regulation_list = ['008|2008|CSU']  # List of regulations

    def cm(self, docx):
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if self.is_affirmative_response_approval_status():
            paragraph.add_run(self.str_council_header + ' ')
            self.cm_answer(docx.add_paragraph(style="List Bullet 2"))
            self.cm_answer_subjects(docx.add_paragraph(style="List Bullet 2"))
            self.cm_answer_subjects_table(docx)
        else:
            paragraph.add_run(self.str_council_header + ' ')
            self.cm_answer(paragraph)

    def cm_answer(self, paragraph):
        paragraph.add_run(
            # pylint: disable=no-member
            self.get_approval_status_display().upper() + ' ').font.bold = True
        paragraph.add_run(self.str_cm[0].format(
            # pylint: disable=no-member
            self.get_calification_display().upper(),
            self.subject_name,
            self.subject_code,
            self.subject_period
        ))

    def cm_answer_subjects(self, paragraph):
        paragraph.add_run(self.str_cm[1].format(
            self.academic_period, self.subject_name
        ))

    def cm_answer_subjects_table(self, docx):
        table_approvals(
            docx,
            SubjectMovility.subjects_to_table_array(
                self.subjects, self.academic_period),
            [self.student_name, self.student_dni,
             self.academic_program, self.institution]
        )

    def pcm(self, docx):
        self.pcm_analysis(docx)
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.str_answer + ': ').bold = True
        paragraph.add_run(self.str_comittee_header + ' ')
        self.pcm_answer(paragraph)

    def pcm_analysis(self, docx):
        analysis_list = []
        analysis_list += [self.str_pcm[0].format(
            self.subject_name,
            self.subject_code,
            self.subject_period,
            self.institution)]
        for subject in self.subjects:
            analysis_list += [self.str_pcm[1].format(
                subject.name_origin,
                subject.grade_origin,
                subject.approval_grade_origin,
                subject.min_grade_origin,
                subject.max_grade_origin,
                subject.name,
                subject.code,
                subject.grade,
                subject.credits,
                subject.get_tipology_display(),
            )]
        analysis_list += self.extra_analysis
        add_analysis_paragraph(docx, analysis_list)

    def pcm_answer(self, paragraph):
        paragraph.add_run(
            # pylint: disable=no-member
            self.get_advisor_response_display().upper()).font.bold = True
        paragraph.add_run(' ')
        paragraph.add_run(self.str_pcm[2].format(
            # pylint: disable=no-member
            self.get_calification_display().lower(),
            self.subject_name,
            self.subject_code,
            self.subject_period,
            self.get_advisor_response_display().upper(),
            self.subject_period))

    def resource_analysis(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.pcm_answer(last_paragraph)
    
    def resource_pre_answer(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.pcm_answer(last_paragraph)

    def resource_answer(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.cm_answer(last_paragraph)
