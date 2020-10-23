from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from mongoengine import StringField, FloatField
from ..models import Request
from .case_utils import add_analysis_paragraph


class EBDA(Request):

    full_name = 'Beca exención de derechos académicos'

    gpa = FloatField(
        display='Promedio obtenido el semestre anterior', default=0.0)
    gpa_period = StringField(
        display='Periodo en el que se obtiene el promedio',
        choices=Request.PERIOD_CHOICES,
        default=Request.PERIOD_DEFAULT)
    target_period = StringField(
        display='Periodo en el que aplica la beca',
        choices=Request.PERIOD_CHOICES,
        default=Request.PERIOD_DEFAULT)

    regulation_list = ['002|2012|CFA']  # List of regulations

    str_case = [
        'Se obtiene el promedio {}/5.0, en el periodo académico {} y se solicita la beca para el ' +
        'periodo académico {}.',
        'La coordinación curricular del programa presenta como beneficiario de la BECA EXENCIÓN D' +
        'E DERECHOS ACADÉMICOS del Acuerdo 2 de 2012 de Consejo de Facultad, por obtener el prome' +
        'dio académico ponderado más alto del semestre en las asignaturas cursadas durante el per' +
        'iodo académico inmediatamente anterior.',
        ' la BECA EXENCIÓN DE DERECHOS ACADÉMICOS en el programa de {} ({}) en el periodo {} y ot' +
        'orgar la exención del 100 % de derechos académicos por obtener el promedio académico pond' +
        'erado más alto del semestre en las asignaturas cursadas durante el periodo académico inm' +
        'ediatamente anterior.'
    ]

    def cm(self, docx):
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.str_council_header + ' ')
        self.cm_answer(paragraph)

    def cm_answer(self, paragraph):
        paragraph.add_run(
            # pylint: disable=no-member
            self.get_approval_status_display().upper() + ' ').font.bold = True
        paragraph.add_run(self.str_case[2].format(
            # pylint: disable=no-member
            self.get_academic_program_display(),
            self.academic_program,
            self.target_period
        ))

    def pcm(self, docx):
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        self.pcm_analysis(docx)
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.str_answer + ': ').font.bold = True
        paragraph.add_run(self.str_council_header + ' ')
        self.cm_answer(paragraph)

    def pcm_answer(self, paragraph):
        paragraph.add_run(
            # pylint: disable=no-member
            self.get_advisor_response_display().upper() + ' ').font.bold = True
        paragraph.add_run(self.str_case[2].format(
            # pylint: disable=no-member
            self.get_academic_program_display(),
            self.academic_program,
            self.target_period
        ))

    def pcm_analysis(self, docx):
        analysis_list = []
        analysis_list += [self.str_case[0].format(
            self.gpa,
            self.gpa_period,
            self.target_period
        )]
        analysis_list += [self.str_case[1].format()]
        analysis_list += self.extra_analysis
        add_analysis_paragraph(docx, analysis_list)

    def resource_analysis(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.pcm_answer(last_paragraph)
    
    def resource_pre_answer(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.pcm_answer(last_paragraph)

    def resource_answer(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.cm_answer(last_paragraph)
