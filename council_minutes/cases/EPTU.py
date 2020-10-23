from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from num2words import num2words
from mongoengine import StringField, IntField, BooleanField
from ..models import Request
from .case_utils import add_analysis_paragraph


class EPTU(Request):

    full_name = 'Exención de pago por cursar tesis como única actividad académica'

    CNA_MAXIMO_EXCEDIDO = 'ME'
    CNA_AVANCE_NO_SATISFACTORIO = 'NS'
    CNA_REINGRESO = 'RE'
    CNA_FUERA_PLAZOS = 'FP'
    CNA_OTRO = 'OT'
    CNA_CHOICES = (
        (CNA_MAXIMO_EXCEDIDO, 'Incentivo aplicado en más de dos/cinco periodos'),
        (CNA_AVANCE_NO_SATISFACTORIO, 'Periodo pasado con avance no satisfactorio'),
        (CNA_REINGRESO, 'Hubo reingreso en el programa'),
        (CNA_FUERA_PLAZOS, 'Solicitud realizada fuera de plazos'),
        (CNA_OTRO, 'Otro')
    )

    target_period = StringField(
        display='Periodo para el que se realiza la solicitud',
        choices=Request.PERIOD_CHOICES, default=Request.PERIOD_DEFAULT)
    points = IntField(display='Cantidad de puntos a eximir', default=0)
    academic_profile = StringField(
        default=Request.PROFILE_INVE, choices=Request.PROFILE_CHOICES,
        display='Perfil de programa curricular')
    right_dates = BooleanField(
        display='Realiza la solicitud en fecha adecuada', default=True)
    periods_in = IntField(display='Periodos de exención aplicada', default=0)
    cna = StringField(
        choices=CNA_CHOICES, default=CNA_OTRO, display='Motivo de rechazo')

    # List of regulations
    regulation_list = ['002|2011|CFA']

    str_cm = [
        'pago de {} ({}) puntos por derechos académicos en el periodo académico {},  condicionado' +
        ' a la inscripción de trabajo final de {} ({}) como única actividad académica en el perio' +
        'do {}',
        'El cálculo de los créditos disponibles se realiza con base en el cupo de créditos establ' +
        'ecido en el {}.'
    ]

    str_pcm = [
        'SIA: {} ({}), perfil de {}.',
        'La solicitud {}fue realizada en fechas adecuadas para ser válida.',
        'El incentivo se aplicará hasta por dos periodos académicos maestría y 5 periodos para do' +
        'ctorado (Literal c, Artículo 16). Ha tenido el incentivo {} periodos.',
        'Para solicitar o renovar este estímulo la evaluación por parte del director, del trabajo' +
        ' desarrollado por el estudiante, en la tesis o en el trabajo final, deberá ser "Avance S' +
        'atisfactorio", si aplica.',
        'Se exceptúa el caso en el cual un estudiante también cursa el último seminario del progr' +
        'ama simultáneamente con la tesis, en cuyo caso se le podrá conceder dicho estímulo.',
        'Estos beneficios no se podrán renovar si el estudiante pierde calidad de estudiante o en' +
        'tra en reserva de cupo.',
        'Las exenciones establecidas en el Artículo 16 de la presente Acuerdo, no podrán ser conc' +
        'edidas a estudiantes que hayan reingresado a programas de posgrado de la Facultad en cua' +
        'lquier época (Artículo 19).',
    ]

    str_pcma_cna = [
        'porque el incentivo sólo se aplicará hasta por dos/cinco periodos académicos.',
        'porque para solicitar o renovar este estímulo la evaluación por parte del director del t' +
        'rabajo desarrollado por el estudiante en la tesis o en el trabajo final deberá ser "Avan' +
        'ce Satisfactorio".',
        'porque las exenciones establecidas en el Artículo 16 del Acuerdo 002 de 2011 del Consejo' +
        ' de Facultad, no podrán ser concedidas a estudiantes que hayan reingresado a programas d' +
        'e posgrado de la Facultad en cualquier época.',
        'porque la solicitud se presenta fuera de los plazos establecidos en el Calendario de Tra' +
        'mites de Fin de Semestre {} - Estudiantes Posgrado.',
        'porque no realiza debidamente la solicitud.',
    ]

    def cm(self, docx):
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.str_council_header + ' ')
        self.cm_answer(paragraph)

    def cm_answer(self, paragraph):
        paragraph.add_run(
            # pylint: disable=no-member
            self.get_approval_status_display().upper() + ' ').font.bold = True
        paragraph.add_run(self.str_cm[0].format(
            # pylint: disable=no-member
            num2words(self.points, lang='es'),
            self.points,
            self.target_period,
            self.get_academic_program_display(),
            self.academic_program,
            self.target_period
        ) + '. ')
        paragraph.add_run(self.str_cm[1].format(
            Request.regulations[self.regulation_list[0]][0]))

    def pcm(self, docx):
        self.pcm_analysis(docx)
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.str_answer + ': ').font.bold = True
        paragraph.add_run(self.str_comittee_header + ' ')
        self.pcm_answer(paragraph)
        if self.is_affirmative_response_advisor_response():
            paragraph.add_run('. ')
        else:
            paragraph.add_run(', ')
            self.pcm_cna(paragraph)
        paragraph.add_run(
            Request.regulations[self.regulation_list[0]][0] + '.')

    def pcm_answer(self, paragraph):
        paragraph.add_run(
            # pylint: disable=no-member
            self.get_advisor_response_display().upper()).font.bold = True
        paragraph.add_run(' ' + self.str_cm[0].format(
            # pylint: disable=no-member
            num2words(self.points, lang='es'),
            self.points,
            self.target_period,
            self.get_academic_program_display(),
            self.academic_program,
            self.target_period
        ))

    def pcm_cna(self, paragraph):
        if self.cna == self.CNA_MAXIMO_EXCEDIDO:
            paragraph.add_run(self.str_pcma_cna[0] + ' ')
        elif self.cna == self.CNA_AVANCE_NO_SATISFACTORIO:
            paragraph.add_run(self.str_pcma_cna[1] + ' ')
        elif self.cna == self.CNA_REINGRESO:
            paragraph.add_run(self.str_pcma_cna[2] + ' ')
        elif self.cna == self.CNA_FUERA_PLAZOS:
            paragraph.add_run(self.str_pcma_cna[3].format(
                self.academic_period) + ' ')
        elif self.cna == self.CNA_OTRO:
            paragraph.add_run(self.council_decision + ' ')

    def pcm_analysis(self, docx):
        analysis_list = []
        analysis_list += [self.str_pcm[0].format(
            # pylint: disable=no-member
            self.get_academic_program_display(),
            self.academic_program,
            self.get_academic_profile_display()
        )]
        analysis_list += [self.str_pcm[1].format(
            '' if self.right_dates else 'no '
        )]
        analysis_list += [self.str_pcm[2].format(
            self.periods_in
        )]
        analysis_list += [self.str_pcm[3]]
        analysis_list += [self.str_pcm[4]]
        analysis_list += [self.str_pcm[5]]
        analysis_list += self.extra_analysis
        add_analysis_paragraph(docx, analysis_list)

    def resource_analysis(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.pcm_answer(last_paragraph)
    
    def resource_pre_answer(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.pcm_answer(last_paragraph)

    def resource_answer(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.cm_answer(last_paragraph)
