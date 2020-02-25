from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from mongoengine import StringField, EmbeddedDocumentListField
from ..models import Request, Professor
from .case_utils import add_analysis_paragraph


class ADIC(Request):

    full_name = 'Designación de codirector'

    node = StringField(required=True, choices=Request.PROFILE_CHOICES,
                       default=Request.PROFILE_INVE, display='Perfil')
    title = StringField(required=True, display='Título de Tesis/Trabajo Final',
                        default='')
    council_number_propuest = StringField(
        required=True, max_length=2, default='00', display='# Acta de aprobación propuesta')
    council_year_propuest = StringField(
        required=True, min_length=4, max_length=4, display='Año del Acta de aprobación propuesta', default='0000')
    proffesors = EmbeddedDocumentListField(Professor, display='Docentes')

    regulation_list = []

    str_cm = [
        'designar codirector de Tesis de {} con título ',
        'aprobado en el acta {} de {}, al(los) profesor(es) '
    ]
    str_pcm = [
        'SIA: {}{}.',
        ', perfil de {}',
        'Aprobación de propuesta y designación de director en el Acta No. {} de {} ' +
        'del Consejo de la Facultad de Ingeniería.'
    ]

    def cm(self, docx):
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        self.cm_answer(paragraph)

    def cm_answer(self, paragraph):
        paragraph.add_run(self.str_council_header + ' ')
        paragraph.add_run(
            # pylint: disable=no-member
            self.get_approval_status_display().upper() + ' ').font.bold = True
        self.add_text(paragraph)

    def pcm(self, docx):
        add_analysis_paragraph(docx, self.fill_analysis())
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        self.pcm_answer(paragraph)

    def pcm_answer(self, paragraph):
        paragraph.add_run(self.str_answer + ' ').font.bold = True
        paragraph.add_run(self.str_comittee_header + ' ')
        paragraph.add_run(
            # pylint: disable=no-member
            self.get_advisor_response_display().upper() + ' ').font.bold = True
        self.add_text(paragraph)

    def fill_analysis(self):
        # pylint: disable=no-member
        if self.node != Request.PROFILE_INVE:
            modifier = self.str_pcm[1].format(self.get_node_display())
        else:
            modifier = ''

        return [
            self.str_pcm[0].format(
                self.get_academic_program_display(), modifier),
            self.str_pcm[2].format(
                self.council_number_propuest, self.council_year_propuest
            )
        ] + self.extra_analysis

    def add_text(self, paragraph):
        paragraph.add_run(self.str_cm[0].format(
            # pylint: disable=no-member
            self.get_academic_program_display()))
        paragraph.add_run('"{}" '.format(self.title)).font.italic = True
        paragraph.add_run(self.str_cm[1].format(
            self.council_number_propuest, self.council_year_propuest))
        self.add_proffesors(paragraph)

    def add_proffesors(self, paragraph):
        for i in range(len(self.proffesors)):
            if self.proffesors[i].department not in (self.DP_EMPTY, self.DP_EXTERNO_FACULTAD):
                mod = self.proffesors[i].get_department_display()
            else:
                mod = self.proffesors[i].institution
                if self.proffesors[i].country != '':
                    mod += ' ({})'.format(self.proffesors[i].country)
            end = ', ' if i + 1 < len(self.proffesors) else '.'
            paragraph.add_run(
                '{} - {}{}'.format(self.proffesors[i].name, mod, end))
