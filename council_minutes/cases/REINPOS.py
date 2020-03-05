from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from mongoengine import StringField, BooleanField, FloatField
from mongoengine import EmbeddedDocumentListField
from ..models import Request, Subject
from .case_utils import add_analysis_paragraph


class REINPOS(Request):

    full_name = 'Reingreso Posgrado'

    node = StringField(
        display='Perfil', choices=Request.PROFILE_CHOICES, default=Request.PROFILE_INVE)
    papa = FloatField(required=True, min_value=0.0,
                      max_value=5.0, display='PAPA', default=0.0)
    first_reing = BooleanField(
        required=True, display='¿Primer reingreso?', default=True)
    reason_of_loss = StringField(
        required=True, display='Razón perdida calidad de estudiante', default='')
    time_limit = BooleanField(
        required=True, default=False, display='Pérdida por tiempo de permanencia')
    remaining_subjects = EmbeddedDocumentListField(
        Subject, display='Asignaturas Pendientes')
    on_time = BooleanField(required=True, default=True,
                           display='Solicitud a tiempo')
    reing_period = StringField(
        required=True, display='Periodo de reingreso', 
        choices=Request.PERIOD_CHOICES, default=Request.PERIOD_DEFAULT)
    grade_option = StringField(
        required=True, choices=Request.GRADE_OPTION_CHOICES,
        display='Opción de Grado', default=Request.GRADE_OPTION_TESIS_MAESTRIA)

    regulation_list = ['008|2008|CSU', '239|2009|VAC', '012|2014|VAC']

    str_cm = [
        'reingreso por única vez en el programa de {}, a partir del periodo {}',
        '. El reingreso del estudiante estará regido por el {}.',
        ', debido a que {}.'
    ]

    str_pcm = [
        'reingreso por única vez al programa {}, a partir del periodo académico {}, ' +
        'el reingreso del estudiante estará regido por el {}. ',
        'Durante el periodo académico adicional otorgado, el estudiante deberá solicitar ' +
        'el nombramiento de jurados de su {}, con el fin de obtener su título, ' +
        'previo cumplimiento de las demás exigencias académicas y administrativas vigentes.' +
        '(Artículo 7 de la {}).'
    ]

    str_analysis = [
        'El estudiante {} ha tenido otro reingreso posterior al 2009-1S, ' +
        '(Artículo 46, {}).',
        'Causa de retiro: {}. Plan de estudios {} - Perfil de {}.',
        '{}iene PAPA superior o igual a 3.5 (literal 3a – Artículo 3, {};' +
        'Artículo 46, {}). SIA PAPA: {}.',
        'En caso de ser por máximo tiempo de permanencia o por tener dos calificaciones ' +
        'NA en su historia académica: las asignaturas que le faltan por aprobar pueden cursarse ' +
        'en un solo periodo académico adicional (literal 5 – Artículo 3, {}; ' +
        'parágrafo 2 Artículo 46, {}).SIA: Le falta por aprobar: {}.',
        'La solicitud {}se hace en fechas de calendario de sede.'
    ]

    def cm(self, docx):
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.str_council_header + ' ')
        self.cm_answer(paragraph)

    def cm_answer(self, paragraph):
        # pylint: disable=no-member
        paragraph.add_run(
            self.get_approval_status_display().upper() + ' ').font.bold = True
        paragraph.add_run(self.str_cm[0].format(
            self.get_academic_program_display(),
            self.reing_period
        ))
        if self.is_affirmative_response_approval_status():
            paragraph.add_run(self.str_cm[1].format(
                self.regulations['008|2008|CSU'][0]))
        else:
            paragraph.add_run(self.str_cm[2].format(self.council_decision))

    def pcm(self, docx):
        add_analysis_paragraph(docx, self.add_analysis())
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.str_answer + ':\n').font.bold = True
        paragraph.add_run(self.str_comittee_header + ' ')
        self.pcm_answer(paragraph)

    def pcm_answer(self, paragraph):
        # pylint: disable=no-member
        paragraph.add_run(
            self.get_advisor_response_display().upper() + ' ').font.bold = True
        paragraph.add_run(self.str_pcm[0].format(
            self.get_academic_program_display(),
            self.reing_period,
            self.regulations['008|2008|CSU'][0]
        ))
        paragraph.add_run(self.str_pcm[1].format(
            self.get_grade_option_display(),
            self.regulations['012|2014|VAC'][0]
        ))

    def add_analysis(self):
        analysis = []

        modifier = 'no' if self.first_reing else 'ya'
        analysis.append(self.str_analysis[0].format(
            modifier, self.regulations['008|2008|CSU'][0]
        ))

        analysis.append(self.str_analysis[1].format(
            self.reason_of_loss,
            # pylint: disable=no-member
            self.get_academic_program_display(),
            self.get_node_display()
        ))

        modifier = 'T' if self.papa >= 3.5 else 'No t'
        analysis.append(self.str_analysis[2].format(
            modifier, self.regulations['239|2009|VAC'][0],
            self.regulations['008|2008|CSU'][0], self.papa
        ))

        subjects = ''
        for s in self.remaining_subjects:
            subjects += s.name + ', '
        analysis.append(self.str_analysis[3].format(
            self.regulations['239|2009|VAC'][0],
            self.regulations['008|2008|CSU'][0],
            subjects[:-2]
        ))

        modifier = '' if self.on_time else 'no '
        analysis.append(self.str_analysis[4].format(modifier))

        return analysis + self.extra_analysis

    def resource_analysis(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.pcm_answer(last_paragraph)
    
    def resource_pre_answer(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.pcm_answer(last_paragraph)

    def resource_answer(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.cm_answer(last_paragraph)
