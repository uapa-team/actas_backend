from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from mongoengine import IntField, FloatField
from ..models import Request
from .case_utils import add_analysis_paragraph


class RDEF(Request):

    full_name = 'Retiro definitivo del programa'

    advance = FloatField(required=True, display='% de Avance', default=0.0)
    enrolled_academic_periods = IntField(
        required=True, display='# Periodos Matriculados', default=0)
    papa = FloatField(required=True, display='PAPA', default=0.0)

    regulation_list = []  # List of regulations

    str_cm = ['presentar con concepto positivo a la División de Registro y Matrícula, ' +
              'el retiro voluntario del programa {} ({})',
              'debido a que']

    list_analysis = [
        'SIA: Porcentaje de avance en el plan: {}%',
        'SIA: Número de matrículas: {}',
        'SIA: P.A.P.A.: {}'
    ]

    def cm(self, docx):
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.str_council_header + ' ')
        self.cm_answer(paragraph)

    def cm_answer(self, paragraph):
        # pylint: disable=no-member
        paragraph.add_run(
            self.get_approval_status_display().upper() + ' ').font.bold = True
        paragraph.add_run(self.str_cm[0].format(
            self.get_academic_program_display(), self.academic_program))
        if not self.is_affirmative_response_approval_status():
            paragraph.add_run(
                ' ' + self.str_cm[1] + ' ' + self.council_decision)
        paragraph.add_run('.')

    def pcm(self, docx):
        self.pcm_analysis(docx)
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.str_answer + ': ').font.bold = True
        paragraph.add_run(self.str_comittee_header + ' ')
        self.pcm_answer(paragraph)

    def pcm_answer(self, paragraph):
        # pylint: disable=no-member
        paragraph.add_run(
            self.get_advisor_response_display().upper() + ' ').font.bold = True
        paragraph.add_run(self.str_cm[0].format(
            self.get_academic_program_display(), self.academic_program))
        if not self.is_affirmative_response_approval_status():
            paragraph.add_run(
                ' ' + self.str_cm[1] + ' ' + self.council_decision)
        paragraph.add_run('.')

    def pcm_analysis(self, docx):
        final_analysis = []
        final_analysis += [self.list_analysis[0].format(self.advance)]
        final_analysis += [self.list_analysis[1]
                           .format(self.enrolled_academic_periods)]
        final_analysis += [self.list_analysis[2].format(self.papa)]
        for extra_a in self.extra_analysis:
            final_analysis += [extra_a]
        add_analysis_paragraph(docx, final_analysis)

    def resource_analysis(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.pcm_answer(last_paragraph)
    
    def resource_pre_answer(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.pcm_answer(last_paragraph)

    def resource_answer(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.cm_answer(last_paragraph)
