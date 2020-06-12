from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from mongoengine import FloatField, IntField, BooleanField
from ..models import Request
from .case_utils import add_analysis_paragraph


class CINF(Request):

    full_name = 'Autorización carga inferior a la mínima'

    papa = FloatField(
        required=True, display='P.A.P.A.', min_value=0.0, max_value=5, default=0.0)
    available_creds = IntField(
        required=True, display='Créditos disponibles', min_value=0, default=0)
    advance_percentage = FloatField(
        required=True, display='Porcentaje de avance', default=0.0, min_value=0.0, max_value=100.0)
    enrolled_academic_periods = IntField(
        required=True, display='Número de matrículas', min_value=0, default=0)

    regulation_list = ['008|2008|CSU']  # List of regulations

    str_cm = ['cursar el periodo académico {} con un número de créditos inferior ' +
              'al mínimo exigido porque ', 'justifica debidamente su solicitud. ', '({}).',
              'Artículo 10 del ']

    list_analysis = ['SIA: Porcentaje de avance en el plan: {}%.',
                     'SIA: Número de matrículas: {}.',
                     'SIA: P.A.P.A.: {}.',
                     'SIA: Créditos disponibles: {}.']

    def cm(self, docx):
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.str_council_header + ' ')
        self.cm_answer(paragraph)

    def cm_answer(self, paragraph):
        paragraph.add_run(
            # pylint: disable=no-member
            self.get_approval_status_display().upper() + ' ').font.bold = True
        paragraph.add_run(self.str_cm[0].format(self.academic_period))
        if self.is_affirmative_response_approval_status():
            self.cm_af(paragraph)
        else:
            self.cm_ng(paragraph)

    def pcm(self, docx):
        self.pcm_analysis(docx)
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.str_answer + ': ').font.bold = True
        paragraph.add_run(self.str_comittee_header + ' ')
        self.pcm_answer(paragraph)

    def pcm_answer(self, paragraph):
        paragraph.add_run(
            # pylint: disable=no-member
            self.get_advisor_response_display().upper()).font.bold = True
        paragraph.add_run(' ' + self.str_cm[0].format(self.academic_period))
        if self.is_affirmative_response_advisor_response():
            self.pcm_answers_af(paragraph)
        else:
            self.pcm_answers_ng(paragraph)

    def cm_af(self, paragraph):
        paragraph.add_run(self.council_decision + '. ' +
                          self.str_cm[2].format(self.str_cm[3] +
                                                self.regulations['008|2008|CSU'][0]))

    def cm_ng(self, paragraph):
        paragraph.add_run(self.council_decision + '. ' +
                          self.str_cm[2].format(self.str_cm[3] +
                                                self.regulations['008|2008|CSU'][0]))

    def pcm_analysis(self, docx):
        final_analysis = []
        final_analysis += [self.list_analysis[0].format(
            self.advance_percentage)]
        final_analysis += [self.list_analysis[1].format(
            self.enrolled_academic_periods)]
        final_analysis += [self.list_analysis[2].format(
            self.papa)]
        final_analysis += [self.list_analysis[3].format(
            self.available_creds)]
        for extra_a in self.extra_analysis:
            final_analysis += [extra_a]
        add_analysis_paragraph(docx, final_analysis)

    def pcm_answers_af(self, paragraph):
        paragraph.add_run(
            self.str_cm[1] + self.str_cm[2].format(self.str_cm[3] +
                                                   self.regulations['008|2008|CSU'][0]))

    def pcm_answers_ng(self, paragraph):
        paragraph.add_run(self.council_decision + '. ' +
                          self.str_cm[2].format(self.str_cm[3] +
                                                self.regulations['008|2008|CSU'][0]))

    def resource_analysis(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.pcm_answer(last_paragraph)

    def resource_pre_answer(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.pcm_answer(last_paragraph)

    def resource_answer(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.cm_answer(last_paragraph)
