import datetime
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from mongoengine import StringField, DateField
from ..models import Request
from .case_utils import add_analysis_paragraph


class AFPD(Request):

    full_name = 'Ampliación de la fecha de pago de recibo de matrícula'
    decision_makers = Request.decision_makers[3]

    justification = StringField(
        required=True, display='Justificación de la decisión', default='')
    limit_date = DateField(
        required=True, display='Fecha Límite', default=datetime.date.today)

    regulation_list = []

    str_cm = [
        'presentar con concepto positivo al Comité de Matriculas de la Sede ' +
        'Bogotá, la expedición de un único recibo correspondiente a los ' +
        'derechos académicos y administrativos para el periodo académico {} ' +
        'debido a que {}.'
    ]
    str_pcm = [
        'El estudiante tiene la historia académica activa.',
        'recomiendar al Consejo de Facultad presentar con concepto positivo al Comité ' +
        'de Matrículas de la Sede Bogotá, la expedición de un único recibo correspondiente ' +
        'a los derechos académicos y administrativos para el periodo académico {} ' +
        'y se le concede como fecha de pago el {}, teniendo en cuenta el estado de pago ' +
        'por parte de {}.'
    ]

    def cm(self, docx):
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.str_council_header + ' ')
        self.cm_answer(paragraph)

    def cm_answer(self, paragraph):
        paragraph.add_run(
            # pylint: disable=no-member
            self.get_approval_status_display().upper() + ' ').font.bold = True
        paragraph.add_run(self.str_cm[0].format(
            self.academic_period, self.justification))

    def pcm(self, docx):
        analysis = [self.str_pcm[0]] + self.extra_analysis
        add_analysis_paragraph(docx, analysis)
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.str_answer + ' ').font.bold = True
        paragraph.add_run(self.str_comittee_header + ' ')
        self.pcm_answer(paragraph)

    def pcm_answer(self, paragraph):
        # pylint: disable=no-member
        paragraph.add_run(
            self.get_advisor_response_display().upper() + ' ').font.bold = True
        paragraph.add_run(self.str_pcm[1].format(
            self.academic_period,
            self.limit_date.strftime('%d/%m/%Y '),
            self.student_name
        ))

    def resource_analysis(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.pcm_answer(last_paragraph)
    
    def resource_pre_answer(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.pcm_answer(last_paragraph)

    def resource_answer(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.cm_answer(last_paragraph)
