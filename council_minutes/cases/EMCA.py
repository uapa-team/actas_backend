from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from mongoengine import FloatField, StringField, BooleanField
from ..models import Request
from .case_utils import add_analysis_paragraph


class EMCA(Request):

    full_name = 'Exención de matrícula por consideración atípica'

    percentage = FloatField(min_value=0.0, max_value=100.0, required=True,
                            display='Porcentaje de exención del valor de la matrícula', default=0.0)
    academic_period_exe = StringField(
        max_length=10, display='Periodo de exención', 
        choices=Request.PERIOD_CHOICES,
        default=Request.PERIOD_DEFAULT)
    ha_active = BooleanField(
        default=True, display='¿Tiene la historia académica activa?')

    regulation_list = ['070|2009|CAC']  # List of regulations

    str_cm = [' otorgrar exención del pago de {}% del valor de la matrícula para el periodo ' +
              'académico {}, ', ' ({}).']

    str_analysis = ['iene la historia académica activa.']

    list_analysis = []

    def cm(self, docx):
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.str_council_header + ' ')
        self.cm_answer(paragraph)

    def cm_answer(self, paragraph):
        # pylint: disable=no-member
        paragraph.add_run(
            self.get_approval_status_display().upper() + ' ').font.bold = True
        paragraph.add_run(self.str_cm[0].format(
            self.percentage, self.academic_period_exe))
        paragraph.add_run(self.council_decision)
        paragraph.add_run(self.str_cm[1].format(
            Request.regulations['070|2009|CAC'][0]))

    def pcm(self, docx):
        self.pcm_analysis(docx)
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.str_answer + ': ').font.bold = True
        paragraph.add_run(self.str_comittee_header + ' ')
        self.pcm_answer(paragraph)

    def pcm_answer(self, paragraph):
            # pylint: disable=no-member
        paragraph.add_run(
            self.get_advisor_response_display().upper()).font.bold = True
        paragraph.add_run(self.str_cm[0].format(
            self.percentage, self.academic_period_exe))
        paragraph.add_run(self.council_decision)
        paragraph.add_run(self.str_cm[1].format(
            Request.regulations['070|2009|CAC'][0]))

    def pcm_analysis(self, docx):
        final_analysis = []
        aux = 'T' if self.ha_active else 'No t'
        final_analysis += [aux + self.str_analysis[0]]
        for extra_a in self.extra_analysis:
            final_analysis += [extra_a]
        add_analysis_paragraph(docx, final_analysis)

    def resource_analysis(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.pcm_answer(last_paragraph)
    
    def resource_pre_answer(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.pcm_answer(last_paragraph)

    def resource_answer(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.cm_answer(last_paragraph)
