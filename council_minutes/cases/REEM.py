from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from num2words import num2words
from mongoengine import IntField, FloatField, ObjectIdField
from ..models import Request
from .case_utils import add_analysis_paragraph


class REEM(Request):

    full_name = 'Reembolso'

    credits_refunded = IntField(display='Créditos Disponibles', default=0)
    percentage = FloatField(
        display='Porcentaje de créditos a cancelar', default=0.0)
    cancelation_case = ObjectIdField(
        display='Código del caso en el que fue cancelado el periodo')

    regulation_list = (
        '032|2010|CSU',
        '1416|2013|RE'
    )

    str_cm = [
        'devolución proporcional del {} por ciento ({} %) del valor pagado por concepto de ' +
        'derechos de matrícula del periodo {}.',
        'Teniendo en cuenta la fecha de presentación de la solicitud y que le fue aprobada la ' +
        'cancelación de periodo en el Acta {} de {} del Consejo de Facultad.'
    ]

    str_pcm = [
        'Fecha de presentación de solicitud: {}.',
        'Créditos a reembolsar: {}.',
        'Cancelación de periodo académico: Caso {}, consecutivo {}, año: {}.'
    ]

    def cm(self, docx):
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.add_run(self.str_council_header + ' ')
        self.cm_answer(paragraph)

    def cm_answer(self, paragraph):
        # pylint: disable=no-member
        cancelation_case = Request.objects.get(id=self.cancelation_case)
        paragraph.add_run(
            # pylint: disable=no-member
            self.get_approval_status_display().upper() + ' ').font.bold = True
        paragraph.add_run(
            self.str_cm[0].format(
                num2words(self.percentage, lang='es'),
                self.percentage,
                cancelation_case.academic_period
            ) + ' '
        )
        paragraph.add_run(
            self.str_cm[1].format(
                cancelation_case.consecutive_minute,
                str(cancelation_case.date)[0:4])
        )

    def pcm(self, docx):
        self.pcm_analysis(docx)
        paragraph = docx.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run('{}: '.format(self.str_answer)).font.bold = True
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.add_run(self.str_comittee_header)
        self.pcm_answer(paragraph)

    def pcm_analysis(self, docx):
        # pylint: disable=no-member
        cancelation_case = Request.objects.get(id=self.cancelation_case)
        analysis_list = []
        analysis_list += [self.str_pcm[0].format(self.date)]
        analysis_list += [self.str_pcm[1].format(self.credits_refunded)]
        analysis_list += [self.str_pcm[2].format(cancelation_case.id,
                                                 cancelation_case.consecutive_minute,
                                                 str(cancelation_case.date)[0:4])]
        analysis_list += self.extra_analysis
        add_analysis_paragraph(docx, analysis_list)

    def pcm_answer(self, paragraph):
        # pylint: disable=no-member
        cancelation_case = Request.objects.get(id=self.cancelation_case)
        paragraph.add_run(
            # pylint: disable=no-member
            ' ' + self.get_advisor_response_display().upper() + ' ').font.bold = True
        paragraph.add_run(self.str_cm[0].format(
            num2words(self.percentage, lang='es'),
            self.percentage,
            cancelation_case.academic_period
        ))

    def resource_analysis(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.pcm_answer(last_paragraph)
    
    def resource_pre_answer(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.pcm_answer(last_paragraph)

    def resource_answer(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.cm_answer(last_paragraph)
