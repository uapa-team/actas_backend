from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from mongoengine import StringField, BooleanField, IntField, ListField, \
    EmbeddedDocument, EmbeddedDocumentListField
from ..models import Request
from .case_utils import add_analysis_paragraph


class APGD(Request):

    class Coadvisor(EmbeddedDocument):
        name_co_advisor = StringField(
            display='Nombre del codirector', required=True, default='')
        inst_co_advisor = StringField(
            display='Departamento de adscripción del codirector',
            default=Request.DP_CIVIL_AGRICOLA, required=True, choices=Request.DP_CHOICES)
        co_advisor_ext = StringField(
            display='Institución externa', default='')

    full_name = 'Aprobación de propuesta de trabajo final de maestría/doctorado' + \
        ' o de proyecto de tesis de maestría y designación de director y co-director'

    GO_TRABAJO_FINAL_MAESTRIA = 'TFM'
    GO_TESIS_MAESTRIA = 'TSM'
    GO_TESIS_DOCTORADO = 'TSD'
    GO_CHOICES = (
        (GO_TRABAJO_FINAL_MAESTRIA, 'Trabajo Final de Maestría'),
        (GO_TESIS_MAESTRIA, 'Tesis de Maestría'),
        (GO_TESIS_DOCTORADO, 'Tesis de Doctorado')
    )

    IG_OWNERSHIP = 'OS'
    IG_NOT_OWNERSHIP = 'NO'
    IG_UNDEFINED = 'UN'
    IG_CHOICES = (
        (IG_OWNERSHIP, 'Sí pertenece'),
        (IG_NOT_OWNERSHIP, 'No pertenece'),
        (IG_UNDEFINED, 'No dice o no se sabe'),
    )

    CP_APROBADA = 'AP'
    CP_NO_APROBADO = 'NA'
    CP_CHOICES = (
        (CP_APROBADA, 'aprobado'),
        (CP_NO_APROBADO, 'no aprobado'),
    )

    grade_option = StringField(
        required=True, display='Tipo de tesis/trabajo final',
        choices=GO_CHOICES, default=GO_TESIS_MAESTRIA)
    enrolled_proyect = BooleanField(
        required=True, default=False, display='¿Tiene inscrita la asignatura ' +
        '(proyecto de tesis)/(propuesta de trabajo final de maestría)?')
    have_signature = BooleanField(required=True, default=False,
                                  display='¿Tiene la firma del (los) director(es)?')
    enrroled_periods = IntField(
        min_value=1, default=1, display='Número de matrícula actual', required=True)
    cd_delivered = BooleanField(
        required=True, default=False, display='¿Entregó CD?')
    general_objetive = StringField(
        required=True, display='Objetivo general', default='')
    specific_objetives = ListField(StringField(),
        display='Objetivos específicos')
    title = StringField(
        required=True, display='Título de la tesis/trabajo final', default='')
    ownership_ig = StringField(
        required=True, choices=IG_CHOICES, default=IG_UNDEFINED,
        display='¿El proyecto hace parte de un grupo de investigación?')
    advisor = StringField(
        display='Director de tesis/trabajo final', default='', required=True)
    advisor_inst = StringField(
        display='Departamento de adscripción del director',
        required=True, choices=Request.DP_CHOICES, default=Request.DP_EMPTY)
    advisor_ext = StringField(
        display='Institución externa', default='')
    co_advisor_list = EmbeddedDocumentListField(
        Coadvisor, display='Codirector(es)')
    grade_proyect = StringField(required=True, display='Calificación de la propuesta/proyecto',
                                choices=CP_CHOICES, default=CP_APROBADA)

    regulation_list = ['040|2017|CFA', '056|2012|CSU']  # List of regulations

    str_cm = ['Calificación {} ({}) a {} de {}, cuyo título es:',
              'Designar director',
              'Designar codirector',
              'al profesor',
              'del',
              'de la institución',
              'de',
              'cuyo título es',
              'Debido a que']

    list_analysis = ['Perfil de {}.',
                     'El estudiante {}tiene inscrita la asignatura {}.',
                     'Estudiante de matrícula número {}.',
                     'ntregó CD.',
                     'iene la firma del (los) director(es) de tesis/trabajo final:',
                     'El proyecto de tesis debe inscribirse y entregarse, antes de alcanzar el 50' +
                     '% de la duración establecida para el programa (Parágrafo Artículo 14 del ' +
                     '{})',
                     'Título:',
                     'Objetivo general:',
                     'Objetivos específicos:']

    def cm(self, docx):
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.str_council_header + ' ')
        self.cm_answer(paragraph)
        self.cm_grade(docx)
        if self.is_affirmative_response_approval_status():
            self.cm_design(docx)
        else:
            paragraph = docx.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.add_run(
                ' ' + self.str_cm[8] + ' ' + self.council_decision + '.')

    def cm_answer(self, paragraph):
        # pylint: disable=no-member
        paragraph.add_run(
            self.get_approval_status_display().upper() + ':').font.bold = True

    def pcm(self, docx):
        self.pcm_analysis(docx)
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.str_answer + ': ').font.bold = True
        paragraph.add_run(self.str_comittee_header + ' ')
        self.pcm_answer(paragraph)
        self.cm_grade(docx)
        if self.is_affirmative_response_advisor_response():
            self.cm_design(docx)
        else:
            paragraph = docx.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.add_run(
                ' ' + self.str_cm[8] + ' ' + self.council_decision + '.')

    def pcm_answer(self, paragraph):
        # pylint: disable=no-member
        paragraph.add_run(
            self.get_advisor_response_display().upper() + ':').font.bold = True

    def pcm_analysis(self, docx):
        # pylint: disable=no-member
        if self.grade_option in [self.GO_TESIS_MAESTRIA, self.GO_TESIS_DOCTORADO]:
            profile = 'investigación'
        else:
            profile = 'profundización'
        type_subject = 'Proyecto' if self.grade_option in [
            self.GO_TESIS_MAESTRIA, self.GO_TESIS_DOCTORADO] else 'Propuesta'
        final_analysis = []
        final_analysis += [self.list_analysis[0].format(profile)]
        ets = '' if self.enrolled_proyect else 'no '
        final_analysis += [self.list_analysis[1].format(
            ets, type_subject + ' ' + self.get_grade_option_display())]
        final_analysis += [self.list_analysis[2].format(
            self.enrroled_periods)]
        cdd = 'E' if self.cd_delivered else 'No e'
        final_analysis += [cdd + self.list_analysis[3]]
        hss = 'T' if self.have_signature else 'No t'
        final_analysis += [hss + self.list_analysis[4] +
                           ' ' + self.advisor + '.']
        final_analysis += [self.list_analysis[5].format(
            Request.regulations['056|2012|CSU'][0]) + '.']
        add_analysis_paragraph(docx, final_analysis)
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.style = 'List Bullet'
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.list_analysis[6] + ' ').font.bold = True
        paragraph.add_run(self.title + '.').font.italic = True
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.style = 'List Bullet'
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.list_analysis[7] + ' ').font.bold = True
        paragraph.add_run(self.general_objetive + '.')
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.style = 'List Bullet'
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.list_analysis[8]).font.bold = True
        for spec_ob in self.specific_objetives:
            paragraph = docx.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph.style = 'List Bullet 2'
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.add_run(spec_ob + '.')
        for ex_an in self.extra_analysis:
            paragraph = docx.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph.style = 'List Bullet'
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.add_run(ex_an + '.')

    def cm_grade(self, docx):
        # pylint: disable=no-member
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.style = 'List Bullet'
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.str_cm[0].format(
            self.grade_proyect, self.get_grade_proyect_display(),
            self.get_grade_option_display(), self.get_academic_program_display()))
        paragraph.add_run(' "{}".'.format(self.title)).font.italic = True

    def cm_design(self, docx):
        # pylint: disable=no-member
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.style = 'List Bullet'
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.str_cm[1] + ' ')
        paragraph.add_run(
            self.str_cm[6] + ' ' + self.get_grade_option_display() + ' ' + self.str_cm[7])
        paragraph.add_run(' "{}" '.format(self.title)).font.italic = True
        paragraph.add_run(self.str_cm[3] + ' ')
        paragraph.add_run(self.advisor)
        if self.advisor_inst == Request.DP_EXTERNO_FACULTAD:
            paragraph.add_run(
                ' ' + self.str_cm[5] + ' ' + self.advisor_ext + '.')
        else:
            paragraph.add_run(
                ' ' + self.str_cm[4] + ' ' + self.get_advisor_inst_display() + '.')
        if self.co_advisor_list != []:
            for co_advc in self.co_advisor_list:
                paragraph = docx.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                paragraph.style = 'List Bullet'
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.add_run(self.str_cm[2] + ' ')
                paragraph.add_run(
                    self.str_cm[6] + ' ' + self.get_grade_option_display() + ' ' + self.str_cm[7])
                paragraph.add_run(' "{}" '.format(
                    self.title)).font.italic = True
                paragraph.add_run(self.str_cm[3] + ' ')
                paragraph.add_run(co_advc.name_co_advisor)
                if co_advc.inst_co_advisor == Request.DP_EXTERNO_FACULTAD:
                    paragraph.add_run(
                        ' ' + self.str_cm[5] + ' ' + co_advc.co_advisor_ext + '.')
                else:
                    paragraph.add_run(
                        ' ' + self.str_cm[4] + ' ' + co_advc.get_inst_co_advisor_display() + '.')

    def resource_analysis(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.pcm_answer(last_paragraph)
    
    def resource_pre_answer(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.pcm_answer(last_paragraph)

    def resource_answer(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.cm_answer(last_paragraph)
