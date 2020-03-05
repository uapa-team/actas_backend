from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from mongoengine import StringField, EmbeddedDocument, EmbeddedDocumentListField
from ..models import Request
from .case_utils import add_analysis_paragraph


class Register(EmbeddedDocument):

    TYPE_EXAMEN = 'E'
    TYPE_TESIS = 'T'
    TYPE_CHOICES = (
        (TYPE_EXAMEN, 'Exámen'),
        (TYPE_TESIS, 'Tesis')
    )
    GRADE_AP = 'AP'
    GRADE_NA = 'NA'
    GRADE_CHOICES = (
        (GRADE_AP, 'Aprobado'),
        (GRADE_NA, 'No Aprobado')
    )

    _type = StringField(
        required=True, display='Tipo de Registro', choices=TYPE_CHOICES)
    code = StringField(required=True, display='Código')
    grade = StringField(
        required=True, display='Calificación', choices=GRADE_CHOICES)
    title = StringField(display='Título')


class RCPE(Request):

    full_name = 'Registro de calificación del proyecto y examen doctoral'

    registers = EmbeddedDocumentListField(
        Register, display='Registros')

    regulation_list = []

    str_cm = []
    str_pcm = []

    str_exam = [
        'Calificar {} ({}) el examen de calificación con código {} en el periodo académico {}.'
    ]

    str_thesis = [
        'Calificar {} ({}) el Proyecto de Tesis de {} con código {} ' +
        'en el periodo académico {}, cuyo título es: '
    ]

    def cm(self, docx):
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.str_council_header + ' ')
        self.cm_answer(paragraph)
        self.add_registers(docx)

    def cm_answer(self, paragraph):
        paragraph.add_run(
            # pylint: disable=no-member
            self.get_approval_status_display().upper() + ':').font.bold = True

    def pcm(self, docx):
        add_analysis_paragraph(docx, self.extra_analysis)
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.str_answer + ' ').font.bold = True
        paragraph.add_run(self.str_comittee_header + ' ')
        self.pcm_answer(paragraph)
        self.add_registers(docx)

    def pcm_answer(self, paragraph):
        paragraph.add_run(
            # pylint: disable=no-member
            self.get_advisor_response_display().upper() + ':').font.bold = True

    def add_registers(self, docx):
        # pylint: disable=no-member, protected-access
        for register in self.registers:
            paragraph = docx.add_paragraph()
            paragraph.style = 'List Bullet 2'
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph.paragraph_format.space_after = Pt(0)
            if register._type == Register.TYPE_EXAMEN:
                paragraph.add_run(self.str_exam[0].format(
                    register.get_grade_display(),
                    register.grade,
                    register.code,
                    self.academic_period
                ))
            elif register._type == Register.TYPE_TESIS:
                paragraph.add_run(self.str_thesis[0].format(
                    register.get_grade_display(),
                    register.grade,
                    self.get_academic_program_display(),
                    register.code,
                    self.academic_period
                ))
                paragraph.add_run('"{}".'.format(
                    register.title)).font.italic = True
            else:
                raise AssertionError(
                    self.assertionerror['CHOICES'].format(register._type))

    def resource_analysis(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.pcm_answer(last_paragraph)
    
    def resource_pre_answer(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.pcm_answer(last_paragraph)

    def resource_answer(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.cm_answer(last_paragraph)
