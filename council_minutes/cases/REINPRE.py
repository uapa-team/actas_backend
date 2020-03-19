import datetime
import functools
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from mongoengine import StringField, IntField, FloatField, BooleanField, DateField
from .case_utils import string_to_date, table_general_data
from .case_utils import table_credits_summary, table_recommend, add_analysis_paragraph
from ..models import Request


class REINPRE(Request):
    RL_ANSWER_RENOV_MATRICULA = 'RM'
    RL_ANSWER_PAPA = 'PA'
    RL_ANSWER_CUPO_CREDITOS = 'CC'
    RL_ANSWER_SANCION = 'SA'
    RL_ANSWER_OTRO = 'OT'
    RL_ANSWER_PAPA_CREDITOS = 'PC'
    RL_ANSWER_CHOICES = (
        (RL_ANSWER_RENOV_MATRICULA, 'No cumplir con los requisitos exigidos para la' +
         ' renovación de la matrícula, en los plazos señalados por la Universidad.'),
        (RL_ANSWER_PAPA,
         'Presentar un Promedio Aritmético Ponderado Acumulado menor que tres punto cero (3.0).'),
        (RL_ANSWER_CUPO_CREDITOS,
         'No disponer de un cupo de créditos suficiente para inscribir las asignaturas' +
         ' del plan de estudios pendientes de aprobación.'),
        (RL_ANSWER_SANCION,
         'Recibir sanción disciplinaria de expulsión o suspensión impuesta de acuerdo' +
         ' con las normas vigentes.'),
        (RL_ANSWER_PAPA_CREDITOS,
         'PAPA menor a 3.0 y cupo de créditos insuficiente.'),
        (RL_ANSWER_OTRO, 'Otro.')
    )

    full_name = 'Reingreso Pregrado'

    regulation_list = ['008|2008|CSU', '239|2009|VAC', '012|2014|VAC']

    reing_period = StringField(
        required=True, display='Periodo de reingreso',
        choices=Request.PERIOD_CHOICES, default=Request.PERIOD_DEFAULT)
    loss_period = StringField(
        required=True, display='Periodo de perdida de calidad de estudiante',
        choices=Request.PERIOD_CHOICES, default=Request.PERIOD_DEFAULT)
    first_reing = BooleanField(
        required=True, display='Primer reingreso', default=True)
    admission_period = StringField(
        required=True, display='Periodo de admisión',
        choices=Request.PERIOD_CHOICES, default=Request.PERIOD_DEFAULT)
    periods_since = IntField(
        required=True, display='periodos desde pérdida de calidad de estudiante', default=0)
    papa = FloatField(required=True, display='PAPA', default=0.0)
    reason_of_loss = StringField(choices=RL_ANSWER_CHOICES,
                                 default=RL_ANSWER_OTRO,
                                 display='Razón pérdida calidad de estudiante')
    credits_bag = IntField(
        required=True, display='Cupo de créditos disponible para inscripción', default=0)
    credits_english = IntField(
        required=True, display='Créditos pendientes inglés', default=0)
    credits_coursed = IntField(required=True,
                               display='Créditos cursados (Aprobados + No Aprobados)' +
                               ' con calificación numérica')

    # Exiged credits
    exi_fund_m = IntField(
        required=True, display='Créditos de fundamentación obligatorios exigidos', default=0)
    exi_fund_o = IntField(
        required=True, display='Créditos de fundamentación optativos exigidos', default=0)
    exi_disc_m = IntField(
        required=True, display='Créditos disciplinares obligatorios exigidos', default=0)
    exi_disc_o = IntField(
        required=True, display='Créditos disciplinares optativos exigidos', default=0)
    exi_free = IntField(
        required=True, display='Créditos de libre elección exigidos', default=0)

    # Approved credits
    app_fund_m = IntField(
        required=True, display='Créditos de fundamentación obligatorios aprobados', default=0)
    app_fund_o = IntField(
        required=True, display='Créditos de fundamentación optativos aprobados', default=0)
    app_disc_m = IntField(
        required=True, display='Créditos disciplinares obligatorios aprobados', default=0)
    app_disc_o = IntField(
        required=True, display='Créditos disciplinares optativos aprobados', default=0)
    app_free = IntField(
        required=True, display='Créditos de libre elección aprobados', default=0)

    comitee_act = StringField(
        required=True, display='Número de acta de comité', default='00')

    comitee_date = DateField(
        required=True, display='Fecha de reunión del comité', default=datetime.date.today
    )

    # Pre-cm variables
    request_in_date = BooleanField(
        display='Solicitud entregada a tiempo', default=True)

    str_pcm_pre = [
        # Used in pcm and cm:
        'reingreso por única vez a partir del periodo académico ',
        '. Si el estudiante no renueva su matrícula en el semestre de reingreso, el acto' +
        ' académico expedido por el Consejo de Facultad queda sin efecto. ',
        '1. Datos Generales:',
        '2. Información Académica:',
        '3. Resumen general de créditos del plan de estudios:',
        '*Sin incluir los créditos correspondientes al cumplimiento del requisito de' +
        ' suficiencia en idioma.',

        # Used only in pcm:
        'El estudiante ',
        ' tiene pendiente por aprobar ',
        ' créditos del plan de estudios de ',
        ' y ',
        ' créditos del requisito de nivelación',
        ' - inglés, con un cupo disponible para inscripción de ',
        ' créditos.',
        'El parágrafo del artículo 11 del ',
        'Superior Universitario establece: ',
        '"Los créditos adicionales que como resultado del ' +
        'proceso de clasificación en la admisión deba aprobar ' +
        'un estudiante de pregrado, se sumarán por única vez al "' +
        'cupo adicional de créditos para inscripción"',
        ', por lo tanto solo es viable otorgar ',
        ' crédito(s) para la inscripción de asignaturas pendientes del plan de estudios de ',

        # Extra credits (optional):
        'y otorga ',
        ' crédito(s) adicional(es) para culminar su plan de estudios. '

    ]

    str_analysis = [
        '{}a tenido otro reingreso después de 2009-01 (Artículo 46, {}). ' +
        'Universitas y SIA: Revisado.',
        'Si perdió calidad antes de 2009-01: Equivalencias incluyendo las asignaturas ' +
        'perdidas. Comité Asesor asigna créditos a las que no tengan equivalencias ' +
        '(Artículo 3, {}). Universitas y SIA: Pérdida de calidad de estudiante al ' +
        'finalizar {} por {}',
        '{}iene PAPA superior o igual a 2.7 (literal 3b – Artículo 3, {}; Artículo 46, ' +
        '{}). SIA: PAPA de {}.',
        '{}ispone de un cupo suficiente de créditos: Cupo adicional de 10 créditos a lo sumo ' +
        '(parágrafo 1 Artículo 46, {}). SIA: {} creditos. En caso de otorgarle un cupo adicional ' +
        'de créditos, éste no podrá ser mayor que el requerido para inscribir las asignaturas ' +
        'pendientes del plan de estudios. (Artículo 6, {}).',
        'La solicitud {}se hace en fechas de calendario de sede.'
    ]

    str_pcm_pre_acadinfo = [
        'Periodo para el cual fue admitido en este plan de estudios',
        '¿Se trata de un primer reingreso?',
        'Si la respuesta es NO, el Comité Asesor no debe recomendar al Consejo ' +
        'de Facultad el reingreso',
        'Es caso de ser primer reingreso en ¿qué periodo académico perdió la ' +
        'calidad de estudiante?',
        'Al momento de presentar la solicitud ¿cuántos periodos académicos (incluido' +
        ' el periodo académico en que presentó la solicitud) han transcurridos a partir' +
        ' del periodo académico en que registró su última matrícula?',
        'En caso que la respuesta sea mayor de 6 periodos académicos no se debe ' +
        'recomendar el reingreso',
        'P.A.P.A.',
        'Causa de la pérdida de la calidad de estudiante',
        'Estudio de créditos',
        'Cupo de créditos menos créditos pendientes',
        'Créditos pendientes por ser aprobados del plan de estudios',
        'Créditos pendientes por ser aprobados de nivelación – Inglés',
        '¿Cuántos créditos adicionales requiere para inscribir asignaturas?',
        # Optional: Grade needed with N credits to keep student condition.
        'Al finalizar el semestre de reingreso para mantener la calidad de estudiante,' +
        ' deberá obtener un promedio semestral mínimo de:',
        'Si inscribe 12 Créditos',
        'Si inscribe 15 Créditos',
        'Si inscribe 18 Créditos',
        'Si inscribe 21 Créditos'
    ]

    str_out_date = [
        'reingreso por única vez a partir del periodo académico {}, porque el estudiante ' +
        'presentó la solicitud fuera de las fechas establecidas en el Calendario Académico ' +
        'de la Sede Bogotá.'
    ]

    def rein_general_data_table(self, docx):
        # pylint: disable=no-member
        general_data = [['Estudiante', self.student_name],
                        ['DNI', self.student_dni],
                        ['Plan de estudios', self.get_academic_program_display()],
                        ['Código del plan de estudios', self.academic_program],
                        ['Fecha de la solicitud', string_to_date(str(self.date))]]

        case = 'REINGRESO'

        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        bullet = paragraph.add_run(self.str_pcm_pre[2])
        bullet.font.bold = True
        bullet.font.size = Pt(8)

        table_general_data(general_data, case, docx)

    def rein_academic_info(self, docx):
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        bullet = paragraph.add_run(self.str_pcm_pre[3])
        bullet.font.bold = True
        bullet.font.size = Pt(8)

        table = docx.add_table(rows=13, cols=3)
        table.style = 'Table Grid'
        table.style.font.size = Pt(8)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for cell in table.columns[0].cells:
            cell.width = 400000
        for cell in table.columns[1].cells:
            cell.width = 3200000
        for cell in table.columns[2].cells:
            cell.width = 1600000
        table.columns[0].width = 400000
        table.columns[1].width = 3200000
        table.columns[2].width = 1600000
        table.cell(0, 0).merge(table.cell(0, 1)).paragraphs[0].add_run(
            self.str_pcm_pre_acadinfo[0]).font.size = Pt(8)
        table.cell(0, 2).paragraphs[0].add_run(
            self.admission_period).font.size = Pt(8)
        table.cell(0, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(1, 0).merge(table.cell(1, 1)).paragraphs[0].add_run(
            self.str_pcm_pre_acadinfo[1]).font.size = Pt(8)

        if self.first_reing:
            table.cell(1, 2).paragraphs[0].add_run('Sí').font.size = Pt(8)
        else:
            table.cell(1, 2).paragraphs[0].add_run('No').font.size = Pt(8)

        table.cell(1, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(2, 0).merge(table.cell(2, 2)).paragraphs[0].add_run(
            self.str_pcm_pre_acadinfo[2]).font.size = Pt(8)
        table.cell(3, 0).merge(table.cell(3, 1)).paragraphs[0].add_run(
            self.str_pcm_pre_acadinfo[3]).font.size = Pt(8)
        table.cell(3, 2).paragraphs[0].add_run(
            self.loss_period).font.size = Pt(8)
        table.cell(3, 2).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        table.cell(3, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(4, 0).merge(table.cell(4, 1)).paragraphs[0].add_run(
            self.str_pcm_pre_acadinfo[4]).font.size = Pt(8)
        table.cell(4, 2).paragraphs[0].add_run(
            str(self.periods_since)).font.size = Pt(8)
        table.cell(4, 2).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        table.cell(4, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(5, 0).merge(table.cell(5, 2)).paragraphs[0].add_run(
            self.str_pcm_pre_acadinfo[5]).font.size = Pt(8)
        table.cell(6, 0).merge(table.cell(6, 1)
                               ).paragraphs[0].add_run(
                                   self.str_pcm_pre_acadinfo[6]).font.size = Pt(8)
        table.cell(6, 2).paragraphs[0].add_run(
            str(self.papa)).font.size = Pt(8)
        table.cell(6, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(7, 0).merge(table.cell(7, 1)).paragraphs[0].add_run(
            self.str_pcm_pre_acadinfo[7]).font.size = Pt(8)
        table.cell(7, 0).merge(table.cell(7, 1)
                               ).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # pylint: disable=no-member
        if self.reason_of_loss == self.RL_ANSWER_PAPA_CREDITOS:
            table.cell(7, 2).paragraphs[0].add_run(
                self.RL_ANSWER_CHOICES[1][1] + '\n' +
                self.RL_ANSWER_CHOICES[2][1]).font.size = Pt(8)
        else:
            table.cell(7, 2).paragraphs[0].add_run(
                self.get_reason_of_loss_display()).font.size = Pt(8)
        table.cell(7, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        mg_cll = table.cell(8, 0).merge(table.cell(8, 2)).paragraphs[0].add_run(
            self.str_pcm_pre_acadinfo[8])
        mg_cll.font.bold = True
        mg_cll.font.size = Pt(8)
        table.cell(9, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(9, 0).paragraphs[0].add_run('1').font.bold = True
        table.cell(10, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(10, 0).paragraphs[0].add_run('2').font.bold = True
        table.cell(11, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(11, 0).paragraphs[0].add_run('3').font.bold = True
        table.cell(12, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(12, 0).paragraphs[0].add_run('4').font.bold = True
        for i in range(9, 13):
            table.cell(
                i, 0).paragraphs[0].runs[0].font.size = Pt(8)
        table.cell(9, 1).paragraphs[0].add_run(
            self.str_pcm_pre_acadinfo[9]).font.size = Pt(8)
        table.cell(10, 1).paragraphs[0].add_run(
            self.str_pcm_pre_acadinfo[10]).font.size = Pt(8)
        table.cell(11, 1).paragraphs[0].add_run(
            self.str_pcm_pre_acadinfo[11]).font.size = Pt(8)
        table.cell(12, 1).paragraphs[0].add_run(
            self.str_pcm_pre_acadinfo[12]).font.size = Pt(8)
        table.cell(9, 2).paragraphs[0].add_run(
            str((self.credits_bag - functools.reduce(
                lambda a, b: a+b, [self.exi_fund_m - self.app_fund_m,
                                   self.exi_fund_o - self.app_fund_o,
                                   self.exi_disc_m - self.app_disc_m,
                                   self.exi_disc_o - self.app_disc_o,
                                   self.exi_free - self.app_free]) -
                 self.credits_english))).font.size = Pt(8)
        table.cell(9, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(10, 2).paragraphs[0].add_run(
            str(functools.reduce(
                lambda a, b: a+b, [self.exi_fund_m - self.app_fund_m,
                                   self.exi_fund_o - self.app_fund_o,
                                   self.exi_disc_m - self.app_disc_m,
                                   self.exi_disc_o - self.app_disc_o,
                                   self.exi_free - self.app_free]))).font.size = Pt(8)
        table.cell(10, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(11, 2).paragraphs[0].add_run(
            str(self.credits_english)).font.size = Pt(8)
        table.cell(11, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        additional_creds_required = self.credits_bag - functools.reduce(
            lambda a, b: a+b, [self.exi_fund_m - self.app_fund_m,
                               self.exi_fund_o - self.app_fund_o,
                               self.exi_disc_m - self.app_disc_m,
                               self.exi_disc_o - self.app_disc_o,
                               self.exi_free - self.app_free]) - self.credits_english
        table.cell(12, 2).paragraphs[0].add_run(
            str(additional_creds_required*(-1) if additional_creds_required < 0 else 0)).font.size = Pt(8)
        table.cell(12, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Optional: Grade needed with N credits to keep student condition.
        if self.reason_of_loss in (self.RL_ANSWER_PAPA, self.RL_ANSWER_PAPA_CREDITOS):
            table = docx.add_table(rows=5, cols=2)
            for col in table.columns:
                for cell in col.cells:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            table.style = 'Table Grid'
            table.style.font.size = Pt(8)
            table.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for cell in table.columns[0].cells:
                cell.width = 3100000
            for cell in table.columns[1].cells:
                cell.width = 2100000
            table.columns[0].width = 3100000
            table.columns[1].width = 2100000
            table.cell(0, 0).merge(table.cell(0, 1)
                                   ).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            table.cell(0, 0).merge(table.cell(0, 1)).paragraphs[0].add_run(
                self.str_pcm_pre_acadinfo[13]).font.size = Pt(8)
            table.cell(1, 0).paragraphs[0].add_run(
                self.str_pcm_pre_acadinfo[14]).font.size = Pt(8)
            table.cell(2, 0).paragraphs[0].add_run(
                self.str_pcm_pre_acadinfo[15]).font.size = Pt(8)
            table.cell(3, 0).paragraphs[0].add_run(
                self.str_pcm_pre_acadinfo[16]).font.size = Pt(8)
            table.cell(4, 0).paragraphs[0].add_run(
                self.str_pcm_pre_acadinfo[17]).font.size = Pt(8)
            table.cell(1, 1).paragraphs[0].add_run(
                str(round(
                    ((3*(
                        self.credits_coursed+12)-self.papa*self.credits_coursed)/12),
                    1))).font.size = Pt(8)
            table.cell(2, 1).paragraphs[0].add_run(
                str(round(
                    ((3*(
                        self.credits_coursed+15)-self.papa*self.credits_coursed)/15),
                    1))).font.size = Pt(8)
            table.cell(3, 1).paragraphs[0].add_run(
                str(round(
                    ((3*(
                        self.credits_coursed+18)-self.papa*self.credits_coursed)/18),
                    1))).font.size = Pt(8)
            table.cell(4, 1).paragraphs[0].add_run(
                str(round(
                    ((3*(
                        self.credits_coursed+21)-self.papa*self.credits_coursed)/21),
                    1))).font.size = Pt(8)

    def rein_credits_summary(self, docx):
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        bullet = paragraph.add_run(self.str_pcm_pre[4])
        bullet.font.bold = True
        bullet.font.size = Pt(8)
        credits_data = [[self.exi_fund_m, self.exi_fund_o, self.exi_disc_m,
                         self.exi_disc_o, self.exi_free],
                        [self.app_fund_m, self.app_fund_o, self.app_disc_m,
                         self.app_disc_o, self.app_free],
                        [self.exi_fund_m - self.app_fund_m,
                         self.exi_fund_o - self.app_fund_o,
                         self.exi_disc_m - self.app_disc_m,
                         self.exi_disc_o - self.app_disc_o,
                         self.exi_free - self.app_free]]
        case = 'REINGRESO'
        table_credits_summary(docx, credits_data, case)

        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        bullet = paragraph.add_run(self.str_pcm_pre[5])
        bullet.font.size = Pt(8)

    def rein_recommends(self, docx):
        details = []
        details.append(
            # pylint: disable=no-member
            self.get_academic_program_display())

        # Migrate to case_utils?
        year = str(self.comitee_date)[0:4]
        month = str(self.comitee_date)[5:7]
        day = str(self.comitee_date)[8:10]
        details.append(day + '-' + month + '-' + year)
        details.append(self.comitee_act)
        details.append(str(self.comitee_date)[0:4])
        if self.advisor_response == self.ARCR_APROBAR:
            details.append(True)
        else:
            details.append(False)
        table_recommend(docx, details)

    def extra_credits(self, paragraph):
        paragraph.add_run(' ' + self.str_pcm_pre[18] +
                          str((self.credits_bag -
                               functools.reduce(
                                   lambda a, b: a+b, [self.exi_fund_m - self.app_fund_m,
                                                      self.exi_fund_o - self.app_fund_o,
                                                      self.exi_disc_m - self.app_disc_m,
                                                      self.exi_disc_o - self.app_disc_o,
                                                      self.exi_free - self.app_free]) -
                               self.credits_english)*(-1)) +
                          self.str_pcm_pre[19])

    def get_analysis(self):
        analysis = []
        modifier = 'No h' if self.first_reing else 'H'
        analysis.append(self.str_analysis[0].format(
            modifier, self.regulations['008|2008|CSU'][0]
        ))
        analysis.append(self.str_analysis[1].format(
            self.regulations['239|2009|VAC'][0],
            # pylint: disable=no-member
            self.loss_period, self.get_reason_of_loss_display()
        ))
        modifier = 'T' if self.papa >= 2.7 else 'No t'
        analysis.append(self.str_analysis[2].format(
            modifier, self.regulations['239|2009|VAC'][0],
            self.regulations['008|2008|CSU'][0], self.papa
        ))
        modifier = 'D' if functools.reduce(
            lambda a, b: a+b, [self.exi_fund_m - self.app_fund_m,
                               self.exi_fund_o - self.app_fund_o,
                               self.exi_disc_m - self.app_disc_m,
                               self.exi_disc_o - self.app_disc_o,
                               self.exi_free - self.app_free]) > 0 else 'No d'
        analysis.append(self.str_analysis[3].format(
            modifier, self.regulations['008|2008|CSU'][0],
            functools.reduce(
                lambda a, b: a+b, [self.exi_fund_m - self.app_fund_m,
                                   self.exi_fund_o - self.app_fund_o,
                                   self.exi_disc_m - self.app_disc_m,
                                   self.exi_disc_o - self.app_disc_o,
                                   self.exi_free -
                                   self.app_free]), self.regulations['012|2014|VAC'][0]
        ))
        modifier = '' if self.request_in_date else 'no '
        analysis.append(self.str_analysis[4].format(modifier))
        return analysis + self.extra_analysis

    def pcm(self, docx):
        add_analysis_paragraph(docx, self.get_analysis())
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.str_answer + ':\n').font.bold = True
        paragraph.add_run(self.str_comittee_header + ' ')
        self.pcm_answer(paragraph)
        # if self.request_in_date: To ommit tables when the request isn't in time
        self.cm_pcm_paragraph(docx)
        self.rein_general_data_table(docx)
        self.rein_academic_info(docx)
        self.rein_credits_summary(docx)
        self.rein_recommends(docx)

    def cm(self, docx):
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        # if self.request_in_date: To ommit tables when the request isn't in time
        paragraph.add_run(self.str_council_header + ' ')
        self.cm_answer(paragraph)
        self.cm_pcm_paragraph(docx)
        self.rein_general_data_table(docx)
        self.rein_academic_info(docx)
        self.rein_credits_summary(docx)
        self.rein_recommends(docx)

    def pcm_answer(self, paragraph):
        paragraph.add_run(
            # pylint: disable=no-member
            self.get_approval_status_display().upper() + ' ').font.bold = True
        if not self.request_in_date:
            self.out_of_date_answer(paragraph)
        else:
            aff = self.is_affirmative_response_advisor_response()
            self.standard_answer(paragraph, aff)

    def cm_pcm_paragraph(self, docx):
        pass
        # if not self.request_in_date:
        #     return  # Skip when it's out of date
        # if self.credits_english == 0:
        #     return
        # para = docx.add_paragraph()
        # para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        # para.paragraph_format.space_after = Pt(0)
        # para.add_run(self.str_pcm_pre[6] + self.student_name +
        #              self.str_pcm_pre[7] + str(functools.reduce(
        #                  lambda a, b: a+b, [self.exi_fund_m - self.app_fund_m,
        #                                     self.exi_fund_o - self.app_fund_o,
        #                                     self.exi_disc_m - self.app_disc_m,
        #                                     self.exi_disc_o - self.app_disc_o,
        #                                     self.exi_free - self.app_free])))
        # para.add_run(self.str_pcm_pre[8] +
        #              # pylint: disable=no-member
        #              self.get_academic_program_display())
        # para.add_run(
        #     self.str_pcm_pre[9] + str(self.credits_english) + self.str_pcm_pre[10])
        # para.add_run(self.str_pcm_pre[11])
        # para.add_run(str(functools.reduce(
        #     lambda a, b: a+b, [self.exi_fund_m - self.app_fund_m,
        #                        self.exi_fund_o - self.app_fund_o,
        #                        self.exi_disc_m - self.app_disc_m,
        #                        self.exi_disc_o - self.app_disc_o,
        #                        self.exi_free - self.app_free]) +
        #                  (self.credits_bag - functools.reduce(
        #                      lambda a, b: a+b, [self.exi_fund_m - self.app_fund_m,
        #                                         self.exi_fund_o - self.app_fund_o,
        #                                         self.exi_disc_m - self.app_disc_m,
        #                                         self.exi_disc_o - self.app_disc_o,
        #                                         self.exi_free - self.app_free]) -
        #                   self.credits_english)))
        # para.add_run(self.str_pcm_pre[12])
        # para = docx.add_paragraph()
        # para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        # para.paragraph_format.space_after = Pt(0)
        # para.add_run(
        #     self.str_pcm_pre[13] + self.regulations['008|2008|CSU'][0])
        # para.add_run(self.str_pcm_pre[14])
        # para.add_run(self.str_pcm_pre[15]).font.italic = True
        # para.add_run(self.str_pcm_pre[16] +
        #              str((self.credits_bag - functools.reduce(
        #                  lambda a, b: a+b, [self.exi_fund_m - self.app_fund_m,
        #                                     self.exi_fund_o - self.app_fund_o,
        #                                     self.exi_disc_m - self.app_disc_m,
        #                                     self.exi_disc_o - self.app_disc_o,
        #                                     self.exi_free - self.app_free]) -
        #                   self.credits_english)*(-1)))
        # para.add_run(self.str_pcm_pre[17])
        # para.add_run(
        #     # pylint: disable=no-member
        #     self.get_academic_program_display() + '.')

    def cm_answer(self, paragraph):
        paragraph.add_run(
            # pylint: disable=no-member
            self.get_approval_status_display().upper() + ' ').font.bold = True
        aff = self.is_affirmative_response_approval_status()
        self.standard_answer(paragraph, aff)

    def standard_answer(self, paragraph, affirmative):

        paragraph.add_run(self.str_pcm_pre[0])
        paragraph.add_run(self.reing_period)

        if ((self.credits_bag - functools.reduce(
                lambda a, b: a+b, [self.exi_fund_m - self.app_fund_m,
                                   self.exi_fund_o - self.app_fund_o,
                                   self.exi_disc_m - self.app_disc_m,
                                   self.exi_disc_o - self.app_disc_o,
                                   self.exi_free - self.app_free]) - self.credits_english)) < 0:
            # Y otorga n créditos adicionales:
            self.extra_credits(paragraph)

        if affirmative:
            paragraph.add_run(self.str_pcm_pre[1])
        else:
            paragraph.add_run('. Debido a que ' + self.council_decision + '.')

        paragraph.add_run('({}).'.format(
            self.regulations['012|2014|VAC'][0] + "; Artículo 46, " +
            self.regulations['008|2008|CSU'][0]))

    def out_of_date_answer(self, paragraph):
        paragraph.add_run(self.str_out_date[0].format(self.reing_period))

    def resource_analysis(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.pcm_answer(last_paragraph)

    def resource_pre_answer(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.pcm_answer(last_paragraph)

    def resource_answer(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.cm_answer(last_paragraph)
