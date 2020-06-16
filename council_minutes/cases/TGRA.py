import datetime
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from mongoengine import StringField, IntField, DateField
from ..models import Request, Subject
from .case_utils import table_subjects, add_analysis_paragraph, num_to_month

class TGRA(Request):

    full_name = 'Inscripción trabajo de grado'

    TGRA_PASANTIA = 'TP'
    TGRA_TRABAJO = 'TT'
    TGRA_CHOICES = (
        (TGRA_PASANTIA, 'Trabajo de grado - Modalidad Pasantía'),
        (TGRA_TRABAJO, 'Trabajo de grado - Modalidad Trabajos Investigativos'),
    )

    period_inscription = StringField(
        display='Periodo de inscripción trabajo de grado', 
        choices=Request.PERIOD_CHOICES, default=Request.PERIOD_DEFAULT)
    type_tgra = StringField(
        choices=TGRA_CHOICES, default=TGRA_PASANTIA, display='Tipo de trabajo de grado')
    title = StringField(default='', display='Título del trabajo de grado')
    organization = StringField(
        default='', display='Empresa donde realiza pasantía')
    professor = StringField(
        default='', display='Profesor director del trabajo')
    dc_approved = IntField(display='Número de créditos aprobados', default=0)
    commite_cm = IntField(default=0, display='Acta de comité')
    commite_cm_date = DateField(
        display='Fecha acta de comité', default=datetime.date.today)

    regulation_list = ['026|2012|CSU', '040|2017|CSU']  # List of regulations

    str_cm = [
        'inscribir la(s) siguiente(s) asignatura(s) en el periodo académico {}, en modalidad {}, ' +
        'bajo la dirección del profesor {}, debido a que {}realiza correctamente la solicitud.'
    ]

    str_pcm = [
        'Formato de registro diligenciado (Artículo 8): Revisado.',
        'Dirección de un profesor de la Universidad, aceptado y formalizado (Artículo 6) en Acta ' +
        '{} de Comité del {}{}{} en modalidad: {}.',
        '{}a cursado por lo menos {} céditos del componente disciplinar. SIA: {} créditos',
        'Tífulo del trabajo de grado: {}.',
        'Institución: {}',
        'Docente encargado: {}'
        'Debido a que formalizó la inscripción de la asignatura Trabajo de Grado en los plazos es' +
        'tablecidos.',
    ]

    def cm(self, docx):
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.str_council_header + ' ')
        self.cm_answer(paragraph)
        # pylint: disable=no-member
        table_subjects(docx,
                       [['2015289' if self.type_tgra == 'TP' else '202599',
                         self.get_type_tgra_display(), '1', Subject.TIP_PRE_TRAB_GRADO[1], '6']])

    def cm_answer(self, paragraph):
        paragraph.add_run(
            # pylint: disable=no-member
            self.get_approval_status_display().upper() + ' ').font.bold = True
        paragraph.add_run(self.str_cm[0].format(
            # pylint: disable=no-member
            self.academic_period, self.get_type_tgra_display(), self.professor, self.council_decision))

    def pcm(self, docx):
        self.pcm_analysis(docx)
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.str_answer + ': ').bold = True
        paragraph.add_run(self.str_comittee_header + ' ')
        self.pcm_answer(paragraph)
        if self.is_affirmative_response_advisor_response():
            # pylint: disable=no-member
            table_subjects(docx,
                           [['2015289' if self.type_tgra == 'TP' else '202599',
                             self.get_type_tgra_display(), '1',
                             Subject.TIP_PRE_TRAB_GRADO[1], '6']])

    def pcm_analysis(self, docx):
        analysis_list = []
        analysis_list += [self.str_pcm[0]]
        # pylint: disable=no-member
        analysis_list += [self.str_pcm[1].format(
            self.commite_cm,
            self.commite_cm_date.day,
            num_to_month(self.commite_cm_date.month),
            self.commite_cm_date.year,
            self.get_type_tgra_display()
        )]
        analysis_list += [self.str_pcm[2].format(
            'H' if self.dc_approved >= self.disciplinar_credits_approved_for_program() else 'No h',
            self.disciplinar_credits_approved_for_program(),
            self.dc_approved
        )]
        analysis_list += [self.str_pcm[3].format(self.title)]
        analysis_list += [self.str_pcm[4].format(self.organization)]
        analysis_list += [self.str_pcm[5].format(self.professor)]
        analysis_list += self.extra_analysis
        add_analysis_paragraph(docx, analysis_list)

    def pcm_answer(self, paragraph):
        paragraph.add_run(
            # pylint: disable=no-member
            self.get_advisor_response_display().upper() + ' ').font.bold = True
        paragraph.add_run(self.str_cm[0].format(
            # pylint: disable=no-member
            self.academic_period, self.get_type_tgra_display(), self.professor, self.council_decision))

    def disciplinar_credits_approved_for_program(self):
        if self.academic_program == self.PI_AGRICOLA:
            return 68
        elif self.academic_program == self.PI_CIVIL:
            return 77
        elif self.academic_program == self.PI_DE_SISTEMAS_Y_COMPUTACION:
            return 60
        elif self.academic_program == self.PI_INDUSTRIAL:
            return 70
        elif self.academic_program == self.PI_ELECTRICA:
            return 57
        elif self.academic_program == self.PI_ELECTRONICA:
            return 63
        elif self.academic_program == self.PI_MECANICA:
            return 69
        elif self.academic_program == self.PI_MECATRONICA:
            return 69
        elif self.academic_program == self.PI_QUIMICA:
            return 60
        else:
            raise AssertionError('TGRA for no PRE academic program!')

    def resource_analysis(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.pcm_answer(last_paragraph)
    
    def resource_pre_answer(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.pcm_answer(last_paragraph)

    def resource_answer(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.cm_answer(last_paragraph)
