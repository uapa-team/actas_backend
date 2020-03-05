from docx.shared import Pt
from num2words import num2words
from docx.enum.text import WD_ALIGN_PARAGRAPH
from mongoengine import StringField, BooleanField, FloatField
from ..models import Request
from .case_utils import add_analysis_paragraph


class CPAC(Request):

    full_name = 'Cancelación periodo académico'

    academic_profile = StringField(
        default=Request.PROFILE_INVE, choices=Request.PROFILE_CHOICES,
        display='Perfil de programa curricular en caso de posgrado')
    percentaje = FloatField(
        display='Porcentaje de devolución ',
        min_value=0.0, max_value=100.0, default=0.0)
    is_fortuitous = BooleanField(
        display='Se considera caso fortuito', default=False)

    regulation_list = ['008|2008|CSU', '032|2010|CSU',
                       '1416|2013|RE']  # List of regulations

    str_cm = [
        'Cancelar la totalidad de las asignaturas en el periodo {}, en el programa de {} ({}), ',
        'Devolución proporcional del {} por ciento ({}%) del valor pagado por ' +
        'concepto de derechos de matrícula del periodo {}, teniendo en cuenta la' +
        'fecha de presentación de la solicitud y que le fue aprobada la cancelación de' +
        'periodo en Acta {} de {} de Consejo de Facultad.'
    ]
    str_cm_reason = [
        'debido a que justifica adecauadamente la fuerza mayor o caso fortuito.',
        'debido a que la situación expuesta no constituye causa extraña (no es una situación ' +
        'intempestiva, insuperable o irresistible), por tanto, no es una situación de fuerza' +
        'mayor o caso fortuito que implique la cancelación del periodo académico. '
    ]

    str_pcm = [
        'SIA: {} ({}).',
        'Perfil de {}.',
        'El comité {}lo considera fuerza mayor o caso fortuito.'
    ]

    def add_paragraph(self, docx):
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        return paragraph

    def add_paragraph_bullet_list(self, docx):
        paragraph = self.add_paragraph(docx)
        paragraph.style = 'List Bullet'
        return paragraph

    def cm(self, docx):
        paragraph = self.add_paragraph(docx)
        paragraph.add_run(self.str_council_header + ' ')
        paragraph.add_run(
            self.get_approval_status_display().upper() + ':').font.bold = True
        self.cm_answer(self.add_paragraph_bullet_list(docx))
        self.cm_answer1(self.add_paragraph_bullet_list(docx))

    def cm_answer(self, paragraph):
        answer = self.str_cm[0].format(
            self.academic_period,
            self.get_academic_program_display(),
            self.academic_program,
        )
        if self.is_affirmative_response_approval_status():
            reason = self.str_cm_reason[0]
        else:
            reason = self.str_cm_reason[1]
        paragraph.add_run(answer + reason)

    def cm_answer1(self, paragraph):
        answer = self.str_cm[1].format(
            num2words(self.percentaje, lang='es'),
            self.percentaje,
            self.academic_period,
            self.consecutive_minute,
            self.year
        )
        paragraph.add_run(answer)

    def pcm(self, docx):
        self.pcm_analysis(docx)
        paragraph = self.add_paragraph(docx)
        paragraph.add_run(self.str_council_header + ' ')
        paragraph.add_run(
            self.get_approval_status_display().upper() + ':').font.bold = True
        self.cm_answer(self.add_paragraph_bullet_list(docx))
        self.cm_answer1(self.add_paragraph_bullet_list(docx))

    def pcm_analysis(self, docx):
        analysis_list = []
        analysis_list += [self.str_pcm[0].format(
            # pylint: disable=no-member
            self.get_academic_program_display(),
            self.academic_program
        )]
        if not self.is_pre():
            analysis_list += [self.str_pcm[1].format(
                # pylint: disable=no-member
                self.get_academic_profile_display()
            )]
        analysis_list += [self.str_pcm[2].format(
            '' if self.is_fortuitous else 'no '
        )]
        analysis_list += self.extra_analysis
        add_analysis_paragraph(docx, analysis_list)

    def pcm_answer(self, paragraph):
        answer = self.str_cm[0].format(
            self.academic_period,
            self.get_academic_program_display(),
            self.academic_program,
        )
        if self.is_affirmative_response_approval_status():
            reason = self.str_cm_reason[0]
        else:
            reason = self.str_cm_reason[1]
        paragraph.add_run(answer + reason)

    def pcm_answer1(self, paragraph):
        answer = self.str_cm[1].format(
            num2words(self.percentaje, lang='es'),
            self.percentaje,
            self.academic_period,
            self.consecutive_minute,
            self.year
        )
        paragraph.add_run(answer)

    def resource_analysis(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.pcm_answer(last_paragraph)
    
    def resource_pre_answer(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.pcm_answer(last_paragraph)

    def resource_answer(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.cm_answer(last_paragraph)
