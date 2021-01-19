from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from mongoengine import StringField, BooleanField, IntField
from mongoengine import EmbeddedDocumentListField, FloatField
from ..models import Request, Subject
from .case_utils import add_analysis_paragraph, table_general_data, string_to_date, table_approvals, indent_table


class TRASPOS(Request):

    class HomologatedSubject(Subject):
        TIP_OBLIGATORIA = 'B'
        TIP_ACTIV_ACADEMICA = 'C'
        TIP_TRAB_GRADO = 'P'
        TIP_ELEGIBLE = 'L'
        TIP_CHOICES = (
            (TIP_OBLIGATORIA, 'Obligatoria'),
            (TIP_ACTIV_ACADEMICA, 'Actividad académica'),
            (TIP_TRAB_GRADO, 'Trabajo de grado'),
            (TIP_ELEGIBLE, 'Elegible'),
        )
        new_name = StringField(
            required=True, display='Nuevo Nombre Asignatura', default='')
        new_code = StringField(
            required=True, display='Nuevo Código Asignatura', default='')
        tipology = StringField(
            required=True, choices=TIP_CHOICES, display='Tipología', default=TIP_OBLIGATORIA)
        group = None
        grade = StringField(required=True, default='', display='Calificación')

    full_name = 'Traslado de programa curricular (Posgrado)'

    TT_INTERCAMPUS = 'TTIC'
    TT_INTERFACULTY = 'TTIF'
    TT_INTRAFACULTY = 'TTRF'
    TT_CHOICES = (
        (TT_INTERCAMPUS, 'Traslado Intersede'),
        (TT_INTERFACULTY, 'Traslado Interfacultad'),
        (TT_INTRAFACULTY, 'Traslado Intrafacultad'),
    )
    TC_BOGOTA = 'BOG'
    TC_MEDELLIN = 'MED'
    TC_MANIZALES = 'MAN'
    TC_PALMIRA = 'PAL'
    TC_LAPAZ = 'PAZ'
    TC_CHOICES = (
        (TC_BOGOTA, 'Bogotá'),
        (TC_MEDELLIN, 'Medellín'),
        (TC_MANIZALES, 'Manizales'),
        (TC_PALMIRA, 'Palmira'),
        (TC_LAPAZ, 'La Paz'),
    )

    PF_INVEST = 'INVE'
    PF_PROFUN = 'PROF'
    PF_CHOICES = (
        (PF_INVEST, 'Investigación'),
        (PF_PROFUN, 'Profundización'),
    )

    at_least_one_period = BooleanField(
        required=True, default=True,
        display='¿Ha cursado por lo menos un periodo académico del primer plan de estudios?')
    finish_first_plan = BooleanField(
        required=True, default=False,
        display='¿Ha culminado el primer plan de estudios?')
    have_entitled_to_enrrol = BooleanField(
        required=True, default=True,
        display='¿Tiene derecho a renovar matrícula?')
    enroled_number = IntField(
        required=True, min_value=1, default=1,
        display='Número de matrículas')
    currently_studying_double_degree = BooleanField(
        required=True, default=False,
        display='¿Está cursando actualmente doble titulación?')
    available_quota_for_transit = BooleanField(
        required=True, default=True,
        display='¿El cupo de créditos para traslado es suficiente?')
    availabe_quota_number = IntField(
        required=True, default=0, min_value=0,
        display='Número de cupos ofertados para traslado')
    campus_origin = StringField(
        required=True, default=TC_BOGOTA, choices=TC_CHOICES,
        display='Sede de origen')
    campus_destination = StringField(
        required=True, default=TC_BOGOTA, choices=TC_CHOICES,
        display='Sede de destino')
    transit_type = StringField(
        required=True, default=TT_INTRAFACULTY, choices=TT_CHOICES,
        display='Tipo de traslado')
    admission_period = StringField(
        required=True, display='Periodo de admisión del estudiante', default='')
    same_degree = BooleanField(
        required=True, default=False,
        display='¿Estos planes de estudios conducen al mismo título?')
    transit_program_code = StringField(
        required=True,
        display='Código del plan de estudios de destino',
        default='')
    transit_program_name = StringField(
        required=True,
        display='Nombre del plan de estudios de destino',
        default='')
    transit_program_profile = StringField(
        required=True,
        display='Perfil del plan de estudios de destino',
        choices=PF_CHOICES, default=PF_INVEST)
    origin_program_code = StringField(
        required=True,
        display='Código del plan de estudios de origen',
        default='')
    origin_program_name = StringField(
        required=True,
        display='Nombre del plan de estudios de origen',
        default='')
    origin_program_profile = StringField(
        required=True,
        display='Perfil del plan de estudios de origen',
        choices=PF_CHOICES, default=PF_INVEST)
    enrroled = BooleanField(
        required=True, default=True,
        display='¿Se encuentra matriculado en el semestre de presentar la solicitud?')
    prev_plan = BooleanField(
        required=True, default=False,
        display='¿Tuvo calidad de estudiante en el plan de estudios destino?')
    completion_percentage = FloatField(
        required=True, default=0.0, min_value=0.0, max_value=100.0,
        display='Porcentaje de créditos aprobados en el plan de estudios origen')
    homologated_subjects = EmbeddedDocumentListField(
        HomologatedSubject, display='Cuadro de equivalencias y convalidaciones')
    agreement_number = StringField(
        required=True, display='Número del acuerdo del CF que reglamenta el plan de estudios',
        default='')
    agreement_year = IntField(
        required=True, display='Año del acuerdo del CF que reglamenta el plan de estudios',
        default=0)

    regulation_list = ['008|2008|CSU', '089|2014|CAC',
                       '155|2014|CSU']  # List of regulations

    str_cm = ['traslado {} del programa {}, plan de estudios de {} al programa {}, plan de ' +
              'estudios de {}, en el periodo académico {}', ', condicionado a conservar la ' +
              'calidad de estudiante al finalizar el periodo académico {}. (Artículo 39 ' +
              'del {} y {}).', 'debido a que']

    str_table = ['Estudiante', 'DNI', 'Plan de estudios de origen (1er plan) - Sede {}',
                 'Código del plan de estudios de origen (1er plan)',
                 'Plan de estudios de destino (2° plan) - Sede {}',
                 'Código del plan de estudios de destino (2° plan)',
                 'Fecha de la solicitud', '¿Estos planes de estudios conducen al mismo título?',
                 'Periodo para el cual fue admitido', '¿El solicitante se encuentra matriculado' +
                 ' en el semestre de presentar la solicitud?', '¿El solicitante tuvo calidad' +
                 ' de estudiante en el plan de estudios de destino (2° plan)?',
                 'Porcentaje de créditos aprobados en el plan de estudios origen (1er plan)',
                 'CUADRO EQUIVALENCIAS Y CONVALIDACIONES DE ASIGNATURAS CURSADAS Y APROBADAS' +
                 ' HASTA LA FECHA DE PRESENTACIÓN DE LA SOLICITUD POR PARTE DEL ESTUDIANTE.',
                 'Universidad Nacional de Colombia plan de estudios de {}',
                 'La oferta de asignaturas en cada una de las agrupaciones y componentes del' +
                 ' plan de estudios del programa de {} - perfil {}, la encuentra en el Acuerdo' +
                 ' No. {} del año {}, expedido por Consejo de Facultad de Ingeniería.']

    srt_titles = ['I) Datos Generales', 'II) Información Académica']

    list_analysis = ['Viene del plan {}, perfil de {} de la sede {}.',
                     'a tenido calidad de estudiante en ese programa previamente ' +
                     '(Parágrafo 1. Artículo 2, {}). Universitas: OK.',
                     'a Culminado el primer plan de estudios.',
                     'iene derecho a renovar la matrícula. Universitas: OK.',
                     'a cursado por lo menos un periodo académico del primer plan ' +
                     'de estudios (Artículo 39, {}). SIA: OK.', 'Ha cursado {} periodos ' +
                     'académicos desde {}.', 'ay cupos disponibles en el plan de estudios ' +
                     'del programa curricular solicitado (Estipulados por Consejo de ' +
                     'Facultad): {} cupos.', 'El estudiante {}cuenta con el suficiente cupo ' +
                     'de créditos para inscribir las asignaturas pendientes de ' +
                     'aprobación en el nuevo plan (Artículo 3, {}).']

    def get_next_period(self, actual_period):
        year = int(actual_period[0:4])
        semester = int(actual_period[5])
        if semester == 1:
            return str(year) + '-' + str(semester + 1) + 'S'
        elif semester == 2:
            return str(year + 1) + '-1S'

    def cm(self, docx):
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.str_council_header + ' ')
        self.cm_answer(paragraph)
        self.add_tables(docx)

    def cm_answer(self, paragraph):
        # pylint: disable=no-member
        paragraph.add_run(
            self.get_approval_status_display().upper() + ' ').font.bold = True
        paragraph.add_run(
            self.str_cm[0].format(
                self.get_transit_type_display().split(
                    ' ')[1].lower(), self.origin_program_name,
                self.get_origin_program_profile_display().lower(), self.transit_program_name,
                self.get_transit_program_profile_display().lower(), self.get_next_period(
                    self.academic_period)))
        if self.is_affirmative_response_approval_status():
            self.cm_af(paragraph)
        else:
            self.cm_ng(paragraph)

    def pcm(self, docx):
        self.pcm_analysis(docx)
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.str_answer + ': ').font.bold = True
        paragraph.add_run(self.str_comittee_header + ' ')
        self.pcm_answer(paragraph)
        self.add_tables(docx)

    def pcm_answer(self, paragraph):
        # pylint: disable=no-member
        paragraph.add_run(
            self.get_advisor_response_display().upper() + ' ').font.bold = True
        paragraph.add_run(
            self.str_cm[0].format(
                self.get_transit_type_display().split(
                    ' ')[1].lower(), self.origin_program_name,
                self.get_origin_program_profile_display().lower(), self.transit_program_name,
                self.get_transit_program_profile_display().lower(), self.get_next_period(
                    self.academic_period)))
        if self.is_affirmative_response_approval_status():
            self.cm_af(paragraph)
        else:
            self.cm_ng(paragraph)

    def cm_af(self, paragraph):
        paragraph.add_run(self.str_cm[1].format(
            self.academic_period, Request.regulations['008|2008|CSU'][0],
            Request.regulations['089|2014|CAC'][0]))

    def cm_ng(self, paragraph):
        paragraph.add_run(
            ', ' + self.str_cm[2] + ' ' + self.council_decision + '.')

    def pcm_analysis(self, docx):
        # pylint: disable=no-member
        final_analysis = []
        final_analysis += [self.list_analysis[0].format(
            self.origin_program_name, self.get_origin_program_profile_display().lower(),
            self.get_campus_origin_display())]
        aux_str = 'H' if self.prev_plan else 'No h'
        final_analysis += [aux_str + self.list_analysis[1]
                           .format(Request.regulations['089|2014|CAC'][0])]
        aux_str = 'H' if self.finish_first_plan else 'No h'
        final_analysis += [aux_str + self.list_analysis[2]]
        aux_str = 'T' if self.have_entitled_to_enrrol else 'No t'
        final_analysis += [aux_str + self.list_analysis[3]]
        aux_str = 'H' if self.at_least_one_period else 'No h'
        final_analysis += [aux_str + self.list_analysis[4]
                           .format(Request.regulations['008|2008|CSU'][0])]
        final_analysis += [self.list_analysis[5].format(
            self.enroled_number, self.admission_period)]
        aux_str = 'H' if self.availabe_quota_number > 0 else 'No h'
        final_analysis += [aux_str +
                           self.list_analysis[6].format(self.availabe_quota_number)]
        aux_str = '' if self.available_quota_for_transit else 'no '
        final_analysis += [self.list_analysis[7].format(
            aux_str, Request.regulations['089|2014|CAC'][0])]
        for extra_a in self.extra_analysis:
            final_analysis += [extra_a]
        add_analysis_paragraph(docx, final_analysis)

    def add_tables(self, docx):
        paragraph = docx.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(self.srt_titles[0])
        run.font.bold = True
        run.font.size = Pt(8)
        # pylint: disable=no-member
        general_data = [
            [self.str_table[0], self.student_name],
            [self.str_table[1], self.student_dni],
            [self.str_table[2].format(
                self.get_campus_origin_display()), self.origin_program_name],
            [self.str_table[3], self.origin_program_code],
            [self.str_table[4].format(
                self.get_campus_destination_display()), self.transit_program_name],
            [self.str_table[5], self.transit_program_code],
            [self.str_table[6], string_to_date(str(self.date))],
            [self.str_table[7], 'Sí' if self.same_degree else 'No'],
        ]
        table_general_data(general_data, 'TRASLADO', docx)
        paragraph = docx.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(self.srt_titles[1])
        run.font.bold = True
        run.font.size = Pt(8)
        table = docx.add_table(rows=4, cols=2, style='Table Grid')
        table.alignment = WD_ALIGN_PARAGRAPH.LEFT
        indent_table(table, 1700)
        table.columns[0].width = 4350000
        table.columns[1].width = 850000
        for cell in table.columns[0].cells:
            cell.width = 4350000
        for cell in table.columns[1].cells:
            cell.width = 850000
        table.cell(0, 0).paragraphs[0].add_run(self.str_table[8])
        table.cell(0, 1).paragraphs[0].add_run(self.admission_period)
        table.cell(1, 0).paragraphs[0].add_run(self.str_table[9])
        table.cell(1, 1).paragraphs[0].add_run('Sí' if self.enrroled else 'No')
        table.cell(2, 0).paragraphs[0].add_run(self.str_table[10])
        table.cell(2, 1).paragraphs[0].add_run(
            'Sí' if self.prev_plan else 'No')
        table.cell(3, 0).paragraphs[0].add_run(self.str_table[11])
        table.cell(3, 1).paragraphs[0].add_run(
            str(self.completion_percentage) + '%')
        paragraph = docx.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(self.str_table[12])
        run.font.bold = True
        run.font.underline = True
        run.font.size = Pt(8)
        subjects = []
        for sbj in self.homologated_subjects:
            subjects.append([self.academic_period, sbj.new_code, sbj.new_name,
                             str(sbj.credits), sbj.tipology, sbj.grade, sbj.name, sbj.grade])
        details = [self.student_name, self.student_dni,
                   self.transit_program_code, self.str_table[13].format(self.origin_program_name)]
        table_approvals(docx, subjects, details)
        paragraph = docx.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = paragraph.add_run(self.str_table[14].format(
            self.transit_program_name,
            self.get_transit_program_profile_display().lower(), self.agreement_number,
            str(self.agreement_year)))
        run.font.underline = True
        run.font.size = Pt(8)

    def resource_analysis(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.pcm_answer(last_paragraph)
    
    def resource_pre_answer(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.pcm_answer(last_paragraph)

    def resource_answer(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.cm_answer(last_paragraph)
