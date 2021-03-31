import datetime
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from num2words import num2words
from mongoengine import DateTimeField, StringField, IntField, BooleanField
from ..models import Request
from .case_utils import num_to_month, add_analysis_paragraph
import datetime


class EPCS(Request):

    full_name = 'Exención de pago por créditos sobrantes de pregrado'

    HC_CESAR = 'CE'
    HC_BOGOTA = 'BO'
    HC_MEDELLIN = 'ME'
    HC_MANIZALES = 'MA'
    HC_PALMIRA = 'PA'
    HC_AMAZONIA = 'AM'
    HC_CARIBE = 'CA'
    HC_ORINOQUIA = 'OR'
    HC_TUMACO = 'TU'
    HEAD_CHOICES = (
        (HC_CESAR, 'Sede Cesar'),
        (HC_BOGOTA, 'Sede Bogotá'),
        (HC_MEDELLIN, 'Sede Medellín'),
        (HC_MANIZALES, 'Sede Manizales'),
        (HC_PALMIRA, 'Sede Palmira'),
        (HC_AMAZONIA, 'Sede Amazonía'),
        (HC_CARIBE, 'Sede Caribe'),
        (HC_ORINOQUIA, 'Sede Orinoquía'),
        (HC_TUMACO, 'Sede Tumaco')
    )

    CJT_ANSWER_DEFAULT = 'DF'
    CJT_ANSWER_JUST_DEB = 'JD'
    CJT_ANSWER_OTRO = 'OT'

    CJT_ANSWER_CHOICES = (
        (CJT_ANSWER_DEFAULT,''),
        (CJT_ANSWER_JUST_DEB,'Justifica debidamente su solicitud'),
        (CJT_ANSWER_OTRO, 'Otro')
    )

    council_decision = StringField(
        max_length=255, choices=CJT_ANSWER_CHOICES,
        default=CJT_ANSWER_DEFAULT, display='Justificación del Consejo')
    academic_profile = StringField(
        default=Request.PROFILE_INVE, choices=Request.PROFILE_CHOICES,
        display='Perfil de programa curricular')
    enrolled_before_preprogram = BooleanField(
        display='Matriculado periodo posterior al pregrado', default=False)
    finalized_period = StringField(
        display='Periodo de culminación de estudios de pregrado',
        choices=Request.PERIOD_CHOICES,
        default=Request.PERIOD_DEFAULT)
    initial_period = StringField(
        display='Periodo de ingreso del posgrado',
        choices=Request.PERIOD_CHOICES,
        default=Request.PERIOD_DEFAULT)
    is_in_right_date = BooleanField(
        display='Solicitud realizada en fechas debidas', default=True)
    right_date = DateTimeField(
        display='Fecha máxima para realizar solicitud.', default=datetime.date.today)
    points = IntField(display='Cantidad de puntos a eximir', default=0)
    # CARE: Choices must be all the programas in the university not only engineering faculty
    bacheilor_program = StringField(
        choices=Request.PLAN_CHOICES, display='Programa de pregrado', default=Request.PI_AGRICOLA)
    headquarters = StringField(
        choices=HEAD_CHOICES, display='Sede donde culminó sus estudios de pregrado', default=HC_BOGOTA)

    # List of regulations
    regulation_list = ['014|2008|CAC']

    str_cm = [
        'otorgar exención del pago de {} ({}) puntos de Derechos Académicos, a partir del periodo' +
        ' académico {}, y durante el siguiente periodo académico, por tener créditos disponibles ' +
        'al finalizar su estudios del programa de pregrado {} ({}), en la {}. ',
        'El cálculo de los créditos disponibles se realiza con base en el cupo de créditos establ' +
        'ecido en el Artículo 2 del {}.'
    ]

    str_pcm = [
        'SIA: Admitido al programa {} ({}) en perfil de {}.',
        'En el año posterior a la culminación de sus estudios de pregrado, el estudiante {}estuvo' +
        ' matriculado en un programa de posgrado. Culminó sus estudios en el periodo {} e ingresó' +
        ' al posgrado en el periodo {}.',
        'La solicitud {}es presentada dentro de las fechas límite establecidas por el reglamento:' +
        ' 2 semanas después de la publicación de resultados de admitidos.',
        'La fecha límite para presentar la solicitud fue el {}.',
    ]

    def cm(self, docx):
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.str_council_header + ' ')
        self.cm_answer(paragraph)
        paragraph.add_run(self.str_cm[1].format(
            Request.regulations[self.regulation_list[0]][0]))

    def cm_answer(self, paragraph):
        paragraph.add_run(
            # pylint: disable=no-member
            self.get_approval_status_display().upper() + ' ').font.bold = True
        paragraph.add_run(self.str_cm[0].format(
            # pylint: disable=no-member
            num2words(self.points, lang='es'),
            self.points,
            self.academic_period,
            self.get_bacheilor_program_display(),
            self.bacheilor_program,
            self.get_headquarters_display()
        ))

    def pcm(self, docx):
        self.pcm_analysis(docx)
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.str_answer + ': ').font.bold = True
        paragraph.add_run(self.str_comittee_header + ' ')
        self.pcm_answer(paragraph)
        paragraph.add_run(self.str_cm[1].format(
            Request.regulations[self.regulation_list[0]][0]))

    def pcm_answer(self, paragraph):
        paragraph.add_run(
            # pylint: disable=no-member
            self.get_advisor_response_display().upper()).font.bold = True
        paragraph.add_run(' ' + self.str_cm[0].format(
            # pylint: disable=no-member
            num2words(self.points, lang='es'),
            self.points,
            self.academic_period,
            self.get_bacheilor_program_display(),
            self.bacheilor_program,
            self.get_headquarters_display()
        ))

    def pcm_analysis(self, docx):
        # pylint: disable=no-member
        analysis_list = []
        analysis_list += [self.str_pcm[0].format(
            self.get_academic_program_display(),
            self.academic_program,
            self.get_academic_profile_display()
        )]
        analysis_list += [self.str_pcm[1].format(
            '' if self.enrolled_before_preprogram else 'no ',
            self.finalized_period,
            self.initial_period
        )]
        analysis_list += [self.str_pcm[2].format(
            '' if self.is_in_right_date else 'no '
        )]
        analysis_list += [self.str_pcm[3].format(
            '{} de {} del {}'.format(
                self.right_date.day,
                num_to_month(self.right_date.month),
                self.right_date.year
            ))]
        analysis_list += self.extra_analysis
        add_analysis_paragraph(docx, analysis_list)

    def resource_analysis(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.pcm_answer(last_paragraph)
    
    def resource_pre_answer(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.pcm_answer(last_paragraph)

    def resource_answer(self, docx):
        last_paragraph = docx.paragraphs[-1]
        self.cm_answer(last_paragraph)
