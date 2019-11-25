
import dateparser
from docx import Document
from docx.shared import RGBColor
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from .models import Request
from .cases.case_utils import header


class CouncilMinuteGenerator():

    def __init__(self):
        self.document = Document()
        for style in self.document.styles:
            try:
                self.document.styles[style.name].font.name = 'Ancizar Sans'
                self.document.styles[style.name].font.color.rgb = RGBColor(
                    0x00, 0x00, 0x00)
            except:  # pylint: disable=bare-except
                pass
        self.case_count = 0

    def add_case_from_request(self, request):
        request.cm(self.document)

    def add_case_from_year_and_council_number(self, council_number, year):
        request_by_number = Request.objects(
            year=year, consecutive_minute=council_number)
        request_by_number_ordered = request_by_number.order_by(
            'academic_program', 'cls')
        requests_pre = [
            request for request in request_by_number_ordered if request.is_pre()]
        requests_pos = [
            request for request in request_by_number_ordered if not request.is_pre()]
        self.__add_cases_from_date_pre_pos(requests_pre, 'PREGRADO')
        self.__add_cases_from_date_pre_pos(requests_pos, 'POSGRADO')

    def add_cases_from_date(self, start_date, end_date):
        # pylint: disable=no-member
        request_by_date = Request.objects(date__gte=dateparser.parse(
            start_date), date__lte=dateparser.parse(end_date))
        request_by_date_ordered = request_by_date.order_by(
            'academic_program', 'cls')
        requests_pre = [
            request for request in request_by_date_ordered if request.is_pre()]
        requests_pos = [
            request for request in request_by_date_ordered if not request.is_pre()]
        self.__add_cases_from_date_pre_pos(requests_pre, 'PREGRADO')
        self.__add_cases_from_date_pre_pos(requests_pos, 'POSGRADO')

    def add_cases_from_array(self, array):
        case_list = Request.objects.filter(id__in=array)
        request_from_array_ordered = case_list.order_by(
            'academic_program', 'cls')

        requests_pre = [
            request for request in request_from_array_ordered if request.is_pre()]
        requests_pos = [
            request for request in request_from_array_ordered if not request.is_pre()]

        self.__add_cases_from_date_pre_pos(requests_pre, 'PREGRADO')
        self.__add_cases_from_date_pre_pos(requests_pos, 'POSGRADO')

    def add_cases_from_date_except_app_status(self, start_date, end_date, app_status):
        # pylint: disable=no-member
        request_by_date = Request.objects(
            date__gte=dateparser.parse(start_date),
            date__lte=dateparser.parse(end_date),
            approval_status__ne=app_status)
        request_by_date_ordered = request_by_date.order_by(
            'academic_program', 'cls')
        requests_pre = [
            request for request in request_by_date_ordered if request.is_pre()]
        requests_pos = [
            request for request in request_by_date_ordered if not request.is_pre()]
        self.__add_cases_from_date_pre_pos(requests_pre, 'PREGRADO')
        self.__add_cases_from_date_pre_pos(requests_pos, 'POSGRADO')

    def __add_cases_from_date_pre_pos(self, requests, pre_pos):
        actual_academic_program = requests[0].academic_program
        para = self.document.add_paragraph(style='Heading 1')
        list_level_1 = 9 if pre_pos == 'PREGRADO' else 10
        list_level_2 = 0
        list_level_3 = 0
        run = para.add_run(
            '{}. ASUNTOS ESTUDIANTILES DE {}'.format(list_level_1, pre_pos))
        run.font.bold = True
        run.font.size = Pt(12)
        actual_academic_program = 'dummy'
        actual_case = 'dummy'
        for request in requests:
            if actual_academic_program != request.academic_program:
                list_level_2 = list_level_2 + 1
                actual_academic_program = request.academic_program
                para = self.document.add_paragraph(style='Heading 2')
                run = para.add_run('{}.{} {}'.format(
                    list_level_1, list_level_2, request.get_academic_program_display().upper()))
                run.font.bold = True
                run.font.size = Pt(12)
                list_level_3 = 0
            if actual_case != request.full_name:
                list_level_3 = list_level_3 + 1
                actual_case = request.full_name
                para = self.document.add_paragraph(style='Heading 2')
                run = para.add_run(request.full_name.upper())
                run.font.bold = True
                run.font.size = Pt(12)
            para = self.document.add_paragraph(style='Heading 3')
            run = para.add_run('{}.{}.{} {} \t DNI. {}'.format(
                list_level_1, list_level_2, list_level_3, request.student_name,
                request.student_dni))
            run.font.bold = True
            run.font.size = Pt(12)
            try:
                request.cm(self.document)
            except NotImplementedError:
                self.document.add_paragraph()
                self.document.add_paragraph(
                    'Not Implemented case {}'.format(request.full_name))
                self.document.add_paragraph()
            except Exception as err:  # pylint: disable=broad-except
                self.document.add_paragraph()
                self.document.add_paragraph(
                    'Error en el acta {}'.format(request.id))
                self.document.add_paragraph('Trace: {}'.format(err))
                self.document.add_paragraph()

    def generate(self, filename):
        self.document.save(filename)


class PreCouncilMinuteGenerator():

    def __init__(self):
        self.document = Document()
        for style in self.document.styles:
            try:
                self.document.styles[style.name].font.name = 'Ancizar Sans'
                self.document.styles[style.name].font.color.rgb = RGBColor(
                    0x00, 0x00, 0x00)
            except:  # pylint: disable=bare-except
                pass
        self.case_count = 0

    def add_case_from_request(self, request):
        hyperlink_style = self.document.styles.add_style(
            # pylint: disable=no-member
            'List Hyperlink', WD_STYLE_TYPE.PARAGRAPH)
        hyperlink_style.base_style = self.document.styles['List Bullet']
        hyperlink_style.font.color.rgb = RGBColor(
            0x00, 0x00, 0xFF)
        hyperlink_style.font.underline = True
        header(request, self.document)
        request.pcm(self.document)

    def add_case_from_year_and_council_number(self, council_number, year):
        request_by_number = Request.objects(
            year=year, consecutive_minute=council_number)
        request_by_number_ordered = request_by_number.order_by(
            'academic_program', 'cls')
        requests_pre = [
            request for request in request_by_number_ordered if request.is_pre()]
        requests_pos = [
            request for request in request_by_number_ordered if not request.is_pre()]
        self.__add_cases_from_date_pre_pos(requests_pre, 'PREGRADO')
        self.__add_cases_from_date_pre_pos(requests_pos, 'POSGRADO')

    def add_cases_from_date(self, start_date, end_date):
        # pylint: disable=no-member
        request_by_date = Request.objects(date__gte=dateparser.parse(
            start_date), date__lte=dateparser.parse(end_date))
        request_by_date_ordered = request_by_date.order_by(
            'academic_program', 'cls')
        requests_pre = [
            request for request in request_by_date_ordered if request.is_pre()]
        requests_pos = [
            request for request in request_by_date_ordered if not request.is_pre()]
        self.__add_cases_from_date_pre_pos(requests_pre, 'PREGRADO')
        self.__add_cases_from_date_pre_pos(requests_pos, 'POSGRADO')

    def add_cases_from_array(self, array):
        case_list = Request.objects.filter(id__in=array)
        request_from_array_ordered = case_list.order_by(
            'academic_program', 'cls')

        requests_pre = [
            request for request in request_from_array_ordered if request.is_pre()]
        requests_pos = [
            request for request in request_from_array_ordered if not request.is_pre()]

        self.__add_cases_from_date_pre_pos(requests_pre, 'PREGRADO')
        self.__add_cases_from_date_pre_pos(requests_pos, 'POSGRADO')

    def __add_cases_from_date_pre_pos(self, requests, pre_pos):
        if len(requests) == 0:
            return
        actual_academic_program = requests[0].academic_program
        para = self.document.add_paragraph(style='Heading 1')
        list_level_1 = 9 if pre_pos == 'PREGRADO' else 10
        list_level_2 = 0
        list_level_3 = 0
        run = para.add_run(
            '{}. ASUNTOS ESTUDIANTILES DE {}'.format(list_level_1, pre_pos))
        run.font.bold = True
        run.font.size = Pt(12)
        actual_academic_program = 'dummy'
        actual_case = 'dummy'

        try:
            hyperlink_style = self.document.styles.add_style(
            # pylint: disable=no-member
            'List Hyperlink', WD_STYLE_TYPE.PARAGRAPH)
            hyperlink_style.base_style = self.document.styles['List Bullet']
            hyperlink_style.font.color.rgb = RGBColor(
                0x00, 0x00, 0xFF)
            hyperlink_style.font.underline = True
        except ValueError:
            pass

        for request in requests:
            if actual_academic_program != request.academic_program:
                list_level_2 = list_level_2 + 1
                actual_academic_program = request.academic_program
                para = self.document.add_paragraph(style='Heading 2')
                run = para.add_run('{}.{} {}'.format(
                    list_level_1, list_level_2, request.get_academic_program_display().upper()))
                run.font.bold = True
                run.font.size = Pt(12)
                list_level_3 = 0
            if actual_case != request.full_name:
                list_level_3 = list_level_3 + 1
                actual_case = request.full_name
                para = self.document.add_paragraph(style='Heading 2')
                run = para.add_run(request.full_name.upper())
                run.font.bold = True
                run.font.size = Pt(12)
            para = self.document.add_paragraph(style='Heading 3')
            run = para.add_run('{}.{}.{} {} \t DNI. {}'.format(
                list_level_1, list_level_2, list_level_3, request.student_name,
                request.student_dni))
            run.font.bold = True
            run.font.size = Pt(12)
            try:

                header(request, self.document)
                request.pcm(self.document)
            except NotImplementedError:
                self.document.add_paragraph()
                self.document.add_paragraph(
                    'Not Implemented case {}'.format(request.full_name))
                self.document.add_paragraph()
            except Exception as err:  # pylint: disable=broad-except
                self.document.add_paragraph()
                self.document.add_paragraph(
                    'Error en el acta {}'.format(request.id))
                self.document.add_paragraph('Trace: {}'.format(err))
                self.document.add_paragraph()

    def generate(self, filename):
        self.document.save(filename)
