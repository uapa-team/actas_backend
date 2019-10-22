from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Pt
from .case_utils import add_hyperlink
from .case_utils import table_general_data
from .case_utils import get_academic_program
from .case_utils import string_to_date
from .case_utils import table_recommend


class REINPRE():

    @staticmethod
    def case_REINGRESO_PREGRADO(request, docx, redirected=False):
        para = docx.add_paragraph()
        para.add_run('Análisis:\t\t\t')
        add_hyperlink(
            para, 'http://www.legal.unal.edu.co/rlunal/home/doc.jsp?d_i=62849',
            'Resolución 012 de 2014'
        )
        para.add_run(', ')
        add_hyperlink(
            para, 'http://www.legal.unal.edu.co/rlunal/home/doc.jsp?d_i=34983',
            'Acuerdo 008 de 2008'
        )
        para = docx.add_paragraph(style='List Number')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if request.detail_cm['first_reingreso'] == 'Sí':
            para.add_run('No h')
        elif request.detail_cm['first_reingreso'] == 'No':
            para.add_run('H')
        para.add_run('a tenido otro reingreso después de 2009-1S ')
        para.add_run(
            '(Artículo 46, Acuerdo 008 de 2008 del Consejo Superior Universitario.).')
        para.add_run(' Universitas y SIA: Revisado.')
        para = docx.add_paragraph(style='List Number')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if float(request.detail_cm['PAPA']) >= 2.7:
            para.add_run('T')
        else:
            para.add_run('No t')
        para.add_run('iene P.A.P.A. superior o igual a 2.7 ')
        para.add_run(
            '(literal 3b - Artículo 3, Resolución 239 de 2009 de Vicerrectoría Académica; ')
        para.add_run(
            'Artículo 46, Acuerdo 008 de 2008 del Consejo Superior Universitario.). SIA: ')
        para.add_run('P.A.P.A. de ' + request.detail_cm['PAPA'] + '.')
        para = docx.add_paragraph(style='List Number')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if int(request.detail_cm['creds_remaining']) >= 0:
            para.add_run('D')
        else:
            para.add_run('No d')
        para.add_run('ispone de un cupo de créditos suficiente: ')
        para.add_run(
            'Cupo adicional de 10 créditos a lo sumo (parágrafo 1 Artículo 46, ')
        para.add_run(
            'Acuerdo 008 de 2008 del Consejo Superior Universitario). ')
        para.add_run(
            'SIA: Revisado. En caso de otorgarle un cupo adicional de créditos, ')
        para.add_run(
            'este no podrá ser mayor que el requerido para inscribir asignaturas ')
        para.add_run(
            'pendientes del plan de estudios. (Artículo 6, Resolución 012 de 2014 ')
        para.add_run('- Vicerrectoría Académica).')
        para = docx.add_paragraph(style='List Number')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run('La solicitud se hace ')
        if request.pre_cm['detail_pre_cm']['request_in_date']:
            para.add_run('en')
        else:
            para.add_run('fuera de las')
        para.add_run(' fechas de calendario de sede (parágrafo Artículo 3).')
        if 'extra_analysis' in request.pre_cm:
            for analysis in request.pre_cm['extra_analysis']:
                para = docx.add_paragraph(style='List Number')
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                para.add_run(analysis)
        para.paragraph_format.space_after = Pt(0)
        para = docx.add_paragraph()
        para.paragraph_format.space_after = Pt(0)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if request.pre_cm['detail_pre_cm']['request_in_date']:
            para.add_run('Concepto: ').font.bold = True
            para.add_run('El Comité Asesor ')
            if request.approval_status == 'RM':
                para.add_run('recomienda')
            elif request.approval_status == 'NM':
                para.add_run('no recomienda')
            para.add_run(' al Consejo de Facultad ')
            para.add_run('APROBAR').font.bold = True
            para.add_run(
                ' reingreso por única vez a partir del periodo académico')
            para.add_run(request.detail_cm['reing_per'])
            if 'grants_cred' in request.pre_cm['detail_pre_cm']:
                para.add_run('y otorga ')
                para.add_run(request.pre_cm['detail_pre_cm']['grants_cred'])
                para.add_run(
                    ' crédito(s) adicional(es) para culminar su plan de estudios')
            para.add_run(
                '. Si el estudiante no renueva su matrícula en el semestre de reingreso el acto ')
            para.add_run(
                'académico expedido por el Consejo de Facultad queda sin efecto. (Resolución 012')
            para.add_run(
                ' de 2014 de Vicerrectoría Académica; Artículo 46, Acuerdo 008 de 2008 del Consejo')
            para.add_run(' Superior Universitario).')
            para = docx.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.space_after = Pt(0)
            para.add_run('El señor ' + request.student_name +
                         ' tiene pendiente por aprobar ' + request.detail_cm['creds_remaining'])
            para.add_run(' créditos del plan de estudios de ' +
                         get_academic_program(request.academic_program))
            para.add_run(
                ' y ' + request.detail_cm['creds_ingl'] + ' créditos del requisito de nivelación')
            para.add_run(
                ' - inglés, con un cupo disponible para inscripción de ')
            para.add_run(str(int(
                request.detail_cm['creds_remaining']) + int(
                    request.detail_cm['creds_minus_remaining'])))
            para.add_run(' créditos.')
            para = docx.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.space_after = Pt(0)
            para.add_run(
                'El parágrafo del artículo 11 del Acuerdo 008 de 2008 de Consejo Superior ')
            para.add_run('Superior Universitario establece: ')
            para.add_run(
                '"Los créditos adicionales que como resultado del ').font.italic = True
            para.add_run(
                'proceso de clasificación en la admisión deba aprobar ').font.italic = True
            para.add_run(
                'un estudiante de pregrado, se sumarán por única vez al "').font.italic = True
            para.add_run(
                'cupo adicional de créditos para inscripción"').font.italic = True
            para.add_run(', por lo tanto solo es viable otorgar ' +
                         request.pre_cm['detail_pre_cm']['grants_cred'])
            para.add_run(
                ' crédito(s) para la inscripción de asignaturas pendientes del plan de estudios ')
            para.add_run(
                ' de ' + get_academic_program(request.academic_program) + '.')
        else:
            para.add_run('Concepto: ').font.bold = True
            para.add_run('El Comité Asesor ')
            if request.approval_status == 'RM':
                para.add_run('recomienda')
            elif request.approval_status == 'NM':
                para.add_run('no recomienda')
            para.add_run(' al Consejo de Facultad ')
            para.add_run('NO APROBAR').font.bold = True
            para.add_run(
                ' reingreso por única vez a partir del periodo académico')
            para.add_run(request.detail_cm['reing_per'])
            para.add_run(
                ', porque el estudiante presentó la solicitud fuera de las fechas establecidas ')
            para.add_run(
                'en el Calendario Académico de la Sede Bogotá. (Resolución 012')
            para.add_run(
                ' de 2014 de Vicerrectoría Académica; Artículo 46, Acuerdo 008 de 2008 del Consejo')
            para.add_run(' Superior Universitario).')
        para.paragraph_format.space_after = Pt(0)
        para = docx.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.space_after = Pt(0)
        general_data = []
        general_data.append(['Estudiante', request.student_name])
        general_data.append(['DNI', request.student_dni])
        general_data.append(['Plan de estudios',
                             get_academic_program(request.academic_program)])
        general_data.append(
            ['Código del plan de estudios', request.academic_program])
        general_data.append(
            ['Fecha de la solicitud', string_to_date(request.detail_cm['solic_date'])])
        bullet = para.add_run('1. Datos Generales')
        bullet.font.bold = True
        bullet.font.size = Pt(8)
        table_general_data(general_data, 'REINGRESO', docx)
        para = docx.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.space_after = Pt(0)
        bullet = para.add_run('2. Información Académica')
        para.paragraph_format.space_after = Pt(0)
        bullet.font.bold = True
        bullet.font.size = Pt(8)
        table = docx.add_table(rows=13, cols=3)
        table.style = 'Table Grid'
        table.style.font.size = Pt(8)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for cell in table.columns[0].cells:
            cell.width = 400000
        for cell in table.columns[1].cells:
            cell.width = 3200000
        for cell in table.columns[2].cells:
            cell.width = 1600000
        table.columns[0].width = 400000
        table.columns[1].width = 3200000
        table.columns[2].width = 1600000
        table.cell(0, 0).merge(table.cell(0, 1)).paragraphs[0].add_run(
            'Periodo para el cual fue admitido en este plan de estudios')
        table.cell(0, 2).paragraphs[0].add_run(request.detail_cm['per_admi'])
        table.cell(0, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(1, 0).merge(table.cell(1, 1)).paragraphs[0].add_run(
            '¿Se trata de un primer reingreso?')
        table.cell(1, 2).paragraphs[0].add_run(
            request.detail_cm['first_reingreso'])
        table.cell(1, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(2, 0).merge(table.cell(2, 2)).paragraphs[0].add_run(
            'Si la respuesta es NO, el Comité Asesor no debe recomendar al Consejo de ' +
            'Facultad el reingreso')
        table.cell(3, 0).merge(table.cell(3, 1)).paragraphs[0].add_run(
            'Es caso de ser primer reingreso en ¿qué periodo académico perdió la calidad ' +
            'de estudiante?')
        table.cell(3, 2).paragraphs[0].add_run(request.detail_cm['per_perd'])
        table.cell(3, 2).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        table.cell(3, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(4, 0).merge(table.cell(4, 1)).paragraphs[0].add_run(
            'Al momento de presentar la solicitud ¿cuántos periodos académicos (incluido ' +
            'el periodo académico en que presentó la solicitud) han transcurridos a partir ' +
            'del periodo académico en que registró su última matrícula?')
        table.cell(4, 2).paragraphs[0].add_run(request.detail_cm['per_transc'])
        table.cell(4, 2).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        table.cell(4, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(5, 0).merge(table.cell(5, 2)).paragraphs[0].add_run(
            'En caso que la respuesta sea mayor de 6 periodos académicos no se debe recomendar ' +
            'el reingreso')
        table.cell(6, 0).merge(table.cell(6, 1)
                               ).paragraphs[0].add_run('P.A.P.A.')
        table.cell(6, 2).paragraphs[0].add_run(request.detail_cm['PAPA'])
        table.cell(6, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(7, 0).merge(table.cell(7, 1)).paragraphs[0].add_run(
            'Causa de la pérdida de la calidad de estudiante')
        table.cell(7, 0).merge(table.cell(7, 1)
                               ).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        table.cell(7, 2).paragraphs[0].add_run(request.detail_cm['causa_perd'])
        table.cell(7, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(8, 0).merge(table.cell(8, 2)).paragraphs[0].add_run(
            'Estudio de créditos').font.bold = True
        table.cell(9, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(9, 0).paragraphs[0].add_run('1').font.bold = True
        table.cell(10, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(10, 0).paragraphs[0].add_run('2').font.bold = True
        table.cell(11, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(11, 0).paragraphs[0].add_run('3').font.bold = True
        table.cell(12, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(12, 0).paragraphs[0].add_run('4').font.bold = True
        table.cell(9, 1).paragraphs[0].add_run(
            'Cupo de créditos menos créditos pendientes')
        table.cell(10, 1).paragraphs[0].add_run(
            'Créditos pendientes por ser aprobados del plan de estudios')
        table.cell(11, 1).paragraphs[0].add_run(
            'Créditos pendientes por ser aprobados de nivelación – Inglés')
        table.cell(12, 1).paragraphs[0].add_run(
            '¿Cuántos créditos adicionales requiere para inscribir asignaturas?')
        table.cell(9, 2).paragraphs[0].add_run(
            request.detail_cm['creds_minus_remaining'])
        table.cell(9, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(10, 2).paragraphs[0].add_run(
            request.detail_cm['creds_remaining'])
        table.cell(10, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(11, 2).paragraphs[0].add_run(
            request.detail_cm['creds_ingl'])
        table.cell(11, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(12, 2).paragraphs[0].add_run(request.detail_cm['creds_add'])
        table.cell(12, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        para = docx.add_paragraph()
        para.paragraph_format.space_after = Pt(0)
        table = docx.add_table(rows=5, cols=2)
        for col in table.columns:
            for cell in col.cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.style = 'Table Grid'
        table.style.font.size = Pt(8)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for cell in table.columns[0].cells:
            cell.width = 3100000
        for cell in table.columns[1].cells:
            cell.width = 2100000
        table.columns[0].width = 3100000
        table.columns[1].width = 2100000
        table.cell(0, 0).merge(table.cell(0, 1)
                               ).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        table.cell(0, 0).merge(table.cell(0, 1)).paragraphs[0].add_run(
            'Al finalizar el semestre de reingreso para mantener la calidad de estudiante, ' +
            'deberá obtener un Promedio Semestral mínimo de:')
        table.cell(1, 0).paragraphs[0].add_run('Si inscribe 12 Créditos')
        table.cell(2, 0).paragraphs[0].add_run('Si inscribe 15 Créditos')
        table.cell(3, 0).paragraphs[0].add_run('Si inscribe 18 Créditos')
        table.cell(4, 0).paragraphs[0].add_run('Si inscribe 21 Créditos')
        table.cell(1, 1).paragraphs[0].add_run(request.detail_cm['12c'])
        table.cell(2, 1).paragraphs[0].add_run(request.detail_cm['15c'])
        table.cell(3, 1).paragraphs[0].add_run(request.detail_cm['18c'])
        table.cell(4, 1).paragraphs[0].add_run(request.detail_cm['21c'])
        para = docx.add_paragraph()
        bullet = para.add_run(
            '3. Resumen general de créditos del plan de estudios')
        para.paragraph_format.space_after = Pt(0)
        bullet.font.bold = True
        bullet.font.size = Pt(8)
        table = docx.add_table(rows=5, cols=7)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for col in table.columns:
            for cell in col.cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.style = 'Table Grid'
        table.style.font.size = Pt(8)
        for cell in table.columns[0].cells:
            cell.width = 1610000
        for cell in table.columns[1].cells:
            cell.width = 690000
        for cell in table.columns[2].cells:
            cell.width = 610000
        for cell in table.columns[3].cells:
            cell.width = 690000
        for cell in table.columns[4].cells:
            cell.width = 610000
        for cell in table.columns[5].cells:
            cell.width = 675000
        for cell in table.columns[5].cells:
            cell.width = 375000
        table.columns[0].width = 1610000
        table.columns[1].width = 690000
        table.columns[2].width = 610000
        table.columns[3].width = 690000
        table.columns[4].width = 610000
        table.columns[5].width = 675000
        table.columns[6].width = 375000
        table.cell(0, 0).merge(table.cell(1, 0)
                               ).paragraphs[0].add_run('Créditos')
        table.cell(0, 0).merge(table.cell(1, 0)
                               ).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        table.cell(0, 1).merge(table.cell(0, 2)
                               ).paragraphs[0].add_run('Fundamentación (B)')
        table.cell(0, 3).merge(table.cell(0, 4)
                               ).paragraphs[0].add_run('Disciplinar (C)')
        table.cell(0, 5).merge(table.cell(1, 5)
                               ).paragraphs[0].add_run('Libre Elección (L)')
        table.cell(0, 6).merge(table.cell(1, 6)).paragraphs[0].add_run('Total')
        table.cell(0, 6).merge(table.cell(1, 6)
                               ).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        table.cell(2, 0).paragraphs[0].add_run('Exigidos*')
        table.cell(3, 0).paragraphs[0].add_run(
            'Aprobados del plan de estudios')
        table.cell(4, 0).paragraphs[0].add_run('Pendientes')
        table.cell(1, 1).paragraphs[0].add_run('Obligatorios')
        table.cell(1, 2).paragraphs[0].add_run('Optativos')
        table.cell(1, 3).paragraphs[0].add_run('Obligatorios')
        table.cell(1, 4).paragraphs[0].add_run('Optativos')
        table.cell(2, 1).paragraphs[0].add_run(
            request.detail_cm['summary']['exiged']['fund_m'])
        table.cell(3, 1).paragraphs[0].add_run(
            request.detail_cm['summary']['approved']['fund_m'])
        table.cell(4, 1).paragraphs[0].add_run(
            request.detail_cm['summary']['remaining']['fund_m'])
        table.cell(2, 2).paragraphs[0].add_run(
            request.detail_cm['summary']['exiged']['fund_o'])
        table.cell(3, 2).paragraphs[0].add_run(
            request.detail_cm['summary']['approved']['fund_o'])
        table.cell(4, 2).paragraphs[0].add_run(
            request.detail_cm['summary']['remaining']['fund_o'])
        table.cell(2, 3).paragraphs[0].add_run(
            request.detail_cm['summary']['exiged']['disc_m'])
        table.cell(3, 3).paragraphs[0].add_run(
            request.detail_cm['summary']['approved']['disc_m'])
        table.cell(4, 3).paragraphs[0].add_run(
            request.detail_cm['summary']['remaining']['disc_m'])
        table.cell(2, 4).paragraphs[0].add_run(
            request.detail_cm['summary']['exiged']['disc_o'])
        table.cell(3, 4).paragraphs[0].add_run(
            request.detail_cm['summary']['approved']['disc_o'])
        table.cell(4, 4).paragraphs[0].add_run(
            request.detail_cm['summary']['remaining']['disc_o'])
        table.cell(2, 5).paragraphs[0].add_run(
            request.detail_cm['summary']['exiged']['free'])
        table.cell(3, 5).paragraphs[0].add_run(
            request.detail_cm['summary']['approved']['free'])
        table.cell(4, 5).paragraphs[0].add_run(
            request.detail_cm['summary']['remaining']['free'])
        exiged = int(request.detail_cm['summary']['exiged']['fund_m']) + int(
            request.detail_cm['summary']['exiged']['fund_o']) + int(
                request.detail_cm['summary']['exiged']['disc_m']) + int(
                    request.detail_cm['summary']['exiged']['disc_o']) + int(
                        request.detail_cm['summary']['exiged']['free'])
        approved = int(request.detail_cm['summary']['approved']['fund_m']) + int(
            request.detail_cm['summary']['approved']['fund_o']) + int(
                request.detail_cm['summary']['approved']['disc_m']) + int(
                    request.detail_cm['summary']['approved']['disc_o']) + int(
                        request.detail_cm['summary']['approved']['free'])
        remaining = int(request.detail_cm['summary']['remaining']['fund_m']) + int(
            request.detail_cm['summary']['remaining']['fund_o']) + int(
                request.detail_cm['summary']['remaining']['disc_m']) + int(
                    request.detail_cm['summary']['remaining']['disc_o']) + int(
                        request.detail_cm['summary']['remaining']['free'])
        table.cell(2, 6).paragraphs[0].add_run(str(exiged))
        table.cell(3, 6).paragraphs[0].add_run(str(approved))
        table.cell(4, 6).paragraphs[0].add_run(str(remaining))
        para = docx.add_paragraph()
        para.paragraph_format.space_after = Pt(0)
        para.add_run(
            '     *Sin incluir los créditos correspondientes al cumplimiento del ' +
            'requisito de suficiencia en idioma.').font.size = Pt(8)
        details = []
        details.append(request.detail_cm['commite_name'])
        details.append(request.detail_cm['comite_date'])
        details.append(request.detail_cm['comite_acta'])
        details.append(request.detail_cm['comite_date'][6:10])
        details.append(request.detail_cm['reccomend'])
        table_recommend(docx, details)
