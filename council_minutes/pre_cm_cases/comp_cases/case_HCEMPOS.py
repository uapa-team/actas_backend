from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from ...models import Request
from .case_utils import *


class HCEMPOS():

    @staticmethod
    def case_HOMOLOGACION_CONVALIDACION_Y_EQUIVALENCIA_POSGRADO(request, docx, redirected=False):
        ### Frequently used ###
        details = request['detail_cm']
        pre_cm = request['pre_cm']
        details_pre = pre_cm['detail_pre_cm']
        is_recommended = request['approval_status'] == 'CR'

        ### Finishing last paragraph ###
        para = docx.paragraphs[-1]
        para.add_run('Análisis:')

        ### Analysis Paragraph ###
        para = docx.add_paragraph(style='List Number')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run('Perfil de {}'.format(details['node']))

        para = docx.add_paragraph(style='List Number')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run('Universitas: Asignaturas asociadas al plan de estudios con tipología L')

        ## Extra Analysis ##
        for analysis in pre_cm['extra_analysis']:
            para = docx.add_paragraph(style='List Number')
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.add_run(analysis)

        para = docx.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run('Concepto: ').bold = True
        para.add_run('El Comité Asesor recomienda al Consejo de Facultad ')
        modifier = 'APROBAR' if is_recommended else 'NO APROBAR'
        para.add_run(modifier).bold = True

        p_aux  = ' homologar, equivaler o convalidar en el programa {} '
        p_aux += 'plan de estudios {}, en el periodo {}, las siguientes asignaturas '
        p_aux += 'cursadas en el programa {}, así:'

        para.add_run(p_aux.format(
            get_academic_program(request['academic_program']),
            details['node'],
            request['academic_period'],
            details_pre['origin_program']
        ))

        options = {
            'homologation'  : ['homologadas', 'homologados'],
            'recognition'   : ['convalidadas', 'convalidados'],
            'equivalence'   : ['equivalidas', 'equivalidos']
        }

        for key in options:
                if key in details:
                    subjects = []
                    for subject in details[key]['subjects']:
                        subjects.append([
                            subject['homologated_period'],
                            subject['code'],
                            subject['subject'],
                            subject['credits'],
                            subject['tipology'],
                            subject['grade'],
                            subject['subject_out'],
                            subject['grade_out']
                        ])
                    
                    draw_table(docx, subjects, options[key])

#Modified table from case_utils.table_approvals
def draw_table(docx, subjects, details):
    '''Add a generated table with approvals subjects

    Params:
        docx (docx): The document to which the table will be added
        subjects (list): A list of list with the subjects in table,
        each list must be a list with following data:
        [0]: Subject's period
        [1]: Subject's SIA code
        [2]: Subject's SIA name
        [3]: Subject's SIA credits
        [4]: Subject's SIA tipology
        [5]: Subject's SIA grade
        [6]: Subject's old name
        [7]: Subject's old grade
        details(list): a list with spanish plural for subjects and credits
        [0]: Subject's plural
        [1]: Credit's plural

    Raises:
        IndexError: All lists must have same size
    '''
    tipology = {}
    periods = []
    for asign in subjects:
        if asign[4] in tipology:
            tipology.update(
                {asign[4]: tipology[asign[4]] + int(asign[3])})
        else:
            tipology.update({asign[4]: int(asign[3])})
        if asign[0] not in periods:
            periods.append(asign[0])
    asign_number = len(subjects)
    tipology_number = len(tipology)
    table = docx.add_table(
        rows=(3+asign_number+tipology_number), cols=8, style='Table Grid')
    table.style.font.size = Pt(8)
    table.columns[0].width = 500000
    table.columns[1].width = 550000
    table.columns[2].width = 1600000
    table.columns[3].width = 300000
    table.columns[4].width = 300000
    table.columns[5].width = 400000
    table.columns[6].width = 1400000
    table.columns[7].width = 400000
    for cell in table.columns[0].cells:
        cell.width = 500000
    for cell in table.columns[1].cells:
        cell.width = 550000
    for cell in table.columns[2].cells:
        cell.width = 1600000
    for cell in table.columns[3].cells:
        cell.width = 300000
    for cell in table.columns[4].cells:
        cell.width = 300000
    for cell in table.columns[5].cells:
        cell.width = 400000
    for cell in table.columns[6].cells:
        cell.width = 1400000
    for cell in table.columns[7].cells:
        cell.width = 400000
    
    cellp = table.cell(0, 0).merge(table.cell(0, 5)).paragraphs[0]
    cellp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cellp.add_run('Asignatura que se {}'.format(details[0])).font.bold = True
    cellp = table.cell(0, 6).merge(table.cell(0, 7)).paragraphs[0]
    cellp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cellp.add_run('Asignatura por la que {}'.format(details[0])).font.bold = True
    for i in range(1, asign_number + 3):
        for j in range(8):
            table.cell(
                i, j).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(1, 0).paragraphs[0].add_run('Periodo').font.bold = True
    table.cell(1, 1).paragraphs[0].add_run('Código').font.bold = True
    table.cell(1, 2).paragraphs[0].add_run('Asignatura').font.bold = True
    table.cell(1, 3).paragraphs[0].add_run('C').font.bold = True
    table.cell(1, 4).paragraphs[0].add_run('T').font.bold = True
    table.cell(1, 5).paragraphs[0].add_run('Nota').font.bold = True
    table.cell(1, 6).paragraphs[0].add_run('Asignatura').font.bold = True
    table.cell(1, 7).paragraphs[0].add_run('Nota').font.bold = True
    count = 2
    for subject in subjects:
        table.cell(count, 0).paragraphs[0].add_run(subject[0])
        table.cell(count, 1).paragraphs[0].add_run(subject[1])
        table.cell(count, 2).paragraphs[0].add_run(
            subject[2])
        table.cell(count, 3).paragraphs[0].add_run(subject[3])
        table.cell(count, 4).paragraphs[0].add_run(subject[4])
        table.cell(count, 5).paragraphs[0].add_run(subject[5])
        table.cell(count, 6).paragraphs[0].add_run(
            subject[6])
        table.cell(count, 7).paragraphs[0].add_run(subject[7])
        count += 1
    total_homologated = 0
    for tip in tipology:
        text = 'Céditos {} '.format(details[1]) + str(tip)
        table.cell(count, 0).merge(table.cell(
            count, 5)).paragraphs[0].add_run(text)
        table.cell(count, 0).merge(table.cell(count, 5)
                                   ).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(count, 6).merge(table.cell(count, 7)
                                   ).paragraphs[0].add_run(str(tipology[tip]))
        table.cell(count, 6).merge(table.cell(count, 7)
                                   ).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        total_homologated += int(tipology[tip])
        count += 1
    table.cell(count, 0).merge(table.cell(count, 5)).paragraphs[0].add_run(
        'Total créditos que se {}'.format(details[1]))
    table.cell(count, 0).merge(table.cell(count, 5)
                               ).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(count, 6).merge(table.cell(count, 7)
                               ).paragraphs[0].add_run(str(total_homologated))
    table.cell(count, 6).merge(table.cell(count, 7)
                               ).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    docx.add_paragraph().paragraph_format.space_after = Pt(0)
