from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from .case_utils import num_to_month
from .case_utils import get_academic_program
from .case_utils import table_subjects


class IASIPRE():

    @staticmethod
    def case_INSCRIPCION_DE_ASIGNATURAS_PREGRADO(request, docx, redirected=False):
        para = docx.add_paragraph()
        para.paragraph_format.space_after = Pt(0)
        para.add_run('Análisis:\t\t')
        add_hyperlink(para, 'Acuerdo 008 de 2008',
                      'http://www.legal.unal.edu.co/rlunal/home/doc.jsp?d_i=34983/')
        para = docx.add_paragraph(style='List Number')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run('SIA: ')
        counts = {'offered': 0, 'not offered': 0}
        for subject in request.detail_cm['subjects']:
            if subject['offered'] == 'true':
                counts.update({'offered': (counts['offered'] + 1)})
            elif subject['offered'] == 'false':
                counts.update({'not offered': (counts['not offered'] + 1)})
            else:
                raise AssertionError(
                    'request.detail_cm["subjects"][i]["offered"] must be "true" or "false".')
        if counts['offered'] > 0:
            line_subjects = ''
            for subject in request.detail_cm['subjects']:
                if subject['offered']:
                    line_subjects += ', '
                    line_subjects += subject['subject']
                    line_subjects += ' (' + subject['cod'] + ')'
            para.add_run(
                'La(s) asignaturas ' + line_subjects[2:] + ' sí se encuentra(n)' +
                ' ofertada(s) para el plan de estudios')
        if counts['offered'] > 0 and counts['not offered'] > 0:
            para.add_run(' y ')
        if counts['not offered'] > 0:
            line_subjects = ''
            for subject in request.detail_cm['subjects']:
                if not subject['offered']:
                    line_subjects += ', '
                    line_subjects += subject['subject']
                    line_subjects += ' (' + subject['cod'] + ')'
            para.add_run(
                'La(s) asignatura(s) ' + line_subjects[2:] + ' no se encuentra(n)' +
                ' ofertada(s) para el plan de estudios.')
        data = []
        for subject in request.detail_cm['subjects']:
            data.append([subject['cod'], subject['subject'],
                         subject['grup'], subject['T'], subject['cre']])
            para = docx.add_paragraph(style='List Number')
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.add_run('SIA: La asignatura ' + subject['subject'])
            para.add_run(' (' + subject['cod'] + '), ')
            para.add_run('tipología ' + subject['T'] + ', cuenta con ')
            para.add_run(subject['quota'] +
                         ' cupos disponibles el día de la revisión ')
            para.add_run(
                '(' + subject['review_date'][0:2] + num_to_month(subject['review_date'][3:5]))
            para.add_run('20' + subject['review_date'][6:8] + ').')
        para = docx.add_paragraph(style='List Number')
        para.paragraph_format.space_after = Pt(0)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if request.pre_cm['detail_pre_cm']['cruces'] == 'false':
            para.add_run('No p')
        elif request.pre_cm['detail_pre_cm']['cruces'] == 'true':
            para.add_run('P')
        para.add_run('resenta cruces con el horario inscrito.')
        if 'extra_analysis' in request.pre_cm:
            for analysis in request.pre_cm['extra_analysis']:
                para = docx.add_paragraph(style='List Number')
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                para.add_run(analysis)
        para.paragraph_format.space_after = Pt(0)
        para = docx.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run('Concepto: ').font.bold = True
        para.add_run('El Comité Asesor ')
        if request.approval_status == 'RM':
            para.add_run('recomienda')
        elif request.approval_status == 'NM':
            para.add_run('no recomienda')
        para.add_run(' inscribir las siguientes asignaturas del programa ')
        para.add_run(get_academic_program(request.academic_program))
        para.add_run(' en el periodo académico de ' +
                     request.academic_period + ':')
        para.paragraph_format.space_after = Pt(0)
        table_subjects(docx, data)
