from docx.shared import Pt


def get_academic_program(cod_program):
    large_program = ''
    for p in Request.PROGRAM_CHOICES:
        if p[0] == cod_program:
            large_program = p[1]
            break
    return large_program


def num_to_month(month):
    if int(month) == 1:
        return ' de enero de '
    elif int(month) == 2:
        return ' de febrero de '
    elif int(month) == 3:
        return ' de marzo de '
    elif int(month) == 4:
        return ' de abril de '
    elif int(month) == 5:
        return ' de mayo de '
    elif int(month) == 6:
        return ' de junio de '
    elif int(month) == 7:
        return ' de julio de '
    elif int(month) == 8:
        return ' de agosto de '
    elif int(month) == 9:
        return ' de septiembre de '
    elif int(month) == 10:
        return ' de octubre de '
    elif int(month) == 11:
        return ' de nomviembre de '
    elif int(month) == 12:
        return ' de diciembre de '


def header(request, docx):
    para = docx.add_paragraph()
    para.add_run('Tipo de solicitud:\t{}\n'.format(request.get_type_display()))
    para.add_run('Justificación:\t\t{}\n'.format(
        request['pre_cm']['justification']))
    para.add_run('Soportes:\t\t{}\n'.format(request['pre_cm']['supports']))
    para.add_run('Fecha radicación:\t{}'.format(request['date']))
    para.paragraph_format.space_after = Pt(0)


def table_subjects():
    raise NotImplementedError


def table_english():
    raise NotImplementedError


def table_approvals():
    raise NotImplementedError


def table_credits_summary():
    raise NotImplementedError


def table_recommend():
    raise NotImplementedError


def table_change_typology():
    raise NotImplementedError
