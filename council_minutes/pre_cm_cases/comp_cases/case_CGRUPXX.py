from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


class CGRUPXX():

    @staticmethod
    def case_CAMBIO_DE_GRUPO(request, docx, redirected=False):
        para = docx.add_paragraph()
        run = para.add_run('Análisis:\t')
        run.font.bold = True
        # add_hyperlink(para, 'http://www.legal.unal.edu.co/rlunal/home/doc.jsp?d_i=34983',
        #              'Acuerdo 008 de 2008 - Consejo Superior Universitario')

        str_1 = '1. El grupo {} de la asignatura {} '
        str_1 += '({}) cuenta con {} cupos.'
        para = docx.add_paragraph()
        para.paragraph_format.left_indent = Pt(36)
        run = para.add_run(str_1.format(request['detail_cm']['subjects'][0]['gd'],
                                        request['detail_cm']['subjects'][0]['subject'],
                                        request['detail_cm']['subjects'][0]['cod'],
                                        request['pre_cm']['free_places'],))

        para = docx.add_paragraph()
        run = para.add_run('Concepto:\t')
        run.font.bold = True
        str_2 = 'El Comité Asesor recomienda'
        str_2 += ' al Consejo de Facultad cambio de grupo de la asignatura/ '
        str_2 += 'actividad {}, código {}, tipología {}, inscrita en el periodo {},'
        str_2 += ' del grupo {} al grupo {} con el profesor {} del Departamento de '
        str_2 += 'Ingeniería {}, debido a que justifica debidamente la solicitud.'
        run = para.add_run(str_2.format(
            request['detail_cm']['subjects'][0]['subject'],
            request['detail_cm']['subjects'][0]['cod'],
            request['detail_cm']['subjects'][0]['tip'],
            request['academic_period'],
            request['detail_cm']['subjects'][0]['gd'],
            request['detail_cm']['subjects'][0]['go'],
            request['pre_cm']['professor'],
            request.get_academic_program_display()))
