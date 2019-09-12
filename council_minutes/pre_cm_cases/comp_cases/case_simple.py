from num2words import num2words
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Pt
from .case_utils import get_academic_program
from .case_utils import table_subjects
from .case_utils import table_change_typology
from .case_utils import table_approvals
from .case_utils import string_to_date
from .case_utils import add_hyperlink
from .case_utils import num_to_month


class simple():

    @staticmethod
    def case_RECURSO_DE_APELACION(request, docx_, redirected=False):
        raise NotImplementedError

    @staticmethod
    def case_RECURSO_DE_REPOSICION(request, docx_, redirected=False):
        raise NotImplementedError

    @staticmethod
    def case_RECURSO_DE_REPOSICION_CON_SUBSIDIO_DE_APELACION(request, docx_, redirected=False):
        raise NotImplementedError

    @staticmethod
    def case_CANCELACION_DE_PERIODO_ACADEMICO_POSGRADO(request, docx_, redirected=False):
        ### Frequently used ###
        details = request['detail_cm']
        pre_cm = request['pre_cm']
        details_pre = pre_cm['detail_pre_cm']
        is_recommended = request['approval_status'] == 'CR'

        ### Finishing last paragraph ###
        para = docx_.paragraphs[-1]
        para.add_run('Análisis:\t\t\t')
        para.add_run('Acuerdo 008 de 2008').underline = True

        ### Analysis Paragraph ###
        para = docx_.add_paragraph(style='List Number')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run(
            'SIA: {} - Perfil de {}.'.format(
                get_academic_program(request['academic_program']), details_pre['academic_profile'])
        )
        para = docx_.add_paragraph(style='List Number')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        modifier = '' if is_recommended else 'no '
        para.add_run(
            'El comité {}lo considera fuerza mayor o caso fortuito.'.format(modifier))
        ## Extra Analysis ##
        for analysis in pre_cm['extra_analysis']:
            para = docx_.add_paragraph(style='List Number')
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.add_run(analysis)
        ### Concept Pragraphs ###
        para = docx_.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run('Concepto: ').bold = True
        para.add_run('El Comité Asesor recomienda al Consejo de Facultad ')
        modifier = 'APROBAR:' if is_recommended else 'NO APROBAR:'
        para.add_run(modifier).bold = True
        ## First Concept Paragraph ##
        para = docx_.add_paragraph(style='List Number 2')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_aux = 'Cancelar la totalidad de asignaturas inscritas en el periodo {}, ' +\
                'en el programa de {}, '
        p_aux += 'teniendo en cuenta que {}justifica documentalmente la fuerza mayor o caso ' + \
            'fortuito '
        p_aux += '(Artículo 18 del Acuerdo 008 del Consejo Superior Universitario).'
        modifier = '' if is_recommended else 'no '
        para.add_run(p_aux.format(
            details['period_cancel'],
            get_academic_program(request['academic_program']),
            modifier
        ))
        ## Subjects Table ##
        subjects = []
        for subject in details_pre['subjects']:
            subjects.append([subject['code'], subject['name'], subject['group'],
                             subject['tipology'], subject['credits']])
        table_subjects(docx_, subjects)
        ## Second Concept Paragraph ##
        para = docx_.add_paragraph(style='List Number 2')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_aux = 'Devolución proporcional del {} por ciento ({}%) del valor pagado por concepto de '
        p_aux += 'derechos de matrícula del periodo {}, teniendo en cuenta la fecha de ' + \
            'presentación '
        p_aux += 'de la solicitud y que le fue aprobada la cancelación de periodo '
        p_aux += 'en el Acta {} de Consejo de Facultad '
        p_aux += '(Acuerdo 032 de 2010 del Consejo Superior Universitario, Artículo 1 ' + \
            'Resolución 1416 de 2013 de Rectoría).'
        para.add_run(p_aux.format(
            num2words(float(details_pre['percentage']), lang='es'),
            details_pre['percentage'],
            details['period_cancel'],
            details_pre['approval_minute']
        ))

    @staticmethod
    def case_CANCELACION_DE_PERIODO_ACADEMICO_PREGRADO(request, docx_, redirected=False):
        analysis_list = simple.case_CANCELACION_DE_PERIODO_ACADEMICO_PREGRADO_Analysis(
            request)
        answers_list = simple.case_CANCELACION_DE_PERIODO_ACADEMICO_PREGRADO_Answers(
            request)
        para = docx_.add_paragraph()
        para.add_run('Analisis:').font.bold = True
        analysis_para = docx_.add_paragraph()
        analysis_para.paragraph_format.left_indent = Pt(36)
        count = 1
        for analysis in analysis_list:
            analysis_para.add_run(str(count) + '. ' + analysis + '\n')
            count = count + 1
        para = docx_.add_paragraph()
        para.add_run('Concepto:').font.bold = True
        answer_para = docx_.add_paragraph()
        answer_para.paragraph_format.left_indent = Pt(36)
        count = 1
        for answer in answers_list:
            answer_para.add_run(str(count) + '. ' + answer + '\n')
            count = count + 1

    @staticmethod
    def case_CANCELACION_DE_PERIODO_ACADEMICO_PREGRADO_Analysis(request):
        a1_f = 'El comité asesor de {}{} lo considera fuerza mayor o caso fortuito documentado.'
        analysis1 = a1_f.format(request['pre_cm']['detail_pre_cm']['advisory_committee'],
                                '' if request['pre_cm']['pre_approval_status'] == 'AP' else ' NO')
        a2_f = 'Información del SIA:\n\t'
        a2_f += 'Porcentaje de avance del plan: {}\n\tNúmero de matrículas{}\n\tPAPA:{}.'
        advance = request['pre_cm']['detail_pre_cm']['advance']
        enrolled_academic_periods = request['pre_cm']['detail_pre_cm']['enrolled_academic_periods']
        papa = request['pre_cm']['detail_pre_cm']['papa']
        analysis2 = a2_f.format(advance, enrolled_academic_periods, papa)
        return [analysis1, analysis2] + request['pre_cm']['extra_analysis']

    @staticmethod
    def case_CANCELACION_DE_PERIODO_ACADEMICO_PREGRADO_Answers(request):
        c1_f1 = '{}ancelar el periodo académico {}, porque {}justifica documentalmente la ' + \
            'fuerza mayor '
        c1_f2 = 'o caso fortuito. (Artículo 18 del Acuerdo 008 del Consejo Superior Universitario).'
        if request['pre_cm']['pre_approval_status'] == 'RC':
            c1 = c1_f1.format('C', request['academic_period'], '') + c1_f2
            c2_f1 = 'Devolución proporcional del {} por ciento ({} %) del valor pagado ' + \
                'por concepto de derechos'
            c2_f2 = ' de matrícula del periodo {}, teniendo en cuenta la fecha de ' + \
                'presentación de la solicitud y '
            c2_f3 = 'que le fue aprobada la cancelación de periodo en el {} de Consejo de Facultad.'
            c2_f1_ = c2_f1.format(num2words(
                request['pre_cm']['devolution'], lang='es'), request['pre_cm']['devolution'])
            c2_f2_ = c2_f2.format(request['academic_period'])
            c2_f2_ = c2_f3.format(request['pre_cm']['cm_cancelation'])
            c2_f4_ = ' (Acuerdo 032 de 2010 del Consejo Superior Universitario,' + \
                ' Artículo 1 Resolución 1416 de 2013 de Rectoría)'
            c2 = c2_f1_ + c2_f2_ + c2_f2_ + c2_f4_
            return [c1, c2]
        else:
            c1 = c1_f1.format(
                'No c', request['academic_period'], 'no ') + c1_f2
            c2 = 'La situación expuesta no constituye causa extraña (no es una ' + \
                'situación intempestiva, insuperable o irresistible), '
            c22 = 'por tanto, no es una situación de fuerza mayor o caso fortuito ' + \
                'que implique la cancelación del periodo académico.'
            return [c1, c2+c22]

    @staticmethod
    def case_CAMBIO_DE_PERFIL_POSGRADO(request, docx_, redirected=False):
        ### Frequently used ###
        details = request['detail_cm']
        pre_cm = request['pre_cm']
        details_pre = pre_cm['detail_pre_cm']
        is_recommended = request['approval_status'] == 'CR'
        ### Finishing last paragraph ###
        para = docx_.paragraphs[-1]
        para.add_run('Análisis:')
        ### Analysis Paragraph ###
        ## Extra Analysis ##
        for analysis in pre_cm['extra_analysis']:
            para = docx_.add_paragraph(style='List Number')
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.add_run(analysis)
        ### Concept Pragraphs ###
        para = docx_.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run('Concepto: ').bold = True
        para.add_run('El Comité Asesor recomienda al Consejo de Facultad ')
        modifier = 'APROBAR:' if is_recommended else 'NO APROBAR:'
        para.add_run(modifier).bold = True
        ## Transfer ##
        para = docx_.add_paragraph(style='List Bullet')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_aux = 'Traslado intrafacultad del estudiante de {}, del plan de estudios {} '
        p_aux += 'al plan de estudios de {} debido a que {}justifica adecuadamente su solicitud.'
        modifier = '' if is_recommended else 'no '
        para.add_run(p_aux.format(
            get_academic_program(request['academic_program']),
            details['from_node'],
            details['to_node'],
            modifier
        ))
        ## Subject's type ##
        para = docx_.add_paragraph(style='List Bullet')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run('Cambiar de componente las siguientes asignaturas del programa {}:'.format(
            get_academic_program(request['academic_program'])
        ))
        subjects = []
        for subject in details_pre['subject_type_mod']:
            subjects.append([
                subject['code'], subject['name'], subject['grade'],
                subject['old_type'], subject['new_type']
            ])
        table_change_typology(docx_, subjects)
        ## Homologation ##
        para = docx_.add_paragraph(style='List Bullet')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_aux = 'Equivaler en el programa {} perfil de {}, '
        p_aux += 'las siguientes asignaturas cursadas en el programa {} perfil de {}:'
        para.add_run(p_aux.format(
            get_academic_program(request['academic_program']),
            details['to_node'],
            get_academic_program(request['academic_program']),
            details['from_node']
        ))
        subjects = []
        for subject in details_pre['subject_homologation']:
            subjects.append([
                subject['period'], subject['code'], subject['name'],
                subject['credits'], subject['tipology'], subject['grade'],
                subject['old_name'], subject['old_grade']
            ])
        table_approvals(docx_, subjects, [
            request['student_name'], request['student_dni'],
            request['academic_program'], 'perfil de {}'.format(
                details['from_node'])
        ])

    @staticmethod
    def case_AMPLIACION_DE_LA_FECHA_DE_PAGO_DE_DERECHOS_ACADEMICOS_POSGRADO(
            request, docx_, redirected=False):
        para = docx_.add_paragraph()
        para.paragraph_format.space_after = Pt(0)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run('Análisis:')
        para = docx_.add_paragraph(style='List Number')
        para.paragraph_format.space_after = Pt(0)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run('El estudiante tiene la historia académica activa.')
        if 'extra_analysis' in request.pre_cm:
            for analysis in request.pre_cm['extra_analysis']:
                para = docx_.add_paragraph(style='List Number')
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                para.paragraph_format.space_after = Pt(0)
                para.add_run(analysis)
        para = docx_.add_paragraph()
        para.paragraph_format.space_after = Pt(0)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run('Concepto:').font.bold = True
        para.add_run(' El comité Asesor ')
        if request.approval_status != 'CR':
            para.add_run('NO').font.bold = True
        para.add_run(
            ' recomienda al Consejo de Facultad presentar con concepto positivo al Comité')
        para.add_run(
            ' de Matrículas de la Sede Bogotá, la expedición de un único recibo correspondiente')
        para.add_run(
            ' a los derechos académicos y administrativos para el periodo académico ')
        para.add_run(request.academic_period)
        para.add_run(' y se le concede como fecha de pago el ')
        para.add_run(string_to_date(
            request.pre_cm['detail_pre_cm']['limit_date']))
        para.add_run(', teniendo en cuenta el estado de pago por parte de ')
        para.add_run(request.student_name)
        para.add_run('.')

    @staticmethod
    def case_REEMBOLSO_PREGRADO(request, docx_, redirected=False):
        raise NotImplementedError

    @staticmethod
    def case_EXENCION_DE_PAGO_POR_CURSAR_TESIS_COMO_UNICA_ACTIVIDAD_ACADEMICA_POSGRADO(
            request, docx_, redirected=False):
        ### Frequently used ###
        details = request['detail_cm']
        pre_cm = request['pre_cm']
        details_pre = pre_cm['detail_pre_cm']
        is_recommended = request['approval_status'] == 'CR'
        ### Finishing last paragraph ###
        para = docx_.paragraphs[-1]
        para.add_run('Análisis:\t\t')
        para.add_run(
            'Acuerdo 002 de 2011 - Consejo de Facultad.').underline = True
        ### Analysis Paragraph ###
        para = docx_.add_paragraph(style='List Number')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run('SIA: {}, perfil de {}.'.format(
            get_academic_program(request['academic_program']), details_pre['node']))
        para = docx_.add_paragraph(style='List Number')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        modifier = 'Está' if details_pre['on_time'] == "true" else 'No está'
        para.add_run('{} en fechas de solicitud.'.format(modifier))
        para = docx_.add_paragraph(style='List Number')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_aux = 'El incentivo se aplicará hasta por dos periodos académicos maestría y 5 periodos '
        p_aux += 'para doctorado (Literal c, Artículo 16). Ha tenido el incentivo {} periodos.'
        para.add_run(p_aux.format(details_pre['total_periods']))
        para = docx_.add_paragraph(style='List Number')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_aux = 'Para solicitar o renovar este estímulo la evaluación por parte del ' + \
            'director del trabajo '
        p_aux += 'desarrollado por el estudiante en la tesis o en el trabajo final ' + \
            'deberá ser "Avance Satisfactorio" si aplica.'
        para.add_run(p_aux)
        para = docx_.add_paragraph(style='List Number')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_aux = 'Se exceptúa el caso en el cual un estudiante también cursa el ' + \
            'último seminario del programa '
        p_aux += 'simultáneamente con la tesis, en cuyo caso se le podrá conceder dicho estímulo.'
        para.add_run(p_aux)
        para = docx_.add_paragraph(style='List Number')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_aux = 'Estos beneficios no se podrán renovar si el estudiante pierde calidad de '
        p_aux += 'estudiante o entra en reserva de cupo.'
        para.add_run(p_aux)
        para = docx_.add_paragraph(style='List Number')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_aux = 'Las exenciones establecidas en el Artículo 16 de la presente Acuerdo,' + \
            ' no podrán ser concedidas '
        p_aux += 'a estudiantes que hayan reingresado a programas de posgrado de la ' + \
            'Facultad en cualquier época (Artículo 19).'
        para.add_run(p_aux)
        ## Extra Analysis ##
        for analysis in pre_cm['extra_analysis']:
            para = docx_.add_paragraph(style='List Number')
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.add_run(analysis)
        ### Concept Pragraphs ###
        para = docx_.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run('Concepto: ').bold = True
        para.add_run('El Comité Asesor recomienda al Consejo de Facultad ')
        modifier = 'APROBAR' if is_recommended else 'NO APROBAR'
        para.add_run(modifier).bold = True

        ## Several options ##
        p_aux = ''
        if is_recommended:
            p_aux += ' pago de {} puntos por derechos académicos en el periodo académico ' + \
                '{}, condicionado a la '
            p_aux += 'inscripción de tesis/trabajo final y/o seminario como única actividad ' + \
                'académica en el '
            p_aux += 'periodo {}. (Literal c, Artículo 16 del Acuerdo 002 de 2011 del Consejo ' + \
                'de Facultad).'
            p_aux = p_aux.format(details['points'], details_pre['target_period'],
                                 details_pre['target_period'])
        elif details_pre['cause'] == "1":
            p_aux += ' pago de {} puntos por derechos académicos del periodo académico {},' + \
                ' condicionado a la '
            p_aux += 'inscripción de tesis/trabajo final y/o seminario como única actividad ' + \
                'académica en el '
            p_aux += 'periodo {}, porque el incentivo sólo se aplicará hasta por dos/cinco ' + \
                'periodos académicos '
            p_aux += '(Artículo 16 del Acuerdo 002 de 2011 del Consejo de Facultad, Literal c).'
            p_aux = p_aux.format(details['points'], details_pre['target_period'],
                                 details_pre['target_period'])
        elif details_pre['cause'] == "2":
            p_aux += ' porque para solicitar o renovar este estímulo la evaluación por parte de' + \
                'l director del trabajo '
            p_aux += 'desarrollado por el estudiante en la tesis o en el trabajo final deberá s' + \
                'er "Avance Satisfactorio". '
            p_aux += '(Literal c, Artículo 16 del Acuerdo 002 de 2011 del Consejo de Facultad).'
        elif details_pre['cause'] == "3":
            p_aux += ' porque las exenciones establecidas en el Artículo 16 del Acuerdo 002 de ' + \
                '2011 del Consejo de '
            p_aux += 'Facultad, no podrán ser concedidas a estudiantes que hayan reingresado a ' + \
                'programas de posgrado '
            p_aux += 'de la Facultad en cualquier época. (Artículo 19 del Acuerdo 002 de 2011 ' + \
                'del Consejo de Facultad).'
        elif details_pre['cause'] == "4":
            p_aux += 'porque la solicitud se presenta fuera de los plazos establecidos en el ' + \
                'Calendario de Tramites de '
            p_aux += 'Fin de Semestre {} - Estudiantes Posgrado (Artículo 58 del Acuerdo 008 de' + \
                ' 2008 del Consejo Superior Universitario).'
            p_aux = p_aux.format(request['academic_period'])
        para.add_run(p_aux)

    @staticmethod
    def case_GENERACION_DE_RECIBO_UNICO_DE_PAGO_POSGRADO(request, docx_, redirected=False):
        raise NotImplementedError

    @staticmethod
    def case_EXENCION_DE_PAGO_POR_CREDITOS_SOBRANTES_DE_PREGRADO_POSGRADO(
            request, docx_, redirected=False):
        para = docx_.add_paragraph()
        para.paragraph_format.space_after = Pt(0)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run('Análisis:\t\t')
        add_hyperlink(para, 'Acuerdo 008 de 2008',
                      'http://www.legal.unal.edu.co/rlunal/home/doc.jsp?d_i=34983/')
        para = docx_.add_paragraph(style='List Number')
        para.paragraph_format.space_after = Pt(0)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run('SIA: Admitido al programa ')
        para.add_run(get_academic_program(request.academic_program))
        para.add_run('. Perfil de ')
        para.add_run(request.pre_cm['detail_pre_cm']['profile'])
        para.add_run('.')
        para = docx_.add_paragraph(style='List Number')
        para.paragraph_format.space_after = Pt(0)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if request.pre_cm['detail_pre_cm']['enrolled'] == 'true':
            para.add_run('S')
        elif request.pre_cm['detail_pre_cm']['enrolled'] == 'false':
            para.add_run('No s')
        else:
            raise AssertionError(
                'request.pre_cm["detail_pre_cm"]["enrolled"] must be ' +
                '"true" or "false"')
        para.add_run(
            'e matriculó en un programa de posgrado de la Universidad en el año siguiente')
        para.add_run(
            ' a la culminación de sus estudios de pregrado. Culminó sus estudios en el periodo ')
        para.add_run(request.pre_cm['detail_pre_cm']['finish_period'])
        para.add_run(' e ingresó al posgrado en ')
        para.add_run(request.pre_cm['detail_pre_cm']['start_period'])
        para.add_run('.')
        para = docx_.add_paragraph(style='List Number')
        para.paragraph_format.space_after = Pt(0)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if request.pre_cm['detail_pre_cm']['limit'] == 'true':
            para.add_run('P')
        elif request.pre_cm['detail_pre_cm']['limit'] == 'false':
            para.add_run('No p')
        else:
            raise AssertionError(
                'request.pre_cm["detail_pre_cm"]["limit"] must be ' +
                '"true" or "false"')
        para.add_run(
            'resenta la solicitud dentro de las fechas límite establecidas por el reglamento:')
        para.add_run(
            ' 2 semanas después de la publicación de resultados de admitidos.')
        para = docx_.add_paragraph(style='List Number')
        para.paragraph_format.space_after = Pt(0)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run('La fecha límite para presentar la solicitud fue el ')
        para.add_run(string_to_date(
            request.pre_cm['detail_pre_cm']['limit_date']))
        para.add_run('.')
        para = docx_.add_paragraph()
        para.paragraph_format.space_after = Pt(0)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run('Concepto:').font.bold = True
        para.add_run(' El comité Asesor ')
        if request.approval_status != 'CR':
            para.add_run('NO').font.bold = True
        para.add_run(
            ' recomienda al Consejo de Facultad otorgar exención del pago de ')
        para.add_run(request.detail_cm['points'])
        para.add_run(' puntos de Derechos Académicos, a partir del periodo ')
        para.add_run(request.pre_cm['detail_pre_cm']['start_period'])
        para.add_run(', y durante los siguientes ')
        para.add_run(request.pre_cm['detail_pre_cm']['num_periods'])
        para.add_run(
            ' periodos académicos, por tener créditos disponibles al finalizar')
        para.add_run(' estudios del programa de pregado ')
        para.add_run(request.detail_cm['program'])
        para.add_run('.')
        if request.approval_status != 'CR' and request.pre_cm['detail_pre_cm']['limit'] == 'false':
            para.add_run(
                ' Debido a que presenta la solicitud fuera de los plazos establecidos')
        para.add_run(
            ' (Artículo 4 de la resolución 121 de 2010) (Artículo 58 del Acuerdo 008 de 2008')
        para.add_run(
            ' del Consejo Superior Universitario.). El cálculo de los créditos disponibles se')
        para.add_run(
            ' realiza con base en el cupo de créditos establecido en el Artículo 2 del Acuerdo')
        para.add_run(' 014 de 2008 del Consejo Académico.')

    @staticmethod
    def case_DEVOLUCION_PROPORCIONAL_DEL_VALOR_PAGADO_POR_CONCEPTO_DERECHOS_DE_MATRICULA_PRE(
            request, docx_, redirected=False):
        raise NotImplementedError

    @staticmethod
    def case_DEVOLUCION_DE_CREDITOS_PREGRADO(request, docx_, redirected=False):
        approval_status = "APROBAR"
        if request.approval_status == "NA":
            approval_status = "NO APROBAR"
        para = docx_.add_paragraph()
        para.add_run("Análisis:\t\t\tAcuerdo 018 de 2014")
        for i in range(0, len(request['pre_cm']['extra_analysis'])):
            para = docx_.add_paragraph(
                request['pre_cm']['extra_analysis'][i], style='List Number')
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para = docx_.add_paragraph()
        para.add_run("Concepto: ").font.bold = True
        para.add_run("El Comité Asesor recomienda al Consejo de Facultad {} reintegrar "
                     .format(approval_status))
        para.add_run(
            "al cupo, los créditos descontados por cancelación de la(s) siguiente(s) asignaturas ")
        para.add_run("en el periodo académico {}. (Circular 001 de 2019 de Vicerrectoría de Sede"
                     .format(request.academic_period))
        para.add_run(
            "Bogotá, Acuerdo 230 de 2016 de Consejo Superior Universitario)")
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        table = docx_.add_table(
            rows=len(request.detail_cm['subjects'])+2, cols=3, style='Table Grid')
        table.style.font.size = Pt(9)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for column in table.columns:
            for cell in column.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for cell in table.columns[0].cells:
            cell.width = 915000
        for cell in table.columns[1].cells:
            cell.width = 3620000
        for cell in table.columns[2].cells:
            cell.width = 915000
        table.cell(0, 0).paragraphs[0].add_run('Código SIA').font.bold = True
        table.cell(0, 1).paragraphs[0].add_run(
            'Nombre Asignatura').font.bold = True
        table.cell(0, 2).paragraphs[0].add_run('Créditos').font.bold = True
        index = 1
        credits_sum = 0
        for subject in request.detail_cm['subjects']:
            credits_sum = credits_sum+int(subject['credits'])
            table.cell(index, 0).paragraphs[0].add_run(subject['code'])
            table.cell(index, 1).paragraphs[0].add_run(subject['name'])
            table.cell(index, 2).paragraphs[0].add_run(subject['credits'])
            index = index + 1
        table.cell(index, 2).paragraphs[0].add_run(str(credits_sum))
        cellp = table.cell(index, 0).merge(table.cell(index, 1)).paragraphs[0]
        cellp.add_run('Total Créditos').font.bold = True

    @staticmethod
    def case_ELIMINACION_DE_LA_HISTORIA_ACADEMICA_BAPI_PREGRADO(request, docx_, redirected=False):
        para = docx_.add_paragraph()
        para.paragraph_format.space_after = Pt(0)
        para.add_run('Analisis:')
        para = docx_.add_paragraph(style='List Number')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run(
            'Modalidad de trabajo de grado: Asignaturas de posgrado. ')
        para.add_run('Acta de comité ' +
                     request.pre_cm['detail_pre_cm']['council_number'])
        para.add_run(
            ' de ' + request.pre_cm['detail_pre_cm']['council_year'] + '.')
        if 'extra_analysis' in request.pre_cm:
            for analysis in request.pre_cm['extra_analysis']:
                para = docx_.add_paragraph(style='List Number')
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                para.add_run(analysis)
        para.paragraph_format.space_after = Pt(0)
        para = docx_.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run('Concepto: ').font.bold = True
        para.add_run('El Comité Asesor ')
        if request.approval_status == 'RM':
            para.add_run('recomienda')
        elif request.approval_status == 'NM':
            para.add_run('no recomienda')
        para.add_run(
            ' al Consejo de Facultad eliminar la historia académica BAPI del periodo ')
        para.add_run(request.academic_period)
        para.add_run(', porque justifica debidamente la solicitud.')

    @staticmethod
    def case_RESERVA_DE_CUPO_ADICIONAL_PREGRADO(request, docx, redirected=False):
        para = docx.add_paragraph()
        run = para.add_run('Analisis:')
        run.font.bold = True
        analysis_para = docx.add_paragraph()
        analysis_para.paragraph_format.left_indent = Pt(36)
        analysis = ''
        answer = ''
        if request['approval_status'] == 'RM' or request['approval_status'] == 'AP':
            analysis_1 = 'El comité de {} considera que la situación personal está debidamente' + \
                ' justificada.'
            answer_1 = 'El Comité Asesor recomienda al Consejo de Facultad aprobar la primera' + \
                ' reserva de cupo adicional'
            answer_2 = ' en el periodo académico {} debido a que justifica debidamente la ' + \
                'solicitud. (Artículo 20 del'
            answer_3 = ' Acuerdo 008 de 2008 del Consejo Superior Universitario.)'
            analysis = analysis_1.format(
                request.get_academic_program_display())
            answer = answer_1 + \
                answer_2.format(request['academic_period']) + answer_3
        elif request['approval_status'] == 'NM' or request['approval_status'] == 'NA':
            analysis_1 = 'El comité de {} NO considera que la situación personal está ' + \
                'debidamente justificada.'
            answer_1 = 'El Comité Asesor recomienda al Consejo de Facultad no aprobar la' + \
                ' primera reserva de cupo adicional para el periodo académico {} teniendo ' + \
                'en cuenta que esta posibilidad es viable a continuación de la segunda reserva ' + \
                'de cupo automática. (Artículo 20 del Acuerdo 008 de 2008 del Consejo ' + \
                    'Superior Universitario)'
            analysis = analysis_1.format(
                request.get_academic_program_display())
            answer = answer_1.format(request['academic_period'])
        else:
            analysis = ' en trámite '
            answer = analysis
        analysis_para.add_run('1. ' + analysis + '\n')
        para = docx.add_paragraph()
        run = para.add_run('Concepto:')
        run.font.bold = True
        answer_para = docx.add_paragraph()
        answer_para.paragraph_format.left_indent = Pt(36)
        answer_para.add_run('1. ' + answer + '\n')
        answer_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    @staticmethod
    def case_REEMBOLSO_POSGRADO(request, docx_, redirected=False):
        raise NotImplementedError

    @staticmethod
    def case_ADMISION_AUTOMATICA_POSGRADO(request, docx_, redirected=False):
        ### Frequently used ###
        details = request['detail_cm']
        pre_cm = request['pre_cm']
        details_pre = pre_cm['detail_pre_cm']
        is_recommended = request['approval_status'] == 'CR'

        ### Finishing last paragraph ###
        para = docx_.paragraphs[-1]
        para.add_run('Análisis:\t')
        para.add_run(
            'Acuerdo 070 de 2009 - Consejo Académico, Acuerdo 008 de 2008').underline = True

        ### Analysis Paragraph ###
        para = docx_.add_paragraph(style='List Number')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run('El estudiante completó plan de estudios en {}.'.format(
            details_pre['grade_period']))

        para = docx_.add_paragraph(style='List Number')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run('Cupo de admisión automática en resolución {}.'.format(
            details_pre['place_resolution']))

        para = docx_.add_paragraph(style='List Number')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run('Solicita admisión al plan de estudios {} - perfil de {}.'.format(
            get_academic_program(request['academic_program']),
            details_pre['academic_profile']
        ))

        ## Extra Analysis ##
        for analysis in pre_cm['extra_analysis']:
            para = docx_.add_paragraph(style='List Number')
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.add_run(analysis)

        para = docx_.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run('Concepto: ').bold = True
        para.add_run('El Comité Asesor recomienda al Consejo de Facultad ')
        modifier = 'APROBAR' if is_recommended else 'NO APROBAR'
        para.add_run(modifier).bold = True
        p_aux = ' admisión automática al programa {} en el plan de estudios de {}'
        p_aux += ' a partir del periodo académico {} (Acuerdo 070 de 2009 de Consejo Académico '
        p_aux += 'y literal c, Artículo 57 del Acuerdo 008 de 2008 del Consejo Superior ' + \
            'Universitario).'
        para.add_run(p_aux.format(
            get_academic_program(request['academic_program']),
            details_pre['academic_profile'],
            details['ing_period']
        ))

    @staticmethod
    def case_REGISTRO_DE_CALIFICACION_DE_MOVILIDAD_PREGRADO(request, docx_, redirected=False):
        raise NotImplementedError

    @staticmethod
    def case_DESIGNACION_DE_JURADOS_CALIFICADORES_DE_TESIS_TRABAJO_FINAL_POSGRADO(
            request, docx_, redirected=False):
        para = docx_.add_paragraph()
        para.paragraph_format.space_after = Pt(0)
        para.add_run('Análisis:\t\t\t')
        add_hyperlink(
            para, 'http://www.legal.unal.edu.co/sisjurun/normas/Norma1.jsp?i=89183',
            'Acuerdo 40 de 2017 - Consejo de Facultad')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.space_after = Pt(0)
        para = docx_.add_paragraph(style='List Number')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run('SIA: Plan de estudios de ')
        if request.detail_cm['tesis_trabajo'] == 'Trabajo Final':
            para.add_run('profundización')
        elif request.detail_cm['tesis_trabajo'] == 'Tesis':
            para.add_run('investigación')
        else:
            raise AssertionError(
                'datail_cm[tesis_trabajo] must be "Trabajo Final" or "Tesis"')
        para.add_run('. La estudiante tiene la asignatura ')
        para.add_run(request.detail_cm['tesis_trabajo'])
        para.add_run(' de ' + request.detail_cm['nivel_pos'])
        para.add_run(
            ' (' + request.pre_cm['detail_pre_cm']['cod_assig'] + ').')
        para = docx_.add_paragraph(style='List Number')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run('Concepto').font.bold = True
        para.add_run(' motivado acerca del trabajo por parte del director ')
        para.add_run(request.pre_cm['detail_pre_cm']['director'])
        para.add_run(' (Artículo 43).')
        para = docx_.add_paragraph(style='List Number')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run('Propuesta de tesis aprobada ')
        para.add_run(
            '(' + request.pre_cm['detail_pre_cm']['date_approval'][0:2])
        para.add_run(' de ')
        para.add_run(num_to_month(
            request.pre_cm['detail_pre_cm']['date_approval'][3:5]))
        para.add_run(
            ' de 20' + request.pre_cm['detail_pre_cm']['date_approval'][6:8])
        para.add_run(
            ' - Acta ' + request.pre_cm['detail_pre_cm']['number_council'])
        para.add_run(' de Consejo de Facultad): ')
        para.add_run(request.detail_cm['tittle']).font.bold = True
        para = docx_.add_paragraph(style='List Number')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run(
            'Copia impresa y versión electrónica en formato pdf (Artículo 43).')
        para = docx_.add_paragraph(style='List Number')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run('Solicitud de nombramiento de jurados (Artículo 44).')
        para = docx_.add_paragraph(style='List Number')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run('Uno o más evaluadores para los trabajos finales, dos o más' +
                     ' jurados para las tesis de Maestría y cuatro jurados para ' +
                     'tesis de Doctorado (Artículo 44).')
        if request.detail_cm['nivel_pos'] == 'Doctorado':
            para = docx_.add_paragraph(style='List Number')
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.add_run('Para las tesis de doctorado, al menos dos de los jurados ' +
                         'deberán ser externos a la Universidad Nacional de Colombia ' +
                         'y preferiblemente laborar en el extranjero (Artículo 44).')
        elif request.detail_cm['nivel_pos'] != 'Maestría':
            raise AssertionError(
                'datail_cm[nivel_pos] must be "Maestría" or "Doctorado"')
        if 'extra_analysis' in request.pre_cm:
            for analysis in request.pre_cm['extra_analysis']:
                para = docx_.add_paragraph(style='List Number')
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                para.add_run(analysis)
        para = docx_.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run('Concepto: ').font.bold = True
        para.add_run('El Comité Asesor ')
        if request.approval_status == 'RM':
            para.add_run('recomienda')
        if request.approval_status == 'NM':
            para.add_run('no recomienda')
        para.add_run(' al Consejo de Facultad')
        para.add_run('designar en el jurado calificador de')
        if request.detail_cm['tesis_trabajo'] == 'Trabajo Final':
            para.add_run('l Trabajo Final de ' +
                         request.detail_cm['nivel_pos'] + ' de ')
        elif request.detail_cm['tesis_trabajo'] == 'Tesis':
            para.add_run(' la Tesis de ' +
                         request.detail_cm['nivel_pos'] + ' de ')
        large_program = get_academic_program(request.academic_program)
        para.add_run(large_program)
        para.add_run(', cuyo título es: "')
        para.add_run(request.detail_cm['tittle']).font.italic = True
        para.add_run('", a los docentes ')
        count = len(request.detail_cm['doc'])
        for doc in request.detail_cm['doc']:
            count = count - 1
            para.add_run(doc['nomb'])
            if 'dep' in doc:
                para.add_run(
                    ' de la Universidad Nacional de Colombia de la dependencia: ')
                para.add_run(doc['dep'])
            elif 'univ' in doc:
                para.add_run(' de la ' + doc['univ'])
            else:
                raise AttributeError
            if count == 0:
                para.add_run('.')
                return
            elif count == 1:
                para.add_run(' y ')
                break
            else:
                para.add_run(', ')
        length = len(request.detail_cm['doc']) - 1
        para.add_run(request.detail_cm['doc'][length]['nomb'])
        if 'dep' in request.detail_cm['doc'][length]:
            para.add_run(
                ' de la Universidad Nacional de Colombia de la dependencia: ')
            para.add_run(request.detail_cm['doc'][length]['dep'])
        elif 'univ' in request.detail_cm['doc'][length]:
            para.add_run(' de la ' + request.detail_cm['doc'][length]['univ'])
            para.add_run('.')
        else:
            raise AttributeError(
                'Each professor must have a "dep" or "univ" key into object')
        para.add_run('.')

    @staticmethod
    def case_MODIFICACION_DE_OBJETIVOS_DE_TESIS_PROPUESTA_POSGRADO(
            request, docx_, redirected=False):
        raise NotImplementedError

    @staticmethod
    def case_CARGA_INFERIOR_A_LA_MINIMA_PREGRADO(request, docx_, redirected=False):
        para = docx_.add_paragraph()
        analysis_str = 'Análisis:'
        run = para.add_run(analysis_str)
        run.font.bold = True
        add_hyperlink(para, 'Acuerdo 008 de 2008',
                      'http://www.legal.unal.edu.co/sisjurun/normas/Norma1.jsp?i=34983')
        para.paragraph_format.left_indent = Pt(36)
        str_in_ana = '\n1. SIA: Porcentaje de avance en el plan: {}. Número de matrículas:' + \
            ' {}. PAPA: {}.\n'
        str_in_ana += '2. SIA: Créditos disponibles: {}.'
        str_in_ans = 'El Comité Asesor recomienda al Consejo de Facultad cursar el periodo'
        str_in_ans += 'académico {} con un número de créditos inferior al mínimo exigido, porque{} '
        str_in_ans += 'justifica debidamente la solicitud (Artículo 10 del Acuerdo 008 de 2008 del '
        str_in_ans += 'Consejo Superior Universitario).'
        para.add_run(str_in_ana.format(request['pre_cm']['advance'],
                                       request['pre_cm']['enrolled_academic_periods'],
                                       request['pre_cm']['papa'],
                                       request['pre_cm']['available']))
        para = docx_.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        answer_str = 'Concepto: '
        run = para.add_run(answer_str)
        run.font.bold = True
        if request['approval_status'] == 'RC':
            para.add_run(str_in_ans.format(request['academic_period'], ' '))
        else:
            para.add_run(str_in_ans.format(request['academic_period'], ' no '))

    @staticmethod
    def case_RETIRO_DEFINITIVO_DEL_PROGRAMA_PREGRADO(request, docx_, redirected=False):
        ### Frequently used ###
        pre_cm = request['pre_cm']
        details_pre = pre_cm['detail_pre_cm']
        is_recommended = request['approval_status'] == 'CR'

        ### Finishing last paragraph ###
        para = docx_.paragraphs[-1]
        para.add_run('Análisis:\t\t\t')
        para.add_run('Acuerdo 026 de 2012').underline = True

        ### Analysis Paragraph ###
        para = docx_.add_paragraph(style='List Number')
        para.add_run(
            'SIA: Porcentaje de avance en el plan: {}%\nNúmero de matriculas: {}\nPAPA: {}.'.format(
                details_pre['percentage'], details_pre['register_num'], details_pre['PAPA']
            ))

        ## Extra Analysis ##
        for analysis in pre_cm['extra_analysis']:
            para = docx_.add_paragraph(style='List Number')
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.add_run(analysis)

        ### Concept Pragraphs ###
        para = docx_.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run('Concepto: ').bold = True
        para.add_run('El Comité Asesor recomienda al Consejo de Facultad ')
        modifier = 'APROBAR ' if is_recommended else 'NO APROBAR '
        para.add_run(modifier).bold = True
        para.add_run(
            'presentar con concepto positivo a la División de Registro y Matrícula, ' + \
                'el retiro voluntario del programa '
        )
        para.add_run(
            '{} ({})'.format(get_academic_program(
                request['academic_program']), request['academic_program'])
        ).bold = True

    @staticmethod
    def case_CREDITOS_EXCEDENTES_MAPI_PREGRADO(request, docx_, redirected=False):
        raise NotImplementedError

    @staticmethod
    def case_CAMBIO_DE_TIPOLOGIA_PREGRADO(request, docx_, redirected=False):
        raise NotImplementedError

    @staticmethod
    def case_TRANSITO_ENTRE_PROGRAMAS_POSGRADO(request, docx_, redirected=False):
        ### Frequently used ###
        details = request['detail_cm']
        pre_cm = request['pre_cm']
        details_pre = pre_cm['detail_pre_cm']
        is_recommended = request['approval_status'] == 'CR'

        ### Finishing last paragraph ###
        para = docx_.paragraphs[-1]
        para.add_run('Análisis:\t')
        para.add_run(
            'Resolución 035 de 2014, Acuerdo 002 de 2011 de Consejo de Facultad').underline = True

        ### Analysis Paragraphs ###
        para = docx_.add_paragraph(style='List Number')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run('En el programa hay {} cupos para tránsito.'.format(
            details_pre['availible_places']))

        para = docx_.add_paragraph(style='List Number')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run('Viene del programa {}.'.format(
            get_academic_program(details['origen'])))

        para = docx_.add_paragraph(style='List Number')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        modifier = '' if details_pre['language_prof'] == 'si' else 'no '
        para.add_run(
            'El estudiante {}cumple con la suficiencia de idioma exigida.'.format(modifier))

        para = docx_.add_paragraph(style='List Number')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_aux = 'La solicitud {}se hace luego de completar el plan de estudios y antes del grado'
        p_aux += ' (a menos que no se haya abierto convocatorio durante el periodo)'
        p_aux += '(Parágrafo 2, Resolución 241 de 2009).'
        modifier = '' if details_pre['on_time'] == 'si' else 'no '
        para.add_run(p_aux.format(modifier))

        ## Extra Analysis ##
        for analysis in pre_cm['extra_analysis']:
            para = docx_.add_paragraph(style='List Number')
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.add_run(analysis)

        ### Concept Paragraph ###
        para = docx_.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run('Concepto: ').bold = True
        para.add_run('El Comité Asesor recomienda al Consejo de Facultad ')
        modifier = 'APROBAR' if is_recommended else 'NO APROBAR'
        para.add_run(modifier).bold = True
        p_aux = ' tránsito del programa {} al programa {}, a partir del periodo académico {} '
        p_aux += '(Artículo 3, Resolución 241 de 2009 de la Vicerrectoría Académica).'
        para.add_run(p_aux.format(
            get_academic_program(details['origen']),
            get_academic_program(request['academic_program']),
            details['desde']
        ))

    @staticmethod
    def case_CAMBIO_DE_DIRECTIOR_CODIRECTOR_JURADO_O_EVALUADOR_POSGRADO(
            request, docx_, redirected=False):
        ### Frequently used ###
        details = request['detail_cm']
        pre_cm = request['pre_cm']
        details_pre = pre_cm['detail_pre_cm']
        is_recommended = request['approval_status'] == 'CR'

        ### Finishing last paragraph ###
        para = docx_.paragraphs[-1]
        para.add_run('Análisis:')

        ### Analysis Paragraph ###
        ## Extra Analysis ##
        for analysis in pre_cm['extra_analysis']:
            para = docx_.add_paragraph(style='List Number')
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.add_run(analysis)

        ### Concept Pragraphs ###
        para = docx_.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run('Concepto: ').bold = True
        para.add_run('El Comité Asesor recomienda al Consejo de Facultad ')
        modifier = 'APROBAR' if is_recommended else 'NO APROBAR'
        para.add_run(modifier).bold = True

        p_aux = ' designar director de {} de {} cuyo título es : "{}", '
        p_aux += 'al profesor {} del {}, en reemplazo del profesor {} '
        p_aux += 'designado en el {}.'

        para.add_run(p_aux.format(
            details['testra'],
            get_academic_program(request['academic_program']),
            details['titulo'],
            details['nuevo'],
            details['depto'],
            details['antiguo'],
            details_pre['minute']
        ))

    @staticmethod
    def case_DESIGNACION_DE_CODIRECTOR_POSGRADO(request, docx_, redirected=False):
        raise NotImplementedError

    @staticmethod
    def case_EVALUADOR_ADICIONAL_POSGRADO(request, docx_, redirected=False):
        ### Frequently used ###
        details = request['detail_cm']
        pre_cm = request['pre_cm']
        is_recommended = request['approval_status'] == 'CR'

        ### Finishing last paragraph ###
        para = docx_.paragraphs[-1]
        para.add_run('Análisis:')

        ### Analysis Paragraph ###
        ## Extra Analysis ##
        for analysis in pre_cm['extra_analysis']:
            para = docx_.add_paragraph(style='List Number')
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.add_run(analysis)

        ### Concept Pragraphs ###
        para = docx_.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run('Concepto: ').bold = True
        para.add_run('El Comité Asesor recomienda al Consejo de Facultad ')
        modifier = 'APROBAR' if is_recommended else 'NO APROBAR'
        para.add_run(modifier).bold = True

        p_aux = ' designar evaluador adicional del trabajo de {}, '
        p_aux += 'cuyo título es {}, al profesor {} '

        if details['professor_university'] != '':
            p_aux += 'de {}, '.format(details['professor_university'])
        else:
            p_aux += 'del departamento {} de la facultad de {}, '.format(
                details['professor_department'],
                details['professor_faculty']
            )

        p_aux += 'quien deberá dirimir la diferencia calificando el trabajo como '
        p_aux += 'aprobado o no aprobado (Acuerdo 56 de 2012 Consejo Superior Universitario).'

        para.add_run(p_aux.format(
            get_academic_program(request['academic_program']),
            details['title'],
            details['professor_name']
        ))

    @staticmethod
    def case_TRABAJO_DE_GRADO_PREGADO(request, docx_, redirected=False):
        raise NotImplementedError

    @staticmethod
    def case_APROBACION_PROYECTO_PROPUESTA_Y_DESIGNACION_DE_DIRECTOR_POSGRADO(
            request, docx_, redirected=False):
        raise NotImplementedError

    @staticmethod
    def case_MODIFICACION_DE_JURADOS_CALIFICADORES_POSGRADO(request, docx_, redirected=False):
        raise NotImplementedError

    @staticmethod
    def case_BECA_MEJOR_PROMEDIO_POSGRADO(request, docx_, redirected=False):
        raise NotImplementedError

    @staticmethod
    def case_EXCENCION_POR_MEJORES_SABER_PRO_POSGRADO(request, docx_, redirected=False):
        ### Frequently used ###
        pre_cm = request['pre_cm']
        details_pre = pre_cm['detail_pre_cm']
        is_recommended = request['approval_status'] == 'CR'

        ### Finishing last paragraph ###
        para = docx_.paragraphs[-1]
        para.add_run('Análisis:')

        ### Analysis Paragraphs ###
        ## Extra Analysis ##
        for analysis in pre_cm['extra_analysis']:
            para = docx_.add_paragraph(style='List Number')
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.add_run(analysis)

        ### Concept Paragraph ###
        para = docx_.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run('Concepto: ').bold = True
        para.add_run('El Comité Asesor recomienda al Consejo de Facultad ')
        modifier = 'APROBAR' if is_recommended else 'NO APROBAR'
        para.add_run(modifier).bold = True
        modifier = '' if details_pre['first_time'] == 'si' else 'la renovación de '
        para.add_run(' {}la exención del {}% del pago de los derechos académicos '.format(
            modifier,
            details_pre['value']
        ))
        modifier = 'l mejor puntaje' if details_pre['value'] == '100'\
             else ' los 10 mejores puntajes'
        para.add_run(
            'y renovación de matrícula a{} del Examen de Estado '.format(modifier))
        para.add_run('de la Calidad de la Educación Superior (SABER PRO) {} - {} '.format(
            details_pre['exam_year'],
            details_pre['degree']
        ))
        para.add_run('para el periodo académico {} '.format(
            request['academic_period']))
        para.add_run(
            '(Literal b, Artículo 16 del Acuerdo No.002 de 2011 del Consejo de Facultad).')

    @staticmethod
    def case_REINGRESO_POSGRADO(request, docx_, redirected=False):
        ### Frequently used ###
        pre_cm = request['pre_cm']
        details_pre = pre_cm['detail_pre_cm']
        is_recommended = request['approval_status'] == 'CR'

        ### Finishing last paragraph ###
        para = docx_.paragraphs[-1]
        para.add_run('Análisis:\t')
        para.add_run(
            'Resolución 239 de 2009,Acuerdo 008 de 2008,Resolución 012 de 2014').underline = True

        ### Analysis Paragraphs ###
        ## Last Reentry ##
        para = docx_.add_paragraph(style='List Number')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_aux = 'El estudiante {} ha tenido otro reingreso posterior al 2009-1S{} '
        p_aux += '(Artículo 46, Acuerdo 008 de 2008 del Consejo Superior Universitario).'
        last = details_pre['last_reentry']
        modifier = ('no', '') if last == '' else (
            'ya', ' en el periodo {}'.format(last))
        para.add_run(p_aux.format(*modifier))

        ## Retirement Cause ##
        para = docx_.add_paragraph(style='List Number')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run(
            '{}. Plan de estudios {} - Perfil de {}.'.format(
                details_pre['retirement_cause'],
                get_academic_program(request['academic_program']),
                details_pre['academic_profile']
            )
        )

        ## P.A.P.A. ##
        para = docx_.add_paragraph(style='List Number')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_aux = '{}iene PAPA superior o igual a 3.5 '
        p_aux += '(literal 3a – Artículo 3, Resolución 239 de 2009 de ' + \
            'Vicerrectoría Académica; Artículo 46, Acuerdo 008 de 2008 ' + \
                'del Consejo Superior Universitario).'
        modifier = 'T' if float(details_pre['PAPA']) >= 3.5 else 'No t'
        p_aux += 'SIA PAPA: '
        para.add_run(p_aux.format(modifier))
        para.add_run('{}.'.format(details_pre['PAPA'])).bold = True

        ## Remaining Subjects ##
        para = docx_.add_paragraph(style='List Number')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_aux = 'En caso de ser por máximo tiempo de permanencia o por tener ' + \
            'dos calificaciones NA en su historia académica:'
        p_aux += 'las asignaturas que le faltan por aprobar pueden cursarse ' + \
            'en un solo periodo académico adicional (literal 5 – Artículo 3, '
        p_aux += 'Resolución 239 de 2009 de Vicerrectoría Académica; parágrafo' + \
            ' 2 Artículo 46, Acuerdo 008 de 2008 del Consejo Superior Universitario).'
        p_aux += 'SIA: Le falta por aprobar '
        para.add_run(p_aux)
        para.add_run('{}.'.format(
            details_pre['remaining_subjects'])).bold = True

        ## On Time ##
        para = docx_.add_paragraph(style='List Number')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_aux = 'La solicitud {}se hace en fechas de calendario de sede ' + \
            '(parágrafo Artículo 3).'
        modifier = '' if details_pre['on_time'] == 'si' else 'no '
        para.add_run(p_aux.format(modifier))

        ## Extra Analysis ##
        for analysis in pre_cm['extra_analysis']:
            para = docx_.add_paragraph(style='List Number')
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.add_run(analysis)

        ### Concept Paragraph ###
        para = docx_.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run('Concepto: ').bold = True
        para.add_run('El Comité Asesor recomienda al Consejo de Facultad ')
        modifier = 'APROBAR' if is_recommended else 'NO APROBAR'
        para.add_run(modifier).bold = True
        para.add_run(' reingreso por única vez al programa {}, '.format(
            get_academic_program(request['academic_program'])))
        if is_recommended:
            para.add_run('a partir del periodo académico {}, '.format(
                details_pre['reentry_period']))

        ## Final Comment ##
        p_aux = 'el reingreso del estudiante estará regido por el Acuerdo ' + \
            '008 de 2008 del Consejo Superior Universitario.'
        p_aux += 'Durante el periodo académico adicional otorgado, el estudiante' + \
            ' deberá solicitar el nombramiento de jurados de su'
        p_aux += ' {}, con el fin de obtener su título, previo cumplimiento de ' + \
            'las demás exigencias académicas y administrativas vigentes.'
        p_aux += '(Artículo 7 de la Resolución 012 de 2014 de la Vicerrectoría Académica).'
        aditional = details_pre['aditional_comments'] + '.'
        modifier = p_aux.format(
            details_pre['grade_option']) if aditional == '.' else aditional
        para.add_run(modifier)

    @staticmethod
    def case_REGISTRO_DE_CALIFICACION_DEL_PROYECTO_Y_EXAMEN_DOCTORAL_POSGRADO(
            request, docx_, redirected=False):
        raise NotImplementedError

    @staticmethod
    def case_CAMBIO_DE_PROYECTO_DE_TESIS(request, docx_, redirected=False):
        ### Frequently used ###
        details = request['detail_cm']
        pre_cm = request['pre_cm']
        details_pre = pre_cm['detail_pre_cm']
        is_recommended = request['approval_status'] == 'CR'

        ### Finishing last paragraph ###
        para = docx_.paragraphs[-1]
        para.add_run('Análisis:  ')
        para.add_run(
            'Acuerdo 002 de 2011 de Consejo de Facultad, ' + \
                'Acuerdo 056 de 2012 C.S.U.').underline = True

        ### Analysis Paragraphs ###

        para = docx_.add_paragraph(style='List Number')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run(
            'Plan de estudios {} - Perfil de {} - Asignatura {}.'.format(
                get_academic_program(request['academic_program']),
                details_pre['academic_profile'],
                details_pre['grade_option']
            )
        )

        para = docx_.add_paragraph(style='List Number')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run('Tiene la firma del (los) director(es) de tesis')

        ## Extra Analysis ##
        for analysis in pre_cm['extra_analysis']:
            para = docx_.add_paragraph(style='List Number')
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.add_run(analysis)
        ### Concept Paragraph ###
        para = docx_.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run('Concepto: ').bold = True
        para.add_run('El Comité Asesor recomienda al Consejo de Facultad ')
        modifier = 'APROBAR' if is_recommended else 'NO APROBAR'
        para.add_run(modifier).bold = True
        para.add_run(' cambiar título de {} a: {}, '.format(
            details_pre['grade_option'], details['titulo']))
        if details_pre['previous_advisor'] == '' \
            or details_pre['previous_advisor'] == details_pre['advisor']:
            para.add_run('ratificando como director al profesor {} del Departamento de {}.'.format(
                details_pre['advisor'],
                details_pre['advisor_department']
            ))
        else:
            para.add_run('designando como nuevo director al profesor' + \
                ' {} del Departamento de {}'.format(
                    details_pre['advisor'],
                    details_pre['advisor_department']
                ))
            para.add_run(', en reemplazo del profesor {} del Departamento de {}.'.format(
                details_pre['previous_advisor'],
                details_pre['previous_advisor_department']
            ))

    @staticmethod
    def case_EXPEDICION_DE_RECIBO_PREGRADO(request, docx_, redirected=False):
        para = docx_.add_paragraph()
        approval = "APRUEBA"
        if request.approval_status == "NA":
            approval = "NO APRUEBA"
        para.add_run("Análisis:\t\t\tResolución 051 de 2003")
        fecha = request.pre_cm['detail_pre_cm']['payment_date'].split("-")
        para = docx_.add_paragraph("Recibo de pago original para cancelar hasta {}{}{}."
                                   .format(fecha[2], num_to_month(fecha[1]),
                                           fecha[0]), style='List Number')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for i in range(0, len(request['pre_cm']['extra_analysis'])):
            para = docx_.add_paragraph(
                request['pre_cm']['extra_analysis'][i], style='List Number')
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para = docx_.add_paragraph()
        para.add_run("Concepto: ").font.bold = True
        para.add_run("El Comité Asesor recomienda al Consejo de Facultad {}'+\
                     ' expedir un nuevo recibo de pago de derechos de matrícula ' + \
                     'con cambio de fecha, para el periodo académico {}."
                     .format(approval, request.academic_period))
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    @staticmethod
    def case_INFORME_DE_AVANCE_DE_TESIS_POSGRADO(request, docx_, redirected=False):
        raise NotImplementedError

    @staticmethod
    def case_ADICION_DE_CODIRECTOR_POSGRADO(request, docx_, redirected=False):
        para = docx_.add_paragraph()
        approval = "APRUEBA"
        if request.approval_status == "CR":
            approval = "NO APRUEBA"
        para = docx_.add_paragraph()
        para.add_run("Análisis:\t\t\tResolución 173 de 2009.")
        para = docx_.add_paragraph("SIA: <no se que va aquí> | <no se que va aquí>: perfil de {}."
                                   .format(request.pre_cm['detail_pre_cm']['profile']),
                                   style='List Number')
        for i in range(0, len(request['pre_cm']['extra_analysis'])):
            para = docx_.add_paragraph(
                request['pre_cm']['extra_analysis'][i], style='List Number')
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para = docx_.add_paragraph()
        para.add_run("Concepto: ").font.bold = True
        para.add_run("El Comité Asesor recomienda al Consejo de Facultad ")
        para.add_run(approval).font.bold = True
        para.add_run(" designar codirector de Tesis de {} - {} con título “{}” "
                     .format(request.pre_cm['detail_pre_cm']['niv_pos'],
                             request.get_academic_program_display(),
                             request.detail_cm['tittle']))
        para.add_run("aprobado en el acta No.{} del año {}, al profesor {} del Departamento de {}."
                     .format(request.detail_cm['council_AP'], request.detail_cm['council_AP_year'],
                             request.detail_cm['professor_name'],
                             request.detail_cm['professor_deparment']))
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    @staticmethod
    def case_TYPE_PROYECTO_DE_TESIS_O_TRABAJO_FINAL_DE_MAESTRIA_POSGRADO(
            request, docx_, redirected=False):
         ### Frequently used ###
        pre_cm = request['pre_cm']
        details_pre = pre_cm['detail_pre_cm']
        is_recommended = request['approval_status'] == 'CR'

        ### Finishing last paragraph ###
        para = docx_.paragraphs[-1]
        para.add_run('Análisis: ')
        para.add_run(
            'Acuerdo 040 de 2017 de Consejo de Facultad, ' + \
                'Acuerdo 056 de 2012 C.S.U.').underline = True

        ### Analysis Paragraph ###
        para = docx_.add_paragraph(style='List Number')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run('El estudiante en el perfil de {}, tiene la asignatura {} ({}).'.format(
            details_pre['node'], details_pre['subject_name'], details_pre['subject_code']))

        para = docx_.add_paragraph(style='List Number')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run('Tiene la firma del director de tesis.')

        if details_pre['research_group'] != '':
            para = docx_.add_paragraph(style='List Number')
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.add_run('El proyecto hace parte del grupo de investigación: {}.'.format(
                details_pre['research_group']))

        para = docx_.add_paragraph(style='List Number')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run('Titulo: ').bold = True
        para.add_run('"{}".'.format(details_pre['title'])).italic = True

        para = docx_.add_paragraph(style='List Number')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run('Objetivo General: ').bold = True
        para.add_run('{}.'.format(details_pre['general_objective']))

        para = docx_.add_paragraph(style='List Number')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run('Objetivos Especificos: ').bold = True
        for objective in details_pre['specific_objectives']:
            para = docx_.add_paragraph(style='List Bullet 2')
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.add_run('{}.'.format(objective))

        para = docx_.add_paragraph(style='List Number')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_aux = 'El proyecto de tesis debe inscribirse y entregarse, antes de alcanzar '
        p_aux += 'el 50% de la duración establecida para el programa (Parágrafo Artículo 14 '
        p_aux += 'del Acuerdo 056 de 2012 del Consejo Superior Universitario).'
        para.add_run(p_aux)

        ## Extra Analysis ##
        for analysis in pre_cm['extra_analysis']:
            para = docx_.add_paragraph(style='List Number')
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.add_run(analysis)

        ### Concept Pragraphs ###
        para = docx_.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run('Concepto: ').bold = True
        para.add_run('El Comité Asesor recomienda al Consejo de Facultad ')
        modifier = 'APROBAR:' if is_recommended else 'NO APROBAR:'
        para.add_run(modifier).bold = True

        para = docx_.add_paragraph(style='List Number 2')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_aux = 'Calificación aprobada (AP) a {} de {}, cuyo título es:'
        para.add_run(p_aux.format(
            details_pre['subject_name'],
            get_academic_program(request['academic_program'])
        ))
        para.add_run('"{}".'.format(details_pre['title'])).italic = True

        para = docx_.add_paragraph(style='List Number 2')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_aux = 'Designar director de {} cuyo título es '
        para.add_run(p_aux.format(details_pre['subject_name']))
        para.add_run('"{}" '.format(details_pre['title'])).italic = True
        p_aux = 'al profesor {}, del departamento de {}.'
        para.add_run(p_aux.format(
            details_pre['advisor_name'], details_pre['advisor_department']))
        para.add_run('.')

    @staticmethod
    def case_HOMOLOGACION_DE_ASIGNATURAS_INTERCAMBIO_ACADEMICO_INTERNACIONAL_PREGRADO(
            request, docx_, redirected=False):
        assign = ['2011183 - Intercambio Académico Internacional',
                  '2014269 - Intercambio Académico Internacional Prórroga',
                  '2026630 - Intercambio Académico Internacional II',
                  '2026631 - Intercambio Académico Internacional II Prórroga']
        para = docx_.add_paragraph()
        para.paragraph_format.space_after = Pt(0)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run('Análisis:\t\t')
        add_hyperlink(para, 'Acuerdo 008 de 2008',
                      'http://www.legal.unal.edu.co/rlunal/home/doc.jsp?d_i=34983/')
        para = docx_.add_paragraph(style='List Number')
        para.paragraph_format.space_after = Pt(0)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run('Solicitud de homologación de ')
        para.add_run(str(len(request.detail_cm['subjects'])))
        para.add_run(' asignaturas del programa ')
        para.add_run(get_academic_program(request.academic_program))
        para.add_run(' de la insitución ')
        para.add_run(request.detail_cm['inst'])
        para.add_run('.')
        para = docx_.add_paragraph(style='List Number')
        para.paragraph_format.space_after = Pt(0)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run('Las asignaturas')
        para.add_run('La(s) asignatura(s) a homologar ')
        if request.pre_cm['detail_pre_cm']['pre_req'] == 'false':
            para.add_run('NO ').font.bold = True
        elif request.pre_cm['detail_pre_cm']['pre_req'] != 'true':
            raise AssertionError('request.pre_cm["detail_pre"]["pre_req"]' +
                                 ' must be string "true" or "false"')
        para.add_run('cumple(n) con los prerrequisitos. ')
        if request.pre_cm['detail_pre_cm']['more_50'] == 'false':
            para.add_run('NO').font.bold = True
            para.add_run(' s')
        elif request.pre_cm['detail_pre_cm']['more_50'] == 'true':
            para.add_run('S')
        para.add_run('e homologan más del 50% de los créditos del plan')
        para.add_run(
            ' (Artículo 38, Acuerdo 008 de 2008 - Consejo Superior Universitario.). ')
        prev = 'antecedente' not in request.pre_cm['detail_pre_cm']
        if prev:
            para.add_run('NO').font.bold = True
            para.add_run(' h')
        else:
            para.add_run('H')
        para.add_run('a tenido homologaciones anteriores.')
        if not prev:
            para = docx_.add_paragraph(style='List Number')
            para.paragraph_format.space_after = Pt(0)
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.add_run(
                'Antecedente de homologación de la institución en Acta ')
            para.add_run(request.pre_cm['detail_pre_cm']
                         ['antecedente']['council_minute_number'])
            para.add_run(' de ')
            para.add_run(request.pre_cm['detail_pre_cm']
                         ['antecedente']['council_minute_year'])
            para.add_run('.')
        for analysis in request.pre_cm['extra_analysis']:
            para = docx_.add_paragraph(style='List Number')
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.space_after = Pt(0)
            para.add_run(analysis)
        para = docx_.add_paragraph()
        para.paragraph_format.space_after = Pt(0)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run('Concepto:').font.bold = True
        para.add_run(' El comité Asesor ')
        if request.approval_status != 'CR':
            para.add_run('NO').font.bold = True
        para.add_run(' recomienda aprobar:')
        para = docx_.add_paragraph(style='List Number 2')
        para.paragraph_format.space_after = Pt(0)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        total_creds = 0
        acum_papa = 0
        for subject in request.detail_cm['subjects']:
            total_creds += int(subject['creds_asig'])
            acum_papa += int(subject['creds_asig']) * \
                float(subject['new_cal_asig'])
        mini_papa = acum_papa / total_creds
        para.add_run('Registrar calificación ')
        if mini_papa > 3:
            para.add_run('aprobada (AP)')
        else:
            para.add_run('no aprobada (NA)')
        para.add_run(' en la asignatura ')
        para.add_run(assign[int(request.pre_cm['detail_pre_cm']['index']) - 1])
        para.add_run(', en el periodo ')
        para.add_run(request.academic_period)
        para.add_run('.')
        para = docx_.add_paragraph(style='List Number 2')
        para.paragraph_format.space_after = Pt(0)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run('Homologar ')
        para.add_run(
            'la(s) siguiente(s) asignatura(s) cursada(s) en Intercambio ')
        para.add_run(
            'Académico Internacional en ')
        if 'old_program' in request.pre_cm['detail_pre_cm']:
            para.add_run(request.pre_cm['detail_pre_cm']['old_program'])
            para.add_run(' de ')
        para.add_run(request.detail_cm['inst'])
        para.add_run(
            ', de la siguiente manera (Artículo 35, Acuerdo 008 de 2008 del ')
        para.add_run(
            'Consejo Superior Universitario y Resolución 105 de 2017 de ')
        para.add_run('Vicerrectoría Académica):')
        details = [request.student_name, request.student_dni,
                   request.academic_program, request.detail_cm['inst']]
        subjects = []
        for sbj in request.detail_cm['subjects']:
            subjects.append([sbj['per_asig'], sbj['cod_asig'], sbj['new_name_asig'],
                             sbj['creds_asig'], sbj['tipo_asig'], sbj['new_cal_asig'],
                             sbj['old_name_asig'], sbj['old_cal_asig']])
        table_approvals(docx_, subjects, details)

    @staticmethod
    def case_HOMOLOGACION_DE_ASIGNATURAS_CONVENIO_CON_UNIVERSIDAD_ANDES_PREGRADO(
            request, docx_, redirected=False):
        assign = ['2011302 - Asignatura Por Convenio Con Universidad De Los Andes I - Pregrado',
                  '2012698 - Asignatura Por Convenio Con Universidad De Los Andes II - Pregrado']
        para = docx_.add_paragraph()
        para.paragraph_format.space_after = Pt(0)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run('Análisis:\t\t')
        add_hyperlink(para, 'Acuerdo 008 de 2008',
                      'http://www.legal.unal.edu.co/rlunal/home/doc.jsp?d_i=34983/')
        para = docx_.add_paragraph(style='List Number')
        para.paragraph_format.space_after = Pt(0)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run('Solicitud de homologación de ')
        para.add_run(str(len(request.detail_cm['subjects'])))
        para.add_run(' asignaturas del programa ')
        para.add_run(get_academic_program(request.academic_program))
        para.add_run(' de la Universidad de los Andes.')
        para = docx_.add_paragraph(style='List Number')
        para.paragraph_format.space_after = Pt(0)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run('La(s) asignatura(s) a homologar ')
        if request.pre_cm['detail_pre_cm']['pre_req'] == 'false':
            para.add_run('NO ').font.bold = True
        elif request.pre_cm['detail_pre_cm']['pre_req'] != 'true':
            raise AssertionError('request.pre_cm["detail_pre"]["pre_req"]' +
                                 ' must be string "true" or "false"')
        para.add_run('cumple(n) con los prerrequisitos.')
        para = docx_.add_paragraph(style='List Number')
        para.paragraph_format.space_after = Pt(0)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if request.pre_cm['detail_pre_cm']['more_50'] == 'false':
            para.add_run('NO').font.bold = True
            para.add_run(' s')
        elif request.pre_cm['detail_pre_cm']['more_50'] == 'true':
            para.add_run('S')
        para.add_run('e homologan más del 50% de los créditos del plan')
        para.add_run(
            ' (Artículo 38, Acuerdo 008 de 2008 - Consejo Superior Universitario.). ')
        prev = 'antecedente' not in request.pre_cm['detail_pre_cm']
        if prev:
            para.add_run('NO').font.bold = True
            para.add_run(' h')
        else:
            para.add_run('H')
        para.add_run('a tenido homologaciones anteriores.')
        if not prev:
            para = docx_.add_paragraph(style='List Number')
            para.paragraph_format.space_after = Pt(0)
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.add_run(
                'Antecedente de homologación de la institución en Acta ')
            para.add_run(request.pre_cm['detail_pre_cm']
                         ['antecedente']['council_minute_number'])
            para.add_run(' de ')
            para.add_run(request.pre_cm['detail_pre_cm']
                         ['antecedente']['council_minute_year'])
            para.add_run('.')
        for analysis in request.pre_cm['extra_analysis']:
            para = docx_.add_paragraph(style='List Number')
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.space_after = Pt(0)
            para.add_run(analysis)
        para = docx_.add_paragraph()
        para.paragraph_format.space_after = Pt(0)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run('Concepto:').font.bold = True
        para.add_run(' El comité Asesor ')
        if request.approval_status != 'CR':
            para.add_run('NO').font.bold = True
        para.add_run(' recomienda aprobar:')
        para = docx_.add_paragraph(style='List Number 2')
        para.paragraph_format.space_after = Pt(0)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run('Registrar calificación ')
        total_creds = 0
        acum_papa = 0
        for subject in request.detail_cm['subjects']:
            total_creds += int(subject['creds_asig'])
            acum_papa += int(subject['creds_asig']) * \
                float(subject['cal_asign'])
        mini_papa = acum_papa / total_creds
        para.add_run('Registrar calificación ')
        if mini_papa > 3:
            para.add_run('aprobada (AP)')
        else:
            para.add_run('no aprobada (NA)')
        para.add_run(' en la asignatura ')
        para.add_run(assign[int(request.detail_cm['index']) - 1])
        para.add_run(', en el periodo ')
        para.add_run(request.academic_period)
        para.add_run('.')
        para = docx_.add_paragraph(style='List Number 2')
        para.paragraph_format.space_after = Pt(0)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run('Homologar, en el periodo académico ')
        para.add_run(request.academic_period)
        para.add_run(
            ', la(s) siguiente(s) asignatura(s) cursada(s) en el Convenio en la ')
        para.add_run(
            'Universidad de los Andes de la siguiente manera (Artículo 35 de Acuerdo')
        para.add_run(' 008 de 2008 del Consejo Superior Universitario):')
        subjects = []
        details = [request.student_name, request.student_dni,
                   request.academic_program, 'Universidad de los Andes']
        for subject in request.detail_cm['subjects']:
            subjects.append(
                [request.academic_period, subject['cod_asig'], subject['name_asig'],
                 subject['creds_asig'], 'L', subject['cal_asign'], subject['name_asig'],
                 subject['cal_asign']])
        table_approvals(docx_, subjects, details)

    @staticmethod
    def case_HOMOLOGACION_DE_ASIGNATURAS_CONVENIO_CON_UNIVERSIDAD_ANDES_POSGRADO(
            request, docx_, redirected=False):
        assign = ['2024944 - Asignatura Por Convenio Con Universidad De Los Andes I - Posgrado',
                  '2024945 - Asignatura Por Convenio Con Universidad De Los Andes II - Posgrado']
        para = docx_.add_paragraph()
        para.paragraph_format.space_after = Pt(0)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run('Análisis:\t\t')
        add_hyperlink(para, 'Acuerdo 008 de 2008',
                      'http://www.legal.unal.edu.co/rlunal/home/doc.jsp?d_i=34983/')
        para = docx_.add_paragraph(style='List Number')
        para.paragraph_format.space_after = Pt(0)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run('Solicitud de homologación de ')
        para.add_run(str(len(request.detail_cm['subjects'])))
        para.add_run(' asignaturas del programa ')
        para.add_run(get_academic_program(request.academic_program))
        para.add_run(' de la Universidad de los Andes.')
        para = docx_.add_paragraph(style='List Number')
        para.paragraph_format.space_after = Pt(0)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run('La(s) asignatura(s) a homologar ')
        if request.pre_cm['detail_pre_cm']['pre_req'] == 'false':
            para.add_run('NO ').font.bold = True
        elif request.pre_cm['detail_pre_cm']['pre_req'] != 'true':
            raise AssertionError('request.pre_cm["detail_pre"]["pre_req"]' +
                                 ' must be string "true" or "false"')
        para.add_run('cumple(n) con los prerrequisitos.')
        para = docx_.add_paragraph(style='List Number')
        para.paragraph_format.space_after = Pt(0)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if request.pre_cm['detail_pre_cm']['more_50'] == 'false':
            para.add_run('NO').font.bold = True
            para.add_run(' s')
        elif request.pre_cm['detail_pre_cm']['more_50'] == 'true':
            para.add_run('S')
        para.add_run('e homologan más del 50% de los créditos del plan')
        para.add_run(
            ' (Artículo 38, Acuerdo 008 de 2008 - Consejo Superior Universitario.). ')
        prev = 'antecedente' not in request.pre_cm['detail_pre_cm']
        if prev:
            para.add_run('NO').font.bold = True
            para.add_run(' h')
        else:
            para.add_run('H')
        para.add_run('a tenido homologaciones anteriores.')
        if not prev:
            para = docx_.add_paragraph(style='List Number')
            para.paragraph_format.space_after = Pt(0)
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.add_run(
                'Antecedente de homologación de la institución en Acta ')
            para.add_run(request.pre_cm['detail_pre_cm']
                         ['antecedente']['council_minute_number'])
            para.add_run(' de ')
            para.add_run(request.pre_cm['detail_pre_cm']
                         ['antecedente']['council_minute_year'])
            para.add_run('.')
        for analysis in request.pre_cm['extra_analysis']:
            para = docx_.add_paragraph(style='List Number')
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.space_after = Pt(0)
            para.add_run(analysis)
        para = docx_.add_paragraph()
        para.paragraph_format.space_after = Pt(0)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run('Concepto:').font.bold = True
        para.add_run(' El comité Asesor ')
        if request.approval_status != 'CR':
            para.add_run('NO').font.bold = True
        para.add_run(' recomienda aprobar:')
        para = docx_.add_paragraph(style='List Number 2')
        para.paragraph_format.space_after = Pt(0)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        total_creds = 0
        acum_papa = 0
        for subject in request.detail_cm['subjects']:
            total_creds += int(subject['creds_asig'])
            acum_papa += int(subject['creds_asig']) * \
                float(subject['cal_asign'])
        mini_papa = acum_papa / total_creds
        para.add_run('Registrar calificación ')
        if mini_papa > 3:
            para.add_run('aprobada (AP)')
        else:
            para.add_run('no aprobada (NA)')
        para.add_run(' en la asignatura ')
        para.add_run(assign[int(request.detail_cm['index']) - 1])
        para.add_run(', en el periodo ')
        para.add_run(request.academic_period)
        para.add_run('.')
        para = docx_.add_paragraph(style='List Number 2')
        para.paragraph_format.space_after = Pt(0)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run('Homologar, en el periodo académico ')
        para.add_run(request.academic_period)
        para.add_run(
            ', la(s) siguiente(s) asignatura(s) cursada(s) en el Convenio en la ')
        para.add_run(
            'Universidad de los Andes de la siguiente manera (Artículo 35 de Acuerdo')
        para.add_run(' 008 de 2008 del Consejo Superior Universitario):')
        subjects = []
        details = [request.student_name, request.student_dni,
                   request.academic_program, 'Universidad de los Andes']
        for subject in request.detail_cm['subjects']:
            subjects.append(
                [request.academic_period, subject['cod_asig'], subject['name_asig'],
                 subject['creds_asig'], 'L', subject['cal_asign'], subject['name_asig'],
                 subject['cal_asign']])
        table_approvals(docx_, subjects, details)

    @staticmethod
    def case_PROYECTO_DE_TESIS_DE_DOCTORADO_POSGRADO(request, docx_, redirected=False):
        ### Frequently used ###
        pre_cm = request['pre_cm']
        details_pre = pre_cm['detail_pre_cm']
        is_recommended = request['approval_status'] == 'CR'

        ### Finishing last paragraph ###
        para = docx_.paragraphs[-1]
        para.add_run('Análisis: ')

        ### Analysis Paragraph ###
        para = docx_.add_paragraph(style='List Number')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_aux = 'SIA: Perfil de {}. El estudiante tiene la asignatura {} ({}).'
        para.add_run(p_aux.format(
            details_pre['node'], details_pre['subject_name'], details_pre['subject_code']
        ))

        para = docx_.add_paragraph(style='List Number')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_aux = 'Jurados evaluadores designados en Acta de Comité {}.'
        para.add_run(p_aux.format(details_pre['approved_jury_minute']))

        para = docx_.add_paragraph(style='List Number')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run(
            'Presenta acta de calificación firmada con jurados evaluadores designados.')

        if details_pre['research_group'] != '':
            para = docx_.add_paragraph(style='List Number')
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.add_run('El proyecto hace parte del grupo de investigación: {}.'.format(
                details_pre['research_group']))

        para = docx_.add_paragraph(style='List Number')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run('Titulo: ').bold = True
        para.add_run('"{}".'.format(details_pre['title'])).italic = True

        para = docx_.add_paragraph(style='List Number')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run('Objetivo General: ').bold = True
        para.add_run('{}.'.format(details_pre['general_objective']))

        para = docx_.add_paragraph(style='List Number')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run('Objetivos Especificos: ').bold = True
        for objective in details_pre['specific_objectives']:
            para = docx_.add_paragraph(style='List Bullet 2')
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.add_run('{}.'.format(objective))

        ## Extra Analysis ##
        for analysis in pre_cm['extra_analysis']:
            para = docx_.add_paragraph(style='List Number')
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.add_run(analysis)

        ### Concept Pragraphs ###
        para = docx_.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run('Concepto: ').bold = True
        para.add_run('El Comité Asesor recomienda al Consejo de Facultad ')
        modifier = 'APROBAR:' if is_recommended else 'NO APROBAR:'
        para.add_run(modifier).bold = True

        para = docx_.add_paragraph(style='List Number 2')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_aux = 'Inscribir y calificar aprobado (AP) el Examen de calificación '
        p_aux += 'con código {} en el periodo académico {}.'
        para.add_run(p_aux.format(
            details_pre['subject_code'],
            details_pre['subject_period']
        ))

        para = docx_.add_paragraph(style='List Number 2')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_aux = 'Calificar aprobado (AP) el proyecto de Tesis de Doctorado en '
        p_aux += 'Ingeniería, cuyo título es: '
        para.add_run(p_aux)
        para.add_run('"{}".'.format(details_pre['title'])).italic = True

        para = docx_.add_paragraph(style='List Number 2')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run(
            'Designar director de la Tesis de Doctorado cuyo título es ')
        para.add_run('"{}" '.format(details_pre['title'])).italic = True
        p_aux = 'al profesor {}, del departamento de {}.'
        para.add_run(p_aux.format(
            details_pre['advisor_name'], details_pre['advisor_department']))
        para.add_run('.')

    @staticmethod
    def case_CANCELACION_DE_ASIGNATURAS_CON_CARGA_INFERIOR_A_LA_MINIMA_PREGRADO(
            request, docx_, redirected=False):
        para = docx_.add_paragraph()
        para.paragraph_format.space_after = Pt(0)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run('Análisis:\t\t')
        add_hyperlink(para, 'Acuerdo 008 de 2008',
                      'http://www.legal.unal.edu.co/rlunal/home/doc.jsp?d_i=34983/')
        para = docx_.add_paragraph(style='List Number')
        para.paragraph_format.space_after = Pt(0)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run('SIA: Porcentaje de avance en el plan: ')
        para.add_run(request.pre_cm['detail_pre_cm']['advance_percentage'])
        para.add_run('%. Número de matrículas: ')
        para.add_run(request.pre_cm['detail_pre_cm']['enrollment_number'])
        para.add_run('. P.A.P.A.: ')
        para.add_run(request.pre_cm['detail_pre_cm']['PAPA'])
        para.add_run('.')
        para = docx_.add_paragraph(style='List Number')
        para.paragraph_format.space_after = Pt(0)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run('SIA: Créditos disponibles: ')
        para.add_run(request.pre_cm['detail_pre_cm']['available_creds'])
        para.add_run('.')
        para = docx_.add_paragraph(style='List Number')
        para.paragraph_format.space_after = Pt(0)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run(
            'SIA: Al aprobar la cancelación de la(s) asignatura(s) ')
        para.add_run('solicitada(s) el estudiante quedaría con ')
        para.add_run(request.pre_cm['detail_pre_cm']['affter_cancel_creds'])
        para.add_run(' créditos inscritos.')
        if 'extra_analysis' in request.pre_cm:
            for analysis in request.pre_cm['extra_analysis']:
                para = docx_.add_paragraph(style='List Number')
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                para.paragraph_format.space_after = Pt(0)
                para.add_run(analysis)
        para = docx_.add_paragraph()
        para.paragraph_format.space_after = Pt(0)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run('Concepto: ').font.bold = True
        para.add_run('El Comité Asesor')
        if request.approval_status != 'CR':
            para.add_run(' NO').font.bold = True
        para.add_run(' recomienda al Consejo de Facultad aprobar:')
        para = docx_.add_paragraph(style='List Number 2')
        para.paragraph_format.space_after = Pt(0)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run('Cursar el periodo académico ')
        para.add_run(request.academic_period)
        para.add_run(
            ' con un número de créditos inferior al mínimo exigido, porque')
        if request.approval_status != 'CR':
            para.add_run(' NO').font.bold = True
        para.add_run(
            ' justifica debidamente la solicitud. (Artículo 10 del Acuerdo 008')
        para.add_run(' de 2008 del Consejo Superior Universitario).')
        para = docx_.add_paragraph(style='List Number 2')
        para.paragraph_format.space_after = Pt(0)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run(
            'Cancelar la(s) siguiente(s) asignatura(s) inscrita(s) del periodo ')
        para.add_run('académico ')
        para.add_run(request.academic_period)
        para.add_run(', porque')
        if request.approval_status != 'CR':
            para.add_run(' NO').font.bold = True
        para.add_run(
            ' justifica debidamente la solicitud. (Artículo 10 del Acuerdo 008')
        para.add_run(' de 2008 del Consejo Superior Universitario).')
        subjects = []
        for sbj in request.detail_cm['subjects']:
            subjects.append([sbj['cod_sia'], sbj['name_asig'],
                             sbj['group_asig'], sbj['tipo_asig'], sbj['creds_asig']])
        table_subjects(docx_, subjects)

    @staticmethod
    def case_APROBACION_PASANTIA_POSGRADO(request, docx_, redirected=False):
        ### Frequently used ###
        pre_cm = request['pre_cm']
        details_pre = pre_cm['detail_pre_cm']
        is_recommended = request['approval_status'] == 'CR'

        ### Finishing last paragraph ###
        para = docx_.paragraphs[-1]
        para.add_run('Análisis: ')

        ### Analysis Paragraph ###
        ## Extra Analysis ##
        for analysis in pre_cm['extra_analysis']:
            para = docx_.add_paragraph(style='List Number')
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.add_run(analysis)

        para = docx_.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run('Concepto: ').bold = True
        para.add_run('El Comité Asesor recomienda al Consejo de Facultad ')
        modifier = 'APROBAR' if is_recommended else 'NO APROBAR:'
        para.add_run(modifier).bold = True

        p_aux = ' realizar pasantía en {} bajo la tutoría del profesor {}, como parte '
        p_aux += 'del desarrollo de su Trabajo de grado, desde {} de {} y por {} meses. '
        p_aux += 'El estudiante deberá estar debidamente matriculado en el periodo {}.'
        para.add_run(p_aux.format(
            details_pre['place'], details_pre['advisor_name'], details_pre['init_month'],
            details_pre['init_year'], details_pre['months'], details_pre['target_period']
        ))
        para.add_run('.')

    @staticmethod
    def case_BECA_EXENSION_DERECHOS_ACADEMICOS(request, docx_, redirected=False):
        ### Frequently used ###
        pre_cm = request['pre_cm']
        details_pre = pre_cm['detail_pre_cm']
        is_recommended = request['approval_status'] == 'CR'

        ### Finishing last paragraph ###
        para = docx_.paragraphs[-1]
        para.add_run('Análisis: ')

        ### Analysis Paragraph ###
        para = docx_.add_paragraph(style='List Number')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_aux = 'La coordinación curricular del programa presenta como beneficiario de la BECA '
        p_aux += 'EXENCIÓN DE DERECHOS ACADÉMICOS del Acuerdo 2 de 2012 de Consejo de Facultad, '
        p_aux += 'por obtener el promedio académico ponderado más alto del semestre en las '
        p_aux += 'asignaturas cursadas durante el periodo académico inmediatamente anterior.'
        para.add_run(p_aux)

        para = docx_.add_paragraph(style='List Number')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run('Promedio: {}/5.0 .'.format(details_pre['average']))

        ## Extra Analysis ##
        for analysis in pre_cm['extra_analysis']:
            para = docx_.add_paragraph(style='List Number')
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.add_run(analysis)

        ### Concept Pragraphs ###
        para = docx_.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run('Concepto: ').bold = True
        para.add_run('El Comité Asesor recomienda al Consejo de Facultad ')
        modifier = 'APROBAR' if is_recommended else 'NO APROBAR'
        para.add_run(modifier).bold = True

        p_aux = ' la BECA EXENCIÓN DE DERECHOS ACADÉMICOS en el programa de {} ({}) '
        p_aux += 'en el periodo {} y otorgar la exención del 100% de derechos académicos '
        p_aux += 'por obtener el promedio académico ponderado más alto del semestre en '
        p_aux += 'las asignaturas cursadas durante el periodo académico inmediatamente '
        p_aux += 'anterior (Artículo 8 Acuerdo 02 de 2012 de Consejo de Facultad).'

        para.add_run(p_aux.format(
            get_academic_program(request['academic_program']),
            request['academic_program'],
            details_pre['target_period']
        ))
