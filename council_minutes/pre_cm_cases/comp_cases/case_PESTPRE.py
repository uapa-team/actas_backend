from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


class PESTPRE():

    @staticmethod
    def case_PRACTICA_ESTUDIANTIL_PREGRADO(request, docx, redirected=False):
        PESTPRE.case_PRACTICA_ESTUDIANTIL_PREGRADO_Analysis(request, docx)
        PESTPRE.case_PRACTICA_ESTUDIANTIL_PREGRADO_Answers(request, docx)

    @staticmethod
    def case_PRACTICA_ESTUDIANTIL_PREGRADO_Analysis(request, docx):
        para = docx.add_paragraph()
        run = para.add_run('Analisis: ')
        run.font.bold = True
        # add_hyperlink(para, 
        # 'Acuerdo 016 de 2011','http://www.legal.unal.edu.co/sisjurun/normas/Norma1.jsp?i=44965')
        para = docx.add_paragraph()
        para.paragraph_format.left_indent = Pt(36)
        PESTPRE.case_PRACTICA_ESTUDIANTIL_PREGRADO_Analysis_1(request, para)
        PESTPRE.case_PRACTICA_ESTUDIANTIL_PREGRADO_Analysis_2(request, para)
        PESTPRE.case_PRACTICA_ESTUDIANTIL_PREGRADO_Analysis_3(request, para)
        PESTPRE.case_PRACTICA_ESTUDIANTIL_PREGRADO_Analysis_4(request, para)

    @staticmethod
    def case_PRACTICA_ESTUDIANTIL_PREGRADO_Analysis_1(request, para):
        str_1 = '\n1. El estudiante{} cumple con el requisito de haber aprobado el'
        str_1 += ' 70% de los créditos del plan de estudios. SIA: {} de avance en '
        str_1 += 'los créditos exigidos del plan de estudios.'
        if int(request['pre_cm']['advance']) >= 70:
            para.add_run(str_1.format(
                ' ', request['pre_cm']['advance'] + '%'))
        else:
            para.add_run(str_1.format(
                ' no ', request['pre_cm']['advance'] + '%'))

    @staticmethod
    def case_PRACTICA_ESTUDIANTIL_PREGRADO_Analysis_2(request, para):
        str_2 = '\n2. El estudiante{}ha cursado otra de las asignaturas con '
        str_2 += 'el nombre Práctica Estudiantil.'
        if request['pre_cm']['another_practice'] == 'true':
            para.add_run(str_2.format(' '))
        else:
            para.add_run(str_2.format(' no '))

    @staticmethod
    def case_PRACTICA_ESTUDIANTIL_PREGRADO_Analysis_3(request, para):
        str_3 = '\n3. Requisitos: Pertinencia, objetivos, alcance, empresa {},'
        str_3 += 'duración: {} horas/semana durante {}, costos, descripción '
        str_3 += 'de actividades (Artículo 3, Acuerdo 016 de 2011 – Consejo '
        str_3 += 'Académico). A cargo de un profesor de la Facultad: {}, '
        str_3 += 'porcentajes de evaluación definidos (Artículo 3, Acuerdo '
        str_3 += '016 de 2011 – Consejo Académico)'
        para.add_run(str_3.format(request['detail_cm']['institution'],
                                        request['pre_cm']['hours_week'],
                                        request['pre_cm']['duration'],
                                        request['detail_cm']['person_un']))

    @staticmethod
    def case_PRACTICA_ESTUDIANTIL_PREGRADO_Analysis_4(request, para):
        str_4 = '\n4. Documentación{}cumple con requisitos: Formato{}'
        str_4 += 'está completamente diligenciado. {}adjunta copia del '
        str_4 += 'Acuerdo firmado. {}adjunta el recibido de la carta de '
        str_4 += 'presentación de la Universidad. {}están fijados los '
        str_4 += 'porcentajes de evaluación.'

        str_ans_1 = ' sí ' if request['pre_cm']['format_right'] == 'true' and \
            request['pre_cm']['agreement_signed'] == 'true' and \
            request['pre_cm']['presentation_letter'] == 'true' and \
            request['pre_cm']['evaluation_percents'] == 'true' else ' no '
        str_ans_2 = ' sí ' if request['pre_cm']['format_right'] == 'true' else ' no '
        str_ans_3 = 'Sí ' if request['pre_cm']['agreement_signed'] == 'true' else 'No '
        str_ans_4 = 'Sí ' if request['pre_cm']['presentation_letter'] == 'true' else 'No '
        str_ans_5 = 'Sí ' if request['pre_cm']['evaluation_percents'] == 'true' else 'No '
        para.add_run(str_4.format(str_ans_1,
                                        str_ans_2,
                                        str_ans_3,
                                        str_ans_4,
                                        str_ans_5))

    @staticmethod
    def case_PRACTICA_ESTUDIANTIL_PREGRADO_Answers(request, docx):
        para = docx.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = para.add_run('Concepto: ')
        run.font.bold = True
        para = docx.add_paragraph()
        para.paragraph_format.left_indent = Pt(36)
        str_a_1 = 'El Comité Asesor recomienda al Consejo de Facultad inscribir la '
        str_a_1 += 'siguiente asignatura en el periodo académico {}, a '
        str_a_1 += 'desarrollar en la empresa {}, a cargo del profesor {}, por '
        str_a_1 += 'parte de la Universidad Nacional de Colombia y el Sr. {} por '
        str_a_1 += 'parte de la entidad.'
        para.add_run(str_a_1.format(request['academic_period'],
                                          request['detail_cm']['institution'],
                                          request['detail_cm']['person_un'],
                                          request['detail_cm']['person_ins']))
