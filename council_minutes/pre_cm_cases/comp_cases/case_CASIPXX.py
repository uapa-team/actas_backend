from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from .case_utils import *


class CASIPXX():

    count = 0

    @staticmethod
    def case_CANCELACION_DE_ASIGNATURAS(request, docx, redirected=False):
        CASIPXX.case_CANCELACION_DE_ASIGNATURAS_Analysis(request, docx)
        CASIPXX.case_CANCELACION_DE_ASIGNATURAS_Answers(request, docx)

    @staticmethod
    def case_CANCELACION_DE_ASIGNATURAS_Analysis(request, docx):
        para = docx.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.left_indent = Pt(36)
        run = para.add_run('Analisis: ')
        run.font.bold = True
        # add_hyperlink(para, 'Acuerdo 008 de 2008',
        # 'http://www.legal.unal.edu.co/rlunal/home/doc.jsp?d_i=34983')
        CASIPXX.case_CANCELACION_DE_ASIGNATURAS_Analysis_1(request, para)
        CASIPXX.case_CANCELACION_DE_ASIGNATURAS_Analysis_2(request, para)
        CASIPXX.case_CANCELACION_DE_ASIGNATURAS_Analysis_3(request, para)
        CASIPXX.case_CANCELACION_DE_ASIGNATURAS_Analysis_extra(request, para)

    @staticmethod
    def case_CANCELACION_DE_ASIGNATURAS_Analysis_1(request, para):
        str_in = '\n1. SIA: Porcentaje de avance en el plan: {}. Número de'
        str_in += 'matrículas: {}. PAPA: {}.'
        para.add_run(str_in.format(request['pre_cm']['advance'],
                           request['pre_cm']['enrolled_academic_periods'],
                           request['pre_cm']['papa']))

    @staticmethod
    def case_CANCELACION_DE_ASIGNATURAS_Analysis_2(request, para):
        str_in = '\n2. SIA: Créditos disponibles: {}.'
        para.add_run(str_in.format(request['pre_cm']['available']))

    @staticmethod
    def case_CANCELACION_DE_ASIGNATURAS_Analysis_S(para, subject):
        str_in = '\n{}. SIA: Al aprobar la cancelación de la asignatura {} ({}) '
        str_in += ' el estudiante quedaría con {} créditos inscritos.'
        para.add_run(str_in.format(subject['number'], subject['code'],
                           subject['subject'], subject['remaining']))

    @staticmethod
    def case_CANCELACION_DE_ASIGNATURAS_Analysis_3(request, docx):
        CASIPXX.count = 2
        for subject in request['detail_cm']['subjects']:
            CASIPXX.count = CASIPXX.count + 1
            subject['number'] = str(CASIPXX.count)
            current_credits = int(request['pre_cm']['current_credits'])
            subject_credits = int(subject['credits'])
            subject['remaining'] = current_credits - subject_credits
            CASIPXX.case_CANCELACION_DE_ASIGNATURAS_Analysis_S(docx, subject)

    @staticmethod
    def case_CANCELACION_DE_ASIGNATURAS_Analysis_extra(request, para):
        for analysis in request['pre_cm']['extra_analysis']:
            CASIPXX.count = CASIPXX.count + 1
            str_in = '\n{}. {}.'
            para.add_run(str_in.format(CASIPXX.count, analysis))

    @staticmethod
    def case_CANCELACION_DE_ASIGNATURAS_Answers(request, docx):
        if request['approval_status'] == 'RC':
            CASIPXX.case_CANCELACION_DE_ASIGNATURAS_Answers_RC(request, docx)
        else:
            CASIPXX.case_CANCELACION_DE_ASIGNATURAS_Answers_NRC(request, docx)

    @staticmethod
    def case_CANCELACION_DE_ASIGNATURAS_Answers_RC(request, docx):
        str_in = 'El Comité Asesor recomienda al Consejo de Facultad'
        str_in += ' cancelar la(s) siguiente(s) asignatura(s) inscrita(s) del '
        str_in += 'periodo académico {}, porque se justifica debidamente '
        str_in += 'la solicitud. (Artículo 15 Acuerdo 008 de 2008 del '
        str_in += 'Consejo Superior Universitario)'
        para = docx.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run(str_in.format(request['academic_period']))
        para.add_run('Concepto: ')
        para.font.bold = True
        data = []
        index = 0
        for subject in request['detail_cm']['subjects']:
            data.append([])
            data[index] += [subject['code']]
            data[index] += [subject['subject']]
            data[index] += [subject['group']]
            data[index] += [subject['tipology']]
            data[index] += [subject['credits']]
            index = index + 1
        table_subjects(docx, data)

    @staticmethod
    def case_CANCELACION_DE_ASIGNATURAS_Answers_NRC(request, docx):
        str_in = 'El Comité Asesor recomienda al Consejo de Facultad'
        str_in += ' NO cancelar la(s) siguiente(s) asignatura(s) inscrita(s) '
        str_in += 'del periodo académico {}, '
        if request['pre_cm']['nrc'] == 'Incoherente o consecuente':
            str_in += 'porque no existe coherencia entre la documentación y '
            str_in += 'justificación que presenta. '
        elif request['pre_cm']['nrc'] == 'No diligente':
            str_in += 'porque lo expuesto es un hecho de su conocimiento '
            str_in += 'desde el inicio del periodo académico; tuvo la '
            str_in += 'oportunidad de resolverlo oportunamente hasta el '
            str_in += '50 % del periodo académico, por tanto, no constituye '
            str_in += 'causa extraña que justifique la cancelación de '
            str_in += 'la(s) asignatura(s). '
        elif request['pre_cm']['nrc'] == 'Motivos Laborales':
            str_in += 'porque de acuerdo con la documentación que presenta, '
            str_in += 'su situación laboral no le impide asistir a las clases '
            str_in += 'y tiene el tiempo suficiente para responder por las '
            str_in += 'actividades académicas de la(s) asignatura(s). '
        elif request['pre_cm']['nrc'] == 'Información Falsa':
            str_in += 'porque verificada la información de los soportes, se '
            str_in += 'encontró que el contenido de los mismos no coincide '
            str_in += 'con lo que en ellos se afirma. '
        elif request['pre_cm']['nrc'] == 'Falta de conocimiento':
            str_in += 'poque es responsabilidad del estudiante indagar sobre '
            str_in += 'el conocimiento requerido y la preparación necesaria '
            str_in += 'para cursar la(s) asignatura(s) antes de inscribir. '
        elif request['pre_cm']['nrc'] == 'Argumentos insuficientes':
            str_in += 'porque lo expuesto no es un hecho que constituya causa '
            str_in += 'extraña que justifique la cancelación de la(s) '
            str_in += 'asignatura(s). '
        elif request['pre_cm']['nrc'] == 'Argumento cuando los soportes no aportan':
            str_in += 'porque de la documentación aportada, se tiene que no hay '
            str_in += 'justificación para acceder a lo pedido. '
        else:
            pass
        str_in += ' (Artículo 15 Acuerdo 008 de 2008 del '
        str_in += 'Consejo Superior Universitario).'
        para = docx.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = para.add_run('Concepto: ')
        run.font.bold = True
        para.add_run(str_in.format(request['academic_period']))