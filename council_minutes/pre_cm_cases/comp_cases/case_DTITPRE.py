from docx import Document
from ...models import Request
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.table import WD_TABLE_ALIGNMENT
from num2words import num2words  ##pip install num2words
from docx.shared import Pt
from .case_utils import *
from ...cm_cases.comp_cases.case_DTITPRE import *


class DTITPRE_pre():
    
    @staticmethod
    def case_DOBLE_TITULACION_PREGRADO(request, docx, redirected=False):
        graduate = "Es"
        enrolled_student = "Es"
        student_quality = "Se"
        quota_credits = "Dispone"
        recommend = "recomendar"
        if request.pre_cm['detail_pre_cm']['graduate_student'] == "No":
            graduate = "No es"
        if request.detail_cm['informacion_academica']['matriculado_solicitud'] == "No":
            enrolled_student = "No es"
        if request.detail_cm['informacion_academica']['calidad_estudiante_seg_plan'] == "No":
            enrolled_student = "No se"
        if request.pre_cm['detail_pre_cm']['quota_credits'] == "No":
            quota_credits = "No ha"
        if request.approval_status == "NA":
            recommend = "no recomendar"
        para = docx.add_paragraph()
        para.add_run("Análisis:\t\t\tAcuerdo 155 de 2014, Acuerdo 008 de 2008")
        para = docx.add_paragraph("{} estudiante de posgrado (Artículo 49 Acuerdo 008 del 2008 Consejo Superior Universitario.). Universitas y SIA: {}"
        .format(graduate, request.pre_cm['detail_pre_cm']['1_SIA_Universitas']), style='List Number')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para = docx.add_paragraph("{} ha matriculado al momento de la solicitud (Artículo 1, Acuerdo 155 de 2014 del Consejo Superior Universitario.). Universitas y SIA: {}"
        .format(enrolled_student, request.pre_cm['detail_pre_cm']['2_SIA_Universitas']), style='List Number')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para = docx.add_paragraph("{} tenido calidad de estudiante en el plan de estudios de doble titulación (Artículo 4, Acuerdo 155 de 2014 del Consejo Superior Universitario.).Universitas: {}"
        .format(student_quality, request.pre_cm['detail_pre_cm']['3_Universitas']), style='List Number')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para = docx.add_paragraph("{} del cupo de créditos necesario para optar por el segundo título luego de convalidar o hacer equivaler todas las asignaturas pertinentes cursadas y aprobadas en el primer plan de estudios (parágrafo 1, Artículo 48 Acuerdo 008 del 2008 Consejo Superior Universitario)."
        .format(quota_credits), style='List Number')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para = docx.add_paragraph("Régimen de convalidaciones y equivalencias PERTINENTES entre el primero y el segundo plan de estudios (Artículo 2, Acuerdo 155 de 2014 del Consejo Superior Universitario.).",
        style='List Number')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para = docx.add_paragraph("No ha perdido la calidad de estudiante por las causales 2, 3, 4 o 5 del Artículo 44 del Acuerdo 008 del 2008 Consejo Superior Universitario. (Artículo 7, Acuerdo 155 de 2014 del Consejo Superior  Universitario.).Universitas: {}"
        .format(request.pre_cm['detail_pre_cm']['6_Universitas']), style='List Number')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para = docx.add_paragraph()
        para.add_run("Concepto: ").font.bold = True
        para.add_run("El Comité Asesor recomienda al Consejo de Facultad APROBAR {} al Consejo de Sede que formalice la admisión y ubicación en el programa de pregrado {} – {}"
        .format(recommend, request.get_academic_program_display(), request.academic_program))
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if request.approval_status == "NA":
            para.add_run(", debido a que tiene un PAPA de {} y no cuenta con el cupo suficiente de créditos para culminar el segundo plan. (Acuerdo 155 de 2014 del Consejo Superior Universitario)"
            .format(request.detail_cm['informacion_academica']['PAPA_plan_origen']))
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        elif request.approval_status == "AP":
            para = docx.add_paragraph("Teniendo en cuenta que el estudiante tiene un Promedio Académico Ponderado Acumulado superior o igual a 4.3. (Acuerdo 155 de 2014 del Consejo Superior Universitario)", style = "List Bullet")
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para = docx.add_paragraph("Teniendo en cuenta que el estudiante cuenta con un cupo de créditos suficiente para culminar el segundo plan de estudios. (Acuerdo 155 de 2014 del Consejo Superior Universitario)", style = "List Bullet")
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para = docx.add_paragraph()
            DTITPRE.case_DOBLE_TITULACION_PREGRADO_TABLES(request, docx, para)
    




