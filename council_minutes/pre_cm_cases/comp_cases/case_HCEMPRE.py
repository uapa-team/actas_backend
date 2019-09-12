from num2words import num2words
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Pt
from ...models import Request
from .case_utils import *


class HCEMPRE():

    @staticmethod
    def case_HOMOLOGACION_CONVALIDACION_EQUIVALENCIA_PREGRADO(request, docx, redirected=False):
        case_d = {'homologation': 'homologa', 'equivalence': 'equivale', 'recognition': 'convalida'}
        # negation = {'negation':'', 'negation_mayus': 'Se'} 
        previous_minute = ""
        previous_approvals = "Ha"
        prerequisites = ""
        credits_50 = "Se"
        subjects = []
        subjects_na = []
        para = docx.add_paragraph() 
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY       
        for i in range (0, len(request.detail_cm['subjects'])):
            if request.detail_cm['subjects'][i]['subject_ap_status'] == "AP":
                temporal = [
                    request.detail_cm['subjects'][i]['period'],
                    request.detail_cm['subjects'][i]['cod'],
                    request.detail_cm['subjects'][i]['name'],
                    request.detail_cm['subjects'][i]['credits'],
                    request.detail_cm['subjects'][i]['typology'],
                    request.detail_cm['subjects'][i]['grade'],
                    request.detail_cm['subjects'][i]['old_name'],
                    request.detail_cm['subjects'][i]['old_grade']
                ]
                subjects.append(temporal)
            else:
                temporal = [
                    request.detail_cm['subjects'][i]['name'],
                    request.detail_cm['subjects'][i]['old_name'],
                    request.detail_cm['subjects'][i]['justification'],
                    request.detail_cm['subjects'][i]['credits'],
                    request.detail_cm['subjects'][i]['old_grade']
                ]
                subjects_na.append(temporal)

        if request.pre_cm['detail_pre_cm']['prerequisites_subjects'] == 'No':
            prerequisites = 'no '
        if request.pre_cm['detail_pre_cm']['50%_credits'] == "No":
            credits_50 = "No se"
        if request.pre_cm['detail_pre_cm']['previous_approvals'] == "No":
            previous_approvals = 'No ha'
        
        para.add_run("Análisis:\t\t\tAcuerdo 008 de 2008, Acuerdo 86 de 2018")
        para = docx.add_paragraph("Solicitud de homologación de {} asignaturas del programa {} de la institución {}."
        .format(request.pre_cm['detail_pre_cm']['requested_subjects'], 
                request.get_academic_program_display(), 
                request.detail_cm['old_university']), style='List Number')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para = docx.add_paragraph("Las asignaturas a homologar {}cumplen con los prerrequisitos."
        .format(prerequisites), style='List Number')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para = docx.add_paragraph("{} homologan/convalidan más del 50% de créditos del plan (Artículo "\
        "38, Acuerdo 008 de 2008 – Consejo Superior Universitario). {} tenido homologaciones/convalidaciones anteriores."
        .format(credits_50, previous_approvals), style='List Number')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        if previous_approvals == "No ha":
            previous_minute = ("Antecedente de homologación de la institución en el {}."
            .format(request.pre_cm['detail_pre_cm']['previous_minute'])) 
            para =  docx.add_paragraph(previous_minute, style = 'List Number')
        
        for i in range (0, len(request['pre_cm']['extra_analysis'])):
            para = docx.add_paragraph(request['pre_cm']['extra_analysis'][i], style = 'List Number')
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para =  docx.add_paragraph()
        para.add_run("Concepto: ").font.bold = True
        # subjects = [
        #     ['2019-1S', '1000024', 'Principios de química', '3', 'B', '3.7', 'Fundamentos químicos y bioquímicos', '3.5'],
        #     ['2019-1S', '1000024', 'Principios de química', '3', 'B', '3.7', 'Termodinámica y fluidos', '3.9'],
        #     ['2019-2S', '2023533', 'Programación orientada a objetos', '3', 'L', '4.1', 'Programación orientada objetos', '4.1'],
        #     ['2019-2S', '1234567', 'Otra materia', '3', 'L', '4.6', 'Programación orientada objetos', '4.1']
        # ]

        # details = ['Laura Molina', '1022431736', '2879', 'Universidad distrital francisco josé de caldas']
        details = [request.student_name, request.student_dni, request.academic_program, request.detail_cm['old_university'], request.detail_cm['case']]
        
        para.add_run("El Comité Asesor recomienda al Consejo de Facultad")
        if len(subjects) > 0 and len(subjects_na) > 0:
            lista = True
        else:
            lista = False

        if len(subjects) > 0:
            cadena_ap = "APROBAR {}r "\
            "la(s) siguiente(s) asignatura(s) cursada(s) en el programa {}, de la siguiente manera "\
            "(Artículo 35 del Acuerdo 008 de 2008 del Consejo Superior Universitario)"
            if lista == True:
                para.add_run(":")
                para = docx.add_paragraph(cadena_ap
                .format(case_d[request.detail_cm['case']], request.academic_period, request.get_academic_program_display()), style='List Bullet')
                table_approvals(docx, subjects, details)
            else:
                para.add_run(" " + cadena_ap
                .format(case_d[request.detail_cm['case']], request.academic_period, request.get_academic_program_display()))
                table_approvals(docx, subjects, details)
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if len(subjects_na) > 0:
            cadena_na = "NO APROBAR {}r "\
            "la(s) siguiente(s) asignatura(s) cursada(s) en el programa {}, de la siguiente manera "\
            "(Artículo 35 del Acuerdo 008 de 2008 del Consejo Superior Universitario)"
            if lista == True:
                para = docx.add_paragraph(cadena_na
                .format(case_d[request.detail_cm['case']], request.academic_period, request.get_academic_program_display()), style='List Bullet')
                table_approvals_na(docx, subjects_na, details)
            else:
                para.add_run(" " + cadena_na
                .format(case_d[request.detail_cm['case']], request.academic_period, request.get_academic_program_display()))
                table_approvals_na(docx, subjects_na, details)   
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


