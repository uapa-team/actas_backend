from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from .case_utils import table_general_data
from .case_utils import table_credits_summary


class TRASPRE():

    @staticmethod
    def case_TRASLADO_PREGRADO(request, docx, redirected=False):
        TRASPRE.case_TRASLADO_PREGRADO_Analysis(request, docx)
        TRASPRE.case_TRASLADO_PREGRADO_Answers(request, docx)

    @staticmethod
    def case_TRASLADO_PREGRADO_Analysis(request, docx):
        para = docx.add_paragraph()
        run = para.add_run('Analisis:')
        run.font.bold = True
        para = docx.add_paragraph()
        para.paragraph_format.left_indent = Pt(36)
        TRASPRE.case_TRASLADO_PREGRADO_Analysis_1(request, para)
        TRASPRE.case_TRASLADO_PREGRADO_Analysis_2(request, para)
        TRASPRE.case_TRASLADO_PREGRADO_Analysis_3(request, para)
        TRASPRE.case_TRASLADO_PREGRADO_Analysis_4(request, para)
        TRASPRE.case_TRASLADO_PREGRADO_Analysis_5(request, para)
        TRASPRE.case_TRASLADO_PREGRADO_Analysis_6_1(request, para)
        TRASPRE.case_TRASLADO_PREGRADO_Analysis_6_2(request, para)
        TRASPRE.case_TRASLADO_PREGRADO_Analysis_7(request, para)
        TRASPRE.case_TRASLADO_PREGRADO_Analysis_8(request, para)

    @staticmethod
    def case_TRASLADO_PREGRADO_Analysis_1(request, para):
        instr = '1. Viene del plan {}.'
        instr = instr.format(request['detail_cm']['origin'])
        para.add_run(instr)

    @staticmethod
    def case_TRASLADO_PREGRADO_Analysis_2(request, para):
        instr_1 = '\n2. {}a tenido calidad estudiante en ese programa previamente '
        instr_2 = '(Parágrafo 1. Artículo 2, Acuerdo 089 de 2014 de Consejo Académico).' + \
            'Universitas: .'
        if request['pre_cm']['was_before'] == 'true':
            instr = instr_1.format('H') + instr_2
        else:
            instr = instr_1.format('No h') + instr_2
        para.add_run(instr)

    @staticmethod
    def case_TRASLADO_PREGRADO_Analysis_3(request, para):
        instr_1 = '\n3. {}a culminado el primer plan de estudios y '
        instr_2 = '{}tiene derecho a renovar la matrícula Universitas: .'
        if request['pre_cm']['has_finished_first'] == 'true':
            instr = instr_1.format('H')
        else:
            instr = instr_1.format('No h')
        if request['pre_cm']['right_to_renovate'] == 'true':
            instr += instr_2.format(" ")
        else:
            instr += instr_2.format("no ")
        para.add_run(instr)

    @staticmethod
    def case_TRASLADO_PREGRADO_Analysis_4(request, para):
        instr_1 = '\n4. {}a cursado por lo menos un periodo académico del primer plan '
        instr_2 = 'de estudios (Artículo 39, Acuerdo 008 de 2008 de Consejo Superior). '
        instr_3 = 'SIA: {} periodos académicos desde {}.'
        if request['pre_cm']['was_before'] == 'true':
            instr = instr_1.format('H') + instr_2
        else:
            instr = instr_1.format('No h') + instr_2
        instr += instr_3.format(request['pre_cm']
                                ['num_periods'], request['detail_cm']['since'])
        para.add_run(instr)

    @staticmethod
    def case_TRASLADO_PREGRADO_Analysis_5(request, para):
        instr_1 = '5. {}stá cursando doble titulación '
        instr_2 = '(Artículo 7. Acuerdo 155 de 2014 – Consejo Superior Universitario.). SIA: OK.'
        if request['pre_cm']['double_titulation'] == 'true':
            instr = instr_1.format('E') + instr_2
        else:
            instr = instr_1.format('No e') + instr_2
        para.add_run(instr)

    @staticmethod
    def case_TRASLADO_PREGRADO_Analysis_6_1(request, para):
        instr_1 = '\n6.1 {}iene un puntaje de admisión igual o superior al puntaje '
        instr_2 = 'del último admitido regular al plan de estudios de destino en el periodo {}.'
        if request['pre_cm']['admission_grade_above'] == 'true':
            instr = instr_1.format(
                'T') + instr_2.format(request['pre_cm']['ing_per'])
        else:
            instr = instr_1.format(
                'No t', request['pre_cm']['ing_per']) + instr_2.format(request['pre_cm']['ing_per'])
        para.add_run(instr)

    @staticmethod
    def case_TRASLADO_PREGRADO_Analysis_6_2(request, para):
        instr_1 = '\n6.2 En el periodo académico {} el PAPA del estudiante fue de {}. '
        instr_2 = 'El PAPA{}encuentra en la franja del 30% de los mejores promedios del ' + \
            'plan de estudios origen.'
        instr = instr_1.format(
            request['pre_cm']['last_per'], request['pre_cm']['last_papa'])
        if request['pre_cm']['pappa_in_30'] == 'true':
            instr += instr_2.format(' se ')
        else:
            instr += instr_2.format(' no se ')
        para.add_run(instr)

    @staticmethod
    def case_TRASLADO_PREGRADO_Analysis_7(request, para):
        instr_1 = '\n7. {}ay cupos disponibles en el plan de estudios del '
        instr_2 = 'programa curricular solicitado. (Estipulados por consejo de facultad). '
        if request['pre_cm']['places'] == 'true':
            instr = instr_1.format('H') + instr_2
        else:
            instr = instr_1.format('No h') + instr_2
        para.add_run(instr)

    @staticmethod
    def case_TRASLADO_PREGRADO_Analysis_8(request, para):
        instr_1 = '\n8. CCPT {} Créditos pendientes de aprobación en nuevo plan. '
        instr_2 = 'El estudiante{}cuenta con el suficiente cupo de créditos para ' + \
            'inscribir las asignaturas pendientes de aprobación en el nuevo plan (Artículo' + \
            ' 3, Acuerdo 089 de 2014 de Consejo Académico).'
        if request['pre_cm']['credits'] == 'true':
            instr = instr_1.format('>=') + instr_2.format(' ')
        else:
            instr = instr_1.format('<') + instr_2.format(' no ')
        para.add_run(instr)

    @staticmethod
    def case_TRASLADO_PREGRADO_Answers(request, docx):
        para = docx.add_paragraph()
        run = para.add_run('Concepto: ')
        run.font.bold = True
        para = docx.add_paragraph()
        para.paragraph_format.left_indent = Pt(36)
        if request['approval_status'] == 'RC':
            TRASPRE.case_TRASLADO_PREGRADO_Answers_AP(request, docx, para)
            TRASPRE.case_TRASLADO_PREGRADO_Answers_AP_table_general_data(
                request, docx)
            TRASPRE.case_TRASLADO_PREGRADO_Answers_AP_table_academic_info(
                request, docx)
            TRASPRE.case_TRASLADO_PREGRADO_Answers_AP_table_credits_summary(
                request, docx)
        else:
            TRASPRE.case_TRASLADO_PREGRADO_Answers_NA(request, docx, para)

    @staticmethod
    def case_TRASLADO_PREGRADO_Answers_AP(request, docx, para):
        para.add_run('El Comité Asesor recomienda al Consejo de ')
        para.add_run('Facultad APROBAR traslado ')
        para.add_run('{} '.format(request['detail_cm']['tras_type']))
        para.add_run('del programa {} ({}) de la '.format(
            request['detail_cm']['origin'], request['detail_cm']['cod_origin']))
        para.add_run('Universidad Nacional de Colombia – Sede Bogotá, ')
        para.add_run('al programa Ingeniería {} ({}) de la '.format(
            request.get_academic_program_display(), request['academic_program']))
        para.add_run('Universidad Nacional de Colombia – Sede Bogotá, ')
        para.add_run('en el periodo académico {}, '.format(
            request['academic_period']))
        para.add_run('condicionado a conservar la calidad de estudiante ')
        para.add_run('al finalizar el periodo académico {}. '.format(
            request['pre_cm']['last_per']))
        para.add_run('(Artículo 39 del Acuerdo 008 de 2008 del ')
        para.add_run('Consejo Superior Universitario y ')
        para.add_run('Acuerdo 089 de 2014 del Consejo Académico)')

    @staticmethod
    def case_TRASLADO_PREGRADO_Answers_AP_table_general_data(request, docx):
        TRASPRE.case_TRASLADO_PREGRADO_Answers_AP_table_listing(
            docx, '1. Datos Generales')
        general_data = [
            ['Estudiante', request['student_name']],
            ['DNI', request['student_dni']],
            ['Plan de estudios origen (1er plan) – Sede {}'.format(
                request['detail_cm']['campus_origin']), request['detail_cm']['origin']],
            ['Código del plan de estudios origen (1er plan)',
             request['detail_cm']['cod_origin']],
            ['Plan de estudios destino (2° plan) – Sede Bogotá',
             request.get_academic_program_display()],
            ['Código del plan de estudios destino (2° plan)',
             request['academic_program']],
            ['Fecha de la Solicitud a través del SIA', str(request['date'])],
            ['¿Estos planes de estudios conducen al mismo título?',
             request['detail_cm']['same_degree']]
        ]
        table_general_data(general_data, 'TRASLADO', docx)

    @staticmethod
    def case_TRASLADO_PREGRADO_Answers_AP_table_listing(docx, str_list):
        para = docx.add_paragraph()
        bullet = para.add_run(str_list)
        para.paragraph_format.space_after = Pt(0)
        bullet.font.bold = True
        bullet.font.size = Pt(8)

    @staticmethod
    def case_TRASLADO_PREGRADO_Answers_AP_table_academic_info(request, docx):
        TRASPRE.case_TRASLADO_PREGRADO_Answers_AP_table_listing(
            docx, '2. Información Académica')
        table = docx.add_table(rows=4, cols=2, style='Table Grid')
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.columns[0].width = 4600000
        table.columns[1].width = 800000

        for cell in table.columns[0].cells:
            cell.width = 4600000
        for cell in table.columns[1].cells:
            cell.width = 850000

        table.cell(0, 0).paragraphs[0].add_run(
            'Periodo para el cual fue admitido')
        table.cell(0, 1).paragraphs[0].add_run(
            request.detail_cm['info_acad'][0]['per_adm'])
        table.cell(1, 0).paragraphs[0].add_run(
            '¿El solicitante se encuentra matriculado en el semestre de presentar la solicitud?')
        table.cell(1, 1).paragraphs[0].add_run(
            request.detail_cm['info_acad'][0]['matr'])
        table.cell(2, 0).paragraphs[0].add_run(
            '¿El solicitante tuvo calidad de estudiante en el plan de estudios destino (2° plan)?')
        table.cell(2, 1).paragraphs[0].add_run(
            request.detail_cm['info_acad'][0]['ant_plan'])
        table.cell(3, 0).paragraphs[0].add_run(
            'Porcentaje de créditos aprobados en el plan de estudios origen (1er plan)')
        table.cell(3, 1).paragraphs[0].add_run(
            request.detail_cm['info_acad'][0]['aprob']+'%')

        if float(request.detail_cm['info_acad'][0]['aprob']) <= 30:
            table = docx.add_table(rows=2, cols=2, style='Table Grid')
            table.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for cell in table.columns[0].cells:
                cell.width = 4600000
            for cell in table.columns[1].cells:
                cell.width = 850000
            table.cell(0, 0).paragraphs[0].add_run(
                '¿Cuál fue el puntaje de admisión del solicitante?')
            table.cell(0, 1).paragraphs[0].add_run(
                request.detail_cm['info_acad'][1]['pun'])
            table.cell(1, 0).paragraphs[0].add_run(
                'Puntaje de admisión del último admitido regular al plan destino ' +
                '(2° plan) en la misma prueba de ingreso del solicitante*')
            table.cell(1, 1).paragraphs[0].add_run(
                request.detail_cm['info_acad'][1]['pun_ult'])

        para = docx.add_paragraph()
        table = docx.add_table(rows=3, cols=2, style='Table Grid')
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.columns[0].width = 4600000
        table.columns[1].width = 800000
        for cell in table.columns[0].cells:
            cell.width = 4600000
        for cell in table.columns[1].cells:
            cell.width = 850000
        table.cell(0, 0).paragraphs[0].add_run(
            'Cupo de créditos menos créditos pendientes en el plan de estudios origen (1er plan)')
        table.cell(0, 1).paragraphs[0].add_run(
            request.detail_cm['info_acad'][2]['cupo_pend'])
        table.cell(1, 0).paragraphs[0].add_run(
            'Cupo de créditos para traslado (literal d del artículo 3 de la resolución' +
            ' 089 de 2015 de Consejo Académico)')
        table.cell(1, 1).paragraphs[0].add_run(
            request.detail_cm['info_acad'][2]['cred_tras'])
        table.cell(2, 0).paragraphs[0].add_run(
            'El cupo de créditos para traslado es igual o mayor al número de créditos' +
            ' pendientes de aprobación en el plan de estudios destino (2° plan)?')
        if (float(request.detail_cm['info_acad'][2]['cupo_pend']) - float(
                request.detail_cm['info_acad'][2]['cred_tras'])) >= 0:
            table.cell(2, 1).paragraphs[0].add_run('Si')
        else:
            table.cell(2, 1).paragraphs[0].add_run('No')
        para = docx.add_paragraph()
        para.add_run(
            '* en caso que el plan de destino tenga convocatoria anual el puntaje' +
            ' será con la anterior convocatoria.')
        para = docx.add_paragraph()

    @staticmethod
    def case_TRASLADO_PREGRADO_Answers_AP_table_credits_summary(request, docx):
        TRASPRE.case_TRASLADO_PREGRADO_Answers_AP_table_listing(
            docx, '3. Resumen General de Créditos del Segundo Plan de Estudios')
        credits_ = [[int(request.detail_cm['resumen'][0][0]),
                     int(request.detail_cm['resumen'][0][1]),
                     int(request.detail_cm['resumen'][0][2]),
                     int(request.detail_cm['resumen'][0][3]),
                     int(request.detail_cm['resumen'][0][4])],
                    [int(request.detail_cm['resumen'][1][0]),
                     int(request.detail_cm['resumen'][1][1]),
                     int(request.detail_cm['resumen'][1][2]),
                     int(request.detail_cm['resumen'][1][3]),
                     int(request.detail_cm['resumen'][1][4])],
                    [int(request.detail_cm['resumen'][2][0]),
                     int(request.detail_cm['resumen'][2][1]),
                     int(request.detail_cm['resumen'][2][2]),
                     int(request.detail_cm['resumen'][2][3]),
                     int(request.detail_cm['resumen'][2][4])]]
        table_credits_summary(docx, credits_, 'Traslado')

    @staticmethod
    def case_TRASLADO_PREGRADO_Answers_NA(request, docx, para):
        str_na = 'El Comité Asesor recomienda al Consejo de Facultad NO APROBAR traslado {} '
        str_na += 'del programa {} ({}) de la Universidad Nacional de Colombia – Sede Bogotá, '
        str_na += 'al programa Ingeniería {} ({}) de la Universidad Nacional de Colombia – Sede'
        str_na += ' Bogotá, en el periodo académico {}.'
        para.add_run(str_na.format(
            request['detail_cm']['tras_type'],
            request['detail_cm']['origin'],
            request['detail_cm']['cod_origin'],
            request.get_academic_program_display(),
            request['academic_program'],
            request['academic_period']
        ))
