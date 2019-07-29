from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL


class CGRUPXX():

    @staticmethod
    def case_CAMBIO_DE_GRUPO(request, docx, redirected=False):
        para = docx.paragraphs[-1]
        if not redirected:
            para = docx.add_paragraph()
            para.add_run('El Consejo de Facultad ')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if request.approval_status == 'AP':
            para.add_run('APRUEBA ').font.bold = True
        else:
            para.add_run('NO APRUEBA ').font.bold = True
        para.add_run('cambio de grupo de la(s) siguiente(s) asignatura(s) ')
        if request.approval_status == 'AP':
            para.add_run(' debido a que justifica debidamente la solicitud.')
        else:
            para.add_run(', debido a que ' + request.justification + '.')

        table = docx.add_table(
            rows=len(request.detail_cm['subjects'])+1, cols=6)
        
        for column in table.columns:
            for cell in column.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.style = 'Table Grid'
        table.style.font.size = Pt(8)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # table.columns[0].width = 200000
        # table.columns[1].width = 800000
        # table.columns[2].width = 900000
        # table.columns[3].width = 700000
        # table.columns[4].width = 500000
        # table.columns[5].width = 500000

        for cell in table.columns[0].cells:
            cell.width = 2050000
        for cell in table.columns[1].cells:
            cell.width = 800000
        for cell in table.columns[2].cells:
            cell.width = 900000
        for cell in table.columns[3].cells:
            cell.width = 700000
        for cell in table.columns[4].cells:
            cell.width = 500000
        for cell in table.columns[5].cells:
            cell.width = 500000

        table.cell(0, 0).paragraphs[0].add_run('Asignatura').font.bold = True
        table.cell(0, 0).paragraphs[0] = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(0, 1).paragraphs[0].add_run('Código').font.bold = True
        table.cell(0, 1).paragraphs[0] = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(0, 2).paragraphs[0].add_run('Tipología').font.bold = True
        table.cell(0, 2).paragraphs[0] = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(0, 3).paragraphs[0].add_run('Periodo').font.bold = True
        table.cell(0, 3).paragraphs[0] = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(0, 4).paragraphs[0].add_run('G.O.').font.bold = True
        table.cell(0, 4).paragraphs[0] = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(0, 5).paragraphs[0].add_run('G.D.').font.bold = True
        table.cell(0, 5).paragraphs[0] = WD_ALIGN_PARAGRAPH.CENTER
        index = 0
        for subject in request.detail_cm['subjects']:
            table.cell(index+1, 0).paragraphs[0].add_run(subject['subject'])
            table.cell(index+1, 1).paragraphs[0].add_run(subject['cod'])
            table.cell(index+1, 2).paragraphs[0].add_run(subject['tip'])
            table.cell(index+1, 3).paragraphs[0].add_run(subject['per'])
            table.cell(index+1, 4).paragraphs[0].add_run(subject['go'])
            table.cell(index+1, 5).paragraphs[0].add_run(subject['gd'])
            index = index + 1
        para = docx.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run('G.O.: Grupo de Origen, G.D.: Grupo de Destino')
