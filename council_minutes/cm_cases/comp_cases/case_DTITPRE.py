from docx import Document
from ...models import Request
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.table import WD_TABLE_ALIGNMENT
from num2words import num2words  ##pip install num2words
from docx.shared import Pt
from .case_REINPRE import REINPRE
from docx.shared import Cm, Inches
from ...pre_cm_cases.comp_cases.case_utils import *
from datetime import date



class DTITPRE():

    @staticmethod
    def GET_PLAN(cod_plan):
        large_program = ''
        for p in Request.PROGRAM_CHOICES:
            if p[0] == cod_plan:
                large_program = p[1]
                break
        return large_program

    @staticmethod
    def case_DOBLE_TITULACION_PREGRADO(request, docx, redirected=False):
        para = docx.paragraphs[-1]
        if not redirected:
            para = docx.add_paragraph()
            para.add_run('El Consejo de Facultad ')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if request.approval_status == 'AP':
            DTITPRE.case_DOBLE_TITULACION_PREGRADO_AP(request, docx, para)
        else:
            DTITPRE.case_DOBLE_TITULACION_PREGRADO_NA(request, docx, para)

    @staticmethod
    def case_DOBLE_TITULACION_PREGRADO_AP(request, docx, paragraph):
        paragraph.add_run('APRUEBA').font.bold = True
        paragraph.add_run(' recomendar al Consejo de Sede que formalice la admisión y ubicación')
        paragraph.add_run('en el programa de pregrado {} – {}, '.format(request.get_academic_program_display(), request.academic_program))
        paragraph.add_run('teniendo en cuenta que el estudiante cuenta con un cupo de créditos suficiente para culminar')
        paragraph.add_run('el segundo plan de estudios. (Acuerdo 155 de 2014 del Consejo Superior Universitario).\n')
        DTITPRE.case_DOBLE_TITULACION_PREGRADO_TABLES(request, docx, paragraph)

    @staticmethod
    def case_DOBLE_TITULACION_PREGRADO_NA(request, docx, paragraph):
        paragraph.add_run('NO APRUEBA').font.bold = True

    @staticmethod
    def case_DOBLE_TITULACION_PREGRADO_TABLES(request, docx, paragraph):
        paragraph.paragraph_format.space_after = Pt(0)
        bullet = paragraph.add_run('1. Datos Generales')
        bullet.font.bold = True
        bullet.font.size = Pt(8)
        DTITPRE.case_DOBLE_TITULACION_PREGRADO_TABLE_DATOS_PERSONALES(request, docx)
        para = docx.add_paragraph()
        para.paragraph_format.space_before = Pt(8)
        bullet = para.add_run('2. Información Académica:')
        para.paragraph_format.space_after = Pt(0)
        bullet.font.bold = True
        bullet.font.size = Pt(8)
        DTITPRE.case_DOBLE_TITULACION_PREGRADO_TABLE_INFORMACION_ACADEMICA(request, docx)
        para = docx.add_paragraph()
        para.paragraph_format.space_before = Pt(8)
        bullet = para.add_run('3. Cuadro equivalencia y convalidaciones de asignaturas cursadas y aprobadas hasta' 
        + 'la fecha de presentación de la solicitud por parte del estudiante.')
        para.paragraph_format.space_after = Pt(0)
        bullet.font.bold = True
        bullet.font.size = Pt(8)
        DTITPRE.case_DOBLE_TITULACION_PREGRADO_TABLE_EQUIVALENCIAS_CONVALIDACIONES(request, docx)
        para = docx.add_paragraph()
        para.paragraph_format.space_before = Pt(8)
        bullet = para.add_run('4. Asignaturas pendientes por cursar en el segundo plan de estudios.')
        para.paragraph_format.space_after = Pt(0)
        bullet.font.bold = True
        bullet.font.size = Pt(8)
        DTITPRE.case_DOBLE_TITULACION_PREGRADO_TABLE_ASIGNATURAS_PENDIENTES(request, docx)
        para = docx.add_paragraph()
        para.paragraph_format.space_before = Pt(8)
        bullet = para.add_run('La oferta de asignaturas optativas en cada una de las agrupaciones y componentes del plan de' + 
        ' estudios del programa curricular de {}, la encuentra en el {} del año {}, expedido'.format(request.get_academic_program_display(), '<acuerdo del programa>', '<año del acuerdo>') 
        + ' por Consejo de facultad de Ingeniería.')
        para.paragraph_format.space_after = Pt(0)
        bullet.font.size = Pt(8)
        bullet.font.underline = True
        bullet.font.italic = True
        
        para = docx.add_paragraph()
        para.paragraph_format.space_before = Pt(8)
        bullet = para.add_run('5. Resumen general de créditos del segundo plan de estudios:')
        para.paragraph_format.space_after = Pt(0)
        bullet.font.size = Pt(8)
        bullet.font.bold = True
        DTITPRE.case_DOBLE_TITULACION_PREGRADO_TABLE_RESUMEN_GENERAL(request, docx)
        para = docx.add_paragraph()
        para.paragraph_format.space_before = Pt(0)
        bullet = para.add_run('*Sin incluir los créditos correspondientes al cumplimiento del requisito de suficiencia en idioma extranjero (Circular 09 de 2013 de la División de Registro).\n**Aprobados del plan de estudios, sin excedentes.')
        para.paragraph_format.space_after = Pt(0)
        bullet.font.size = Pt(6)
        # table = docx.add_table(rows=1, cols=5, style='Table Grid')
        # table.style.font.size = Pt(8)
        # table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # for cell in table.columns[0].cells:
        #     cell.width = 3000000
        # for cell in table.columns[1].cells:
        #     cell.width = 650000
        # for cell in table.columns[2].cells:
        #     cell.width = 450000
        # for cell in table.columns[3].cells:
        #     cell.width = 650000
        # for cell in table.columns[4].cells:
        #     cell.width = 450000
        # table.cell(0, 0).paragraphs[0].add_run('El Consejo de la Facultad de Ingeniería en sesión del día {} Acta {}'.format(str(request.detail_cm['fecha_sesion'].day) + REINPRE.num_to_month(request.detail_cm['fecha_sesion'].month) + str(request.detail_cm['fecha_sesion'].year), request.detail_cm['acta']))
        # table.cell(0, 1).paragraphs[0].add_run('Recomienda')
        # table.cell(0, 3).paragraphs[0].add_run('No recomienda')
        # if request.detail_cm['recomienda'] == 'Si':
        #     table.cell(0, 2).paragraphs[0].add_run('X')
        #     table.cell(0, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # else:
        #     table.cell(0, 4).paragraphs[0].add_run('X')
        #     table.cell(0, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        fecha_sesion = request.detail_cm['fecha_sesion'].split('-')
        details = [
            "Consejo de la Facultad de Ingeniería",
            fecha_sesion[2] + "-" + fecha_sesion[1] + "-" + fecha_sesion[0],
            request.detail_cm['acta'],
            fecha_sesion[0],
            True
        ]
        table_recommend(docx, details)
    
    @staticmethod
    def case_DOBLE_TITULACION_PREGRADO_TABLE_DATOS_PERSONALES(request, docx):
        
        general_data = [    
            ['Nombre Estudiante', request.student_name], 
            ['DNI', request.student_dni], 
            ['Plan de estudios origen (1er plan) – Sede', get_academic_program(request.detail_cm['datos_generales']['plan_origen'])], 
            ['Código del plan de estudios', request.detail_cm['datos_generales']['plan_origen']], 
            ['Plan de estudios doble titulación (2do plan)', request.get_academic_program_display()], 
            ['Código del plan de estudios doble titulación', request.academic_program], 
            ['Fecha de la Solicitud', str(request.detail_cm['datos_generales']['fecha_solicitud'].day) + REINPRE.num_to_month(request.detail_cm['datos_generales']['fecha_solicitud'].month) + str(request.detail_cm['datos_generales']['fecha_solicitud'].year)]
        ]
        table_general_data(general_data, "DOBLE TITULACIÓN", docx)

    @staticmethod
    def case_DOBLE_TITULACION_PREGRADO_TABLE_INFORMACION_ACADEMICA(request, docx):
        table = docx.add_table(rows=4, cols=5, style='Table Grid')
        table.style.font.size = Pt(8)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER  
        for cell in table.columns[0].cells:
            cell.width = 4000000
        for cell in table.columns[1].cells:
            cell.width = 300000
        for cell in table.columns[2].cells:
            cell.width = 300000
        for cell in table.columns[3].cells:
            cell.width = 300000
        for cell in table.columns[4].cells:
            cell.width = 300000
        table.cell(0, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(0, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(1, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(1, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(0, 0).paragraphs[0].add_run('¿Tuvo calidad de estudiante en el 2do plan?')
        table.cell(0, 1).paragraphs[0].add_run('Si')
        table.cell(0, 3).paragraphs[0].add_run('No')
        if request.detail_cm['informacion_academica']['calidad_estudiante_seg_plan'] == 'Si':
            table.cell(0, 2).paragraphs[0].add_run('X')
        else:
            table.cell(0, 4).paragraphs[0].add_run('X')
        
        table.cell(1, 0).paragraphs[0].add_run('Se encuentra matriculado al momento de presentar la solicitud')
        table.cell(1, 1).paragraphs[0].add_run('Si')
        table.cell(1, 3).paragraphs[0].add_run('No')
        if request.detail_cm['informacion_academica']['matriculado_solicitud'] == 'Si':
            table.cell(1, 2).paragraphs[0].add_run('X')
        else:
            table.cell(1, 4).paragraphs[0].add_run('X')
        
        table.cell(2, 0).paragraphs[0].add_run('PAPA en el primer plan de estudio')
        cellp = table.cell(2, 1).merge(table.cell(2, 4)).paragraphs[0]
        cellp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cellp.add_run(request.detail_cm['informacion_academica']['PAPA_plan_origen'])
        table.cell(3, 0).paragraphs[0].add_run('Créditos estudio doble titulación')
        cellp = table.cell(3, 1).merge(table.cell(3, 4)).paragraphs[0]
        cellp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cellp.add_run(request.detail_cm['informacion_academica']['cred_estudio_doble_titulacion'])
    
    @staticmethod
    def case_DOBLE_TITULACION_PREGRADO_TABLE_EQUIVALENCIAS_CONVALIDACIONES(request, docx):
    #TIPOLOGÍA B    
        table = docx.add_table(rows=len(request.detail_cm['equivalencias_convalidaciones']['T_B'])+2, cols=9, style='Table Grid')
        table.style.font.size = Pt(8)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER  
        for cell in table.columns[0].cells:
            cell.width = 720000
        for cell in table.columns[1].cells:
            cell.width = 900000
        for cell in table.columns[2].cells:
            cell.width = 600000
        for cell in table.columns[3].cells:
            cell.width = 900000
        for cell in table.columns[4].cells:
            cell.width = 300000
        for cell in table.columns[5].cells:
            cell.width = 300000
        for cell in table.columns[6].cells:
            cell.width = 900000
        for cell in table.columns[7].cells:
            cell.width = 200000
        for cell in table.columns[8].cells:
            cell.width = 300000
        for column in table.columns:
            for cell in column.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cellp = table.cell(0, 0).merge(table.cell(0, 1)).paragraphs[0]
        cellp.add_run('PLAN DE ESTUDIOS (1)').font.bold = True
        cellp = table.cell(0, 2).merge(table.cell(0, 8)).paragraphs[0]   
        cellp.add_run('PLAN DE ESTUDIOS (2)\nCOMPONENTE DE FUNDAMENTACIÓN (Tipología B)').font.bold = True
        table.cell(1, 0).paragraphs[0].add_run('Código').font.bold = True
        table.cell(1, 1).paragraphs[0].add_run('Asignatura').font.bold = True
        table.cell(1, 2).paragraphs[0].add_run('Código').font.bold = True
        table.cell(1, 3).paragraphs[0].add_run('Asignatura').font.bold = True
        table.cell(1, 4).paragraphs[0].add_run('Ob*').font.bold = True
        table.cell(1, 5).paragraphs[0].add_run('Op*').font.bold = True
        table.cell(1, 6).paragraphs[0].add_run('Agrupación').font.bold = True
        table.cell(1, 7).paragraphs[0].add_run('C*').font.bold = True
        table.cell(1, 8).paragraphs[0].add_run('Nota').font.bold = True
        row = 2
        column = 0
        for i in range (0, len(request.detail_cm['equivalencias_convalidaciones']['T_B']) - 1):
            table.cell(row, column).paragraphs[0].add_run(request.detail_cm['equivalencias_convalidaciones']['T_B'][i]['codigo_origen'])
            table.cell(row, column+1).paragraphs[0].add_run(request.detail_cm['equivalencias_convalidaciones']['T_B'][i]['asignatura_origen'])
            table.cell(row, column+2).paragraphs[0].add_run(request.detail_cm['equivalencias_convalidaciones']['T_B'][i]['codigo_seg'])
            table.cell(row, column+3).paragraphs[0].add_run(request.detail_cm['equivalencias_convalidaciones']['T_B'][i]['asignatura_seg'])
            if request.detail_cm['equivalencias_convalidaciones']['T_B'][0]['Ob/Op'] == 'Ob':
                table.cell(row, column+4).paragraphs[0].add_run('X')
            else:
                table.cell(row, column+5).paragraphs[0].add_run('X')
            table.cell(row, column+6).paragraphs[0].add_run(request.detail_cm['equivalencias_convalidaciones']['T_B'][i]['agrupacion'])
            table.cell(row, column+7).paragraphs[0].add_run(request.detail_cm['equivalencias_convalidaciones']['T_B'][i]['C'])
            table.cell(row, column+8).paragraphs[0].add_run(request.detail_cm['equivalencias_convalidaciones']['T_B'][i]['nota'])
            row = row + 1
            column = 0
        cellp = table.cell(row, 0).merge(table.cell(row, 6)).paragraphs[0]
        cellp.add_run('Total créditos convalidados/equivalentes en el componente')
        cellp = table.cell(row, 7).merge(table.cell(row, 8)).paragraphs[0]
        cellp.add_run(request.detail_cm['equivalencias_convalidaciones']['T_B'][-1])
        
        #TIPOLOGÍA C
        para = docx.add_paragraph()
        para.paragraph_format.space_before = Pt(0)
        para.add_run(' ').font.size = Pt(8)
        para.paragraph_format.space_after = Pt(0)
        table = docx.add_table(rows=len(request.detail_cm['equivalencias_convalidaciones']['T_C'])+2, cols=9, style='Table Grid')
        table.style.font.size = Pt(8)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER  
        for cell in table.columns[0].cells:
            cell.width = 720000
        for cell in table.columns[1].cells:
            cell.width = 900000
        for cell in table.columns[2].cells:
            cell.width = 600000
        for cell in table.columns[3].cells:
            cell.width = 900000
        for cell in table.columns[4].cells:
            cell.width = 300000
        for cell in table.columns[5].cells:
            cell.width = 300000
        for cell in table.columns[6].cells:
            cell.width = 900000
        for cell in table.columns[7].cells:
            cell.width = 200000
        for cell in table.columns[8].cells:
            cell.width = 300000
        for column in table.columns:
            for cell in column.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        cellp = table.cell(0, 0).merge(table.cell(0, 1)).paragraphs[0]
        cellp.add_run('PLAN DE ESTUDIOS (1)').font.bold = True
        cellp = table.cell(0, 2).merge(table.cell(0, 8)).paragraphs[0]   
        cellp.add_run('PLAN DE ESTUDIOS (2)\nCOMPONENTE DISCIPLINAR/PROFESIONAL (Tipología C)').font.bold = True
        table.cell(1, 0).paragraphs[0].add_run('Código').font.bold = True
        table.cell(1, 1).paragraphs[0].add_run('Asignatura').font.bold = True
        table.cell(1, 2).paragraphs[0].add_run('Código').font.bold = True
        table.cell(1, 3).paragraphs[0].add_run('Asignatura').font.bold = True
        table.cell(1, 4).paragraphs[0].add_run('Ob*').font.bold = True
        table.cell(1, 5).paragraphs[0].add_run('Op*').font.bold = True
        table.cell(1, 6).paragraphs[0].add_run('Agrupación').font.bold = True
        table.cell(1, 7).paragraphs[0].add_run('C*').font.bold = True
        table.cell(1, 8).paragraphs[0].add_run('Nota').font.bold = True
        row = 2
        column = 0
        for i in range (0, len(request.detail_cm['equivalencias_convalidaciones']['T_C']) - 1):
            table.cell(row, column).paragraphs[0].add_run(request.detail_cm['equivalencias_convalidaciones']['T_C'][i]['codigo_origen'])
            table.cell(row, column+1).paragraphs[0].add_run(request.detail_cm['equivalencias_convalidaciones']['T_C'][i]['asignatura_origen'])
            table.cell(row, column+2).paragraphs[0].add_run(request.detail_cm['equivalencias_convalidaciones']['T_C'][i]['codigo_seg'])
            table.cell(row, column+3).paragraphs[0].add_run(request.detail_cm['equivalencias_convalidaciones']['T_C'][i]['asignatura_seg'])
            if request.detail_cm['equivalencias_convalidaciones']['T_C'][0]['Ob/Op'] == 'Ob':
                table.cell(row, column+4).paragraphs[0].add_run('X')
            else:
                table.cell(row, column+5).paragraphs[0].add_run('X')
            table.cell(row, column+6).paragraphs[0].add_run(request.detail_cm['equivalencias_convalidaciones']['T_C'][i]['agrupacion'])
            table.cell(row, column+7).paragraphs[0].add_run(request.detail_cm['equivalencias_convalidaciones']['T_C'][i]['C'])
            table.cell(row, column+8).paragraphs[0].add_run(request.detail_cm['equivalencias_convalidaciones']['T_C'][i]['nota'])
            row = row + 1
            column = 0
        cellp = table.cell(row, 0).merge(table.cell(row, 6)).paragraphs[0]
        cellp.add_run('Total créditos convalidados/equivalentes en el componente')
        cellp = table.cell(row, 7).merge(table.cell(row, 8)).paragraphs[0]
        cellp.add_run(request.detail_cm['equivalencias_convalidaciones']['T_C'][-1])

        #TIPOLOGÍA L 
        para = docx.add_paragraph()
        para.paragraph_format.space_before = Pt(0)
        para.add_run(' ').font.size = Pt(8)
        para.paragraph_format.space_after = Pt(0)
        table = docx.add_table(rows=len(request.detail_cm['equivalencias_convalidaciones']['T_L'])+2, cols=6, style='Table Grid')
        table.style.font.size = Pt(8)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER  
        for cell in table.columns[0].cells:
            cell.width = 900000
        for cell in table.columns[1].cells:
            cell.width = 1400000
        for cell in table.columns[2].cells:
            cell.width = 700000
        for cell in table.columns[3].cells:
            cell.width = 1400000
        for cell in table.columns[4].cells:
            cell.width = 400000
        for cell in table.columns[5].cells:
            cell.width = 400000
        for column in table.columns:
            for cell in column.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        cellp = table.cell(0, 0).merge(table.cell(0, 1)).paragraphs[0]
        cellp.add_run('PLAN DE ESTUDIOS (1)').font.bold = True
        cellp = table.cell(0, 2).merge(table.cell(0, 5)).paragraphs[0]   
        cellp.add_run('PLAN DE ESTUDIOS (2)\nCOMPONENTE DE LIBRE ELECCIÓN (Tipología L)').font.bold = True
        table.cell(1, 0).paragraphs[0].add_run('Código').font.bold = True
        table.cell(1, 1).paragraphs[0].add_run('Asignatura').font.bold = True
        table.cell(1, 2).paragraphs[0].add_run('Código').font.bold = True
        table.cell(1, 3).paragraphs[0].add_run('Asignatura').font.bold = True
        table.cell(1, 4).paragraphs[0].add_run('C*').font.bold = True
        table.cell(1, 5).paragraphs[0].add_run('Nota').font.bold = True
        row = 2
        column = 0
        for i in range (0, len(request.detail_cm['equivalencias_convalidaciones']['T_L']) - 1):
            table.cell(row, column).paragraphs[0].add_run(request.detail_cm['equivalencias_convalidaciones']['T_L'][i]['codigo_origen'])
            table.cell(row, column+1).paragraphs[0].add_run(request.detail_cm['equivalencias_convalidaciones']['T_L'][i]['asignatura_origen'])
            table.cell(row, column+2).paragraphs[0].add_run(request.detail_cm['equivalencias_convalidaciones']['T_L'][i]['codigo_seg'])
            table.cell(row, column+3).paragraphs[0].add_run(request.detail_cm['equivalencias_convalidaciones']['T_L'][i]['asignatura_seg'])
            table.cell(row, column+4).paragraphs[0].add_run(request.detail_cm['equivalencias_convalidaciones']['T_L'][i]['C'])
            table.cell(row, column+5).paragraphs[0].add_run(request.detail_cm['equivalencias_convalidaciones']['T_L'][i]['nota'])
            row = row + 1
            column = 0
        cellp = table.cell(row, 0).merge(table.cell(row, 3)).paragraphs[0]
        cellp.add_run('Total créditos convalidados/equivalentes en el componente')
        cellp = table.cell(row, 4).merge(table.cell(row, 5)).paragraphs[0]
        cellp.add_run(request.detail_cm['equivalencias_convalidaciones']['T_L'][-1])

    @staticmethod
    def case_DOBLE_TITULACION_PREGRADO_TABLE_ASIGNATURAS_PENDIENTES(request, docx):
        total_pendientes = 0
        cant_rows = 0
        for i in range (0, len(request.detail_cm['asignaturas_pendientes']['T_B']['obligatorias']['categorias'])):
            for m in request.detail_cm['asignaturas_pendientes']['T_B']['obligatorias']['categorias'][i]['materias']:
                cant_rows = cant_rows + 1

        table = docx.add_table(rows=cant_rows + 3, cols=5, style='Table Grid')
        table.style.font.size = Pt(8)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER 

        for cell in table.columns[0].cells:
            cell.width = 970000
        for cell in table.columns[1].cells:
            cell.width = 900000
        for cell in table.columns[2].cells:
            cell.width = 1300000
        for cell in table.columns[3].cells:
            cell.width = 1000000
        for cell in table.columns[4].cells:
            cell.width = 1000000

        for column in table.columns:
            for cell in column.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER   
        cellp = table.cell(0, 0).merge(table.cell(0, 4)).paragraphs[0]
        cellp.add_run('Componente de Fundamentación (B)').font.bold = True
        cellp = table.cell(1, 0).merge(table.cell(1, 4)).paragraphs[0]

        #OBLIGATORIAS T_B
        cellp.add_run('Obligatorias').font.bold = True
        table.cell(2, 0).paragraphs[0].add_run('Agrupación').font.bold = True
        table.cell(2, 1).paragraphs[0].add_run('Código').font.bold = True
        table.cell(2, 2).paragraphs[0].add_run('Asignatura').font.bold = True
        table.cell(2, 3).paragraphs[0].add_run('Créditos asignatura').font.bold = True
        table.cell(2, 4).paragraphs[0].add_run('Créditos pendientes por cursar por el estudiante').font.bold = True
        row_a = 3
        row_m = 3
        for i in range (0, len(request.detail_cm['asignaturas_pendientes']['T_B']['obligatorias']['categorias'])):
            cellp = table.cell(row_a, 0).merge(table.cell(row_a + len(request.detail_cm['asignaturas_pendientes']['T_B']['obligatorias']['categorias'][i]['materias']) - 1,0)).paragraphs[0]
            cellp.add_run(request.detail_cm['asignaturas_pendientes']['T_B']['obligatorias']['categorias'][i]['agrupacion'])
            cellp = table.cell(row_a, 4).merge(table.cell(row_a + len(request.detail_cm['asignaturas_pendientes']['T_B']['obligatorias']['categorias'][i]['materias']) - 1,4)).paragraphs[0]
            cellp.add_run(request.detail_cm['asignaturas_pendientes']['T_B']['obligatorias']['categorias'][i]['creditos_pendientes'])
            total_pendientes = total_pendientes + int(request.detail_cm['asignaturas_pendientes']['T_B']['obligatorias']['categorias'][i]['creditos_pendientes'])
            for j in range (0, len(request.detail_cm['asignaturas_pendientes']['T_B']['obligatorias']['categorias'][i]['materias'])):
                table.cell(row_m, 1).paragraphs[0].add_run(request.detail_cm['asignaturas_pendientes']['T_B']['obligatorias']['categorias'][i]['materias'][j]['codigo'])
                table.cell(row_m, 2).paragraphs[0].add_run(request.detail_cm['asignaturas_pendientes']['T_B']['obligatorias']['categorias'][i]['materias'][j]['asignatura'])
                table.cell(row_m, 3).paragraphs[0].add_run(request.detail_cm['asignaturas_pendientes']['T_B']['obligatorias']['categorias'][i]['materias'][j]['creditos'])
                row_m = row_m + 1
            row_a = row_m
        
        #OPTATIVAS T_B
        cant_rows = 0
        for i in request.detail_cm['asignaturas_pendientes']['T_B']['optativas']:
            cant_rows = cant_rows + 1
        table = docx.add_table(rows=cant_rows + 3, cols=3, style='Table Grid')
        table.style.font.size = Pt(8)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER  
        
        for cell in table.columns[0].cells:
            cell.width = 1900000
        for cell in table.columns[1].cells:
            cell.width = 1630000
        for cell in table.columns[2].cells:
            cell.width = 1630000


        for column in table.columns:
            for cell in column.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER   
        
        cellp = table.cell(0, 0).merge(table.cell(0, 2)).paragraphs[0] 
        cellp.add_run('Optativas').font.bold = True
        table.cell(1, 0).paragraphs[0].add_run('Nombre de la Agrupación').font.bold = True
        table.cell(1, 1).paragraphs[0].add_run('Créditos Requeridos').font.bold = True
        table.cell(1, 2).paragraphs[0].add_run('Créditos pendientes por cursar por el estudiante').font.bold = True
        row_m = 2
        for i in range (0, len(request.detail_cm['asignaturas_pendientes']['T_B']['optativas'])):
            table.cell(row_m, 0).paragraphs[0].add_run(request.detail_cm['asignaturas_pendientes']['T_B']['optativas'][j]['agrupacion'])
            table.cell(row_m, 1).paragraphs[0].add_run(request.detail_cm['asignaturas_pendientes']['T_B']['optativas'][j]['cred_requeridos'])
            table.cell(row_m, 2).paragraphs[0].add_run(request.detail_cm['asignaturas_pendientes']['T_B']['optativas'][j]['cred_pendientes'])
            total_pendientes =  total_pendientes + int(request.detail_cm['asignaturas_pendientes']['T_B']['optativas'][j]['cred_pendientes'])
            row_m = row_m + 1
        
        
        cellp = table.cell(row_m, 0).merge(table.cell(row_m, 1)).paragraphs[0]
        cellp.add_run('Total créditos pendientes').font.bold = True
        table.cell(row_m, 2).paragraphs[0].add_run(str(total_pendientes))

        cant_rows = 0
        for i in range (0, len(request.detail_cm['asignaturas_pendientes']['T_B']['obligatorias']['categorias'])):
            for m in request.detail_cm['asignaturas_pendientes']['T_B']['obligatorias']['categorias'][i]['materias']:
                cant_rows = cant_rows + 1

        #T_C
        total_pendientes = 0
        para = docx.add_paragraph()
        para.paragraph_format.space_before = Pt(0)
        para.add_run(' ').font.size = Pt(8)
        para.paragraph_format.space_after = Pt(0)
        table = docx.add_table(rows=cant_rows + 3, cols=5, style='Table Grid')
        table.style.font.size = Pt(8)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER  

        for cell in table.columns[0].cells:
            cell.width = 1000000
        for cell in table.columns[1].cells:
            cell.width = 900000
        for cell in table.columns[2].cells:
            cell.width = 1300000
        for cell in table.columns[3].cells:
            cell.width = 1000000
        for cell in table.columns[4].cells:
            cell.width = 1000000

        for column in table.columns:
            for cell in column.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER   
        cellp = table.cell(0, 0).merge(table.cell(0, 4)).paragraphs[0]
        cellp.add_run('Componente  Disciplinar/Profesional  (C)').font.bold = True
        cellp = table.cell(1, 0).merge(table.cell(1, 4)).paragraphs[0]

        #OBLIGATORIAS T_C
        cellp.add_run('Obligatorias').font.bold = True
        table.cell(2, 0).paragraphs[0].add_run('Agrupación').font.bold = True
        table.cell(2, 1).paragraphs[0].add_run('Código').font.bold = True
        table.cell(2, 2).paragraphs[0].add_run('Asignatura').font.bold = True
        table.cell(2, 3).paragraphs[0].add_run('Créditos asignatura').font.bold = True
        table.cell(2, 4).paragraphs[0].add_run('Créditos pendientes por cursar por el estudiante').font.bold = True
        row_a = 3
        row_m = 3

        for i in range (0, len(request.detail_cm['asignaturas_pendientes']['T_C']['obligatorias']['categorias'])):
            cellp = table.cell(row_a, 0).merge(table.cell(row_a + len(request.detail_cm['asignaturas_pendientes']['T_C']['obligatorias']['categorias'][i]['materias']) - 1,0)).paragraphs[0]
            cellp.add_run(request.detail_cm['asignaturas_pendientes']['T_C']['obligatorias']['categorias'][i]['agrupacion'])
            cellp = table.cell(row_a, 4).merge(table.cell(row_a + len(request.detail_cm['asignaturas_pendientes']['T_C']['obligatorias']['categorias'][i]['materias']) - 1,4)).paragraphs[0]
            cellp.add_run(request.detail_cm['asignaturas_pendientes']['T_C']['obligatorias']['categorias'][i]['creditos_pendientes'])
            total_pendientes = total_pendientes + int(request.detail_cm['asignaturas_pendientes']['T_C']['obligatorias']['categorias'][i]['creditos_pendientes'])
            for j in range (0, len(request.detail_cm['asignaturas_pendientes']['T_C']['obligatorias']['categorias'][i]['materias'])):
                table.cell(row_m, 1).paragraphs[0].add_run(request.detail_cm['asignaturas_pendientes']['T_C']['obligatorias']['categorias'][i]['materias'][j]['codigo'])
                table.cell(row_m, 2).paragraphs[0].add_run(request.detail_cm['asignaturas_pendientes']['T_C']['obligatorias']['categorias'][i]['materias'][j]['asignatura'])
                table.cell(row_m, 3).paragraphs[0].add_run(request.detail_cm['asignaturas_pendientes']['T_C']['obligatorias']['categorias'][i]['materias'][j]['creditos'])
                row_m = row_m + 1
            row_a = row_m

        
        #OPTATIVAS T_C
        cant_rows = 0
        for i in request.detail_cm['asignaturas_pendientes']['T_C']['optativas']:
            cant_rows = cant_rows + 1

        table = docx.add_table(rows=cant_rows + 3, cols=3, style='Table Grid')
        table.style.font.size = Pt(8)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER  

        for cell in table.columns[0].cells:
            cell.width = 1900000
        for cell in table.columns[1].cells:
            cell.width = 1650000
        for cell in table.columns[2].cells:
            cell.width = 1650000


        for column in table.columns:
            for cell in column.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER   
        
        cellp = table.cell(0, 0).merge(table.cell(0, 2)).paragraphs[0] 
        cellp.add_run('Optativas').font.bold = True
        table.cell(1, 0).paragraphs[0].add_run('Nombre de la Agrupación').font.bold = True
        table.cell(1, 1).paragraphs[0].add_run('Créditos Requeridos').font.bold = True
        table.cell(1, 2).paragraphs[0].add_run('Créditos pendientes por cursar por el estudiante').font.bold = True
        row_m = 2
        for i in range (0, len(request.detail_cm['asignaturas_pendientes']['T_C']['optativas'])):
            table.cell(row_m, 0).paragraphs[0].add_run(request.detail_cm['asignaturas_pendientes']['T_C']['optativas'][j]['agrupacion'])
            table.cell(row_m, 1).paragraphs[0].add_run(request.detail_cm['asignaturas_pendientes']['T_C']['optativas'][j]['cred_requeridos'])
            table.cell(row_m, 2).paragraphs[0].add_run(request.detail_cm['asignaturas_pendientes']['T_C']['optativas'][j]['cred_pendientes'])
            total_pendientes =  total_pendientes + int(request.detail_cm['asignaturas_pendientes']['T_C']['optativas'][j]['cred_pendientes'])
            row_m = row_m + 1
        
        cellp = table.cell(row_m, 0).merge(table.cell(row_m, 1)).paragraphs[0]
        cellp.add_run('Total créditos pendientes').font.bold = True
        table.cell(row_m, 2).paragraphs[0].add_run(str(total_pendientes))

        #ELECTIVAS
        para = docx.add_paragraph()
        para.paragraph_format.space_before = Pt(0)
        para.add_run(' ').font.size = Pt(8)
        para.paragraph_format.space_after = Pt(0)
        table = docx.add_table(rows=1, cols=2, style='Table Grid')
        table.style.font.size = Pt(8)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER  
        for cell in table.columns[0].cells:
            cell.width = 4000000
        for cell in table.columns[1].cells:
            cell.width = 1200000
        table.cell(0, 0).paragraphs[0].add_run('Componente de Libre Elección (L) (Créditos pendientes)').font.bold = True
        table.cell(0, 1).paragraphs[0].add_run(request.detail_cm['asignaturas_pendientes']['electivas']).font.bold = True

    @staticmethod
    def case_DOBLE_TITULACION_PREGRADO_TABLE_RESUMEN_GENERAL(request, docx):
        summary_credits = [
            [
                int(request.detail_cm['resumen_general']['exigidos']['B']['ob']),
                int(request.detail_cm['resumen_general']['exigidos']['B']['op']),
                int(request.detail_cm['resumen_general']['exigidos']['C']['ob']),
                int(request.detail_cm['resumen_general']['exigidos']['C']['op']),
                int(request.detail_cm['resumen_general']['exigidos']['L'])
            ],
            [
                int(request.detail_cm['resumen_general']['equiv_conv']['B']['ob']),
                int(request.detail_cm['resumen_general']['equiv_conv']['B']['op']),
                int(request.detail_cm['resumen_general']['equiv_conv']['C']['ob']),
                int(request.detail_cm['resumen_general']['equiv_conv']['C']['op']),
                int(request.detail_cm['resumen_general']['equiv_conv']['L'])
            ],
            [
                int(request.detail_cm['resumen_general']['pendientes']['B']['ob']),
                int(request.detail_cm['resumen_general']['pendientes']['B']['op']),
                int(request.detail_cm['resumen_general']['pendientes']['C']['ob']),
                int(request.detail_cm['resumen_general']['pendientes']['C']['op']),
                int(request.detail_cm['resumen_general']['pendientes']['L'])
            ]
        ]

        table_credits_summary(docx, summary_credits, "DOBLE TITULACIÓN")
        
