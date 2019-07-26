from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from ...models import Request


class TRASPOS():

    @staticmethod
    def case_TRASLADO_POSGRADO(request, docx, redirected=False):
        if redirected:
            para = docx.paragraphs[-1]
        else:
            para = docx.add_paragraph()
            para.add_run('El Consejo de Facultad ')
        large_program = ''
        for p in Request.PROGRAM_CHOICES:
            if p[0] == request.academic_program:
                large_program = p[1]
                break
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if request.approval_status == 'AP':
            para.add_run('APRUEBA:').font.bold = True
            para = docx.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.add_run('1. T')
        else:
            para.add_run('NO APRUEBA').font.bold = True
            para.add_run(' t')
        para.add_run('raslado ' + request.detail_cm['tras_type'] + ' del programa de ' +
                     request.detail_cm['origin'] + ' del plan de estudios de ')
        para.add_run(request.detail_cm['or_inv_prof'] +
                     ' de la Universidad Nacional de Colombia - Sede ' + request.detail_cm['campus_origin'])
        para.add_run(', al programa ' + large_program + ' del plan de estudios de ' +
                     request.detail_cm['des_inv_prof'] + ' de la Universidad Nacional de Colombia - Sede Bogotá, debido a que ' + request.justification + '.')
        if request.approval_status == 'AP':
            para = docx.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.add_run('2. Equivaler en el programa ' + large_program +
                         ', plan de estudios de ' + request.detail_cm['des_inv_prof'] + ', ')
            para.add_run('en el periodo académico ' +
                         request.detail_cm['period'] + ', las asignaturas cursadas en el programa ' + request.detail_cm['origin'] + '.')
            table = docx.add_table(
                rows=len(request.detail_cm['equivalencia'])+2, cols=6)
            table.style = 'Table Grid'
            table.style.font.size = Pt(8)
            table.alignment = WD_ALIGN_PARAGRAPH.CENTER
            table.columns[0].width = 1800000
            table.columns[1].width = 600000
            table.columns[2].width = 1800000
            table.columns[3].width = 600000
            table.columns[4].width = 200000
            table.columns[5].width = 200000
            cellp = table.cell(0, 1).merge(table.cell(0, 5)).paragraphs[0]
            cellp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cellp.add_run('ASIGNATURA POR LA QUE SE EQUIVALE').font.bold = True
            table.cell(0, 0).paragraphs[0].add_run(
                'ASIGNATURA QUE SE EQUIVALE').font.bold = True
            table.cell(1, 0).paragraphs[0].add_run('NOMBRE').font.bold = True
            table.cell(1, 1).paragraphs[0].add_run('CÓDIGO').font.bold = True
            table.cell(1, 2).paragraphs[0].add_run('NOMBRE').font.bold = True
            table.cell(1, 3).paragraphs[0].add_run('NOTA').font.bold = True
            table.cell(1, 4).paragraphs[0].add_run('C').font.bold = True
            table.cell(1, 5).paragraphs[0].add_run('T').font.bold = True
            index = 0
            for subject in request.detail_cm['equivalencia']:
                table.cell(index+2, 0).paragraphs[0].add_run(subject['nom'])
                table.cell(index+2, 1).paragraphs[0].add_run(subject['cod'])
                table.cell(index+2, 2).paragraphs[0].add_run(subject['nom1'])
                table.cell(index+2, 3).paragraphs[0].add_run(subject['nota'])
                table.cell(index+2, 4).paragraphs[0].add_run(subject['c'])
                table.cell(index+2, 5).paragraphs[0].add_run(subject['t'])
                index = index+1
