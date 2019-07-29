from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Pt


class HOIDPXX():

    @staticmethod
    def case_HOMOLOGACION_INGLES(request, docx, redirected=False):
        para = docx.paragraphs[-1]
        if not redirected:
            para = docx.add_paragraph()
            para.add_run('El Consejo de Facultad ')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if request.approval_status == 'AP':
            HOIDPXX.case_HOMOLOGACION_INGLES_AP(request, docx, para)
        else:
            HOIDPXX.case_HOMOLOGACION_INGLES_NP(request, docx, para)

    @staticmethod
    def case_HOMOLOGACION_INGLES_AP(request, docx, paragraph):
        paragraph.add_run('APRUEBA').font.bold = True
        paragraph.add_run(
            ' homologar en el periodo académico ' + request.academic_period)
        paragraph.add_run(
            ', el requisito de idioma inglés por obtener una calificación de ')
        paragraph.add_run(
            request.detail_cm['grade_got'] + ' en el examen ' + request.detail_cm['institution'])
        paragraph.add_run(
            ', siendo ' + request.detail_cm['min_grade']+' el mínimo exigido.')
        paragraph.add_run(
            ' (Acuerdo 102 de 2013 del Consejo Superior Universitario).')
        HOIDPXX.case_HOMOLOGACION_INGLES_TABLE(request, docx)

    @staticmethod
    def case_HOMOLOGACION_INGLES_NP(request, docx, paragraph):
        paragraph.add_run('NO APRUEBA').font.bold = True
        paragraph.add_run(
            ' homologar en el periodo académico ' + request.academic_period)
        paragraph.add_run(', el requisito de idioma inglés debido a que {}'.format(
            request.justification))
        paragraph.add_run(
            ' (Acuerdo 102 de 2013 del Consejo Superior Universitario).')

    @staticmethod
    def case_HOMOLOGACION_INGLES_TABLE(request, docx):
        table = docx.add_table(
            rows=len(request.detail_cm['subjects'])+5, cols=7)
        for column in table.columns:
            for cell in column.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.style = 'Table Grid'
        table.style.font.size = Pt(8)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.columns[0].width = 600000
        table.columns[1].width = 1800000
        table.columns[2].width = 300000
        table.columns[3].width = 300000
        table.columns[4].width = 400000
        table.columns[5].width = 1400000
        table.columns[6].width = 400000

        for cell in table.columns[0].cells:
            cell.width = 850000
        for cell in table.columns[1].cells:
            cell.width = 1800000
        for cell in table.columns[2].cells:
            cell.width = 300000
        for cell in table.columns[3].cells:
            cell.width = 300000
        for cell in table.columns[4].cells:
            cell.width = 400000
        for cell in table.columns[5].cells:
            cell.width = 1400000
        for cell in table.columns[6].cells:
            cell.width = 400000

        cell = table.cell(0, 0).merge(table.cell(0, 6)).paragraphs[0]
        cell.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        cell.add_run(request.student_name + '\t\tDNI. ' +
                     request.student_dni).font.bold = True
        cell = table.cell(1, 0).merge(table.cell(1, 4)).paragraphs[0]
        str_prog = 'Asignaturas a homologar en el plan de estudios {} ({})'
        cell.add_run(str_prog.format(request.get_academic_program_display(
        ), request.academic_program)).font.bold = True

        cellp = table.cell(1, 5).merge(table.cell(2, 5)).paragraphs[0]
        cellp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(1, 5).merge(table.cell(2, 5)).paragraphs[0].add_run(
            'Examen de inglés presentado').font.bold = True
        table.cell(1, 5).merge(table.cell(2, 5)
                               ).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        cellp = table.cell(1, 6).merge(table.cell(2, 6)).paragraphs[0]
        cellp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(1, 6).merge(table.cell(2, 6)).paragraphs[0].add_run(
            'Nota').font.bold = True
        table.cell(1, 6).merge(table.cell(2, 6)
                               ).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        cell = table.cell(3, 5).merge(table.cell(
            len(request.detail_cm['subjects'])+2, 5)).paragraphs[0]
        cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.add_run(request.detail_cm['institution'])
        table.cell(3, 5).merge(table.cell(len(
            request.detail_cm['subjects'])+2, 5)).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        cell = table.cell(3, 6).merge(table.cell(
            len(request.detail_cm['subjects'])+2, 6)).paragraphs[0]
        cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.add_run(request.detail_cm['grade_got'])
        table.cell(3, 6).merge(table.cell(len(
            request.detail_cm['subjects'])+2, 6)).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        table.cell(2, 0).paragraphs[0].add_run('Código').font.bold = True
        table.cell(2, 1).paragraphs[0].add_run('Asignatura').font.bold = True
        table.cell(2, 2).paragraphs[0].add_run('C').font.bold = True
        table.cell(2, 3).paragraphs[0].add_run('T').font.bold = True
        table.cell(2, 4).paragraphs[0].add_run('Nota').font.bold = True
        index = 0
        credits_sum = 0
        for subject in request.detail_cm['subjects']:
            credits_sum = credits_sum + int(subject['credits'])
            table.cell(index+3, 0).paragraphs[0].add_run(subject['code'])
            table.cell(index+3, 1).paragraphs[0].add_run(subject['subject'])
            table.cell(index+3, 2).paragraphs[0].add_run(subject['credits'])
            table.cell(index+3, 3).paragraphs[0].add_run(subject['tipology'])
            table.cell(index+3, 4).paragraphs[0].add_run(subject['grade'])
            index = index + 1
        cellp = table.cell(index+3, 0).merge(table.cell(index+3, 3)).paragraphs[0]
        cellp.add_run('Créditos homologados P')
        cellp = table.cell(index+3, 4).merge(table.cell(index+3, 6)).paragraphs[0]
        cellp.add_run(str(credits_sum))
        cellp = table.cell(index+4, 0).merge(table.cell(index+4, 3)).paragraphs[0]
        cellp.add_run('Total créditos que se homologan')
        cellp = table.cell(index+4, 4).merge(table.cell(index+4, 6)).paragraphs[0]
        cellp.add_run(str(credits_sum))
          
