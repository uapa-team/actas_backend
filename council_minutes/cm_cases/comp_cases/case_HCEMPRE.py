from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Pt


class HCEMPRE():

    @staticmethod
    def case_HOMOLOGACION_CONVALIDACION_EQUIVALENCIA_PREGRADO(request, docx, redirected=False):
        if redirected:
            para = docx.paragraphs[-1]
        else:
            para = docx.add_paragraph()
            para.add_run('El Consejo de Facultad ')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if request.approval_status == 'AP':
            para.add_run('APRUEBA:').font.bold = True
            para.paragraph_format.space_after = Pt(0)
            HCEMPRE.case_HOMOLOGACION_CONVALIDACION_EQUIVALENCIA_PREGRADO_AP(
                request, docx, para)
        else:
            para.add_run('NO APRUEBA:').font.bold = True
            para.paragraph_format.space_after = Pt(0)
            HCEMPRE.case_HOMOLOGACION_CONVALIDACION_EQUIVALENCIA_PREGRADO_NA(
                request, docx, para)

    @staticmethod
    def case_HOMOLOGACION_CONVALIDACION_EQUIVALENCIA_PREGRADO_AP(request, docx, paragraph):
        if 'homologation' in request.detail_cm:
            Item1 = 'Homologar en el periodo académico {}, la(s) siguiente(s) asignatura(s) cursada(s)'
            Item2 = ' en el programa {} de la {}, de la siguiente manera (Artículo 35 del Acuerdo 008 de 2008 del Consejo Superior Universitario)'
            Item = Item1 + Item2
            para = docx.add_paragraph(Item.format(
                request.academic_period, request.detail_cm['homologation'][-1], request.detail_cm['homologation'][-2]), style='List Number')
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.space_after = Pt(0)
            HCEMPRE.case_HOMOLOGACION_CONVALIDACION_EQUIVALENCIA_PREGRADO_TABLE(
                request, docx, 'homologation')
        if 'equivalence' in request.detail_cm:
            Item1 = 'Equivaler en el periodo académico {}, la(s) siguiente(s) asignatura(s) cursada(s)'
            Item2 = ' en el programa {}, de la siguiente manera (Artículo 35 del Acuerdo 008 de 2008 del Consejo Superior Universitario)'
            Item = Item1 + Item2
            para = docx.add_paragraph(Item.format(
                request.academic_period, request.detail_cm['equivalence'][-1], request.detail_cm['equivalence'][-2]), style='List Number')
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.space_after = Pt(0)
            HCEMPRE.case_HOMOLOGACION_CONVALIDACION_EQUIVALENCIA_PREGRADO_TABLE(
                request, docx, 'equivalence')
        if 'recognition' in request.detail_cm:
            Item1 = 'Convalidar en el periodo académico {}, la(s) siguiente(s) asignatura(s) cursada(s)'
            Item2 = ' en el programa {} de la {}, de la siguiente manera'
            Item = Item1 + Item2
            para = docx.add_paragraph(Item.format(
                request.academic_period, request.detail_cm['recognition'][-1], request.detail_cm['recognition'][-2]), style='List Number')
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.space_after = Pt(0)
            HCEMPRE.case_HOMOLOGACION_CONVALIDACION_EQUIVALENCIA_PREGRADO_TABLE(
                request, docx, 'recognition')

    @staticmethod
    def case_HOMOLOGACION_CONVALIDACION_EQUIVALENCIA_PREGRADO_NA(request, docx, paragraph):
        if 'homologation' in request.detail_cm:
            Item1 = 'Homologar en el periodo académico {}, la(s) siguiente(s) asignatura(s)'
            Item2 = ' cursada(s) en el programa {} de la {}, debido a que {}. (Artículo 35 del Acuerdo 008 del 2008 del Consejo Superior Universitario)'
            Item = Item1 + Item2
            para = docx.add_paragraph(Item.format(
                request.academic_period, request.detail_cm['homologation'][-1], request.detail_cm['homologation'][-2], request.justification), style='List Number')
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.space_after = Pt(0)
            HCEMPRE.case_HOMOLOGACION_CONVALIDACION_EQUIVALENCIA_PREGRADO_TABLE_NA(
                request, docx, 'homologation')
        if 'equivalence' in request.detail_cm:
            Item1 = 'Equivaler en el periodo académico {}, la(s) siguiente(s) asignatura(s)'
            Item2 = ' cursada(s) en el programa {} de la {}, debido a que {}. (Artículo 35 del Acuerdo 008 del 2008 del Consejo Superior Universitario)'
            Item = Item1 + Item2
            para = docx.add_paragraph(Item.format(
                request.academic_period, request.detail_cm['equivalence'][-1], request.detail_cm['equivalence'][-2], request.justification), style='List Number')
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.space_after = Pt(0)
            HCEMPRE.case_HOMOLOGACION_CONVALIDACION_EQUIVALENCIA_PREGRADO_TABLE_NA(
                request, docx, 'equivalence')
        if 'recognition' in request.detail_cm:
            Item1 = 'Convalidar en el periodo académico {}, la(s) siguiente(s) asignatura(s)'
            Item2 = ' cursada(s) en el programa {} de la {}, debido a que {}.'
            Item = Item1 + Item2
            para = docx.add_paragraph(Item.format(
                request.academic_period, request.detail_cm['recognition'][-1], request.detail_cm['recognition'][-2], request.justification), style='List Number')
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.space_after = Pt(0)
            HCEMPRE.case_HOMOLOGACION_CONVALIDACION_EQUIVALENCIA_PREGRADO_TABLE_NA(
                request, docx, 'recognition')

    @staticmethod
    def case_HOMOLOGACION_CONVALIDACION_EQUIVALENCIA_PREGRADO_TABLE(request, docx, case):
        case_d = {'homologation': 'homologar',
                  'equivalence': 'equivaler', 'recognition': 'convalidar'}
        num_credits = {'L': 0, 'B': 0, 'C': 0, 'total': 0}
        cant_rows = 0
        for i in range(0, len(request.detail_cm[case]) - 2):
            if len(request.detail_cm[case][i]['origin_subject']) > len(request.detail_cm[case][i]['subject_out']):
                cant_rows += len(request.detail_cm[case][i]['origin_subject'])
            else:
                cant_rows += len(request.detail_cm[case][i]['subject_out'])
        table = docx.add_table(rows=cant_rows+3, cols=7, style='Table Grid')
        table.style.font.size = Pt(8)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cellp = table.cell(0, 0).merge(table.cell(0, 6)).paragraphs[0]
        cellp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cellp.add_run('{}\t\t\tDNI.{}'.format(
            request.student_name, request.student_dni)).font.bold = True
        cellp = table.cell(1, 0).merge(table.cell(1, 4)).paragraphs[0]
        cellp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cellp.add_run('Asignaturas a {} en el plan de estudios de {}({})'.format(
            case_d[case], request.get_academic_program_display(), request.academic_program)).font.bold = True
        cellp = table.cell(1, 5).merge(table.cell(1, 6)).paragraphs[0]
        cellp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cellp.add_run('Asignaturas cursadas en el plan de estudios {} de la {}'.format(
            request.detail_cm[case][-1], request.detail_cm[case][-2])).font.bold = True
        table.cell(2, 0).paragraphs[0].add_run('Código').font.bold = True
        table.cell(2, 1).paragraphs[0].add_run('Asignatura').font.bold = True
        table.cell(2, 2).paragraphs[0].add_run('C').font.bold = True
        table.cell(2, 3).paragraphs[0].add_run('T').font.bold = True
        table.cell(2, 4).paragraphs[0].add_run('Nota').font.bold = True
        table.cell(2, 5).paragraphs[0].add_run('Asignatura').font.bold = True
        table.cell(2, 6).paragraphs[0].add_run('Nota').font.bold = True
        row_a = 3
        for i in range(0, len(request.detail_cm[case]) - 2):
            if len(request.detail_cm[case][i]['origin_subject']) > len(request.detail_cm[case][i]['subject_out']):
                row_m = row_a + \
                    len(request.detail_cm[case][i]['origin_subject']) - 1
                cellp = table.cell(row_a, 0).merge(
                    table.cell(row_m, 0)).paragraphs[0]
                cellp.add_run(
                    request.detail_cm[case][i]['subject_out'][0]['code'])
                cellp = table.cell(row_a, 1).merge(
                    table.cell(row_m, 1)).paragraphs[0]
                cellp.add_run(
                    request.detail_cm[case][i]['subject_out'][0]['name'])
                cellp = table.cell(row_a, 2).merge(
                    table.cell(row_m, 2)).paragraphs[0]
                cellp.add_run(
                    request.detail_cm[case][i]['subject_out'][0]['credits'])
                cellp = table.cell(row_a, 3).merge(
                    table.cell(row_m, 3)).paragraphs[0]
                cellp.add_run(
                    request.detail_cm[case][i]['subject_out'][0]['tipology'])
                cellp = table.cell(row_a, 4).merge(
                    table.cell(row_m, 4)).paragraphs[0]
                cellp.add_run(
                    request.detail_cm[case][i]['subject_out'][0]['grade'])
                credits = num_credits[request.detail_cm[case]
                                      [i]['subject_out'][0]['tipology']]
                credits += int(request.detail_cm[case]
                               [i]['subject_out'][0]['credits'])
                num_credits[request.detail_cm[case][i]
                            ['subject_out'][0]['tipology']] = credits
                num_credits['total'] += credits
                for j in range(0, len(request.detail_cm[case][i]['origin_subject'])):
                    table.cell(row_a, 5).paragraphs[0].add_run(
                        request.detail_cm[case][i]['origin_subject'][j]['name'])
                    table.cell(row_a, 6).paragraphs[0].add_run(
                        request.detail_cm[case][i]['origin_subject'][j]['grade'])
                    row_a += 1
            else:
                row_m = row_a + \
                    len(request.detail_cm[case][i]['subject_out']) - 1
                cellp = table.cell(row_a, 5).merge(
                    table.cell(row_m, 5)).paragraphs[0]
                cellp.add_run(
                    request.detail_cm[case][i]['origin_subject'][0]['name'])
                cellp = table.cell(row_a, 6).merge(
                    table.cell(row_m, 6)).paragraphs[0]
                cellp.add_run(
                    request.detail_cm[case][i]['origin_subject'][0]['grade'])
                for j in range(0, len(request.detail_cm[case][i]['subject_out'])):
                    table.cell(row_a, 0).paragraphs[0].add_run(
                        request.detail_cm[case][i]['subject_out'][j]['code'])
                    table.cell(row_a, 1).paragraphs[0].add_run(
                        request.detail_cm[case][i]['subject_out'][j]['name'])
                    table.cell(row_a, 2).paragraphs[0].add_run(
                        request.detail_cm[case][i]['subject_out'][j]['credits'])
                    table.cell(row_a, 3).paragraphs[0].add_run(
                        request.detail_cm[case][i]['subject_out'][j]['tipology'])
                    table.cell(row_a, 4).paragraphs[0].add_run(
                        request.detail_cm[case][i]['subject_out'][j]['grade'])
                    credits = num_credits[request.detail_cm[case]
                                          [i]['subject_out'][j]['tipology']]
                    credits += int(request.detail_cm[case]
                                   [i]['subject_out'][j]['credits'])
                    num_credits[request.detail_cm[case][i]
                                ['subject_out'][j]['tipology']] = credits
                    num_credits['total'] += credits
                    row_a += 1
        cant_rows = 0
        appear = []
        if num_credits['L'] != 0:
            cant_rows += 1
            appear.append('L')
        if num_credits['B'] != 0:
            cant_rows += 1
            appear.append('B')
        if num_credits['C'] != 0:
            cant_rows += 1
            appear.append('C')
        for i in range(0, cant_rows + 1):
            table.add_row()
        # table = docx.add_table(rows=cant_rows+1, cols=3, style='Table Grid')
        # table.style.font.size = Pt(8)
        # table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cellp = table.cell(row_a, 0).merge(table.cell(row_a, 1)).paragraphs[0]
        cellp.add_run('Total créditos')
        table.cell(row_a, 2).paragraphs[0].add_run(str(num_credits['total']))
        row_a += 1
        for i in range(0, cant_rows):
            cellp = table.cell(
                i + row_a, 0).merge(table.cell(i + row_a, 1)).paragraphs[0]
            cellp.add_run('Créditos homologados ' + appear[i-1] + ':')
            table.cell(
                i + row_a, 2).paragraphs[0].add_run(str(num_credits[appear[i-1]]))

        cellp = table.cell(
            row_a - 1, 3).merge(table.cell(cant_rows - 1 + row_a, 6)).paragraphs[0]
        for column in table.columns:
            for cell in column.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    @staticmethod
    def case_HOMOLOGACION_CONVALIDACION_EQUIVALENCIA_PREGRADO_TABLE_NA(request, docx, case):
        case_d = {'homologation': 'homologan',
                  'equivalence': 'equivalen', 'recognition': 'convalidan'}
        cant_rows = 0
        for i in range(0, len(request.detail_cm[case]) - 3):
            if len(request.detail_cm[case][i]['origin_subject']) > len(request.detail_cm[case][i]['subject_out']):
                cant_rows += len(request.detail_cm[case][i]['origin_subject'])
            else:
                cant_rows += len(request.detail_cm[case][i]['subject_out'])
        table = docx.add_table(rows=cant_rows+3, cols=5, style='Table Grid')
        table.style.font.size = Pt(8)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for column in table.columns:
            for cell in column.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cellp = table.cell(0, 0).merge(table.cell(0, 4)).paragraphs[0]
        cellp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cellp.add_run('{}\t\t\tDNI.{}'.format(
            request.student_name, request.student_dni)).font.bold = True
        cellp = table.cell(1, 0).merge(table.cell(1, 4)).paragraphs[0]
        cellp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cellp.add_run('Asignaturas que no se {} en el plan de estudios de {}({})'.format(
            case_d[case], request.get_academic_program_display(), request.academic_program)).font.bold = True
        table.cell(2, 0).paragraphs[0].add_run('Asignatura Universidad Nacional de Colombia ({})'.format(
            request.academic_program)).font.bold = True
        table.cell(2, 1).paragraphs[0].add_run('Asignatura ({})'.format(
            request.detail_cm[case][-2])).font.bold = True
        table.cell(2, 2).paragraphs[0].add_run(
            'Justificación').font.bold = True
        table.cell(2, 3).paragraphs[0].add_run('C').font.bold = True
        table.cell(2, 4).paragraphs[0].add_run('Nota').font.bold = True
        row_a = 3
        for i in range(0, len(request.detail_cm[case]) - 3):
            if len(request.detail_cm[case][i]['origin_subject']) > len(request.detail_cm[case][i]['subject_out']):
                row_m = row_a + \
                    len(request.detail_cm[case][i]['origin_subject']) - 1
                cellp = table.cell(row_a, 0).merge(
                    table.cell(row_m, 0)).paragraphs[0]
                cellp.add_run(
                    request.detail_cm[case][i]['subject_out'][0]['name'])
                cellp = table.cell(row_a, 3).merge(
                    table.cell(row_m, 3)).paragraphs[0]
                cellp.add_run(
                    request.detail_cm[case][i]['subject_out'][0]['credits'])
                cellp = table.cell(row_a, 4).merge(
                    table.cell(row_m, 4)).paragraphs[0]
                cellp.add_run(
                    request.detail_cm[case][i]['subject_out'][0]['grade'])
                for j in range(0, len(request.detail_cm[case][i]['origin_subject'])):
                    table.cell(row_a, 1).paragraphs[0].add_run(
                        request.detail_cm[case][i]['origin_subject'][j]['name'])
                    table.cell(row_a, 2).paragraphs[0].add_run(
                        request.detail_cm[case][i]['origin_subject'][j]['justification'])
                    row_a += 1
            else:
                row_m = row_a + \
                    len(request.detail_cm[case][i]['subject_out']) - 1
                cellp = table.cell(row_a, 1).merge(
                    table.cell(row_m, 1)).paragraphs[0]
                cellp.add_run(
                    request.detail_cm[case][i]['origin_subject'][0]['name'])
                cellp = table.cell(row_a, 2).merge(
                    table.cell(row_m, 2)).paragraphs[0]
                cellp.add_run(
                    request.detail_cm[case][i]['origin_subject'][0]['justification'])
                for j in range(0, len(request.detail_cm[case][i]['subject_out'])):
                    table.cell(row_a, 0).paragraphs[0].add_run(
                        request.detail_cm[case][i]['subject_out'][j]['name'])
                    table.cell(row_a, 3).paragraphs[0].add_run(
                        request.detail_cm[case][i]['subject_out'][j]['credits'])
                    table.cell(row_a, 4).paragraphs[0].add_run(
                        request.detail_cm[case][i]['subject_out'][j]['grade'])
                    row_a += 1
