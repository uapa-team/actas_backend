from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from ...models import Request


class HCEMPOS():

    @staticmethod
    def case_HOMOLOGACION_CONVALIDACION_Y_EQUIVALENCIA_POSGRADO(request, docx, redirected=False):
        if redirected:
            para = docx.paragraphs[-1]
        else:
            para = docx.add_paragraph()
            para.add_run('El Consejo de Facultad ')
        large_program = ''
        for p in Request.PROGRAM_CHOICES:
            if p[0] == request.academic_program:
                large_program = p[1]
                break
        para.paragraph_format.space_after = Pt(4)
        if request.approval_status == 'AP':
            para.add_run('APRUEBA:').font.bold = True
        else:
            para.add_run('NO APRUEBA:').font.bold = True
        if 'homologation' in request.detail_cm:
            para2 = docx.add_paragraph(style='List Number')
            para2.paragraph_format.space_after = Pt(0)
            # para2.add_run(cases_n + '.').font.bold = True
            # cases_n = cases_n + 1
            para2.add_run('Homologar en el programa ' +
                          large_program + ' plan de estudios ')
            para2.add_run(request.detail_cm['node'])
            para2.add_run(', las siguientes asignaturas cursadas en ')
            para2.add_run(
                request.detail_cm['homologation']['institution'])
            para2.add_run(', debido a que ' + request.justification)
            if request.approval_status == 'AP':
                para2.add_run(', así:')
            else:
                para2.add_run('.')
            table_rows = len(request.detail_cm['homologation']['subjects'])+2
            table = docx.add_table(rows=table_rows, cols=7)
            table.style = 'Table Grid'
            table.style.font.size = Pt(8)
            table.alignment = WD_ALIGN_PARAGRAPH.CENTER
            table.columns[0].width = 1400000
            table.columns[1].width = 530000
            table.columns[2].width = 1400000
            table.columns[3].width = 440000
            table.columns[4].width = 440000
            table.columns[5].width = 440000
            table.columns[6].width = 550000

            for cell in table.columns[0].cells:
                cell.width = 1400000
            for cell in table.columns[1].cells:
                cell.width = 580000
            for cell in table.columns[2].cells:
                cell.width = 1400000
            for cell in table.columns[3].cells:
                cell.width = 490000
            for cell in table.columns[4].cells:
                cell.width = 490000
            for cell in table.columns[5].cells:
                cell.width = 490000
            for cell in table.columns[6].cells:
                cell.width = 600000
            
            for col in table.columns:
                for cell in col.cells:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            table.cell(0, 0).paragraphs[0].add_run(
                'Asignatura que se homologa').font.bold = True
            merge1 = table.cell(0, 1).merge(table.cell(0, 6))
            merge1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            merge1.paragraphs[0].add_run(
                'Asignatura por la que se homologa').font.bold = True
            table.cell(1, 0).paragraphs[0].add_run('Nombre').font.bold = True
            table.cell(1, 1).paragraphs[0].add_run('Código').font.bold = True
            table.cell(1, 2).paragraphs[0].add_run('Nombre').font.bold = True
            table.cell(1, 3).paragraphs[0].add_run('Nota').font.bold = True
            table.cell(1, 4).paragraphs[0].add_run('C').font.bold = True
            table.cell(1, 5).paragraphs[0].add_run('T').font.bold = True
            table.cell(1, 6).paragraphs[0].add_run('Periodo').font.bold = True
            row = 2
            for subjet in request.detail_cm['homologation']['subjects']:
                table.cell(row, 0).paragraphs[0].add_run(subjet['subject'])
                table.cell(row, 1).paragraphs[0].add_run(subjet['code'])
                table.cell(row, 2).paragraphs[0].add_run(subjet['subject_out'])
                table.cell(row, 3).paragraphs[0].add_run(subjet['grade'])
                table.cell(row, 4).paragraphs[0].add_run(subjet['credits'])
                table.cell(row, 5).paragraphs[0].add_run(subjet['tipology'])
                table.cell(row, 6).paragraphs[0].add_run(
                    subjet['homologated_period'])
                row = row + 1
        if 'recognition' in request.detail_cm:
            para3 = docx.add_paragraph(style='List Number')
            para3.paragraph_format.space_after = Pt(0)
            para3.add_run('Convalidar en el programa ' +
                          large_program + ' plan de estudios ')
            para3.add_run(request.detail_cm['node'])
            para3.add_run(', las siguientes asignaturas')
            if 'program' in request.detail_cm["recognition"]:
                para3.add_run(' cursadas en el programa ' + large_program)
                if 'node' in request.detail_cm["recognition"]:
                    large_program = ''
                    for p in Request.PROGRAM_CHOICES:
                        if p[0] == request.detail_cm["recognition"]['program']:
                            large_program = p[1]
                            break
                    para3.add_run(' plan de estudios ')
                    para3.add_run(request.detail_cm["recognition"]['node'])
            if 'institution' in request.detail_cm["recognition"]:
                para3.add_run(' en el(la) ' +
                              request.detail_cm['recognition']['institution'])
            para3.add_run(', debido a que ' + request.justification)
            if request.approval_status == 'AP':
                para3.add_run(', así:')
            else:
                para3.add_run('.')
            table_rows = len(request.detail_cm['recognition']['subjects']) + 2
            table = docx.add_table(rows=table_rows, cols=7)
            table.style = 'Table Grid'
            table.style.font.size = Pt(8)
            table.alignment = WD_ALIGN_PARAGRAPH.CENTER
            table.columns[0].width = 1400000
            table.columns[1].width = 530000
            table.columns[2].width = 1400000
            table.columns[3].width = 440000
            table.columns[4].width = 440000
            table.columns[5].width = 440000
            table.columns[6].width = 550000

            for cell in table.columns[0].cells:
                cell.width = 1400000
            for cell in table.columns[1].cells:
                cell.width = 580000
            for cell in table.columns[2].cells:
                cell.width = 1400000
            for cell in table.columns[3].cells:
                cell.width = 490000
            for cell in table.columns[4].cells:
                cell.width = 490000
            for cell in table.columns[5].cells:
                cell.width = 490000
            for cell in table.columns[6].cells:
                cell.width = 600000
            
            for col in table.columns:
                for cell in col.cells:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            table.cell(0, 0).paragraphs[0].add_run(
                'Asignatura que se convalida').font.bold = True
            merge1 = table.cell(0, 1).merge(table.cell(0, 6))
            merge1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            merge1.paragraphs[0].add_run(
                'Asignatura por la que se convalida').font.bold = True
            table.cell(1, 0).paragraphs[0].add_run('Nombre').font.bold = True
            table.cell(1, 1).paragraphs[0].add_run('Código').font.bold = True
            table.cell(1, 2).paragraphs[0].add_run('Nombre').font.bold = True
            table.cell(1, 3).paragraphs[0].add_run('Nota').font.bold = True
            table.cell(1, 4).paragraphs[0].add_run('C').font.bold = True
            table.cell(1, 5).paragraphs[0].add_run('T').font.bold = True
            table.cell(1, 6).paragraphs[0].add_run('Periodo').font.bold = True
            row = 2
            for subjet in request.detail_cm['recognition']['subjects']:
                print(request.detail_cm['recognition']['subjects'])
                table.cell(row, 0).paragraphs[0].add_run(subjet['subject'])
                table.cell(row, 1).paragraphs[0].add_run(subjet['code'])
                table.cell(row, 2).paragraphs[0].add_run(subjet['subject_out'])
                table.cell(row, 3).paragraphs[0].add_run(subjet['grade'])
                table.cell(row, 4).paragraphs[0].add_run(subjet['credits'])
                table.cell(row, 5).paragraphs[0].add_run(subjet['tipology'])
                table.cell(row, 6).paragraphs[0].add_run(
                    subjet['homologated_period'])
                row = row + 1
        if 'equivalence' in request.detail_cm:
            para4 = docx.add_paragraph(style='List Number')
            para4.paragraph_format.space_after = Pt(0)
            para4.add_run('Equivaler en el programa ' +
                          large_program + ' plan de estudios ')
            para4.add_run(request.detail_cm['node'])
            para4.add_run(', las siguientes asignaturas cursadas en ')
            large_program = ''
            for p in Request.PROGRAM_CHOICES:
                if p[0] == request.detail_cm["equivalence"]['program']:
                    large_program = p[1]
                    break
            para4.add_run('el programa ' + large_program)
            if 'node' in request.detail_cm["equivalence"]:
                para4.add_run(' plan de estudios ')
                para4.add_run(
                    request.detail_cm["equivalence"]['node'])
            para4.add_run(', debido a que ' + request.justification)
            if request.approval_status == 'AP':
                para4.add_run(', así:')
            else:
                para4.add_run('.')
            table_rows = len(request.detail_cm['equivalence']['subjects'])+2
            table = docx.add_table(rows=table_rows, cols=7)
            table.style = 'Table Grid'
            table.style.font.size = Pt(8)
            table.alignment = WD_ALIGN_PARAGRAPH.CENTER
            table.columns[0].width = 1400000
            table.columns[1].width = 530000
            table.columns[2].width = 1400000
            table.columns[3].width = 440000
            table.columns[4].width = 440000
            table.columns[5].width = 440000
            table.columns[6].width = 550000
        
            for cell in table.columns[0].cells:
                cell.width = 1400000
            for cell in table.columns[1].cells:
                cell.width = 580000
            for cell in table.columns[2].cells:
                cell.width = 1400000
            for cell in table.columns[3].cells:
                cell.width = 490000
            for cell in table.columns[4].cells:
                cell.width = 490000
            for cell in table.columns[5].cells:
                cell.width = 490000
            for cell in table.columns[6].cells:
                cell.width = 600000
            
            for col in table.columns:
                for cell in col.cells:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            table.cell(0, 0).paragraphs[0].add_run(
                'Asignatura que se equivale').font.bold = True
            merge1 = table.cell(0, 1).merge(table.cell(0, 6))
            merge1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            merge1.paragraphs[0].add_run(
                'Asignatura por la que se equivale').font.bold = True
            table.cell(1, 0).paragraphs[0].add_run('Nombre').font.bold = True
            table.cell(1, 1).paragraphs[0].add_run('Código').font.bold = True
            table.cell(1, 2).paragraphs[0].add_run('Nombre').font.bold = True
            table.cell(1, 3).paragraphs[0].add_run('Nota').font.bold = True
            table.cell(1, 4).paragraphs[0].add_run('C').font.bold = True
            table.cell(1, 5).paragraphs[0].add_run('T').font.bold = True
            table.cell(1, 6).paragraphs[0].add_run('Periodo').font.bold = True
            row = 2
            for subjet in request.detail_cm['equivalence']['subjects']:
                table.cell(row, 0).paragraphs[0].add_run(subjet['subject'])
                table.cell(row, 1).paragraphs[0].add_run(subjet['code'])
                table.cell(row, 2).paragraphs[0].add_run(subjet['subject_out'])
                table.cell(row, 3).paragraphs[0].add_run(subjet['grade'])
                table.cell(row, 4).paragraphs[0].add_run(subjet['credits'])
                table.cell(row, 5).paragraphs[0].add_run(subjet['tipology'])
                table.cell(row, 6).paragraphs[0].add_run(
                    subjet['homologated_period'])
                row = row + 1
