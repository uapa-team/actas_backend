from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from ...models import Request


class IASIPRE():

    @staticmethod
    def case_INSCRIPCION_DE_ASIGNATURAS_PREGRADO(request, docx, redirected=False):
        para = docx.paragraphs[-1]
        if not redirected:
            para = docx.add_paragraph()
            para.add_run('El Consejo de Facultad ')
        large_program = ''
        for p in Request.PROGRAM_CHOICES:
            if p[0] == request.academic_program:
                large_program = p[1]
                break
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if request.approval_status == 'AP':
            para.add_run('APRUEBA ').font.bold = True
        else:
            para.add_run('NO APRUEBA ').font.bold = True
        para.add_run(
            'inscribir la(s) siguiente(s) asignatura(s) del programa ' + large_program)
        para.add_run(', en el periodo académico ' + request.academic_period)
        if request.approval_status == 'AP':
            para.add_run(':')
        else:
            para.add_run(', debido a que ' + request.justification + '.')
        table = docx.add_table(
            rows=len(request.detail_cm['subjects'])+1, cols=5)
        table.style = 'Table Grid'
        table.style.font.size = Pt(8)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # table.columns[0].width = 700000
        # table.columns[1].width = 2300000
        # table.columns[2].width = 800000
        # table.columns[3].width = 800000
        # table.columns[4].width = 800000

        for column in table.columns:
            for cell in column.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        for cell in table.columns[0].cells:
            cell.width = 750000
        for cell in table.columns[1].cells:
            cell.width = 2300000
        for cell in table.columns[2].cells:
            cell.width = 800000
        for cell in table.columns[3].cells:
            cell.width = 800000
        for cell in table.columns[4].cells:
            cell.width = 800000

        table.cell(0, 0).paragraphs[0].add_run('Código').font.bold = True
        table.cell(0, 1).paragraphs[0].add_run('Asignatura').font.bold = True
        table.cell(0, 2).paragraphs[0].add_run('Grupo').font.bold = True
        table.cell(0, 3).paragraphs[0].add_run('Tipología').font.bold = True
        table.cell(0, 4).paragraphs[0].add_run('Créditos').font.bold = True
        index = 0
        for subject in request.detail_cm['subjects']:
            table.cell(index+1, 0).paragraphs[0].add_run(subject['cod'])
            table.cell(index+1, 1).paragraphs[0].add_run(subject['subject'])
            table.cell(index+1, 2).paragraphs[0].add_run(subject['grup'])
            table.cell(index+1, 3).paragraphs[0].add_run(subject['T'])
            table.cell(index+1, 4).paragraphs[0].add_run(subject['cre'])
            index = index + 1
        para = docx.add_paragraph()
