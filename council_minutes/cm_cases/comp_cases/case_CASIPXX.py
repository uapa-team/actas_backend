from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL


class CASIPXX():

    @staticmethod
    def case_CANCELACION_DE_ASIGNATURAS(request, docx, redirected=False):
        if redirected:
            para = docx.paragraphs[-1]
        else:
            para = docx.add_paragraph()
            para.add_run('El Consejo de Facultad ')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if request.approval_status == 'AP':
            CASIPXX.case_CANCELACION_DE_ASIGNATURAS_AP(request, docx, para)
        else:
            CASIPXX.case_CANCELACION_DE_ASIGNATURAS_NAP(request, docx, para)

    @staticmethod
    def case_CANCELACION_DE_ASIGNATURAS_AP(request, docx, paragraph):
        paragraph.add_run('APRUEBA').font.bold = True
        paragraph.add_run(
            ' cancelar la(s) siguiente(s) asignatura(s) inscrita(s) del periodo académico ')
        paragraph.add_run(request.academic_period +
                          ', porque justifica debidamente la solicitud.')
        paragraph.add_run(
            ' (Artículo 15 Acuerdo 008 de 2008 del Consejo Superior Universitario).')
        CASIPXX.case_CANCELACION_DE_ASIGNATURAS_TABLE(request, docx)

    @staticmethod
    def case_CANCELACION_DE_ASIGNATURAS_NAP(request, docx, paragraph):
        paragraph.add_run('NO APRUEBA').font.bold = True
        paragraph.add_run(
            ' cancelar la(s) siguiente(s) asignatura(s) inscrita(s) del periodo académico')
        paragraph.add_run(request.academic_period +
                          ', porque ' + request.justification + '. ')
        paragraph.add_run(
            '(Artículo 15 Acuerdo 008 de 2008 del Consejo Superior Universitario).')
        CASIPXX.case_CANCELACION_DE_ASIGNATURAS_TABLE(request, docx)

    @staticmethod
    def case_CANCELACION_DE_ASIGNATURAS_TABLE(request, docx):
        table = docx.add_table(
            rows=len(request.detail_cm['subjects'])+1, cols=5)
        for column in table.columns:
            for cell in column.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.style = 'Table Grid'
        table.style.font.size = Pt(8)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.columns[0].width = 700000
        table.columns[1].width = 2000000
        table.columns[2].width = 900000
        table.columns[3].width = 900000
        table.columns[4].width = 900000

        for cell in table.columns[0].cells:
            cell.width = 750000
        for cell in table.columns[1].cells:
            cell.width = 2000000
        for cell in table.columns[2].cells:
            cell.width = 900000
        for cell in table.columns[3].cells:
            cell.width = 900000
        for cell in table.columns[4].cells:
            cell.width = 900000

        cellp = table.cell(0, 0).paragraphs[0]
        cellp.add_run('Código SIA').font.bold = True
        cellp.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cellp = table.cell(0, 1).paragraphs[0]
        cellp.add_run('Nombre Asignatura').font.bold = True
        cellp.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cellp = table.cell(0, 2).paragraphs[0]
        cellp.add_run('Grupo').font.bold = True
        cellp.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cellp = table.cell(0, 3).paragraphs[0]
        cellp.add_run('Tipología').font.bold = True
        cellp.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cellp = table.cell(0, 4).paragraphs[0]
        cellp.add_run('Créditos').font.bold = True
        cellp.alignment = WD_ALIGN_PARAGRAPH.CENTER

        index = 0
        for subject in request.detail_cm['subjects']:
            table.cell(index+1, 0).paragraphs[0].add_run(subject['code'])
            table.cell(index+1, 1).paragraphs[0].add_run(subject['subject'])
            table.cell(index+1, 4).paragraphs[0].add_run(subject['group'])
            table.cell(index+1, 3).paragraphs[0].add_run(subject['tipology'])
            table.cell(index+1, 2).paragraphs[0].add_run(subject['credits'])
            index = index + 1
