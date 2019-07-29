from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from ...models import Request

class TRASPRE():


    @staticmethod
    def case_TRASLADO_PREGRADO(request, docx, redirected=False):
        para = docx.paragraphs[-1]
        if not redirected:
            para = docx.add_paragraph()
            para.add_run('El Consejo de Facultad ')
        large_program = ''
        for p in Request.PROGRAM_CHOICES:
            if p[0] == request.academic_program:
                large_program = p[1]
                break
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if request.approval_status == 'AP':
            para.add_run('APRUEBA').font.bold = True
            para.add_run(' traslado ' + request.detail_cm['tras_type'] + ' del programa ' + request.detail_cm['origin'] +'(')
            para.add_run(request.detail_cm['cod_origin'] + ') de la Universidad Nacional de Colombia - Sede ' + request.detail_cm['campus_origin'])
            para.add_run(', al programa ' + large_program + ' (' + request.academic_program)
            para.add_run(') de la Universidad Nacional de Colombia - Sede Bogotá, en el periodo académico ' + request.detail_cm['since'])
            para.add_run(' condicionado a consevar la calidad de estudiante al finalizar el periodo académico ' + request.academic_period)
            para.add_run('. (Artículo 39 del Acuerdo 008 de 2008 del Consejo Superior Universitario y Acuerdo 089 de 2014 del Consejo Académico). ')
            para = docx.add_paragraph()
        else:
            para.add_run('NO APRUEBA').font.bold = True
            para.add_run(' traslado ' + request.detail_cm['tras_type'] + ' del programa ' + request.detail_cm['origin'] +'(')
            para.add_run(request.detail_cm['cod_origin'] + ') de la Universidad Nacional de Colombia - Sede ' + request.detail_cm['campus_origin'])
            para.add_run(', al programa ' + large_program + ' (' + request.academic_program)
            para.add_run(') de la Universidad Nacional de Colombia - Sede Bogotá, en el periodo académico ' + request.detail_cm['since'])
            para.add_run(', porque ' + request.justification +'.')
            para.add_run('(Artículo 39 del Acuerdo 008 de 2008 del Consejo Superior Universitario y Acuerdo 089 de 2014 del Consejo Académico). ')
            para = docx.add_paragraph()

        para.paragraph_format.space_after = Pt(0)
        bullet = para.add_run('1. Datos Generales')
        bullet.font.bold = True
        bullet.font.size = Pt(8)
        table = docx.add_table(rows=9, cols=3)
        table.style='Table Grid'
        table.alignment=WD_ALIGN_PARAGRAPH.CENTER
        # table.columns[0].width = 200000
        # table.columns[1].width = 2600000
        # table.columns[2].width = 2600000
        table.style.font.size = Pt(8)
        for cell in table.columns[0].cells:
            cell.width = 450000
        for cell in table.columns[1].cells:
            cell.width = 2500000
        for cell in table.columns[2].cells:
            cell.width = 2500000
        
        for i in range(1,9):
            table.cell(i, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        cellp = table.cell(0, 0).merge(table.cell(0, 2)).paragraphs[0]
        cellp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cellp.add_run('TRASLADO').font.bold = True
        cellp.add_run('\nNormativa Asociada: Artículo 39 del Acuerdo 008 de 2008 del CSU y Acuerdo 089 de 2014 del C.A.')
        table.cell(1, 0).paragraphs[0].add_run('1').font.bold = True
        table.cell(1, 1).paragraphs[0].add_run('Estudiante')
        table.cell(1, 2).paragraphs[0].add_run(request.student_name)
        table.cell(2, 0).paragraphs[0].add_run('2').font.bold = True
        table.cell(2, 1).paragraphs[0].add_run('DNI')
        table.cell(2, 2).paragraphs[0].add_run(request.student_dni)
        table.cell(3, 0).paragraphs[0].add_run('3').font.bold = True
        table.cell(3, 1).paragraphs[0].add_run('Plan de estudios origen (1er plan)')
        table.cell(3, 2).paragraphs[0].add_run(request.detail_cm['origin'])
        table.cell(4, 0).paragraphs[0].add_run('4').font.bold = True
        table.cell(4, 1).paragraphs[0].add_run('Código del plan de estudios origen (1er plan)')
        table.cell(4, 2).paragraphs[0].add_run(request.detail_cm['cod_origin'])
        table.cell(5, 0).paragraphs[0].add_run('5').font.bold = True
        table.cell(5, 1).paragraphs[0].add_run('Plan de estudios destino (2° plan)')
        table.cell(5, 2).paragraphs[0].add_run(large_program)
        table.cell(6, 0).paragraphs[0].add_run('6').font.bold = True
        table.cell(6, 1).paragraphs[0].add_run('Código del plan de estudios destino (2° plan)')
        table.cell(6, 2).paragraphs[0].add_run(request.academic_program)
        table.cell(7, 0).paragraphs[0].add_run('7').font.bold = True
        table.cell(7, 1).paragraphs[0].add_run('Fecha de la solicitud a través del SIA')
        table.cell(7, 2).paragraphs[0].add_run(request.detail_cm['sia_request'])
        table.cell(8, 0).paragraphs[0].add_run('8').font.bold = True
        table.cell(8, 1).paragraphs[0].add_run('¿Estos planes de estudio conducen al mismo título?')
        table.cell(8, 2).paragraphs[0].add_run(request.detail_cm['same_degree'])
        para = docx.add_paragraph()
        para.paragraph_format.space_before = Pt(8)
        bullet = para.add_run('2. Información Académica:')
        para.paragraph_format.space_after = Pt(0)
        bullet.font.bold = True
        bullet.font.size = Pt(8)
        table = docx.add_table(rows=4, cols=2, style='Table Grid')
        table.alignment=WD_ALIGN_PARAGRAPH.CENTER
        # table.columns[0].width = 4600000
        # table.columns[1].width = 800000

        for cell in table.columns[0].cells:
            cell.width = 4600000
        for cell in table.columns[1].cells:
            cell.width = 850000

        table.cell(0, 0).paragraphs[0].add_run('Periodo para el cual fue admitido')
        table.cell(0, 1).paragraphs[0].add_run(request.detail_cm['info_acad'][0]['per_adm'])
        table.cell(1, 0).paragraphs[0].add_run('¿El solicitante se encuentra matriculado en el semestre de presentar la solicitud?')            
        table.cell(1, 1).paragraphs[0].add_run(request.detail_cm['info_acad'][0]['matr'])
        table.cell(2, 0).paragraphs[0].add_run('¿El solicitante tuvo calidad de estudiante en el plan de estudios destino (2° plan)?')            
        table.cell(2, 1).paragraphs[0].add_run(request.detail_cm['info_acad'][0]['ant_plan'])
        table.cell(3, 0).paragraphs[0].add_run('Porcentaje de créditos aprobados en el plan de estudios origen (1er plan)')            
        table.cell(3, 1).paragraphs[0].add_run(request.detail_cm['info_acad'][0]['aprob']+'%')

        if float(request.detail_cm['info_acad'][0]['aprob']) <= 30:
            table = docx.add_table(rows=2, cols=2, style='Table Grid')
            table.alignment=WD_ALIGN_PARAGRAPH.CENTER
            # table.columns[0].width = 4600000
            # table.columns[1].width = 800000

            for cell in table.columns[0].cells:
                cell.width = 4600000
            for cell in table.columns[1].cells:
                cell.width = 850000

            table.cell(0, 0).paragraphs[0].add_run('¿Cuál fue el puntaje de admisión del solicitante?')
            table.cell(0, 1).paragraphs[0].add_run(request.detail_cm['info_acad'][1]['pun'])
            table.cell(1, 0).paragraphs[0].add_run('Puntaje de admisión del último admitido regular al plan destino (2° plan) en la misma prueba de ingreso del solicitante*')
            table.cell(1, 1).paragraphs[0].add_run(request.detail_cm['info_acad'][1]['pun_ult'])

        if (float(request.detail_cm['info_acad'][1]['pun'])>=float(request.detail_cm['info_acad'][1]['pun_ult'])) or float(request.detail_cm['info_acad'][0]['aprob'])>=30:
            para = docx.add_paragraph()
            table = docx.add_table(rows=3, cols=2, style='Table Grid')
            table.alignment=WD_ALIGN_PARAGRAPH.CENTER
            # table.columns[0].width = 4600000
            # table.columns[1].width = 800000
            for cell in table.columns[0].cells:
                cell.width = 4600000
            for cell in table.columns[1].cells:
                cell.width = 850000
            table.cell(0, 0).paragraphs[0].add_run('Cupo de créditos menos créditos pendientes en el plan de estudios origen (1er plan)')
            table.cell(0, 1).paragraphs[0].add_run(request.detail_cm['info_acad'][2]['cupo_pend'])
            table.cell(1, 0).paragraphs[0].add_run('Cupo de créditos para traslado (literal d del artículo 3 de la resolución 089 de 2015 de Consejo Académico)')
            table.cell(1, 1).paragraphs[0].add_run(request.detail_cm['info_acad'][2]['cred_tras'])
            table.cell(2, 0).paragraphs[0].add_run('El cupo de créditos para traslado es igual o mayor al número de créditos pendientes de aprobación en el plan de estudios destino (2° plan)?')
            if (float(request.detail_cm['info_acad'][2]['cupo_pend'])-float(request.detail_cm['info_acad'][2]['cred_tras']))>=0:
                table.cell(2, 1).paragraphs[0].add_run('Si')
            else:
                table.cell(2, 1).paragraphs[0].add_run('No')
            para = docx.add_paragraph()
            para.add_run('* en caso que el plan de destino tenga convocatoria anual el puntaje será con la anterior convocatoria.')
            para = docx.add_paragraph()
            
        if request.approval_status == 'AP':
            para = docx.add_paragraph() 
            para.paragraph_format.space_before = Pt(8)
            bullet = para.add_run('3. Resumen General de Créditos del Segundo Plan de Estudios')
            para.paragraph_format.space_after = Pt(0)
            bullet.font.bold = True
            bullet.font.size = Pt(8)
            table = docx.add_table(rows=5, cols=7)
            table.style='Table Grid'
            table.alignment=WD_ALIGN_PARAGRAPH.CENTER
            # table.columns[0].width = 950000
            # table.columns[1].width = 820000
            # table.columns[2].width = 780000
            # table.columns[3].width = 820000
            # table.columns[4].width = 780000
            # table.columns[5].width = 675000
            # table.columns[6].width = 575000
            for column in table.columns:
                for cell in column.cells:
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            for cell in table.columns[0].cells:
                cell.width = 1250000
            for cell in table.columns[1].cells:
                cell.width = 700000
            for cell in table.columns[2].cells:
                cell.width = 700000
            for cell in table.columns[3].cells:
                cell.width = 700000
            for cell in table.columns[4].cells:
                cell.width = 700000
            for cell in table.columns[5].cells:
                cell.width = 850000
            for cell in table.columns[6].cells:
                cell.width = 550000
            cellp = table.cell(0, 0).merge(table.cell(1, 0)).paragraphs[0]
            cellp.vertical_alignement = WD_ALIGN_VERTICAL.CENTER
            cellp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cellp.add_run('Créditos').font.bold = True
            cellp = table.cell(0, 1).merge(table.cell(0, 2)).paragraphs[0]
            cellp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cellp.add_run('Fundamentación (B)').font.bold = True
            cellp = table.cell(0, 3).merge(table.cell(0, 4)).paragraphs[0]
            cellp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cellp.add_run('Disciplinar (C)').font.bold = True
            cellp = table.cell(0, 5).merge(table.cell(1, 5)).paragraphs[0]
            cellp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cellp.add_run('Libre Elección (L)').font.bold = True
            cellp = table.cell(0, 6).merge(table.cell(1, 6)).paragraphs[0]
            cellp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cellp.add_run('Total').font.bold = True
            table.cell(2, 0).paragraphs[0].add_run('Exigidos*').font.bold = True
            table.cell(3, 0).paragraphs[0].add_run('Convalidados/Equivalentes').font.bold = True
            table.cell(4, 0).paragraphs[0].add_run('Pendientes').font.bold = True
            table.cell(1, 1).paragraphs[0].add_run('Obligatorios').font.bold = True
            table.cell(1, 2).paragraphs[0].add_run('Optativos').font.bold = True
            table.cell(1, 3).paragraphs[0].add_run('Obligatorios').font.bold = True
            table.cell(1, 4).paragraphs[0].add_run('Optativos').font.bold = True
            table.cell(2, 1).paragraphs[0].add_run(request.detail_cm['resumen'][0][0])
            table.cell(2, 2).paragraphs[0].add_run(request.detail_cm['resumen'][0][1])
            table.cell(2, 3).paragraphs[0].add_run(request.detail_cm['resumen'][0][2])
            table.cell(2, 4).paragraphs[0].add_run(request.detail_cm['resumen'][0][3])
            table.cell(2, 5).paragraphs[0].add_run(request.detail_cm['resumen'][0][4])
            table.cell(2, 6).paragraphs[0].add_run(str(int(request.detail_cm['resumen'][0][0])+int(request.detail_cm['resumen'][0][1])+int(request.detail_cm['resumen'][0][2])+int(request.detail_cm['resumen'][0][3])+int(request.detail_cm['resumen'][0][4])))
            table.cell(3, 1).paragraphs[0].add_run(request.detail_cm['resumen'][1][0])
            table.cell(3, 2).paragraphs[0].add_run(request.detail_cm['resumen'][1][1])
            table.cell(3, 3).paragraphs[0].add_run(request.detail_cm['resumen'][1][2])
            table.cell(3, 4).paragraphs[0].add_run(request.detail_cm['resumen'][1][3])
            table.cell(3, 5).paragraphs[0].add_run(request.detail_cm['resumen'][1][4])
            table.cell(3, 6).paragraphs[0].add_run(str(int(request.detail_cm['resumen'][1][0])+int(request.detail_cm['resumen'][1][1])+int(request.detail_cm['resumen'][1][2])+int(request.detail_cm['resumen'][1][3])+int(request.detail_cm['resumen'][1][4])))
            table.cell(4, 1).paragraphs[0].add_run(request.detail_cm['resumen'][2][0])
            table.cell(4, 2).paragraphs[0].add_run(request.detail_cm['resumen'][2][1])
            table.cell(4, 3).paragraphs[0].add_run(request.detail_cm['resumen'][2][2])
            table.cell(4, 4).paragraphs[0].add_run(request.detail_cm['resumen'][2][3])
            table.cell(4, 5).paragraphs[0].add_run(request.detail_cm['resumen'][2][4])
            table.cell(4, 6).paragraphs[0].add_run(str(int(request.detail_cm['resumen'][2][0])+int(request.detail_cm['resumen'][2][1])+int(request.detail_cm['resumen'][2][2])+int(request.detail_cm['resumen'][2][3])+int(request.detail_cm['resumen'][2][4])))
            
            para = docx.add_paragraph() 
            para.paragraph_format.space_before = Pt(0)
            bullet = para.add_run('*Sin incluir los créditos correspondientes al cumplimiento del requisito de suficiencia en idioma extranjero.')
            para.paragraph_format.space_after = Pt(0)
            bullet.font.size = Pt(8)
            para = docx.add_paragraph()
            para.paragraph_format.space_before = Pt(8)
            para.alignment=WD_ALIGN_PARAGRAPH.CENTER
            bullet = para.add_run('CUADRO EQUIVALENCIA Y CONVALIDACIONES DE ASIGNATURAS CURSADAS Y APROBADAS HASTA LA FECHA DE PRESENTACIÓN DE LA SOLICITUD POR PARTE DEL ESTUDIANTE')
            para.paragraph_format.space_after = Pt(0)
            bullet.font.bold = True
            bullet.font.size = Pt(8)
            bullet.font.underline = True
            table = docx.add_table(rows=len(request.detail_cm['equivalencia'])+3, cols=10, style='Table Grid')
            table.alignment=WD_ALIGN_PARAGRAPH.CENTER
            # table.columns[0].width = 700000
            # table.columns[1].width = 900000
            # table.columns[2].width = 700000
            # table.columns[3].width = 900000
            # table.columns[4].width = 225000
            # table.columns[5].width = 225000
            # table.columns[6].width = 225000
            # table.columns[7].width = 900000
            # table.columns[8].width = 225000
            # table.columns[9].width = 400000
            for column in table.columns:
                for cell in column.cells:
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            table.allow_autofit = False
            for cell in table.columns[0].cells:
                cell.width = 650000
            for cell in table.columns[1].cells:
                cell.width = 900000
            for cell in table.columns[2].cells:
                cell.width = 650000
            for cell in table.columns[3].cells:
                cell.width = 900000
            for cell in table.columns[4].cells:
                cell.width = 250000
            for cell in table.columns[5].cells:
                cell.width = 250000
            for cell in table.columns[6].cells:
                cell.width = 250000
            for cell in table.columns[7].cells:
                cell.width = 900000
            for cell in table.columns[8].cells:
                cell.width = 250000
            for cell in table.columns[9].cells:
                cell.width = 300000
            
            cellp = table.cell(0, 0).merge(table.cell(0, 1)).paragraphs[0]
            cellp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cellp.add_run('Plan de estudios 1').font.bold = True
            cellp = table.cell(0, 2).merge(table.cell(0, 9)).paragraphs[0]
            cellp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cellp.add_run('Plan de estudios 2').font.bold = True
            cellp = table.cell(len(request.detail_cm['equivalencia'])+2, 0).merge(table.cell(len(request.detail_cm['equivalencia'])+2, 7)).paragraphs[0]
            cellp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cellp.add_run('Total créditos convalidados/equivalentes').font.bold = True
            table.cell(1, 0).paragraphs[0].add_run('Código').font.bold = True
            table.cell(1, 1).paragraphs[0].add_run('Asignatura').font.bold = True
            table.cell(1, 2).paragraphs[0].add_run('Código').font.bold = True
            table.cell(1, 3).paragraphs[0].add_run('Asignatura').font.bold = True
            table.cell(1, 4).paragraphs[0].add_run('T*').font.bold = True
            table.cell(1, 5).paragraphs[0].add_run('Ob*').font.bold = True
            table.cell(1, 6).paragraphs[0].add_run('Op*').font.bold = True
            table.cell(1, 7).paragraphs[0].add_run('Agrupación').font.bold = True
            table.cell(1, 8).paragraphs[0].add_run('C*').font.bold = True
            table.cell(1, 9).paragraphs[0].add_run('Nota').font.bold = True         
            credits_sum=0
            index=0
            for subject in request.detail_cm['equivalencia']:
                credits_sum = credits_sum + int(subject['C'])
                table.cell(index+2, 0).paragraphs[0].add_run(subject['cod'])
                table.cell(index+2, 1).paragraphs[0].add_run(subject['asig'])
                table.cell(index+2, 2).paragraphs[0].add_run(subject['cod2'])
                table.cell(index+2, 3).paragraphs[0].add_run(subject['asig2'])
                table.cell(index+2, 4).paragraphs[0].add_run(subject['t'])
                table.cell(index+2, 5).paragraphs[0].add_run(subject['ob'])   
                table.cell(index+2, 6).paragraphs[0].add_run(subject['op']) 
                table.cell(index+2, 7).paragraphs[0].add_run(subject['agr']) 
                table.cell(index+2, 8).paragraphs[0].add_run(subject['C']) 
                table.cell(index+2, 9).paragraphs[0].add_run(subject['nota'])            
                index = index + 1
            cellp = table.cell(len(request.detail_cm['equivalencia'])+2, 8).merge(table.cell(len(request.detail_cm['equivalencia'])+2, 9)).paragraphs[0]
            cellp.add_run(str(credits_sum)) 
    
            para = docx.add_paragraph()
            para.paragraph_format.space_before = Pt(8)
            para.alignment=WD_ALIGN_PARAGRAPH.CENTER
            bullet = para.add_run('CUADRO EQUIVALENCIA Y CONVALIDACIONES DE ASIGNATURAS CURSADAS Y NO APROBADAS HASTA LA FECHA DE PRESENTACIÓN DE LA SOLICITUD POR PARTE DEL ESTUDIANTE')
            para.paragraph_format.space_after = Pt(0)
            bullet.font.bold = True
            bullet.font.size = Pt(8)
            bullet.font.underline = True
            table = docx.add_table(rows=len(request.detail_cm['eq_na'])+3, cols=10, style='Table Grid')
            table.alignment=WD_ALIGN_PARAGRAPH.CENTER
            # table.columns[0].width = 700000
            # table.columns[1].width = 900000
            # table.columns[2].width = 700000
            # table.columns[3].width = 900000
            # table.columns[4].width = 225000
            # table.columns[5].width = 225000
            # table.columns[6].width = 225000
            # table.columns[7].width = 900000
            # table.columns[8].width = 225000
            # table.columns[9].width = 400000
            for column in table.columns:
                for cell in column.cells:
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            table.allow_autofit = False
            for cell in table.columns[0].cells:
                cell.width = 650000
            for cell in table.columns[1].cells:
                cell.width = 900000
            for cell in table.columns[2].cells:
                cell.width = 650000
            for cell in table.columns[3].cells:
                cell.width = 900000
            for cell in table.columns[4].cells:
                cell.width = 250000
            for cell in table.columns[5].cells:
                cell.width = 250000
            for cell in table.columns[6].cells:
                cell.width = 250000
            for cell in table.columns[7].cells:
                cell.width = 900000
            for cell in table.columns[8].cells:
                cell.width = 250000
            for cell in table.columns[9].cells:
                cell.width = 300000

            cellp = table.cell(0, 0).merge(table.cell(0, 1)).paragraphs[0]
            cellp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cellp.add_run('Plan de estudios 1').font.bold = True
            cellp = table.cell(0, 2).merge(table.cell(0, 9)).paragraphs[0]
            cellp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cellp.add_run('Plan de estudios 2').font.bold = True
            cellp = table.cell(len(request.detail_cm['eq_na'])+2, 0).merge(table.cell(len(request.detail_cm['eq_na'])+2, 7)).paragraphs[0]
            cellp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cellp.add_run('Total créditos convalidados/equivalentes').font.bold = True
            table.cell(1, 0).paragraphs[0].add_run('Código').font.bold = True
            table.cell(1, 1).paragraphs[0].add_run('Asignatura').font.bold = True
            table.cell(1, 2).paragraphs[0].add_run('Código').font.bold = True
            table.cell(1, 3).paragraphs[0].add_run('Asignatura').font.bold = True
            table.cell(1, 4).paragraphs[0].add_run('T*').font.bold = True
            table.cell(1, 5).paragraphs[0].add_run('Ob*').font.bold = True
            table.cell(1, 6).paragraphs[0].add_run('Op*').font.bold = True
            table.cell(1, 7).paragraphs[0].add_run('Agrupación').font.bold = True
            table.cell(1, 8).paragraphs[0].add_run('C*').font.bold = True
            table.cell(1, 9).paragraphs[0].add_run('Nota').font.bold = True          
            credits_sum=0
            index=0
            for subject in request.detail_cm['eq_na']:
                credits_sum = credits_sum + int(subject['C'])
                table.cell(index+2, 0).paragraphs[0].add_run(subject['cod'])
                table.cell(index+2, 1).paragraphs[0].add_run(subject['asig'])
                table.cell(index+2, 2).paragraphs[0].add_run(subject['cod2'])
                table.cell(index+2, 3).paragraphs[0].add_run(subject['asig2'])
                table.cell(index+2, 4).paragraphs[0].add_run(subject['t'])
                table.cell(index+2, 5).paragraphs[0].add_run(subject['ob'])   
                table.cell(index+2, 6).paragraphs[0].add_run(subject['op']) 
                table.cell(index+2, 7).paragraphs[0].add_run(subject['agr']) 
                table.cell(index+2, 8).paragraphs[0].add_run(subject['C']) 
                table.cell(index+2, 9).paragraphs[0].add_run(subject['nota'])            
                index = index + 1
            table.cell(len(request.detail_cm['eq_na'])+2, 8).paragraphs[0].add_run(str(credits_sum)) 

            para = docx.add_paragraph()
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(8)
            bullet = para.add_run('*T:tipología(C/B/L). Ob: obligatoria. Op: optativa. C: créditos')
            bullet.font.underline = True
            bullet.font.italic = True
            bullet.font.size =  Pt(8)
            para = docx.add_paragraph()
            para.alignment=WD_ALIGN_PARAGRAPH.CENTER
            bullet = para.add_run('ASIGNATURAS PENDIENTES POR CURSAR EN EL SEGUNDO PLAN DE ESTUDIOS')
            bullet.font.size = Pt(8)
            para.paragraph_format.space_after = Pt(0)
            bullet.font.bold = True
            bullet.font.underline = True
            sum=0
            for i in request.detail_cm['pen_funda_obl']:
                sum = sum+len(i['mat'])
            sum2=len(request.detail_cm['pen_funda_opt'])
            table = docx.add_table(rows=sum+7+sum2, cols=5, style='Table Grid')
            table.alignment=WD_ALIGN_PARAGRAPH.CENTER
            # table.columns[0].width = 1300000
            # table.columns[1].width = 700000
            # table.columns[2].width = 1400000
            # table.columns[3].width = 1000000
            # table.columns[4].width = 1000000

            for cell in table.columns[0].cells:
                cell.width = 1000000
            for cell in table.columns[1].cells:
                cell.width = 1050000
            for cell in table.columns[2].cells:
                cell.width = 1400000
            for cell in table.columns[3].cells:
                cell.width = 1000000
            for cell in table.columns[4].cells:
                cell.width = 1000000
            
            for column in table.columns:
                for cell in column.cells:
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            cellp = table.cell(0, 0).merge(table.cell(0, 4)).paragraphs[0]
            cellp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cellp.add_run('Componente de fundamentación (B)').font.bold = True
            cellp = table.cell(1, 0).merge(table.cell(1, 4)).paragraphs[0]
            cellp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cellp.add_run('Obligatorias').font.bold = True
            cellp = table.cell(sum+3, 0).merge(table.cell(sum+3, 3)).paragraphs[0]
            cellp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cellp.add_run('Total créditos pendientes').font.bold = True
            cellp = table.cell(sum+4, 0).merge(table.cell(sum+4, 4)).paragraphs[0]
            cellp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cellp.add_run('Optativas').font.bold = True
            cellp = table.cell(sum+5, 0).merge(table.cell(sum+5, 1)).paragraphs[0]
            cellp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cellp.add_run('Nombre de la Agrupación').font.bold = True
            cellp = table.cell(sum+5, 2).merge(table.cell(sum+5, 3)).paragraphs[0]
            cellp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cellp.add_run('Créditos Requeridos').font.bold = True
            cellp = table.cell(sum+sum2+6, 0).merge(table.cell(sum+sum2+6, 3)).paragraphs[0]
            cellp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cellp.add_run('Total créditos pendientes').font.bold = True
            table.cell(2, 0).paragraphs[0].add_run('Agrupación').font.bold = True
            table.cell(2, 1).paragraphs[0].add_run('Código').font.bold = True
            table.cell(2, 2).paragraphs[0].add_run('Asignatura').font.bold = True
            table.cell(2, 3).paragraphs[0].add_run('Créditos asignatura').font.bold = True
            table.cell(2, 4).paragraphs[0].add_run('Créditos pendientes por cursar por el estudiante').font.bold = True
            table.cell(sum+5, 4).paragraphs[0].add_run('Créditos pendientes por cursar por el estudiante').font.bold = True
            index=0
            creditosgrandes=0
            for agrupacion in request.detail_cm['pen_funda_obl']:
                if len(agrupacion['mat'])>=2:
                    cellp = table.cell(index+3, 0).merge(table.cell(index+2+len(agrupacion['mat']), 0)).paragraphs[0]
                    cellp.add_run(agrupacion['grup'])
                    index1=0
                    creditos=0
                    for materia in agrupacion['mat']:
                        table.cell(index+index1+3, 1).paragraphs[0].add_run(materia['cod'])
                        table.cell(index+index1+3, 2).paragraphs[0].add_run(materia['mat'])
                        table.cell(index+index1+3, 3).paragraphs[0].add_run(materia['cre'])
                        index1=index1+1
                        creditos=creditos+int(materia['cre'])
                    cellp = table.cell(index+3, 4).merge(table.cell(index+2+len(agrupacion['mat']), 4)).paragraphs[0]
                    cellp.add_run(str(creditos))
                    creditosgrandes=creditosgrandes+creditos
                else:
                    table.cell(index+3, 0).paragraphs[0].add_run(agrupacion['grup'])
                    table.cell(index+3, 1).paragraphs[0].add_run(agrupacion['mat'][0]['cod'])
                    table.cell(index+3, 2).paragraphs[0].add_run(agrupacion['mat'][0]['mat'])
                    table.cell(index+3, 3).paragraphs[0].add_run(agrupacion['mat'][0]['cre'])
                    table.cell(index+3, 4).paragraphs[0].add_run(agrupacion['mat'][0]['cre'])
                    creditosgrandes=creditosgrandes+int(agrupacion['mat'][0]['cre'])
                index=index+len(agrupacion['mat'])
            table.cell(sum+3, 4).paragraphs[0].add_run(str(creditosgrandes))
            index=0
            creditosgrandes=0
            for agrupacion in request.detail_cm['pen_funda_opt']:
                cellp = table.cell(sum+6+index, 0).merge(table.cell(sum+6+index, 1)).paragraphs[0]
                cellp.add_run(agrupacion['agr'])
                cellp = table.cell(sum+6+index, 2).merge(table.cell(sum+6+index, 3)).paragraphs[0]
                cellp.add_run(agrupacion['req'])
                table.cell(sum+6+index, 4).paragraphs[0].add_run(agrupacion['pen'])
                index=index+1
                creditosgrandes=creditosgrandes+int(agrupacion['pen'])
            table.cell(sum+6+sum2, 4).paragraphs[0].add_run(str(creditosgrandes))
            para = docx.add_paragraph()

            sum=0
            for i in request.detail_cm['pen_dis_obl']:
                sum = sum+len(i['mat'])
            sum2=len(request.detail_cm['pen_dis_opt'])
            para.paragraph_format.space_before = Pt(0)
            para.add_run(' ').font.size = Pt(8)
            para.paragraph_format.space_after = Pt(0)
            table = docx.add_table(rows=sum+7+sum2, cols=5, style='Table Grid')
            table.alignment=WD_ALIGN_PARAGRAPH.CENTER
            # table.columns[0].width = 1300000
            # table.columns[1].width = 700000
            # table.columns[2].width = 1400000
            # table.columns[3].width = 1000000
            # table.columns[4].width = 1000000

            for cell in table.columns[0].cells:
                cell.width = 1000000
            for cell in table.columns[1].cells:
                cell.width = 1050000
            for cell in table.columns[2].cells:
                cell.width = 1400000
            for cell in table.columns[3].cells:
                cell.width = 1000000
            for cell in table.columns[4].cells:
                cell.width = 1000000
            
            for column in table.columns:
                for cell in column.cells:
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  

            cellp = table.cell(0, 0).merge(table.cell(0, 4)).paragraphs[0]
            cellp.add_run('Componente Disciplinar/Profesional (C)').font.bold = True
            cellp = table.cell(1, 0).merge(table.cell(1, 4)).paragraphs[0]
            cellp.add_run('Obligatorias').font.bold = True
            cellp = table.cell(sum+3, 0).merge(table.cell(sum+3, 3)).paragraphs[0]
            cellp.add_run('Total créditos pendientes').font.bold = True
            cellp = table.cell(sum+4, 0).merge(table.cell(sum+4, 4)).paragraphs[0]
            cellp.add_run('Optativas').font.bold = True
            cellp = table.cell(sum+5, 0).merge(table.cell(sum+5, 1)).paragraphs[0]
            cellp.add_run('Nombre de la Agrupación').font.bold = True
            cellp = table.cell(sum+5, 2).merge(table.cell(sum+5, 3)).paragraphs[0]
            cellp.add_run('Créditos Requeridos').font.bold = True
            cellp = table.cell(sum+sum2+6, 0).merge(table.cell(sum+sum2+6, 3)).paragraphs[0]
            cellp.add_run('Total créditos pendientes').font.bold = True
            table.cell(2, 0).paragraphs[0].add_run('Agrupación').font.bold = True
            table.cell(2, 1).paragraphs[0].add_run('Código').font.bold = True
            table.cell(2, 2).paragraphs[0].add_run('Asignatura').font.bold = True
            table.cell(2, 3).paragraphs[0].add_run('Créditos asignatura').font.bold = True
            table.cell(2, 4).paragraphs[0].add_run('Créditos pendientes por cursar por el estudiante').font.bold = True
            table.cell(sum+5, 4).paragraphs[0].add_run('Créditos pendientes por cursar por el estudiante').font.bold = True
            index=0
            creditosgrandes=0
            for agrupacion in request.detail_cm['pen_dis_obl']:
                if len(agrupacion['mat'])>=2:
                    cellp = table.cell(index+3, 0).merge(table.cell(index+2+len(agrupacion['mat']), 0)).paragraphs[0]
                    cellp.add_run(agrupacion['grup'])
                    index1=0
                    creditos=0
                    for materia in agrupacion['mat']:
                        table.cell(index+index1+3, 1).paragraphs[0].add_run(materia['cod'])
                        table.cell(index+index1+3, 2).paragraphs[0].add_run(materia['mat'])
                        table.cell(index+index1+3, 3).paragraphs[0].add_run(materia['cre'])
                        index1=index1+1
                        creditos=creditos+int(materia['cre'])
                    cellp = table.cell(index+3, 4).merge(table.cell(index+2+len(agrupacion['mat']), 4)).paragraphs[0]
                    cellp.add_run(str(creditos))
                    creditosgrandes=creditosgrandes+creditos
                else:
                    table.cell(index+3, 0).paragraphs[0].add_run(agrupacion['grup'])
                    table.cell(index+3, 1).paragraphs[0].add_run(agrupacion['mat'][0]['cod'])
                    table.cell(index+3, 2).paragraphs[0].add_run(agrupacion['mat'][0]['mat'])
                    table.cell(index+3, 3).paragraphs[0].add_run(agrupacion['mat'][0]['cre'])
                    table.cell(index+3, 4).paragraphs[0].add_run(agrupacion['mat'][0]['cre'])
                    creditosgrandes=creditosgrandes+int(agrupacion['mat'][0]['cre'])
                index=index+len(agrupacion['mat'])
            table.cell(sum+3, 4).paragraphs[0].add_run(str(creditosgrandes))
            index=0
            creditosgrandes=0
            for agrupacion in request.detail_cm['pen_dis_opt']:
                cellp = table.cell(sum+6+index, 0).merge(table.cell(sum+6+index, 1)).paragraphs[0]
                cellp.add_run(agrupacion['agr'])
                cellp = table.cell(sum+6+index, 2).merge(table.cell(sum+6+index, 3)).paragraphs[0]
                cellp.add_run(agrupacion['req'])
                table.cell(sum+6+index, 4).paragraphs[0].add_run(agrupacion['pen'])
                index=index+1
                creditosgrandes=creditosgrandes+int(agrupacion['pen'])
            table.cell(sum+6+sum2, 4).paragraphs[0].add_run(str(creditosgrandes))
            para = docx.add_paragraph()
            para.paragraph_format.space_before = Pt(0)
            para.add_run(' ').font.size = Pt(8)
            para.paragraph_format.space_after = Pt(0)
            table = docx.add_table(rows=1, cols=2, style='Table Grid')
            table.alignment=WD_ALIGN_PARAGRAPH.CENTER
            # table.columns[0].width = 4300000
            # table.columns[1].width = 1000000
            for cell in table.columns[0].cells:
                cell.width = 4000000
            for cell in table.columns[1].cells:
                cell.width = 1450000
            table.cell(0, 0).paragraphs[0].add_run('Componente de Libre Elección (L) (Créditos pendientes)').font.bold = True
            table.cell(0, 1).paragraphs[0].add_run(request.detail_cm['pen_lib'])
            para = docx.add_paragraph()
            para.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
            para.add_run('La oferta de asignaturas optativas en cada una de las agrupaciones y componentes ')
            para.add_run('del plan de estudios del programa de ' + request.detail_cm['origin'] + ', la ')
            para.add_run('encuentra en el Acuerdo ' + request.detail_cm['normativa'] + ', expedido por el ')
            para.add_run('consejo de la Facultad de Ingeniería.')

