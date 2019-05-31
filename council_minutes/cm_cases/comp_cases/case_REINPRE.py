from docx import Document
from docx.shared import Pt
from ...models import Request
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT

class REINPRE():

    @staticmethod
    def num_to_month(month):
        if int(month) == 1:
            return ' de enero de '
        elif int(month) == 2:
            return ' de febrero de '
        elif int(month) == 3:
            return ' de marzo de '
        elif int(month) == 4:
            return ' de abril de '
        elif int(month) == 5:
            return ' de mayo de '
        elif int(month) == 6:
            return ' de junio de '
        elif int(month) == 7:
            return ' de julio de '
        elif int(month) == 8:
            return ' de agosto de '
        elif int(month) == 9:
            return ' de septiembre de '
        elif int(month) == 10:
            return ' de octubre de '
        elif int(month) == 11:
            return ' de nomviembre de '
        elif int(month) == 12:
            return ' de diciembre de '

    @staticmethod
    def case_REINGRESO_PREGRADO(request, docx):
        large_program = ''
        for p in Request.PROGRAM_CHOICES:
            if p[0] == request.academic_program:
                large_program = p[1]
                break
        para = docx.add_paragraph()
        para.paragraph_format.space_after = Pt(0)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run('El Consejo de Facultad ')
        if request.approval_status == 'AP':
            para.add_run('APRUEBA ').font.bold = True
        else:
            para.add_run('NO APRUEBA ').font.bold = True
        para.add_run('reingreso por única vez a partir del periodo académico ')
        para.add_run(request.detail_cm['reing_per'])
        para.add_run('. Si el estudiante no renueva su matrícula en el semestre de reingreso, ')
        para.add_run('el acto académico expedido por el Consejo de Facultad queda ')
        para.add_run('sin efecto. (Resolución 012 de 2014 de Vicerrectoría Académica; Artículo ')
        para.add_run('46, Acuerdo 008 de 2008 del Consejo Superior Universitario).\n')
        bullet = para.add_run('1. Datos Generales')
        bullet.font.bold = True
        bullet.font.size = Pt(8)
        table = docx.add_table(rows=6, cols=3)
        table.style='Table Grid'
        table.style.font.size = Pt(8)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.columns[0].width = 400000
        table.columns[1].width = 2400000
        table.columns[2].width = 2400000
        cellp = table.cell(0, 0).merge(table.cell(0, 2)).paragraphs[0]
        cellp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cellp.add_run('REINGRESO\n').font.bold = True
        cellp.add_run('Normativa Asociada: Articulo 46 del Acuerdo 008 de 2008 del CSU')
        cellp.add_run(' y Resolución 012 de 2014 de VRA')
        table.cell(1, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(2, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(3, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(4, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(5, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(1, 0).paragraphs[0].add_run('1').font.bold = True
        table.cell(1, 1).paragraphs[0].add_run('Estudiante')
        table.cell(1, 2).paragraphs[0].add_run(request.student_name)
        table.cell(2, 0).paragraphs[0].add_run('2').font.bold = True
        table.cell(2, 1).paragraphs[0].add_run('DNI')
        table.cell(2, 2).paragraphs[0].add_run(request.student_dni)
        table.cell(3, 0).paragraphs[0].add_run('3').font.bold = True
        table.cell(3, 1).paragraphs[0].add_run('Plan de estudios')
        table.cell(3, 2).paragraphs[0].add_run(large_program)
        table.cell(4, 0).paragraphs[0].add_run('4').font.bold = True
        table.cell(4, 1).paragraphs[0].add_run('Código del plan de estudios')
        table.cell(4, 2).paragraphs[0].add_run(request.academic_program)
        table.cell(5, 0).paragraphs[0].add_run('5').font.bold = True
        table.cell(5, 1).paragraphs[0].add_run('Fecha solicitud')
        table.cell(5, 2).paragraphs[0].add_run(str(request.detail_cm['solic_date'].day) + REINPRE.num_to_month(request.detail_cm['solic_date'].month) + str(request.detail_cm['solic_date'].year))
        para = docx.add_paragraph()
        bullet = para.add_run('2. Información Académica')
        para.paragraph_format.space_after = Pt(0)
        bullet.font.bold = True
        bullet.font.size = Pt(8)
        table = docx.add_table(rows=13, cols=3)
        table.style='Table Grid'
        table.style.font.size = Pt(8)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.columns[0].width = 400000
        table.columns[1].width = 3200000
        table.columns[2].width = 1600000
        table.cell(0, 0).merge(table.cell(0, 1)).paragraphs[0].add_run('Periodo para el cual fue admitido en este plan de estudios')
        table.cell(0, 2).paragraphs[0].add_run(request.detail_cm['per_admi'])
        table.cell(0, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(1, 0).merge(table.cell(1, 1)).paragraphs[0].add_run('¿Se trata de un primer reingreso?')
        table.cell(1, 2).paragraphs[0].add_run(request.detail_cm['first_reingreso'])
        table.cell(1, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(2, 0).merge(table.cell(2, 2)).paragraphs[0].add_run('Si la respuesta es NO, el Comité Asesor no debe recomendar al Consejo de Facultad el reingreso')
        table.cell(3, 0).merge(table.cell(3, 1)).paragraphs[0].add_run('Es caso de ser primer reingreso en ¿qué periodo académico perdió la calidad de estudiante?')
        table.cell(3, 2).paragraphs[0].add_run(request.detail_cm['per_perd'])
        table.cell(3, 2).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        table.cell(3, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(4, 0).merge(table.cell(4, 1)).paragraphs[0].add_run('Al momento de presentar la solicitud ¿cuántos periodos académicos (incluido el periodo académico en que presentó la solicitud) han transcurridos a partir del periodo académico en que registró su última matrícula?')
        table.cell(4, 2).paragraphs[0].add_run(request.detail_cm['per_transc'])
        table.cell(4, 2).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        table.cell(4, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(5, 0).merge(table.cell(5, 2)).paragraphs[0].add_run('En caso que la respuesta sea mayor de 6 periodos académicos no se debe recomendar el reingreso')
        table.cell(6, 0).merge(table.cell(6, 1)).paragraphs[0].add_run('P.A.P.A.')
        table.cell(6, 2).paragraphs[0].add_run(request.detail_cm['PAPA'])
        table.cell(6, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(7, 0).merge(table.cell(7, 1)).paragraphs[0].add_run('Causa de la pérdida de la calidad de estudiante')
        table.cell(7, 0).merge(table.cell(7, 1)).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        table.cell(7, 2).paragraphs[0].add_run(request.detail_cm['causa_perd'])
        table.cell(7, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(8, 0).merge(table.cell(8, 2)).paragraphs[0].add_run('Estudio de créditos').font.bold = True
        table.cell(9, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(9, 0).paragraphs[0].add_run('1').font.bold = True
        table.cell(10, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(10, 0).paragraphs[0].add_run('2').font.bold = True
        table.cell(11, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(11, 0).paragraphs[0].add_run('3').font.bold = True
        table.cell(12, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(12, 0).paragraphs[0].add_run('4').font.bold = True
        table.cell(9, 1).paragraphs[0].add_run('Cupo de créditos menos créditos pendientes')
        table.cell(10, 1).paragraphs[0].add_run('Créditos pendientes por ser aprobados del plan de estudios')
        table.cell(11, 1).paragraphs[0].add_run('Créditos pendientes por ser aprobados de nivelación – Inglés')
        table.cell(12, 1).paragraphs[0].add_run('¿Cuántos créditos adicionales requiere para inscribir asignaturas?')
        table.cell(9, 2).paragraphs[0].add_run(request.detail_cm['creds_minus_remaining'])
        table.cell(9, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(10, 2).paragraphs[0].add_run(request.detail_cm['creds_remaining'])
        table.cell(10, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(11, 2).paragraphs[0].add_run(request.detail_cm['creds_ingl'])
        table.cell(11, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(12, 2).paragraphs[0].add_run(request.detail_cm['creds_add'])
        table.cell(12, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        para = docx.add_paragraph()
        para.paragraph_format.space_after = Pt(0)
        table = docx.add_table(rows=5, cols=2)
        for col in table.columns:
            for cell in col.cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.style = 'Table Grid'
        table.style.font.size = Pt(8)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.columns[0].width = 3100000
        table.columns[1].width = 2100000
        table.cell(0, 0).merge(table.cell(0, 1)).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        table.cell(0, 0).merge(table.cell(0, 1)).paragraphs[0].add_run('Al finalizar el semestre de reingreso para mantener la calidad de estudiante, deberá obtener un Promedio Semestral mínimo de:')
        table.cell(1, 0).paragraphs[0].add_run('Si inscribe 12 Créditos')
        table.cell(2, 0).paragraphs[0].add_run('Si inscribe 15 Créditos')
        table.cell(3, 0).paragraphs[0].add_run('Si inscribe 18 Créditos')
        table.cell(4, 0).paragraphs[0].add_run('Si inscribe 21 Créditos')
        table.cell(1, 1).paragraphs[0].add_run(request.detail_cm['12c'])
        table.cell(2, 1).paragraphs[0].add_run(request.detail_cm['15c'])
        table.cell(3, 1).paragraphs[0].add_run(request.detail_cm['18c'])
        table.cell(4, 1).paragraphs[0].add_run(request.detail_cm['21c'])
        para = docx.add_paragraph()
        bullet = para.add_run('3. Resumen general de créditos del plan de estudios')
        para.paragraph_format.space_after = Pt(0)
        bullet.font.bold = True
        bullet.font.size = Pt(8)
        table = docx.add_table(rows=5, cols=7)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for col in table.columns:
            for cell in col.cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.style = 'Table Grid'
        table.style.font.size = Pt(8)
        table.columns[0].width = 1610000
        table.columns[1].width = 690000
        table.columns[2].width = 610000
        table.columns[3].width = 690000
        table.columns[4].width = 610000
        table.columns[5].width = 675000
        table.columns[6].width = 375000
        table.cell(0, 0).merge(table.cell(1, 0)).paragraphs[0].add_run('Créditos')
        table.cell(0, 0).merge(table.cell(1, 0)).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        table.cell(0, 1).merge(table.cell(0, 2)).paragraphs[0].add_run('Fundamentación (B)')
        table.cell(0, 3).merge(table.cell(0, 4)).paragraphs[0].add_run('Disciplinar (C)')
        table.cell(0, 5).merge(table.cell(1, 5)).paragraphs[0].add_run('Libre Elección (L)')
        table.cell(0, 6).merge(table.cell(1, 6)).paragraphs[0].add_run('Total')
        table.cell(0, 6).merge(table.cell(1, 6)).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        table.cell(2, 0).paragraphs[0].add_run('Exigidos*')
        table.cell(3, 0).paragraphs[0].add_run('Aprobados del plan de estudios')
        table.cell(4, 0).paragraphs[0].add_run('Pendientes')
        table.cell(1, 1).paragraphs[0].add_run('Obligatorios')
        table.cell(1, 2).paragraphs[0].add_run('Optativos')
        table.cell(1, 3).paragraphs[0].add_run('Obligatorios')
        table.cell(1, 4).paragraphs[0].add_run('Optativos')
        table.cell(2, 1).paragraphs[0].add_run(request.detail_cm['summary']['exiged']['fund_m'])
        table.cell(3, 1).paragraphs[0].add_run(request.detail_cm['summary']['approved']['fund_m'])
        table.cell(4, 1).paragraphs[0].add_run(request.detail_cm['summary']['remaining']['fund_m'])
        table.cell(2, 2).paragraphs[0].add_run(request.detail_cm['summary']['exiged']['fund_o'])
        table.cell(3, 2).paragraphs[0].add_run(request.detail_cm['summary']['approved']['fund_o'])
        table.cell(4, 2).paragraphs[0].add_run(request.detail_cm['summary']['remaining']['fund_o'])
        table.cell(2, 3).paragraphs[0].add_run(request.detail_cm['summary']['exiged']['disc_m'])
        table.cell(3, 3).paragraphs[0].add_run(request.detail_cm['summary']['approved']['disc_m'])
        table.cell(4, 3).paragraphs[0].add_run(request.detail_cm['summary']['remaining']['disc_m'])
        table.cell(2, 4).paragraphs[0].add_run(request.detail_cm['summary']['exiged']['disc_o'])
        table.cell(3, 4).paragraphs[0].add_run(request.detail_cm['summary']['approved']['disc_o'])
        table.cell(4, 4).paragraphs[0].add_run(request.detail_cm['summary']['remaining']['disc_o'])
        table.cell(2, 5).paragraphs[0].add_run(request.detail_cm['summary']['exiged']['free'])
        table.cell(3, 5).paragraphs[0].add_run(request.detail_cm['summary']['approved']['free'])
        table.cell(4, 5).paragraphs[0].add_run(request.detail_cm['summary']['remaining']['free'])
        exiged = int(request.detail_cm['summary']['exiged']['fund_m']) + int(request.detail_cm['summary']['exiged']['fund_o']) + int(request.detail_cm['summary']['exiged']['disc_m']) + int(request.detail_cm['summary']['exiged']['disc_o']) + int(request.detail_cm['summary']['exiged']['free'])
        approved = int(request.detail_cm['summary']['approved']['fund_m']) + int(request.detail_cm['summary']['approved']['fund_o']) + int(request.detail_cm['summary']['approved']['disc_m']) + int(request.detail_cm['summary']['approved']['disc_o']) + int(request.detail_cm['summary']['approved']['free'])
        remaining = int(request.detail_cm['summary']['remaining']['fund_m']) + int(request.detail_cm['summary']['remaining']['fund_o']) + int(request.detail_cm['summary']['remaining']['disc_m']) + int(request.detail_cm['summary']['remaining']['disc_o']) + int(request.detail_cm['summary']['remaining']['free'])
        table.cell(2, 6).paragraphs[0].add_run(str(exiged))
        table.cell(3, 6).paragraphs[0].add_run(str(approved))
        table.cell(4, 6).paragraphs[0].add_run(str(remaining))
        para = docx.add_paragraph()
        para.paragraph_format.space_after = Pt(0)
        para.add_run('     *Sin incluir los créditos correspondientes al cumplimiento del requisito de suficiencia en idioma.').font.size = Pt(8)
        table = docx.add_table(rows=1, cols=5)
        table.style='Table Grid'
        table.style.font.size = Pt(8)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.columns[0].width = 3000000
        table.columns[1].width = 800000
        table.columns[2].width = 300000
        table.columns[3].width = 800000
        table.columns[4].width = 300000
        table.cell(0, 0).paragraphs[0].add_run('El Comité Asesor de ' + request.detail_cm['commite_name'] + ' en sesión del día ')
        table.cell(0, 0).paragraphs[0].add_run(str(request.detail_cm['comite_date'].day) + REINPRE.num_to_month(request.detail_cm['comite_date'].month) + str(request.detail_cm['comite_date'].year))
        table.cell(0, 0).paragraphs[0].add_run('. Acta ' + request.detail_cm['comite_acta'] + '.')
        table.cell(0, 1).paragraphs[0].add_run('Recomienda')
        table.cell(0, 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        table.cell(0, 3).paragraphs[0].add_run('No Recomienda')
        table.cell(0, 3).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if request.detail_cm['reccomend'] == 'true':
            table.cell(0, 2).paragraphs[0].add_run('X')
        else:
            table.cell(0, 4).paragraphs[0].add_run('X')
        table.cell(0, 2).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        table.cell(0, 4).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        