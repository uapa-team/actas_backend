from docx.enum.text import WD_ALIGN_PARAGRAPH


class HOIDPRE():

    @staticmethod
    def case_HOMOLOGACION_INGLES_PREGRADO(request, docx, redirected=False):
        if redirected:
            para = docx.paragraphs[-1]
        else:
            para = docx.add_paragraph()
            para.add_run('El Consejo de Facultad ')
        para.add_run('APRUEBA').font.bold = True
        para.add_run(' homologar en el periodo académico ' +
                     request.academic_period)
        para.add_run(
            ', el requisito de idioma inglés por obtener una calificación de ')
        para.add_run(request.detail_cm['min_grade'] +
                     ' en el examen ' + request.detail_cm['institution'])
        para.add_run(
            ', siendo ' + request.detail_cm['grade_got']+' el mínimo exigido.')

        table = docx.add_table(
            rows=len(request.detail_cm['subjects'])+5, cols=7, style='Table Grid')

        cellp = table.cell(0, 0).merge(table.cell(0, 6)).paragraphs[0]
        cellp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        cellp.add_run(request.student_name + '\t\tDNI. ' +
                      request.student_dni).font.bold = True

        cellp = table.cell(1, 0).merge(table.cell(1, 4)).paragraphs[0]
        cellp.add_run(
            'Asignaturas a homologar en el plan de estudios ' + 'Ing. Sistemas')

        cellp = table.cell(1, 5).merge(table.cell(2, 5)).paragraphs[0]
        cellp.add_run('Examen de inglés presentado')

        cellp = table.cell(1, 6).merge(table.cell(2, 6)).paragraphs[0]
        cellp.add_run('Nota')

        cellp = table.cell(3, 5).merge(table.cell(
            len(request.detail_cm['subjects'])+3, 5)).paragraphs[0]
        cellp.add_run(request.detail_cm['institution'])

        cellp = table.cell(3, 6).merge(table.cell(
            len(request.detail_cm['subjects'])+3, 6)).paragraphs[0]
        cellp.add_run(request.detail_cm['grade_got'])

        table.cell(2, 0).paragraphs[0].add_run('Código')
        table.cell(2, 1).paragraphs[0].add_run('Asignatura')
        table.cell(2, 2).paragraphs[0].add_run('C')
        table.cell(2, 3).paragraphs[0].add_run('T')
        table.cell(2, 4).paragraphs[0].add_run('Nota')

        index = 0
        credits_sum = 0
        for subject in request.detail_cm['subjects']:
            credits_sum = credits_sum + int(subject['credits'])
            table.cell(index+3, 0).paragraphs[0].add_run(subject['code'])
            table.cell(index+3, 1).paragraphs[0].add_run(subject['subject'])
            table.cell(index+3, 2).paragraphs[0].add_run(subject['credits'])
            table.cell(index+3, 3).paragraphs[0].add_run(subject['tipology'])
            table.cell(index+3, 4).paragraphs[0].add_run(subject['grade'])
            index = index + 1

        table.cell(index+3, 1).paragraphs[0].add_run('Créditos homologados P')
        table.cell(index+3, 2).paragraphs[0].add_run(str(credits_sum))
        table.cell(
            index+4, 1).paragraphs[0].add_run('Total créditos que se homologan')
        table.cell(index+4, 2).paragraphs[0].add_run(str(credits_sum))
