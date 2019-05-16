from docx import Document
from ..models import Program

class HCEMPRE():

    def __init__(self):
        self.subtypes = {
            'HING' : self.subtype_HOMOLOGACION_INGLES,
            'HNOR' : self.subtype_HOMOLOGACIONES_NORMALES,
            'EQUI' : self.subtype_EQUIVALENCIAS,
            'CONV' : self.subtype_CONVALIDACION
        }

    def case_HOMOLOGACION_CONVALIDACION_EQUIVALENCIA_PREGRADO(self, request, docx):
        return self.subtypes[request.detail_cm['subtype']](request, docx)

    def subtype_HOMOLOGACION_INGLES(self, request, docx):
        curr_para = docx.add_paragraph()
        curr_para.add_run('El Consejo de Facultad ')
        curr_para.add_run('APRUEBA').font.bold = True
        curr_para.add_run(' homologar en el periodo académico ')
        curr_para.add_run(request.detail_cm['academic_period'])
        curr_para.add_run(', el requisito de inioma inglés por obtener un puntaje de ')
        curr_para.add_run(request.detail_cm['minimal_calification'])
        curr_para.add_run(' en el examen ')
        curr_para.add_run(request.detail_cm['institution'])
        curr_para.add_run(', siendo ')
        curr_para.add_run(request.detail_cm['grade_got'])
        curr_para.add_run(' el mínimo exigido.')
        
        curr_tabl = docx.add_table(rows=len(request.detail_cm['subjects']), cols=7, style='Table Grid')
        curr_cell_1 = curr_tabl.cell(0, 0)
        curr_cell_2 = curr_tabl.cell(0, 6)
        curr_cell = curr_cell_1.merge(curr_cell_2)
        curr_cell.add_paragraph(request.student_name + '\t' + request.student_dni)
        curr_cell_1 = curr_tabl.cell(1, 0)
        curr_cell_2 = curr_tabl.cell(1, 5)
        curr_cell = curr_cell_1.merge(curr_cell_2)
        curr_cell.add_paragraph('Asignaturas a homologar en el plan de estudios ' + 'hola')

    def subtype_HOMOLOGACIONES_NORMALES(self, request, docx):
        raise NotImplementedError

    def subtype_EQUIVALENCIAS(self, request, docx):
        raise NotImplementedError

    def subtype_CONVALIDACION(self, request, docx):
        raise NotImplementedError